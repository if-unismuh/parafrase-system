# enhanced_turnitin_gemini.py
"""
Enhanced Turnitin-Aware Paraphraser with Advanced Gemini Integration
Multi-level AI processing for maximum Turnitin avoidance
Author: DevNoLife
Version: 3.1
"""

import os
import json
import re
import time
import random
from datetime import datetime
from pathlib import Path
import docx
from docx.shared import RGBColor
from docx.enum.text import WD_COLOR_INDEX
import shutil
from collections import defaultdict
import hashlib

# Load environment variables from .env file
try:
    from dotenv import load_dotenv
    load_dotenv()
    print("✅ Environment variables loaded from .env file")
except ImportError:
    print("⚠️  python-dotenv not found. Install with: pip install python-dotenv")
    print("💡 Attempting to load from system environment variables...")

# Updated Gemini AI imports
try:
    from google import genai
    from google.genai import types
    GEMINI_AVAILABLE = True
except ImportError:
    print("⚠️  google-genai package not found. Install with: pip install google-genai")
    GEMINI_AVAILABLE = False

# Import sistem parafrase lokal
from main_paraphrase_system import IndonesianParaphraseSystem

class EnhancedTurnitinGeminiParaphraser:
    def __init__(self, synonym_file='sinonim.json', gemini_api_key=None):
        print("🚀 Initializing Enhanced Turnitin-Aware + Gemini Paraphraser...")
        
        # Initialize local paraphrase system
        self.local_paraphraser = IndonesianParaphraseSystem(synonym_file)
        
        # Enhanced Gemini AI configuration with new API
        # Get API key from parameter, environment, or fallback
        self.gemini_api_key = gemini_api_key or os.environ.get("GEMINI_API_KEY")
        self.gemini_client = None
        
        if self.gemini_api_key and GEMINI_AVAILABLE:
            try:
                # Configure the client properly with explicit API key
                os.environ["GEMINI_API_KEY"] = self.gemini_api_key
                self.gemini_client = genai.Client(api_key=self.gemini_api_key)
                # Verify the configuration worked
                print(f"🔑 API Key configured: {self.gemini_api_key[:10]}...")
                print("✅ Gemini 2.0 Flash configured with enhanced Turnitin integration")
            except Exception as e:
                print(f"⚠️  Error initializing Gemini client: {e}")
                print(f"🔍 API Key length: {len(self.gemini_api_key) if self.gemini_api_key else 'None'}")
                self.gemini_client = None
        elif not self.gemini_api_key:
            print("⚠️  No Gemini API key provided")
        elif not GEMINI_AVAILABLE:
            print("⚠️  Google GenAI package not available")
        
        # Enhanced cost tracking
        self.cost_tracker = {
            'local_calls': 0,
            'ai_calls_level1': 0,    # Basic AI paraphrasing
            'ai_calls_level2': 0,    # Advanced pattern targeting
            'ai_calls_level3': 0,    # Deep restructuring
            'ai_tokens_used': 0,
            'estimated_cost_usd': 0.0,
            'turnitin_patterns_detected': 0,
            'high_risk_patterns_fixed': 0,
            'ai_vs_local_wins': {'ai': 0, 'local': 0}
        }
        
        # AI result cache with different levels
        self.ai_cache = {
            'level1': {},  # Basic paraphrasing
            'level2': {},  # Pattern-specific
            'level3': {}   # Deep restructuring
        }
        
        # Enhanced Turnitin-specific configuration
        self.config = {
            'turnitin_threshold': 0.7,           # Lower threshold for more AI usage
            'academic_template_threshold': 0.85,  
            'technical_definition_threshold': 0.8,
            'ai_level1_threshold': 0.4,          # Basic AI usage
            'ai_level2_threshold': 0.6,          # Advanced AI usage  
            'ai_level3_threshold': 0.8,          # Deep AI restructuring
            'model_name': 'gemini-2.0-flash',
            'max_batch_size': 1,                 # Individual processing for quality
            'ai_retry_attempts': 3,
            'multi_level_processing': True       # Enable multi-level AI
        }
        
        # EXPANDED: More comprehensive Turnitin patterns from your PDF
        self.turnitin_detected_patterns = {
            # Academic boilerplate (RED ZONES in Turnitin)
            'academic_templates': [
                r'berdasarkan latar belakang yang telah diuraikan',
                r'maka rumusan masalah dalam penelitian ini adalah sebagai berikut',
                r'penelitian ini bertujuan untuk',
                r'menambah wawasan dan pemahaman penulis mengenai',
                r'menjadi referensi bagi peneliti.*akademisi.*pengembang',
                r'secara garis besar penulisan.*laporan.*terbagi menjadi',
                r'bab ini menerangkan secara singkat dan jelas mengenai',
                r'pada bab ini membahas tentang teori-teori yang melandasi',
                r'membahas tentang metode penelitian dan alat yang digunakan',
                r'fokus penelitian adalah pada',
                r'penelitian ini dibatasi pada',
                r'tidak termasuk dalam ruang lingkup penelitian'
            ],
            
            # Technical definitions (HEAVILY FLAGGED in your PDF)
            'technical_definitions': [
                r'least significant bit.*lsb.*adalah teknik steganografi yang umum digunakan',
                r'cara kerja metode lsb yaitu mengubah bit terakhir',
                r'bit yang paling tidak signifikan.*dari data piksel citra penampung',
                r'cover image.*yang tidak berpengaruh signifikan dengan bit',
                r'metode ini banyak digunakan dalam invisible watermarking',
                r'implementasinya yang relatif mudah dan cepat',
                r'kemampuan menyisipkan data dalam jumlah cukup besar',
                r'quick response.*qr.*code telah diadopsi secara luas',
                r'fleksibilitas dan kapasitas penyimpanan qr code',
                r'steganografi dan watermarking bertujuan menyembunyikan informasi',
                r'invisible watermarking menawarkan solusi dengan menyisipkan data'
            ],
            
            # Research methodology patterns
            'methodology_patterns': [
                r'penelitian ini menggunakan metode',
                r'teknik pengumpulan data yang digunakan',
                r'populasi dalam penelitian ini adalah',
                r'sampel penelitian ini berjumlah',
                r'instrumen penelitian yang digunakan'
            ],
            
            # Citation patterns (common academic phrases)
            'citation_patterns': [
                r'menurut.*\(\d{4}\)',
                r'berdasarkan.*\(\d{4}\)',
                r'penelitian.*et al.*\(\d{4}\)',
                r'hasil penelitian menunjukkan',
                r'studi.*menunjukkan bahwa'
            ],
            
            # Domain-specific terminology (from your document)
            'domain_terms': [
                'steganografi', 'watermarking', 'least significant bit', 'lsb',
                'cover image', 'citra penampung', 'invisible watermarking', 
                'qr code', 'quick response', 'imperceptibility', 'ketidakterlihatan',
                'bit terakhir', 'piksel citra', 'data rahasia', 'penyembunyian informasi'
            ]
        }
        
        # ADVANCED: Multi-level paraphrase strategies
        self.gemini_prompts = {
            'level1_basic': """Kamu adalah ahli paraphrasing akademik Indonesia. Parafrase teks berikut dengan natural tapi tetap formal:

TEKS: {text}

INSTRUKSI:
✅ Ganti dengan sinonim yang tepat
✅ Ubah struktur kalimat 
✅ Pertahankan makna akademik
✅ Hasil harus natural dan mudah dibaca

PARAFRASE:""",

            'level2_pattern_specific': """Kamu adalah expert anti-plagiarisme yang paham sistem deteksi Turnitin. 

DETEKSI: Teks ini mengandung pattern berisiko tinggi yang sering di-flag Turnitin:
PATTERN: {detected_patterns}

TEKS BERISIKO: {text}

MISI: Restructure teks ini agar menghindari deteksi pattern recognition Turnitin:
✅ UBAH STRUKTUR academic template yang standar
✅ VARIASIKAN technical definition dengan approach berbeda  
✅ PAKAI terminology alternatif untuk domain terms
✅ PECAH atau GABUNG kalimat untuk mengubah pattern

TARGET: Similarity reduction 40-60%
PARAFRASE ANTI-TURNITIN:""",

            'level3_deep_restructuring': """Kamu adalah master paraphrasing yang sangat expert dalam menghindari detection system plagiarisme tingkat enterprise seperti Turnitin.

ANALISIS RISIKO: {risk_analysis}
PATTERN TERDETEKSI: {detected_patterns}
SIMILARITY RISK: {risk_level}%

TEKS HIGH-RISK: {text}

MISI CRITICAL: Deep restructuring untuk completely bypass Turnitin detection:

STRATEGI ADVANCED:
🔥 COMPLETE sentence restructuring (aktif↔pasif, inverse order)
🔥 ADVANCED terminology substitution (gunakan sinonim yang jarang dipakai)
🔥 BREAKDOWN complex sentences atau COMBINE simple ones
🔥 REORDER information flow tanpa ubah makna
🔥 ELIMINATE academic boilerplate templates
🔥 TRANSFORM technical definitions jadi explanatory approach

PANTANGAN CRITICAL:
❌ Jangan pertahankan original sentence structure
❌ Jangan gunakan common academic phrases  
❌ Jangan biarkan technical terms in original form

TARGET AGGRESSIVE: Turunkan similarity 60-80%
HASIL DEEP RESTRUCTURING:"""
        }
        
        # Enhanced academic replacements
        self.enhanced_academic_replacements = {
            'berdasarkan latar belakang yang telah diuraikan': [
                'mengacu pada konteks yang telah dijelaskan sebelumnya',
                'merujuk pada uraian kondisi yang telah dipaparkan',
                'berlandaskan pemaparan situasi yang dikemukakan di atas',
                'sesuai dengan gambaran kondisi yang telah disampaikan'
            ],
            'penelitian ini bertujuan untuk': [
                'kajian ini dimaksudkan untuk', 'studi ini berupaya untuk',
                'riset ini dirancang untuk', 'investigasi ini ditujukan untuk',
                'eksplorasi ini bermaksud untuk'
            ],
            'cara kerja metode lsb yaitu': [
                'prinsip operasional teknik LSB melibatkan',
                'mekanisme kerja pendekatan LSB adalah',
                'sistem operasi metode LSB berfungsi dengan',
                'prosedur teknik LSB bekerja melalui'
            ],
            'menambah wawasan dan pemahaman penulis': [
                'memperluas pengetahuan dan insight peneliti',
                'mengembangkan kompetensi dan perspektif penulis', 
                'meningkatkan kapasitas dan pemahaman mendalam penulis',
                'memperkaya khazanah pengetahuan dan analisis penulis'
            ]
        }
        
        print(f"✅ Local system ready with {len(self.local_paraphraser.synonyms)} synonyms")
        print(f"🎯 Enhanced Turnitin patterns: {sum(len(v) for v in self.turnitin_detected_patterns.values())} patterns")
        print(f"🤖 Multi-level AI: {'3 Levels Enabled' if self.gemini_client else 'Disabled'}")
        print(f"💎 Enhanced academic replacements: {len(self.enhanced_academic_replacements)}")
        print("✅ Enhanced Turnitin-Gemini Paraphraser ready!")
    
    def detect_enhanced_turnitin_risk(self, text):
        """Enhanced risk detection with more granular analysis"""
        risk_score = 0.0
        detected_patterns = []
        pattern_details = {}
        
        text_lower = text.lower()
        
        # Academic templates (highest risk - from your PDF)
        academic_matches = 0
        for pattern in self.turnitin_detected_patterns['academic_templates']:
            if re.search(pattern, text_lower):
                risk_score += 0.35  # Higher weight
                detected_patterns.append(('academic_template', pattern))
                academic_matches += 1
        
        # Technical definitions (very high risk - heavily flagged in your PDF)
        technical_matches = 0
        for pattern in self.turnitin_detected_patterns['technical_definitions']:
            if re.search(pattern, text_lower):
                risk_score += 0.3
                detected_patterns.append(('technical_definition', pattern))
                technical_matches += 1
        
        # Citation patterns
        citation_matches = 0
        for pattern in self.turnitin_detected_patterns['citation_patterns']:
            if re.search(pattern, text_lower):
                risk_score += 0.2
                detected_patterns.append(('citation', pattern))
                citation_matches += 1
        
        # Methodology patterns
        methodology_matches = 0
        for pattern in self.turnitin_detected_patterns['methodology_patterns']:
            if re.search(pattern, text_lower):
                risk_score += 0.25
                detected_patterns.append(('methodology', pattern))
                methodology_matches += 1
        
        # Domain terms density
        domain_count = sum(1 for term in self.turnitin_detected_patterns['domain_terms'] 
                          if term in text_lower)
        if domain_count > 0:
            domain_risk = min(domain_count * 0.08, 0.4)
            risk_score += domain_risk
            detected_patterns.append(('domain_terms', f"{domain_count} specialized terms"))
        
        # Normalize and categorize
        risk_level = min(risk_score, 1.0)
        
        # Determine AI processing level needed
        ai_level_needed = 0
        if risk_level >= self.config['ai_level3_threshold']:
            ai_level_needed = 3  # Deep restructuring
        elif risk_level >= self.config['ai_level2_threshold']:
            ai_level_needed = 2  # Pattern-specific
        elif risk_level >= self.config['ai_level1_threshold']:
            ai_level_needed = 1  # Basic AI
        
        return {
            'risk_level': risk_level,
            'risk_category': self.get_enhanced_risk_category(risk_level),
            'detected_patterns': detected_patterns,
            'pattern_counts': {
                'academic': academic_matches,
                'technical': technical_matches,
                'citation': citation_matches,
                'methodology': methodology_matches,
                'domain_terms': domain_count
            },
            'ai_level_needed': ai_level_needed,
            'needs_ai': ai_level_needed > 0
        }
    
    def get_enhanced_risk_category(self, risk_level):
        """Enhanced risk categorization"""
        if risk_level >= 0.9:
            return "🔴 CRITICAL - Multiple Turnitin flags detected"
        elif risk_level >= 0.7:
            return "🔴 VERY HIGH - Turnitin will definitely flag this"
        elif risk_level >= 0.5:
            return "🟠 HIGH - Very likely to be flagged"
        elif risk_level >= 0.3:
            return "🟡 MEDIUM - May be flagged"
        elif risk_level >= 0.1:
            return "🟢 LOW - Probably safe"
        else:
            return "✅ VERY LOW - Safe"
    
    def call_gemini_multi_level(self, text, risk_analysis, ai_level):
        """Multi-level Gemini AI processing"""
        if not self.gemini_client:
            return None
        
        if not self.gemini_api_key:
            print("❌ No API key available")
            return None
        
        try:
            # Select appropriate prompt based on AI level
            if ai_level == 1:
                prompt = self.gemini_prompts['level1_basic'].format(text=text)
                cache_key = f"l1_{hashlib.md5(text.encode()).hexdigest()[:12]}"
                cache_level = 'level1'
                self.cost_tracker['ai_calls_level1'] += 1
            elif ai_level == 2:
                detected_patterns_str = ', '.join([p[1][:50] for p in risk_analysis['detected_patterns'][:3]])
                prompt = self.gemini_prompts['level2_pattern_specific'].format(
                    text=text,
                    detected_patterns=detected_patterns_str
                )
                cache_key = f"l2_{hashlib.md5(text.encode()).hexdigest()[:12]}"
                cache_level = 'level2'
                self.cost_tracker['ai_calls_level2'] += 1
            else:  # ai_level == 3
                detected_patterns_str = '\n'.join([f"- {p[1]}" for p in risk_analysis['detected_patterns'][:5]])
                prompt = self.gemini_prompts['level3_deep_restructuring'].format(
                    text=text,
                    risk_analysis=risk_analysis['risk_category'],
                    detected_patterns=detected_patterns_str,
                    risk_level=int(risk_analysis['risk_level'] * 100)
                )
                cache_key = f"l3_{hashlib.md5(text.encode()).hexdigest()[:12]}"
                cache_level = 'level3'
                self.cost_tracker['ai_calls_level3'] += 1
            
            # Check cache first
            if cache_key in self.ai_cache[cache_level]:
                return self.ai_cache[cache_level][cache_key]
            
            # Configure generation based on AI level
            temperature = 0.3 + (ai_level * 0.1)  # Higher creativity for higher levels
            max_tokens = 2048 + (ai_level * 1024)  # More tokens for complex restructuring
            
            contents = [
                types.Content(
                    role="user",
                    parts=[types.Part.from_text(text=prompt)],
                ),
            ]
            
            generate_content_config = types.GenerateContentConfig(
                temperature=temperature,
                top_k=40 + (ai_level * 20),
                top_p=0.9,
                max_output_tokens=max_tokens,
                candidate_count=1,
            )
            
            for attempt in range(self.config['ai_retry_attempts']):
                try:
                    response_text = ""
                    for chunk in self.gemini_client.models.generate_content_stream(
                        model=self.config['model_name'],
                        contents=contents,
                        config=generate_content_config,
                    ):
                        if hasattr(chunk, 'text') and chunk.text:
                            response_text += chunk.text
                    
                    if response_text.strip():
                        # Clean response
                        cleaned_response = self.clean_ai_response(response_text)
                        
                        # Track usage
                        tokens_used = len(prompt + response_text) // 4
                        self.cost_tracker['ai_tokens_used'] += tokens_used
                        self.cost_tracker['estimated_cost_usd'] += tokens_used * 0.000001
                        
                        # Cache result
                        self.ai_cache[cache_level][cache_key] = cleaned_response
                        
                        return cleaned_response
                    
                except Exception as e:
                    error_msg = str(e)
                    print(f"⚠️  Level {ai_level} AI attempt {attempt + 1} failed: {e}")
                    
                    # Special handling for API key errors
                    if "API_KEY_INVALID" in error_msg or "API Key not found" in error_msg:
                        print("🔍 API Key issue detected. Trying to reconfigure...")
                        try:
                            os.environ["GEMINI_API_KEY"] = self.gemini_api_key
                            self.gemini_client = genai.Client(api_key=self.gemini_api_key)
                            print("✅ API key reconfigured")
                        except Exception as reconfig_error:
                            print(f"❌ Failed to reconfigure API key: {reconfig_error}")
                            return None
                    
                    if attempt < self.config['ai_retry_attempts'] - 1:
                        time.sleep(2 ** attempt)
            
            return None
            
        except Exception as e:
            print(f"❌ Multi-level AI error: {e}")
            return None
    
    def clean_ai_response(self, response_text):
        """Clean and normalize AI response"""
        # Remove common prefixes/suffixes
        cleaned = response_text.strip()
        
        # Remove markdown formatting
        cleaned = re.sub(r'\*\*(.*?)\*\*', r'\1', cleaned)
        cleaned = re.sub(r'\*(.*?)\*', r'\1', cleaned)
        
        # Remove any instructional text
        lines = cleaned.split('\n')
        content_lines = []
        for line in lines:
            line = line.strip()
            if (not line.startswith('PARAFRASE') and 
                not line.startswith('HASIL') and
                not line.startswith('TARGET') and
                line and
                not line.isupper()):
                content_lines.append(line)
        
        result = ' '.join(content_lines)
        
        # Final cleanup
        result = re.sub(r'\s+', ' ', result).strip()
        
        return result
    
    def process_paragraph_enhanced_turnitin(self, paragraph_text, paragraph_index):
        """Enhanced processing with multi-level AI"""
        
        # Step 1: Enhanced risk analysis
        risk_analysis = self.detect_enhanced_turnitin_risk(paragraph_text)
        self.cost_tracker['turnitin_patterns_detected'] += len(risk_analysis['detected_patterns'])
        
        # Step 2: Apply local enhanced replacements first
        local_modified = paragraph_text
        for original, alternatives in self.enhanced_academic_replacements.items():
            if original in local_modified.lower():
                replacement = random.choice(alternatives)
                pattern = re.compile(re.escape(original), re.IGNORECASE)
                local_modified = pattern.sub(replacement, local_modified, count=1)
        
        # Step 3: Get local system result
        local_result = self.local_paraphraser.generate_paraphrase(local_modified, aggressiveness=0.8)
        local_result['method'] = 'enhanced_local'
        local_result['risk_analysis'] = risk_analysis
        self.cost_tracker['local_calls'] += 1
        
        # Step 4: Multi-level AI processing based on risk
        if risk_analysis['needs_ai'] and self.gemini_client:
            ai_level = risk_analysis['ai_level_needed']
            
            print(f"    🤖 Level {ai_level} AI processing: {risk_analysis['risk_category']}")
            
            ai_result_text = self.call_gemini_multi_level(paragraph_text, risk_analysis, ai_level)
            
            if ai_result_text:
                # Calculate AI result metrics
                ai_similarity = self.local_paraphraser.calculate_similarity(paragraph_text, ai_result_text)
                
                ai_result = {
                    'paraphrase': ai_result_text,
                    'similarity': round(ai_similarity, 2),
                    'plagiarism_reduction': round(100 - ai_similarity, 2),
                    'changes_made': 1,
                    'method': f'enhanced_ai_level_{ai_level}',
                    'status': self.local_paraphraser._get_plagiarism_status(ai_similarity),
                    'risk_analysis': risk_analysis,
                    'ai_level_used': ai_level
                }
                
                # Compare AI vs Local
                if ai_result['plagiarism_reduction'] > local_result['plagiarism_reduction']:
                    self.cost_tracker['high_risk_patterns_fixed'] += 1
                    self.cost_tracker['ai_vs_local_wins']['ai'] += 1
                    print(f"      🏆 AI Level {ai_level} wins: {ai_result['plagiarism_reduction']:.1f}% vs {local_result['plagiarism_reduction']:.1f}%")
                    return ai_result
                else:
                    self.cost_tracker['ai_vs_local_wins']['local'] += 1
                    print(f"      🏆 Local wins: {local_result['plagiarism_reduction']:.1f}% vs {ai_result['plagiarism_reduction']:.1f}%")
        
        return local_result
    
    def is_section_header(self, text):
        """Check if text is a section header"""
        # Section headers are typically short and may contain numbers or keywords
        if len(text.split()) <= 6:
            # Common section patterns
            section_patterns = [
                r'bab\s+\d+', r'chapter\s+\d+', r'\d+\.\s*\w+',
                r'pendahuluan', r'tinjauan pustaka', r'metodologi',
                r'hasil dan pembahasan', r'kesimpulan', r'daftar pustaka',
                r'abstract', r'abstrak', r'kata pengantar'
            ]
            
            text_lower = text.lower().strip()
            for pattern in section_patterns:
                if re.search(pattern, text_lower):
                    return True
        
        return False
    
    def is_paragraph_suitable_for_paraphrasing(self, text):
        """Check if paragraph is suitable for paraphrasing"""
        # Filter out unsuitable paragraphs
        if len(text.split()) < 5:  # Too short
            return False
        
        # Skip citations, references, tables, figures
        unsuitable_patterns = [
            r'^\s*\d+\.\s*$',  # Just numbers
            r'^\s*[a-z]\.\s*$',  # Just letters
            r'gambar\s+\d+', r'tabel\s+\d+', r'figure\s+\d+',
            r'^\s*\([^)]+\)\s*$',  # Just citations
            r'^\s*sumber\s*:', r'^\s*source\s*:',
            r'^\s*\w+\s*:\s*\w+\s*$'  # Simple key-value pairs
        ]
        
        text_lower = text.lower().strip()
        for pattern in unsuitable_patterns:
            if re.search(pattern, text_lower):
                return False
        
        return True
    
    def print_enhanced_summary(self, total_stats, processing_time, backup_folder):
        """Print comprehensive processing summary"""
        print("\n" + "=" * 80)
        print("📊 ENHANCED PROCESSING SUMMARY")
        print("=" * 80)
        
        print(f"📄 Documents processed: {total_stats['documents_processed']}")
        print(f"📝 Paragraphs processed: {total_stats['paragraphs_processed']}")
        print(f"⏱️  Processing time: {processing_time:.1f} seconds")
        
        print(f"\n🎯 TURNITIN RISK ANALYSIS:")
        print(f"  🔴 Critical risk: {total_stats['critical_risk']}")
        print(f"  🔴 Very high risk: {total_stats['very_high_risk']}")
        print(f"  🟠 High risk: {total_stats['high_risk']}")
        print(f"  🟡 Medium risk: {total_stats['medium_risk']}")
        print(f"  🟢 Low risk: {total_stats['low_risk']}")
        
        print(f"\n🤖 AI PROCESSING LEVELS:")
        print(f"  Level 3 (Deep): {total_stats['ai_level_3']}")
        print(f"  Level 2 (Pattern): {total_stats['ai_level_2']}")
        print(f"  Level 1 (Basic): {total_stats['ai_level_1']}")
        print(f"  Local Enhanced: {total_stats['local_enhanced']}")
        
        print(f"\n💰 COST TRACKING:")
        print(f"  Local calls: {self.cost_tracker['local_calls']}")
        print(f"  AI calls L1: {self.cost_tracker['ai_calls_level1']}")
        print(f"  AI calls L2: {self.cost_tracker['ai_calls_level2']}")
        print(f"  AI calls L3: {self.cost_tracker['ai_calls_level3']}")
        print(f"  Estimated cost: ${self.cost_tracker['estimated_cost_usd']:.4f}")
        
        if backup_folder:
            print(f"\n💾 Backup folder: {backup_folder}")
        
        print("=" * 80)
    
    def generate_enhanced_report(self, input_folder, total_stats, processing_time):
        """Generate detailed processing report"""
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        report_file = f"enhanced_turnitin_report_{timestamp}.json"
        
        report_data = {
            'timestamp': timestamp,
            'input_folder': input_folder,
            'processing_time_seconds': processing_time,
            'statistics': total_stats,
            'cost_tracker': self.cost_tracker,
            'configuration': self.config
        }
        
        try:
            with open(report_file, 'w', encoding='utf-8') as f:
                json.dump(report_data, f, indent=2, ensure_ascii=False)
            print(f"📊 Detailed report saved: {report_file}")
        except Exception as e:
            print(f"⚠️  Could not save report: {e}")

    def process_batch_documents(self, input_folder, create_backup=True):
        """Enhanced batch processing"""
        start_time = datetime.now()
        
        print("=" * 80)
        print("🎯 ENHANCED TURNITIN-AWARE + MULTI-LEVEL GEMINI PARAPHRASER")
        print("=" * 80)
        
        if not os.path.exists(input_folder):
            print(f"❌ Input folder not found: {input_folder}")
            return
        
        docx_files = [f for f in os.listdir(input_folder) if f.endswith('.docx') and not f.startswith('~')]
        
        if not docx_files:
            print(f"❌ No Word documents found in: {input_folder}")
            return
        
        print(f"📁 Input folder: {input_folder}")
        print(f"📄 Found {len(docx_files)} documents")
        print(f"🎯 Multi-level AI: {'3 Levels Enabled' if self.gemini_client else 'Local Only'}")
        
        # Create backup
        backup_folder = None
        if create_backup:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            backup_folder = f"{input_folder}_backup_enhanced_{timestamp}"
            
            try:
                if not os.path.exists(backup_folder):
                    os.makedirs(backup_folder)
                
                for file in docx_files:
                    src = os.path.join(input_folder, file)
                    dst = os.path.join(backup_folder, file)
                    shutil.copy2(src, dst)
                
                print(f"✅ Backup created: {backup_folder}")
            except Exception as e:
                print(f"❌ Error creating backup: {e}")
                return
        
        print(f"\n🚀 Starting enhanced multi-level processing...")
        
        # Process documents
        total_stats = {
            'documents_processed': 0,
            'paragraphs_processed': 0,
            'critical_risk': 0,
            'very_high_risk': 0,
            'high_risk': 0,
            'medium_risk': 0,
            'low_risk': 0,
            'ai_level_1': 0,
            'ai_level_2': 0,
            'ai_level_3': 0,
            'local_enhanced': 0,
            'total_changes': 0
        }
        
        for i, filename in enumerate(docx_files, 1):
            file_path = os.path.join(input_folder, filename)
            
            print(f"\n{'='*20} Document {i}/{len(docx_files)} {'='*20}")
            
            doc_stats = self.process_single_document_enhanced(file_path)
            
            if doc_stats:
                total_stats['documents_processed'] += 1
                for key in doc_stats:
                    if key in total_stats:
                        total_stats[key] += doc_stats[key]
        
        # Calculate processing time
        end_time = datetime.now()
        processing_time = (end_time - start_time).total_seconds()
        
        # Print comprehensive summary
        self.print_enhanced_summary(total_stats, processing_time, backup_folder)
        
        # Generate detailed report
        self.generate_enhanced_report(input_folder, total_stats, processing_time)
    
    def process_single_document_enhanced(self, file_path):
        """Enhanced single document processing"""
        try:
            print(f"\n📄 Processing: {os.path.basename(file_path)}")
            
            doc = docx.Document(file_path)
            
            doc_stats = {
                'filename': os.path.basename(file_path),
                'total_paragraphs': 0,
                'paragraphs_processed': 0,
                'critical_risk': 0,
                'very_high_risk': 0,
                'high_risk': 0,
                'medium_risk': 0,
                'low_risk': 0,
                'ai_level_1': 0,
                'ai_level_2': 0, 
                'ai_level_3': 0,
                'local_enhanced': 0,
                'total_changes': 0
            }
            
            current_section = 'UNKNOWN'
            processed_paragraphs = []
            
            # Collect suitable paragraphs
            for paragraph in doc.paragraphs:
                para_text = paragraph.text.strip()
                doc_stats['total_paragraphs'] += 1
                
                if not para_text:
                    continue
                
                if self.is_section_header(para_text):
                    current_section = para_text
                    print(f"  📍 Section: {current_section}")
                    continue
                
                if self.is_paragraph_suitable_for_paraphrasing(para_text):
                    processed_paragraphs.append({
                        'paragraph': paragraph,
                        'text': para_text,
                        'index': len(processed_paragraphs)
                    })
            
            total_suitable = len(processed_paragraphs)
            print(f"  📝 Found {total_suitable} suitable paragraphs")
            
            # Enhanced processing
            for para_data in processed_paragraphs:
                paragraph = para_data['paragraph']
                para_text = para_data['text']
                para_index = para_data['index']
                
                # Enhanced processing
                result = self.process_paragraph_enhanced_turnitin(para_text, para_index)
                
                # Track risk levels and methods
                risk_level = result.get('risk_analysis', {}).get('risk_level', 0)
                if risk_level >= 0.9:
                    doc_stats['critical_risk'] += 1
                elif risk_level >= 0.7:
                    doc_stats['very_high_risk'] += 1
                elif risk_level >= 0.5:
                    doc_stats['high_risk'] += 1
                elif risk_level >= 0.3:
                    doc_stats['medium_risk'] += 1
                else:
                    doc_stats['low_risk'] += 1
                
                # Track AI levels used
                if 'level_1' in result['method']:
                    doc_stats['ai_level_1'] += 1
                elif 'level_2' in result['method']:
                    doc_stats['ai_level_2'] += 1
                elif 'level_3' in result['method']:
                    doc_stats['ai_level_3'] += 1
                else:
                    doc_stats['local_enhanced'] += 1
                
                # Apply changes
                if result['plagiarism_reduction'] > 10:  # Lower threshold
                    paragraph.clear()
                    paragraph.add_run(result['paraphrase'])
                    
                    # Enhanced color coding
                    if 'level_3' in result['method']:
                        # Purple for Level 3 (Deep restructuring)
                        for run in paragraph.runs:
                            run.font.highlight_color = WD_COLOR_INDEX.VIOLET
                    elif 'level_2' in result['method']:
                        # Orange for Level 2 (Pattern-specific)
                        for run in paragraph.runs:
                            run.font.highlight_color = WD_COLOR_INDEX.BRIGHT_GREEN
                    elif 'level_1' in result['method']:
                        # Yellow for Level 1 (Basic AI)
                        for run in paragraph.runs:
                            run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                    else:
                        # Light blue for local enhanced
                        for run in paragraph.runs:
                            run.font.highlight_color = WD_COLOR_INDEX.TURQUOISE
                    
                    doc_stats['total_changes'] += 1
                    doc_stats['paragraphs_processed'] += 1
            
            # Save the document
            doc.save(file_path)
            print(f"  ✅ Document processed: {doc_stats['paragraphs_processed']} paragraphs modified")
            return doc_stats
            
        except Exception as e:
            print(f"❌ Error processing document: {e}")
            return None


def main():
    """Main function for enhanced Turnitin-aware processing"""
    print("🎯 ENHANCED TURNITIN-AWARE + MULTI-LEVEL GEMINI PARAPHRASER")
    print("Multi-level AI processing untuk maximum Turnitin avoidance")
    print("=" * 80)
    
    # Check dependencies
    if not GEMINI_AVAILABLE:
        print("❌ Required package missing!")
        print("📦 Install with: pip install google-genai")
        return
    
    # Configuration
    SYNONYM_FILE = 'contoh.json'            # Your synonym database
    INPUT_FOLDER = 'documents'              # Folder with Word documents
    CREATE_BACKUP = True                    # Always backup first
    
    # Get API key from environment or set directly
    GEMINI_API_KEY = os.environ.get("GEMINI_API_KEY")
    
    if not GEMINI_API_KEY:
        print("⚠️  WARNING: GEMINI_API_KEY not found!")
        print("🔑 Set your API key:")
        print("   Method 1: export GEMINI_API_KEY='your-api-key-here'")
        print("   Method 2: Set in environment variables")
        print("   Method 3: Uncomment and set directly in code below")
        print()
        # Uncomment and set your API key here if needed:
        # GEMINI_API_KEY = "your-gemini-api-key-here"
        
        if not GEMINI_API_KEY:
            print("❌ Cannot proceed without API key")
            print("📖 Get your free API key from: https://makersuite.google.com/app/apikey")
            return
    
    print(f"✅ API key configured: {GEMINI_API_KEY[:20]}...")
    
    # Initialize enhanced processor
    processor = EnhancedTurnitinGeminiParaphraser(
        synonym_file=SYNONYM_FILE,
        gemini_api_key=GEMINI_API_KEY
    )
    
    # Process batch with enhanced approach
    processor.process_batch_documents(
        input_folder=INPUT_FOLDER,
        create_backup=CREATE_BACKUP
    )
    
    print("\n🎉 Enhanced Turnitin-aware processing completed!")
    print("💡 Color coding:")
    print("   🟡 Yellow = AI Level 1 (Basic)")
    print("   🟢 Green = AI Level 2 (Pattern-specific)")
    print("   🟣 Purple = AI Level 3 (Deep restructuring)")
    print("   🔵 Blue = Local enhanced")
    print("\n📊 Check the detailed report for Turnitin risk analysis!")
    print("🚀 Powered by Multi-level Gemini AI processing!")


if __name__ == "__main__":
    main()
