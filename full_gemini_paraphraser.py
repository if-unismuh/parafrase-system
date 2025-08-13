# full_gemini_paraphraser.py
"""
Full Gemini AI Paraphraser - Enhanced Quality System
Focus: High-quality paraphrasing using full Gemini AI with content protection
Features:
- Full Gemini AI for superior sentence quality
- Protection for titles, author names, and important headings
- Academic document optimization
- Smart content preservation

Author: DevNoLife
Version: 1.0 - Full AI Edition
"""

import os
import json
import re
import time
from datetime import datetime
from pathlib import Path
import docx
from collections import defaultdict
import requests
from difflib import SequenceMatcher
import glob

# Load environment variables
try:
    from dotenv import load_dotenv
    load_dotenv()
    print("✅ Environment variables loaded from .env file")
except ImportError:
    print("⚠️  python-dotenv not found. Install with: pip install python-dotenv")
    print("💡 Using system environment variables...")

# Gemini AI imports
try:
    import google.generativeai as genai
    GEMINI_AVAILABLE = True
except ImportError:
    try:
        # Fallback to requests-based API
        GEMINI_AVAILABLE = "requests"
    except:
        GEMINI_AVAILABLE = False

class FullGeminiParaphraser:
    def __init__(self, synonym_file='sinonim.json', gemini_api_key=None, chunk_size=5000, max_retries=5):
        print("🚀 Initializing Full Gemini Paraphraser v2.0...")
        print("🎯 Focus: High-Quality AI Paraphrasing with Large File Support")
        
        # API Configuration
        self.gemini_api_key = gemini_api_key or os.getenv('GEMINI_API_KEY')
        if not self.gemini_api_key:
            print("❌ GEMINI_API_KEY not found! Set it in environment variables or .env file")
            raise ValueError("Gemini API key is required")
        
        # Large file processing configuration
        self.chunk_size = chunk_size  # Maximum characters per chunk
        self.max_retries = max_retries
        self.rate_limit_delay = 1.0  # Base delay between requests
        self.exponential_backoff = True
        
        # Progress tracking for large files
        self.progress_file = None
        self.resume_enabled = True
        
        # Initialize Gemini client
        self._initialize_gemini_client()
        
        # Load synonym database for reference (not primary paraphrasing)
        self.synonyms = self.load_synonyms(synonym_file)
        print(f"📚 Loaded {len(self.synonyms)} synonym groups")
        
        # Content protection patterns
        self.protection_patterns = {
            'titles': [
                r'^BAB\s+[IVX]+.*',  # Chapter titles
                r'^CHAPTER\s+\d+.*',
                r'^Bab\s+\d+.*',
                r'^\d+\.\s*[A-Z][^.]*$',  # Numbered titles without ending punctuation
                r'^[A-Z][A-Z\s]{5,}$',  # ALL CAPS titles
            ],
            'author_names': [
                r'Author:\s*(.+)',
                r'Penulis:\s*(.+)', 
                r'Oleh:\s*(.+)',
                r'^([A-Z][a-z]+\s+[A-Z][a-z]+(\s+[A-Z][a-z]+)*)$',  # Proper names
            ],
            'headings': [
                r'^\d+\.\d+.*',  # 1.1 Heading format
                r'^\d+\.\d+\.\d+.*',  # 1.1.1 Subheading format
                r'^[A-Z]\.\s+[A-Z].*',  # A. Heading format
                r'^ABSTRAK$|^ABSTRACT$|^PENDAHULUAN$|^KESIMPULAN$',
                r'^DAFTAR\s+(ISI|PUSTAKA|GAMBAR|TABEL)',
            ],
            'citations': [
                r'\(.+?\d{4}.+?\)',  # (Author, 2024)
                r'\[.+?\]',  # [1], [Author, 2024]
                r'(Menurut|According to).+?\(\d{4}\)',
            ],
            'formulas': [
                r'[A-Za-z]\s*=\s*[^.]{1,50}',  # Mathematical formulas
                r'\$.*?\$',  # LaTeX formulas
            ],
            'tables_figures': [
                r'^Tabel\s+\d+.*',
                r'^Gambar\s+\d+.*',
                r'^Figure\s+\d+.*',
                r'^Table\s+\d+.*',
            ]
        }
        
        # Cost tracking
        self.cost_tracker = {
            'total_api_calls': 0,
            'total_tokens_used': 0,
            'total_cost_usd': 0.0,
            'paragraphs_processed': 0,
            'protected_content': 0
        }
        
        # Processing statistics
        self.stats = {
            'start_time': datetime.now(),
            'total_similarity_reduction': 0,
            'protected_items': [],
            'quality_scores': []
        }
        
        print("✅ Full Gemini Paraphraser initialized successfully!")

    def _initialize_gemini_client(self):
        """Initialize Gemini AI client"""
        if GEMINI_AVAILABLE == True:
            try:
                genai.configure(api_key=self.gemini_api_key)
                # Use the most advanced model available
                self.gemini_client = genai.GenerativeModel('gemini-2.0-flash-thinking-exp')
                self.api_method = 'genai'
                print(f"✅ Gemini client initialized (gemini-2.0-flash-thinking-exp)")
            except Exception as e:
                try:
                    # Fallback to standard model
                    self.gemini_client = genai.GenerativeModel('gemini-1.5-pro')
                    self.api_method = 'genai'
                    print(f"✅ Gemini client initialized (gemini-1.5-pro fallback)")
                except Exception as e2:
                    print(f"⚠️  Google GenAI failed: {e2}")
                    self.gemini_client = None
                    self.api_method = 'requests'
        elif GEMINI_AVAILABLE == "requests":
            self.api_method = 'requests'
            self.gemini_api_url = "https://generativelanguage.googleapis.com/v1beta/models/gemini-2.0-flash:generateContent"
        else:
            raise ValueError("Gemini AI not available. Please install google-generativeai")

    def load_synonyms(self, filename):
        """Load synonym database from JSON file"""
        if not filename or not os.path.exists(filename):
            print("⚠️  No synonym file provided or file not found")
            return {}
        
        try:
            # Try components directory first
            if not os.path.exists(filename):
                components_path = os.path.join('components', filename)
                if os.path.exists(components_path):
                    filename = components_path
            
            with open(filename, 'r', encoding='utf-8') as f:
                data = json.load(f)
            
            synonyms = defaultdict(list)
            count = 0
            
            for word, info in data.items():
                if isinstance(info, dict) and 'sinonim' in info:
                    synonym_list = info['sinonim']
                    if isinstance(synonym_list, list) and len(synonym_list) > 0:
                        all_words = [word] + synonym_list
                        for w in all_words:
                            clean_w = w.lower().strip()
                            synonyms[clean_w] = [s.lower().strip() for s in all_words if s.lower().strip() != clean_w]
                        count += 1
            
            return dict(synonyms)
            
        except Exception as e:
            print(f"❌ Error loading synonyms: {e}")
            return {}

    def is_protected_content(self, text):
        """Check if content should be protected from paraphrasing"""
        text = text.strip()
        
        # Check each protection category
        for category, patterns in self.protection_patterns.items():
            for pattern in patterns:
                if re.match(pattern, text, re.IGNORECASE):
                    self.stats['protected_items'].append({
                        'text': text[:100],  # First 100 chars for logging
                        'category': category,
                        'pattern': pattern
                    })
                    return True, category
        
        return False, None

    def create_full_gemini_prompt(self, text, context="academic"):
        """Create comprehensive prompt for full Gemini AI paraphrasing"""
        
        prompt = f"""Anda adalah expert akademis dalam paraphrasing bahasa Indonesia dengan spesialisasi pada teks ilmiah dan penelitian.

MISI UTAMA: Parafrase teks berikut untuk menghasilkan kualitas tulisan akademis yang superior dengan pengurangan plagiarisme maksimal sambil mempertahankan makna dan konteks yang tepat.

STRATEGI PARAPHRASING ADVANCED:
✅ SINONIM KONTEKSTUAL: Gunakan sinonim yang tepat sesuai konteks akademis
✅ TRANSFORMASI STRUKTURAL: Ubah struktur kalimat secara mendalam (aktif↔pasif, subordinasi↔koordinasi)
✅ VARIASI FRASA AKADEMIS: Ganti terminologi akademis dengan variasi yang setara namun berbeda
✅ REKONSTRUKSI GRAMATIKAL: Ubah konstruksi gramatikal tanpa mengubah substansi
✅ OPTIMASI FLOW: Pastikan alur kalimat natural dan mudah dibaca
✅ TONE AKADEMIS: Pertahankan registrer formal dan scientific writing style
✅ COHERENCE: Jaga koherensi antar kalimat dan paragraf

TARGET KUALITAS:
- Pengurangan similarity: 50-80%
- Readability score: Tinggi (mudah dipahami)
- Academic tone: Dipertahankan
- Meaning preservation: 100%

KONTEKS: {context}

ATURAN KHUSUS:
- JANGAN ubah nama orang, tempat, atau istilah teknis spesifik
- PERTAHANKAN angka, tanggal, dan data statistik
- JAGA konsistensi terminologi dalam satu dokumen
- Hasil harus terdengar natural bagi pembaca Indonesia

TEKS YANG AKAN DIPARAFRASE:
"{text}"

INSTRUKSI RESPONSE:
Berikan HANYA hasil parafrase tanpa penjelasan tambahan. Hasil harus berupa paragraf yang flow dan natural."""

        return prompt

    def call_gemini_ai(self, prompt, max_retries=None):
        """Enhanced Gemini AI API call with intelligent rate limiting and retry logic"""
        if max_retries is None:
            max_retries = self.max_retries
            
        # Dynamic delay based on content size and current load
        content_size = len(prompt)
        base_delay = self.rate_limit_delay
        
        if content_size > 3000:
            base_delay *= 1.5  # Longer delay for larger content
        
        for attempt in range(max_retries):
            try:
                # Apply rate limiting delay before each request
                if attempt > 0:
                    if self.exponential_backoff:
                        delay = base_delay * (2 ** attempt) + (attempt * 0.5)
                    else:
                        delay = base_delay + (attempt * 1.0)
                    
                    print(f"⏳ Rate limiting: waiting {delay:.1f}s before retry {attempt + 1}")
                    time.sleep(delay)
                
                if self.api_method == 'genai':
                    response = self.gemini_client.generate_content(
                        prompt,
                        generation_config=genai.types.GenerationConfig(
                            temperature=0.7,  # Balanced creativity
                            max_output_tokens=4000,
                            top_p=0.8,
                            top_k=40
                        )
                    )
                    
                    if response and response.text:
                        self.cost_tracker['total_api_calls'] += 1
                        # Estimate token usage (rough calculation)
                        estimated_tokens = len(prompt.split()) + len(response.text.split())
                        self.cost_tracker['total_tokens_used'] += estimated_tokens
                        
                        # Adaptive rate limiting based on success
                        if attempt == 0:
                            self.rate_limit_delay = max(0.5, self.rate_limit_delay * 0.95)  # Decrease delay on success
                        
                        return response.text.strip()
                
                elif self.api_method == 'requests':
                    headers = {
                        'Content-Type': 'application/json',
                    }
                    
                    payload = {
                        "contents": [{
                            "parts": [{"text": prompt}]
                        }],
                        "generationConfig": {
                            "temperature": 0.7,
                            "maxOutputTokens": 4000,
                            "topP": 0.8,
                            "topK": 40
                        }
                    }
                    
                    response = requests.post(
                        f"{self.gemini_api_url}?key={self.gemini_api_key}",
                        headers=headers,
                        json=payload,
                        timeout=60
                    )
                    
                    if response.status_code == 200:
                        data = response.json()
                        if 'candidates' in data and len(data['candidates']) > 0:
                            text = data['candidates'][0]['content']['parts'][0]['text']
                            self.cost_tracker['total_api_calls'] += 1
                            estimated_tokens = len(prompt.split()) + len(text.split())
                            self.cost_tracker['total_tokens_used'] += estimated_tokens
                            
                            # Adaptive rate limiting based on success
                            if attempt == 0:
                                self.rate_limit_delay = max(0.5, self.rate_limit_delay * 0.95)
                            
                            return text.strip()
                    elif response.status_code == 429:  # Rate limit exceeded
                        print(f"⚠️  Rate limit exceeded, increasing delay")
                        self.rate_limit_delay = min(10.0, self.rate_limit_delay * 2)  # Increase delay
                        raise Exception(f"Rate limit exceeded (429)")
                    elif response.status_code == 503:  # Service unavailable
                        print(f"⚠️  Service temporarily unavailable")
                        raise Exception(f"Service unavailable (503)")
                    else:
                        print(f"❌ API Error {response.status_code}: {response.text}")
                        raise Exception(f"API Error {response.status_code}")
                        
            except Exception as e:
                print(f"⚠️  Attempt {attempt + 1}/{max_retries} failed: {e}")
                
                # Increase delay on repeated failures
                if "rate limit" in str(e).lower() or "429" in str(e):
                    self.rate_limit_delay = min(10.0, self.rate_limit_delay * 1.5)
                
                if attempt < max_retries - 1:
                    continue  # Will apply delay at start of next iteration
                else:
                    print(f"❌ All {max_retries} API attempts failed")
                    return None
        
        return None

    def calculate_similarity(self, text1, text2):
        """Enhanced similarity calculation"""
        # Method 1: Word-based Jaccard similarity
        words1 = set(re.findall(r'\w+', text1.lower()))
        words2 = set(re.findall(r'\w+', text2.lower()))
        
        if not words1 and not words2:
            return 100.0
        if not words1 or not words2:
            return 0.0
        
        jaccard_sim = len(words1.intersection(words2)) / len(words1.union(words2)) * 100
        
        # Method 2: SequenceMatcher for structural similarity
        sequence_sim = SequenceMatcher(None, text1.lower(), text2.lower()).ratio() * 100
        
        # Method 3: Phrase-level similarity
        phrases1 = set(text1.lower().split())
        phrases2 = set(text2.lower().split())
        phrase_sim = len(phrases1.intersection(phrases2)) / len(phrases1.union(phrases2)) * 100
        
        # Weighted combination: 40% Jaccard, 40% Sequence, 20% Phrase
        final_similarity = (jaccard_sim * 0.4) + (sequence_sim * 0.4) + (phrase_sim * 0.2)
        
        return round(final_similarity, 2)

    def paraphrase_text(self, text, context="academic"):
        """Main paraphrasing method using full Gemini AI"""
        
        # Check if content should be protected
        is_protected, category = self.is_protected_content(text)
        if is_protected:
            print(f"🛡️  Protected content detected ({category}): {text[:50]}...")
            self.cost_tracker['protected_content'] += 1
            return {
                'paraphrase': text,  # Return original
                'original': text,
                'similarity': 100.0,
                'plagiarism_reduction': 0.0,
                'method': 'protected_content',
                'category': category,
                'status': 'PROTECTED'
            }
        
        # Create comprehensive prompt
        prompt = self.create_full_gemini_prompt(text, context)
        
        # Call Gemini AI
        print(f"🤖 Processing with Full Gemini AI...")
        paraphrased = self.call_gemini_ai(prompt)
        
        if not paraphrased:
            print("❌ AI processing failed, returning original text")
            return {
                'paraphrase': text,
                'original': text,
                'similarity': 100.0,
                'plagiarism_reduction': 0.0,
                'method': 'ai_failed',
                'status': 'FAILED'
            }
        
        # Calculate similarity and quality metrics
        similarity = self.calculate_similarity(text, paraphrased)
        plagiarism_reduction = 100 - similarity
        
        # Update statistics
        self.cost_tracker['paragraphs_processed'] += 1
        self.stats['total_similarity_reduction'] += plagiarism_reduction
        
        # Quality assessment
        quality_score = self.assess_quality(text, paraphrased)
        self.stats['quality_scores'].append(quality_score)
        
        result = {
            'paraphrase': paraphrased,
            'original': text,
            'similarity': similarity,
            'plagiarism_reduction': plagiarism_reduction,
            'method': 'full_gemini_ai',
            'quality_score': quality_score,
            'status': self._get_status(similarity),
            'word_count_original': len(text.split()),
            'word_count_paraphrase': len(paraphrased.split())
        }
        
        return result

    def assess_quality(self, original, paraphrased):
        """Assess the quality of paraphrasing"""
        score = 0
        
        # Length similarity (good paraphrase should have similar length)
        len_ratio = min(len(paraphrased), len(original)) / max(len(paraphrased), len(original))
        score += len_ratio * 25
        
        # Word diversity (how many different words used)
        orig_words = set(original.lower().split())
        para_words = set(paraphrased.lower().split())
        diversity = len(para_words - orig_words) / len(orig_words) if orig_words else 0
        score += min(diversity, 1) * 25
        
        # Sentence structure variety (different sentence patterns)
        orig_sentences = len(re.findall(r'[.!?]+', original))
        para_sentences = len(re.findall(r'[.!?]+', paraphrased))
        structure_score = 25 if orig_sentences == para_sentences else 15
        score += structure_score
        
        # Readability (avoid overly complex sentences)
        avg_word_len = sum(len(word) for word in paraphrased.split()) / len(paraphrased.split()) if paraphrased.split() else 0
        readability = max(0, 25 - (avg_word_len - 6) * 2)  # Penalize very long words
        score += readability
        
        return round(score, 2)

    def _get_status(self, similarity):
        """Get processing status based on similarity"""
        if similarity >= 70:
            return 'HIGH_SIMILARITY'
        elif similarity >= 40:
            return 'MEDIUM_SIMILARITY'
        elif similarity >= 20:
            return 'LOW_SIMILARITY'
        else:
            return 'VERY_LOW_SIMILARITY'

    def save_progress(self, progress_data):
        """Save processing progress to file"""
        if self.progress_file:
            try:
                with open(self.progress_file, 'w', encoding='utf-8') as f:
                    json.dump(progress_data, f, indent=2, ensure_ascii=False, default=str)
                print(f"💾 Progress saved: {self.progress_file}")
            except Exception as e:
                print(f"⚠️  Failed to save progress: {e}")

    def load_progress(self, progress_file):
        """Load processing progress from file"""
        if os.path.exists(progress_file):
            try:
                with open(progress_file, 'r', encoding='utf-8') as f:
                    progress_data = json.load(f)
                print(f"📂 Progress loaded: {progress_file}")
                return progress_data
            except Exception as e:
                print(f"⚠️  Failed to load progress: {e}")
        return None

    def detect_chapter_boundaries(self, doc):
        """Detect BAB I (start) and end of main content, skip appendices"""
        start_index = None
        end_index = None
        
        chapter_patterns = [
            r'^BAB\s+I\b.*PENDAHULUAN',
            r'^BAB\s+1\b.*PENDAHULUAN', 
            r'^CHAPTER\s+I\b.*PENDAHULUAN',
            r'^CHAPTER\s+1\b.*PENDAHULUAN',
            r'^1\.\s*PENDAHULUAN',
            r'^BAB\s+I\b',
            r'^BAB\s+1\b'
        ]
        
        end_patterns = [
            r'^BAB\s+V\b.*KESIMPULAN',
            r'^BAB\s+5\b.*KESIMPULAN',
            r'^CHAPTER\s+V\b.*KESIMPULAN',
            r'^CHAPTER\s+5\b.*KESIMPULAN',
            r'^5\.\s*KESIMPULAN',
            r'^BAB\s+V\b',
            r'^BAB\s+5\b',
            r'^DAFTAR\s+PUSTAKA',
            r'^REFERENCES',
            r'^BIBLIOGRAPHY',
            r'^LAMPIRAN',
            r'^APPENDIX',
            r'^APPENDICES'
        ]
        
        # Find start of main content (BAB I)
        for i, paragraph in enumerate(doc.paragraphs):
            text = paragraph.text.strip().upper()
            for pattern in chapter_patterns:
                if re.match(pattern, text, re.IGNORECASE):
                    start_index = i
                    print(f"📖 Found start at paragraph {i}: {text[:50]}...")
                    break
            if start_index is not None:
                break
        
        # Find end of main content 
        if start_index is not None:
            for i, paragraph in enumerate(doc.paragraphs[start_index:], start_index):
                text = paragraph.text.strip().upper()
                for pattern in end_patterns:
                    if re.match(pattern, text, re.IGNORECASE):
                        # If it's KESIMPULAN, include it. If it's DAFTAR PUSTAKA/LAMPIRAN, stop before it
                        if 'KESIMPULAN' in pattern:
                            # Look for the end of KESIMPULAN chapter
                            for j in range(i + 1, len(doc.paragraphs)):
                                next_text = doc.paragraphs[j].text.strip().upper()
                                if any(re.match(p, next_text, re.IGNORECASE) for p in end_patterns if 'KESIMPULAN' not in p):
                                    end_index = j
                                    break
                            if end_index is None:
                                end_index = len(doc.paragraphs)
                        else:
                            end_index = i
                        print(f"📖 Found end at paragraph {end_index}: {text[:50]}...")
                        break
                if end_index is not None:
                    break
        
        return start_index, end_index

    def split_large_text(self, text):
        """Split large text into manageable chunks while preserving sentence boundaries"""
        if len(text) <= self.chunk_size:
            return [text]
        
        chunks = []
        current_chunk = ""
        sentences = re.split(r'([.!?]+\s*)', text)
        
        for i in range(0, len(sentences), 2):
            sentence = sentences[i] if i < len(sentences) else ""
            punct = sentences[i + 1] if i + 1 < len(sentences) else ""
            full_sentence = sentence + punct
            
            if len(current_chunk) + len(full_sentence) <= self.chunk_size:
                current_chunk += full_sentence
            else:
                if current_chunk.strip():
                    chunks.append(current_chunk.strip())
                current_chunk = full_sentence
        
        if current_chunk.strip():
            chunks.append(current_chunk.strip())
        
        return chunks if chunks else [text]

    def process_document(self, input_path, output_path=None, resume=True):
        """Enhanced document processing with progress tracking and resume capability"""
        print(f"🎯 Processing document: {input_path}")
        print("🤖 Using Full Gemini AI with large file support...")
        
        if not os.path.exists(input_path):
            print(f"❌ File not found: {input_path}")
            return None
        
        # Generate output path if not provided
        if not output_path:
            base_name = Path(input_path).stem
            ext = Path(input_path).suffix
            # Save to completed folder
            completed_dir = "completed"
            os.makedirs(completed_dir, exist_ok=True)
            output_path = os.path.join(completed_dir, f"{base_name}_paraphrased{ext}")
        
        # Setup progress tracking
        progress_dir = "progress"
        os.makedirs(progress_dir, exist_ok=True)
        self.progress_file = os.path.join(progress_dir, f"{Path(input_path).stem}_progress.json")
        
        # Check for existing progress
        progress_data = None
        if resume and self.resume_enabled:
            progress_data = self.load_progress(self.progress_file)
        
        try:
            # Load document
            doc = docx.Document(input_path)
            
            # Detect chapter boundaries (BAB I to BAB V/KESIMPULAN)
            start_index, end_index = self.detect_chapter_boundaries(doc)
            
            if start_index is None:
                print("⚠️  Could not find BAB I PENDAHULUAN. Processing entire document.")
                start_index = 0
                end_index = len(doc.paragraphs)
            
            if end_index is None:
                print("⚠️  Could not find end marker. Processing until end of document.")
                end_index = len(doc.paragraphs)
            
            # Filter paragraphs to main content only
            main_paragraphs = []
            for i in range(start_index, end_index):
                if i < len(doc.paragraphs) and doc.paragraphs[i].text.strip():
                    main_paragraphs.append((i, doc.paragraphs[i]))
            
            total_paragraphs = len(main_paragraphs)
            
            print(f"📄 Processing main content: paragraphs {start_index} to {end_index}")
            print(f"📄 Found {total_paragraphs} paragraphs to process (excluding cover, appendices, etc.)")
            
            # Initialize or restore progress
            if progress_data and 'completed_paragraphs' in progress_data:
                completed_paragraphs = set(progress_data['completed_paragraphs'])
                results = progress_data.get('results', [])
                print(f"🔄 Resuming from paragraph {len(completed_paragraphs) + 1}/{total_paragraphs}")
            else:
                completed_paragraphs = set()
                results = []
                progress_data = {
                    'input_path': input_path,
                    'output_path': output_path,
                    'total_paragraphs': total_paragraphs,
                    'completed_paragraphs': [],
                    'results': [],
                    'start_time': datetime.now().isoformat(),
                    'last_update': datetime.now().isoformat()
                }
            
            # Process paragraphs (main content only)
            for paragraph_index, paragraph in main_paragraphs:
                if paragraph_index in completed_paragraphs:
                    continue  # Skip already processed paragraphs
                
                paragraph_text = paragraph.text.strip()
                print(f"🔄 Processing paragraph {len(completed_paragraphs) + 1}/{total_paragraphs}")
                
                # Handle large paragraphs by splitting into chunks
                text_chunks = self.split_large_text(paragraph_text)
                
                if len(text_chunks) > 1:
                    print(f"📝 Large paragraph split into {len(text_chunks)} chunks")
                
                chunk_results = []
                paraphrased_chunks = []
                
                for chunk_idx, chunk in enumerate(text_chunks):
                    print(f"   🔸 Processing chunk {chunk_idx + 1}/{len(text_chunks)}")
                    
                    # Process chunk
                    result = self.paraphrase_text(chunk)
                    chunk_results.append(result)
                    
                    # Collect paraphrased text
                    if result['method'] != 'protected_content':
                        paraphrased_chunks.append(result['paraphrase'])
                    else:
                        paraphrased_chunks.append(chunk)
                    
                    # Save progress after each chunk for very large files
                    if len(text_chunks) > 5:
                        progress_data['last_update'] = datetime.now().isoformat()
                        self.save_progress(progress_data)
                
                # Combine chunks back into paragraph
                combined_result = {
                    'paraphrase': ' '.join(paraphrased_chunks),
                    'original': paragraph_text,
                    'chunks': chunk_results,
                    'chunk_count': len(text_chunks),
                    'paragraph_index': paragraph_index,
                    'method': 'chunked_processing' if len(text_chunks) > 1 else chunk_results[0]['method'],
                    'status': 'COMPLETED'
                }
                
                # Calculate combined metrics
                if chunk_results:
                    combined_result['similarity'] = sum(r.get('similarity', 0) for r in chunk_results) / len(chunk_results)
                    combined_result['plagiarism_reduction'] = 100 - combined_result['similarity']
                    combined_result['quality_score'] = sum(r.get('quality_score', 0) for r in chunk_results) / len(chunk_results)
                
                results.append(combined_result)
                
                # Update paragraph with paraphrased text
                paragraph.text = combined_result['paraphrase']
                
                # Update progress
                completed_paragraphs.add(paragraph_index)
                progress_data['completed_paragraphs'] = list(completed_paragraphs)
                progress_data['results'] = results
                progress_data['last_update'] = datetime.now().isoformat()
                
                # Save progress regularly
                if len(completed_paragraphs) % 5 == 0:  # Every 5 paragraphs
                    self.save_progress(progress_data)
                
                # Apply base delay between paragraphs
                time.sleep(self.rate_limit_delay)
            
            # Save processed document
            doc.save(output_path)
            print(f"✅ Document saved: {output_path}")
            
            # Clean up progress file on successful completion
            if os.path.exists(self.progress_file):
                os.remove(self.progress_file)
                print("🧹 Progress file cleaned up")
            
            # Generate detailed report
            report = self.generate_report(results, input_path, output_path)
            
            return {
                'output_path': output_path,
                'results': results,
                'report': report,
                'statistics': self.get_statistics()
            }
            
        except KeyboardInterrupt:
            print("\n⏸️  Processing interrupted by user")
            print(f"💾 Progress saved to: {self.progress_file}")
            print("🔄 Use resume=True to continue processing from where you left off")
            return None
        except Exception as e:
            print(f"❌ Error processing document: {e}")
            print(f"💾 Progress saved to: {self.progress_file}")
            return None

    def generate_report(self, results, input_path, output_path):
        """Generate detailed processing report"""
        report = {
            'processing_info': {
                'input_file': input_path,
                'output_file': output_path,
                'processing_time': str(datetime.now() - self.stats['start_time']),
                'timestamp': datetime.now().isoformat()
            },
            'summary': {
                'total_paragraphs': len(results),
                'protected_content': sum(1 for r in results if r.get('method') == 'protected_content'),
                'successfully_paraphrased': sum(1 for r in results if r.get('method') == 'full_gemini_ai'),
                'average_similarity_reduction': round(
                    sum(r.get('plagiarism_reduction', 0) for r in results) / len(results), 2
                ) if results else 0,
                'average_quality_score': round(
                    sum(r.get('quality_score', 0) for r in results if 'quality_score' in r) / 
                    len([r for r in results if 'quality_score' in r]), 2
                ) if any('quality_score' in r for r in results) else 0
            },
            'costs': self.cost_tracker,
            'protected_items': self.stats['protected_items'],
            'quality_distribution': self._get_quality_distribution(results)
        }
        
        # Save report to JSON
        report_path = f"report_{Path(output_path).stem}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json"
        with open(report_path, 'w', encoding='utf-8') as f:
            json.dump(report, f, indent=2, ensure_ascii=False)
        
        print(f"📊 Report saved: {report_path}")
        return report

    def _get_quality_distribution(self, results):
        """Get quality score distribution"""
        quality_scores = [r.get('quality_score', 0) for r in results if 'quality_score' in r]
        if not quality_scores:
            return {}
        
        return {
            'excellent': sum(1 for q in quality_scores if q >= 80),
            'good': sum(1 for q in quality_scores if 60 <= q < 80),
            'fair': sum(1 for q in quality_scores if 40 <= q < 60),
            'poor': sum(1 for q in quality_scores if q < 40),
            'average': round(sum(quality_scores) / len(quality_scores), 2),
            'highest': max(quality_scores),
            'lowest': min(quality_scores)
        }

    def get_statistics(self):
        """Get processing statistics"""
        return {
            'runtime': str(datetime.now() - self.stats['start_time']),
            'total_api_calls': self.cost_tracker['total_api_calls'],
            'total_tokens_used': self.cost_tracker['total_tokens_used'],
            'paragraphs_processed': self.cost_tracker['paragraphs_processed'],
            'protected_content': self.cost_tracker['protected_content'],
            'average_quality': round(
                sum(self.stats['quality_scores']) / len(self.stats['quality_scores']), 2
            ) if self.stats['quality_scores'] else 0,
            'total_similarity_reduction': round(
                self.stats['total_similarity_reduction'] / max(1, self.cost_tracker['paragraphs_processed']), 2
            )
        }
    
    def process_core_chapters_only(self, input_path, output_path=None):
        """Process only core chapters (BAB I - BAB V/KESIMPULAN), skip preliminaries and appendices"""
        print(f"🎯 Processing CORE CHAPTERS ONLY: {input_path}")
        print("📖 Focus: BAB I PENDAHULUAN → BAB V KESIMPULAN")
        print("⏭️  Skipping: Cover pages, appendices, references")
        
        if not os.path.exists(input_path):
            print(f"❌ File not found: {input_path}")
            return None
        
        # Generate output path if not provided
        if not output_path:
            base_name = Path(input_path).stem
            ext = Path(input_path).suffix
            completed_dir = "completed"
            os.makedirs(completed_dir, exist_ok=True)
            output_path = os.path.join(completed_dir, f"{base_name}_core_chapters_paraphrased{ext}")
        
        try:
            # Load document
            doc = docx.Document(input_path)
            
            # Detect chapter boundaries
            start_index, end_index = self.detect_chapter_boundaries(doc)
            
            if start_index is None:
                print("❌ Could not find BAB I PENDAHULUAN in the document!")
                print("💡 Make sure your document has 'BAB I' or 'BAB 1' followed by 'PENDAHULUAN'")
                return None
            
            if end_index is None:
                print("⚠️  Could not find specific end marker. Processing until end of document.")
                end_index = len(doc.paragraphs)
            
            print(f"📍 Processing range: paragraph {start_index} to {end_index}")
            
            # Create new document with only core chapters
            new_doc = docx.Document()
            
            # Copy styles from original document
            try:
                new_doc.styles = doc.styles
            except:
                pass  # If style copying fails, continue with default styles
            
            results = []
            processed_count = 0
            
            # Process and copy core chapters only
            for i in range(start_index, end_index):
                if i >= len(doc.paragraphs):
                    break
                    
                paragraph = doc.paragraphs[i]
                if not paragraph.text.strip():
                    # Copy empty paragraphs as-is
                    new_doc.add_paragraph("")
                    continue
                
                processed_count += 1
                print(f"🔄 Processing paragraph {processed_count}: {paragraph.text.strip()[:60]}...")
                
                # Process paragraph
                result = self.paraphrase_text(paragraph.text.strip())
                results.append(result)
                
                # Add processed paragraph to new document
                new_paragraph = new_doc.add_paragraph()
                if result['method'] != 'protected_content':
                    new_paragraph.text = result['paraphrase']
                else:
                    new_paragraph.text = paragraph.text  # Keep protected content as-is
                
                # Copy paragraph formatting if possible
                try:
                    if paragraph.style:
                        new_paragraph.style = paragraph.style
                except:
                    pass  # If style copying fails, use default
                
                # Apply rate limiting
                time.sleep(self.rate_limit_delay)
            
            # Save processed document
            new_doc.save(output_path)
            print(f"✅ Core chapters saved: {output_path}")
            print(f"📊 Processed {processed_count} paragraphs from core chapters only")
            
            # Generate report
            report = self.generate_report(results, input_path, output_path)
            
            return {
                'output_path': output_path,
                'results': results,
                'report': report,
                'statistics': self.get_statistics(),
                'core_chapters_only': True,
                'start_paragraph': start_index,
                'end_paragraph': end_index,
                'processed_paragraphs': processed_count
            }
            
        except Exception as e:
            print(f"❌ Error processing core chapters: {e}")
            return None

    def resume_document_processing(self, input_path, output_path=None):
        """Resume processing of a specific document from saved progress"""
        return self.process_document(input_path, output_path, resume=True)

    def check_pending_documents(self, progress_dir="progress"):
        """Check for documents with incomplete processing"""
        if not os.path.exists(progress_dir):
            return []
        
        pending_files = []
        for progress_file in glob.glob(os.path.join(progress_dir, "*_progress.json")):
            try:
                with open(progress_file, 'r', encoding='utf-8') as f:
                    progress_data = json.load(f)
                
                total = progress_data.get('total_paragraphs', 0)
                completed = len(progress_data.get('completed_paragraphs', []))
                
                pending_files.append({
                    'input_path': progress_data.get('input_path', ''),
                    'output_path': progress_data.get('output_path', ''),
                    'progress_file': progress_file,
                    'completion_percentage': (completed / total * 100) if total > 0 else 0,
                    'completed_paragraphs': completed,
                    'total_paragraphs': total,
                    'last_update': progress_data.get('last_update', '')
                })
            except Exception as e:
                print(f"⚠️  Error reading progress file {progress_file}: {e}")
        
        return pending_files

    def resume_all_pending(self):
        """Resume processing of all pending documents"""
        pending = self.check_pending_documents()
        
        if not pending:
            print("✅ No pending documents found")
            return []
        
        print(f"🔄 Found {len(pending)} pending documents")
        
        results = []
        for doc_info in pending:
            print(f"\n📄 Resuming: {Path(doc_info['input_path']).name}")
            print(f"   Progress: {doc_info['completed_paragraphs']}/{doc_info['total_paragraphs']} paragraphs ({doc_info['completion_percentage']:.1f}%)")
            
            result = self.resume_document_processing(
                doc_info['input_path'],
                doc_info['output_path']
            )
            
            if result:
                results.append(result)
                print(f"✅ Completed: {Path(doc_info['output_path']).name}")
            else:
                print(f"❌ Failed to complete: {Path(doc_info['input_path']).name}")
        
        return results

    def process_all_core_chapters(self, input_dir="documents", output_dir="completed"):
        """Process core chapters (BAB I - BAB V) for all documents, skip preliminaries and appendices"""
        print(f"🚀 Processing CORE CHAPTERS for all documents from: {input_dir}")
        print(f"📁 Output directory: {output_dir}")
        print("📖 Focus: BAB I PENDAHULUAN → BAB V KESIMPULAN only")
        
        # Create output directory if it doesn't exist
        os.makedirs(output_dir, exist_ok=True)
        
        # Find all Word documents
        doc_patterns = ["*.docx", "*.doc"]
        doc_files = []
        
        for pattern in doc_patterns:
            doc_files.extend(glob.glob(os.path.join(input_dir, pattern)))
        
        if not doc_files:
            print(f"❌ No documents found in {input_dir}")
            return []
        
        print(f"📄 Found {len(doc_files)} documents to process (core chapters only)")
        
        results = []
        for i, doc_path in enumerate(doc_files, 1):
            print(f"\n🔄 Processing document {i}/{len(doc_files)}: {Path(doc_path).name}")
            
            # Generate output path
            base_name = Path(doc_path).stem
            ext = Path(doc_path).suffix
            output_path = os.path.join(output_dir, f"{base_name}_core_chapters_paraphrased{ext}")
            
            # Skip if output already exists
            if os.path.exists(output_path):
                print(f"⏭️  Skipping (output exists): {Path(output_path).name}")
                continue
            
            # Process core chapters only
            result = self.process_core_chapters_only(doc_path, output_path)
            
            if result:
                result['input_file'] = doc_path
                result['document_number'] = i
                results.append(result)
                print(f"✅ Completed: {Path(output_path).name}")
            else:
                print(f"❌ Failed to process: {Path(doc_path).name}")
            
            # Add delay between documents
            if i < len(doc_files):
                delay = self.rate_limit_delay * 2
                print(f"⏳ Waiting {delay}s before next document...")
                time.sleep(delay)
        
        # Generate summary report
        if results:
            summary_report = self.generate_batch_summary(results, input_dir, output_dir)
            print(f"\n🎉 Core chapters processing completed!")
            print(f"✅ Processed: {len(results)} documents")
            print(f"📊 Summary report: {summary_report['report_path']}")
        
        return results

    def process_all_documents(self, input_dir="documents", output_dir="completed", resume_pending=True):
        """Enhanced batch processing with resume capability"""
        print(f"🚀 Processing all documents from: {input_dir}")
        print(f"📁 Output directory: {output_dir}")
        
        # Create output directory if it doesn't exist
        os.makedirs(output_dir, exist_ok=True)
        
        results = []
        
        # First, resume any pending documents if enabled
        if resume_pending:
            print("\n🔍 Checking for pending documents...")
            pending_results = self.resume_all_pending()
            results.extend(pending_results)
        
        # Find all Word documents
        doc_patterns = ["*.docx", "*.doc"]
        doc_files = []
        
        for pattern in doc_patterns:
            doc_files.extend(glob.glob(os.path.join(input_dir, pattern)))
        
        if not doc_files:
            print(f"❌ No new documents found in {input_dir}")
            if results:
                print(f"✅ Resumed {len(results)} pending documents")
            return results
        
        # Filter out documents that are already completed
        completed_files = set()
        for result in results:
            if 'input_file' in result:
                completed_files.add(result['input_file'])
        
        new_doc_files = [f for f in doc_files if f not in completed_files]
        
        if new_doc_files:
            print(f"📄 Found {len(new_doc_files)} new documents to process")
        else:
            print("✅ All documents already processed or resumed")
            return results
        
        # Process new documents
        for i, doc_path in enumerate(new_doc_files, 1):
            print(f"\n🔄 Processing document {i}/{len(new_doc_files)}: {Path(doc_path).name}")
            
            # Generate output path
            base_name = Path(doc_path).stem
            ext = Path(doc_path).suffix
            output_path = os.path.join(output_dir, f"{base_name}_paraphrased{ext}")
            
            # Skip if output already exists and is not from a resumed document
            if os.path.exists(output_path) and doc_path not in completed_files:
                print(f"⏭️  Skipping (output exists): {Path(output_path).name}")
                continue
            
            # Process document
            result = self.process_document(doc_path, output_path, resume=True)
            
            if result:
                result['input_file'] = doc_path
                result['document_number'] = len(results) + 1
                results.append(result)
                print(f"✅ Completed: {Path(output_path).name}")
            else:
                print(f"❌ Failed to process: {Path(doc_path).name}")
                print(f"💡 You can resume this document later using resume_document_processing()")
            
            # Add delay between documents to avoid rate limiting
            if i < len(new_doc_files):
                delay = self.rate_limit_delay * 2  # Longer delay between documents
                print(f"⏳ Waiting {delay}s before next document...")
                time.sleep(delay)
        
        # Generate summary report
        if results:
            summary_report = self.generate_batch_summary(results, input_dir, output_dir)
            print(f"\n🎉 Batch processing completed!")
            print(f"✅ Processed: {len(results)} documents total")
            print(f"📊 Summary report: {summary_report['report_path']}")
        
        return results
    
    def generate_batch_summary(self, results, input_dir, output_dir):
        """Generate summary report for batch processing"""
        summary = {
            'batch_info': {
                'input_directory': input_dir,
                'output_directory': output_dir,
                'total_documents': len(results),
                'processing_date': datetime.now().isoformat(),
                'total_processing_time': str(datetime.now() - self.stats['start_time'])
            },
            'aggregate_statistics': {
                'total_paragraphs_processed': sum(r['report']['summary']['total_paragraphs'] for r in results if 'report' in r),
                'total_api_calls': sum(r['report']['costs']['total_api_calls'] for r in results if 'report' in r),
                'total_tokens_used': sum(r['report']['costs']['total_tokens_used'] for r in results if 'report' in r),
                'total_protected_content': sum(r['report']['summary']['protected_content'] for r in results if 'report' in r),
                'average_similarity_reduction': round(
                    sum(r['report']['summary']['average_similarity_reduction'] for r in results if 'report' in r) / len(results), 2
                ) if results else 0,
                'average_quality_score': round(
                    sum(r['report']['summary']['average_quality_score'] for r in results if 'report' in r) / len(results), 2
                ) if results else 0
            },
            'document_results': [
                {
                    'filename': Path(r['input_file']).name,
                    'output_file': Path(r['output_path']).name,
                    'paragraphs': r['report']['summary']['total_paragraphs'],
                    'similarity_reduction': r['report']['summary']['average_similarity_reduction'],
                    'quality_score': r['report']['summary']['average_quality_score']
                } for r in results if 'report' in r
            ]
        }
        
        # Save batch summary
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        report_path = os.path.join(output_dir, f"batch_summary_{timestamp}.json")
        
        with open(report_path, 'w', encoding='utf-8') as f:
            json.dump(summary, f, indent=2, ensure_ascii=False)
        
        return {
            'summary': summary,
            'report_path': report_path
        }

def demo_large_file_processing():
    """Demonstrate large file processing capabilities"""
    print("🚀 Full Gemini AI Paraphraser System v2.0")
    print("🎯 Large File Processing Demo")
    print("=" * 60)
    
    # Initialize with custom settings for large files
    paraphraser = FullGeminiParaphraser(
        synonym_file='sinonim.json',
        chunk_size=3000,  # Smaller chunks for better API handling
        max_retries=5     # More retries for reliability
    )
    
    # Check for pending documents first
    print("\n🔍 Checking for interrupted processes...")
    pending = paraphraser.check_pending_documents()
    
    if pending:
        print(f"📋 Found {len(pending)} pending documents:")
        for doc in pending:
            print(f"   📄 {Path(doc['input_path']).name}: {doc['completion_percentage']:.1f}% complete")
        
        response = input("\n❓ Resume pending documents? (y/n): ").lower().strip()
        if response == 'y':
            paraphraser.resume_all_pending()
    
    # Process all documents
    results = paraphraser.process_all_documents(
        input_dir="documents",
        output_dir="completed",
        resume_pending=True
    )
    
    if results:
        print(f"\n🎉 Processing completed! {len(results)} documents processed")
        
        # Display final statistics
        final_stats = paraphraser.get_statistics()
        print("\n📈 Final Statistics:")
        for key, value in final_stats.items():
            print(f"   {key}: {value}")
    else:
        print("❌ No documents were processed")

def process_single_large_file(file_path, chunk_size=5000):
    """Process a single large file with custom chunk size"""
    print(f"🎯 Processing large file: {file_path}")
    print(f"📏 Chunk size: {chunk_size} characters")
    
    paraphraser = FullGeminiParaphraser(
        synonym_file='sinonim.json',
        chunk_size=chunk_size,
        max_retries=5
    )
    
    result = paraphraser.process_document(file_path, resume=True)
    
    if result:
        print(f"✅ Successfully processed: {result['output_path']}")
        return result
    else:
        print(f"❌ Failed to process: {file_path}")
        return None

def main():
    """Main processing function with menu options"""
    import sys
    
    if len(sys.argv) > 1:
        command = sys.argv[1].lower()
        
        if command == "demo":
            demo_large_file_processing()
        elif command == "resume":
            paraphraser = FullGeminiParaphraser(synonym_file='sinonim.json')
            paraphraser.resume_all_pending()
        elif command == "check":
            paraphraser = FullGeminiParaphraser(synonym_file='sinonim.json')
            pending = paraphraser.check_pending_documents()
            if pending:
                print(f"📋 Found {len(pending)} pending documents:")
                for doc in pending:
                    print(f"   📄 {Path(doc['input_path']).name}: {doc['completion_percentage']:.1f}% complete")
            else:
                print("✅ No pending documents found")
        elif command == "core":
            paraphraser = FullGeminiParaphraser(synonym_file='sinonim.json', chunk_size=4000)
            paraphraser.process_all_core_chapters()
        elif command == "core-single" and len(sys.argv) > 2:
            file_path = sys.argv[2]
            paraphraser = FullGeminiParaphraser(synonym_file='sinonim.json', chunk_size=4000)
            result = paraphraser.process_core_chapters_only(file_path)
            if result:
                print(f"✅ Core chapters processed: {result['output_path']}")
                print(f"📊 Processed {result['processed_paragraphs']} paragraphs")
        elif command == "single" and len(sys.argv) > 2:
            file_path = sys.argv[2]
            chunk_size = int(sys.argv[3]) if len(sys.argv) > 3 else 5000
            process_single_large_file(file_path, chunk_size)
        else:
            print("Usage:")
            print("  python full_gemini_paraphraser.py demo    # Run demo with all documents")
            print("  python full_gemini_paraphraser.py core    # Process core chapters only (BAB I-V)")
            print("  python full_gemini_paraphraser.py core-single <file>  # Process single file core chapters")
            print("  python full_gemini_paraphraser.py resume  # Resume pending documents")
            print("  python full_gemini_paraphraser.py check   # Check pending documents")
            print("  python full_gemini_paraphraser.py single <file> [chunk_size]  # Process single file")
    else:
        # Default behavior
        try:
            # Initialize paraphraser with large file support
            print("🚀 Full Gemini AI Paraphraser System v2.0")
            print("🎯 Enhanced with Large File Support & Resume Capability")
            print("=" * 60)
            
            paraphraser = FullGeminiParaphraser(
                synonym_file='sinonim.json',
                chunk_size=4000,  # 4KB chunks for balance
                max_retries=5
            )
            
            # Process all documents with resume capability
            results = paraphraser.process_all_documents(
                input_dir="documents",
                output_dir="completed",
                resume_pending=True
            )
            
            if results:
                print(f"\n🎉 All documents processed successfully! ({len(results)} total)")
                
                # Display final statistics
                final_stats = paraphraser.get_statistics()
                print("\n📈 Final Statistics:")
                for key, value in final_stats.items():
                    print(f"   {key}: {value}")
            else:
                print("❌ No documents were processed successfully")
            
        except KeyboardInterrupt:
            print("\n⏸️  Processing interrupted by user")
            print("💡 You can resume processing by running the script again")
        except Exception as e:
            print(f"❌ Processing error: {e}")
            import traceback
            traceback.print_exc()

if __name__ == "__main__":
    main()