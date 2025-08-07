# integrated_smart_paraphrase_system.py
"""
Integrated Smart Paraphrase System - Complete Solution
Combines: Online Plagiarism Detection + Smart Paraphrasing + Auto Processing
Features: 
- Online plagiarism detection with internet sources
- Smart paraphrasing based on risk levels
- Automated workflow from detection to paraphrasing
- Comprehensive reporting and analytics

Author: DevNoLife
Version: 1.0 - Integrated Edition
"""

import os
import json
import time
from datetime import datetime
from pathlib import Path
import docx
from docx.shared import RGBColor
import shutil

# Import our existing modules
from plagiarism_detector import PlagiarismDetector
from ultimate_hybrid_paraphraser import UltimateHybridParaphraser
from smart_plagiarism_checker import SmartPlagiarismChecker

class IntegratedSmartParaphraseSystem:
    def __init__(self, gemini_api_key=None, mode='smart'):
        print("🚀 Initializing Integrated Smart Paraphrase System...")
        print("🎯 Complete Solution: Detection → Analysis → Paraphrasing → Reporting")
        
        # Initialize components
        self.plagiarism_detector = PlagiarismDetector()
        self.pattern_checker = SmartPlagiarismChecker()
        self.paraphraser = UltimateHybridParaphraser(
            synonym_file='sinonim.json',
            gemini_api_key=gemini_api_key,
            mode=mode
        )
        
        # System configuration
        self.config = {
            'high_risk_threshold': 80,      # >= 80% similarity = HIGH RISK
            'medium_risk_threshold': 50,    # >= 50% similarity = MEDIUM RISK
            'low_risk_threshold': 20,       # >= 20% similarity = LOW RISK
            'auto_paraphrase_threshold': 70, # Auto paraphrase if similarity >= 70%
            'pattern_risk_threshold': 50,   # Pattern analysis threshold
            'min_paragraph_length': 15,     # Minimum words to process
            'backup_original': True,        # Create backup before processing
            'detailed_logging': True        # Enable detailed logs
        }
        
        # Processing statistics
        self.stats = {
            'total_paragraphs': 0,
            'online_checked_paragraphs': 0,
            'pattern_checked_paragraphs': 0,
            'paraphrased_paragraphs': 0,
            'high_risk_found': 0,
            'medium_risk_found': 0,
            'low_risk_found': 0,
            'processing_start_time': None,
            'processing_end_time': None,
            'total_processing_time': 0,
            'online_detection_time': 0,
            'pattern_analysis_time': 0,
            'paraphrasing_time': 0
        }
        
        # Results storage
        self.analysis_results = []
        self.paraphrasing_results = []
        
        print("✅ Online plagiarism detector: Ready")
        print("✅ Pattern-based checker: Ready")
        print("✅ Hybrid paraphraser: Ready")
        print(f"✅ Processing mode: {mode.upper()}")
        print("✅ Integrated system ready!")
    
    def create_backup(self, file_path):
        """Create backup of original document"""
        if not self.config['backup_original']:
            return None
            
        try:
            backup_dir = Path(file_path).parent / "backups"
            backup_dir.mkdir(exist_ok=True)
            
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = Path(file_path).stem
            extension = Path(file_path).suffix
            backup_path = backup_dir / f"{filename}_backup_{timestamp}{extension}"
            
            shutil.copy2(file_path, backup_path)
            print(f"📋 Backup created: {backup_path}")
            return str(backup_path)
            
        except Exception as e:
            print(f"⚠️  Could not create backup: {e}")
            return None
    
    def analyze_document_comprehensive(self, file_path):
        """Comprehensive document analysis: Online + Pattern detection"""
        print("=" * 80)
        print("🔍 COMPREHENSIVE DOCUMENT ANALYSIS")
        print("=" * 80)
        
        if not os.path.exists(file_path):
            print(f"❌ File not found: {file_path}")
            return None
        
        self.stats['processing_start_time'] = datetime.now()
        
        print(f"📄 Analyzing: {os.path.basename(file_path)}")
        
        # Step 1: Online plagiarism detection
        print(f"\n🌐 STEP 1: Online Plagiarism Detection")
        online_start = time.time()
        
        online_results = self.plagiarism_detector.scan_document_for_plagiarism(file_path)
        
        online_end = time.time()
        self.stats['online_detection_time'] = online_end - online_start
        
        if not online_results:
            print("❌ Online detection failed")
            return None
        
        # Step 2: Pattern-based analysis
        print(f"\n🔍 STEP 2: Pattern-Based Risk Analysis")
        pattern_start = time.time()
        
        pattern_results = self.pattern_checker.scan_document(file_path)
        
        pattern_end = time.time()
        self.stats['pattern_analysis_time'] = pattern_end - pattern_start
        
        if not pattern_results:
            print("❌ Pattern analysis failed")
            return None
        
        # Step 3: Combine and analyze results
        print(f"\n🔄 STEP 3: Combining Analysis Results")
        combined_results = self.combine_analysis_results(online_results, pattern_results)
        
        # Update statistics
        self.update_analysis_statistics(combined_results)
        
        # Save combined analysis
        self.save_combined_analysis(combined_results)
        
        print(f"\n✅ Comprehensive analysis completed!")
        return combined_results
    
    def combine_analysis_results(self, online_results, pattern_results):
        """Combine online and pattern analysis results"""
        combined = {
            'filename': online_results['filename'],
            'analysis_timestamp': datetime.now().isoformat(),
            'total_paragraphs': online_results['total_paragraphs'],
            'analysis_summary': {
                'online_analysis': online_results['summary'],
                'pattern_analysis': pattern_results['summary']
            },
            'paragraph_analysis': [],
            'recommendations': [],
            'priority_actions': []
        }
        
        # Create comprehensive paragraph analysis
        for i in range(online_results['total_paragraphs']):
            para_analysis = {
                'paragraph_index': i + 1,
                'online_analysis': None,
                'pattern_analysis': None,
                'combined_risk_level': 'unknown',
                'combined_risk_score': 0,
                'recommendation': 'no_action',
                'priority': 'low'
            }
            
            # Find online analysis for this paragraph
            all_online_paras = (
                online_results.get('high_risk_paragraphs', []) +
                online_results.get('medium_risk_paragraphs', []) +
                online_results.get('low_risk_paragraphs', [])
            )
            
            for online_para in all_online_paras:
                if online_para['paragraph_index'] == i + 1:
                    para_analysis['online_analysis'] = online_para
                    break
            
            # Find pattern analysis for this paragraph
            all_pattern_paras = (
                pattern_results.get('very_high_risk', []) +
                pattern_results.get('high_risk', []) +
                pattern_results.get('medium_risk', []) +
                pattern_results.get('low_risk', [])
            )
            
            for pattern_para in all_pattern_paras:
                if pattern_para['paragraph_index'] == i + 1:
                    para_analysis['pattern_analysis'] = pattern_para
                    break
            
            # Calculate combined risk
            combined_risk = self.calculate_combined_risk(para_analysis)
            para_analysis.update(combined_risk)
            
            combined['paragraph_analysis'].append(para_analysis)
        
        # Generate recommendations
        combined['recommendations'] = self.generate_recommendations(combined)
        
        return combined
    
    def calculate_combined_risk(self, para_analysis):
        """Calculate combined risk score from online and pattern analysis"""
        online_score = 0
        pattern_score = 0
        
        # Get online similarity score
        if para_analysis['online_analysis']:
            online_score = para_analysis['online_analysis'].get('similarity', 0)
        
        # Get pattern risk score
        if para_analysis['pattern_analysis']:
            pattern_score = para_analysis['pattern_analysis']['detailed_analysis']['overall_risk_score']
        
        # Combined scoring algorithm
        # Online detection is weighted more heavily (70%) as it checks actual sources
        # Pattern analysis provides supporting evidence (30%)
        combined_score = (online_score * 0.7) + (pattern_score * 0.3)
        
        # Determine risk level and recommendation
        if combined_score >= self.config['high_risk_threshold']:
            risk_level = 'very_high'
            recommendation = 'immediate_paraphrase'
            priority = 'critical'
        elif combined_score >= self.config['medium_risk_threshold']:
            risk_level = 'high'
            recommendation = 'paraphrase_required' 
            priority = 'high'
        elif combined_score >= self.config['low_risk_threshold']:
            risk_level = 'medium'
            recommendation = 'paraphrase_recommended'
            priority = 'medium'
        else:
            risk_level = 'low'
            recommendation = 'monitor_only'
            priority = 'low'
        
        return {
            'combined_risk_level': risk_level,
            'combined_risk_score': round(combined_score, 2),
            'online_score': online_score,
            'pattern_score': pattern_score,
            'recommendation': recommendation,
            'priority': priority
        }
    
    def generate_recommendations(self, combined_results):
        """Generate actionable recommendations based on combined analysis"""
        recommendations = []
        
        # Count paragraphs by risk level
        risk_counts = {'very_high': 0, 'high': 0, 'medium': 0, 'low': 0}
        
        for para in combined_results['paragraph_analysis']:
            risk_level = para['combined_risk_level']
            if risk_level in risk_counts:
                risk_counts[risk_level] += 1
        
        # Generate specific recommendations
        if risk_counts['very_high'] > 0:
            recommendations.append({
                'priority': 'critical',
                'action': 'immediate_paraphrase',
                'description': f"🚨 {risk_counts['very_high']} paragraphs require immediate paraphrasing",
                'paragraphs': risk_counts['very_high']
            })
        
        if risk_counts['high'] > 0:
            recommendations.append({
                'priority': 'high',
                'action': 'paraphrase_required',
                'description': f"⚠️ {risk_counts['high']} paragraphs need paraphrasing",
                'paragraphs': risk_counts['high']
            })
        
        if risk_counts['medium'] > 0:
            recommendations.append({
                'priority': 'medium',
                'action': 'paraphrase_recommended',
                'description': f"💡 {risk_counts['medium']} paragraphs would benefit from paraphrasing",
                'paragraphs': risk_counts['medium']
            })
        
        # Overall document recommendation
        total_problematic = risk_counts['very_high'] + risk_counts['high']
        total_paragraphs = sum(risk_counts.values())
        
        if total_paragraphs > 0:
            risk_percentage = (total_problematic / total_paragraphs) * 100
            
            if risk_percentage > 30:
                recommendations.insert(0, {
                    'priority': 'critical',
                    'action': 'document_revision',
                    'description': f"🔴 Document has {risk_percentage:.1f}% high-risk content - Extensive revision needed",
                    'risk_percentage': risk_percentage
                })
            elif risk_percentage > 15:
                recommendations.insert(0, {
                    'priority': 'high',
                    'action': 'selective_paraphrasing',
                    'description': f"🟡 Document has {risk_percentage:.1f}% problematic content - Targeted paraphrasing needed",
                    'risk_percentage': risk_percentage
                })
        
        return recommendations
    
    def auto_paraphrase_document(self, file_path, combined_analysis):
        """Automatically paraphrase high-risk paragraphs"""
        print("\n" + "=" * 80)
        print("🤖 AUTO-PARAPHRASING HIGH-RISK CONTENT")
        print("=" * 80)
        
        if not os.path.exists(file_path):
            print(f"❌ File not found: {file_path}")
            return None
        
        paraphrase_start = time.time()
        
        try:
            doc = docx.Document(file_path)
            
            paraphrase_results = {
                'filename': os.path.basename(file_path),
                'paraphrase_timestamp': datetime.now().isoformat(),
                'total_paragraphs': len(doc.paragraphs),
                'paragraphs_processed': 0,
                'paragraphs_paraphrased': 0,
                'skipped_paragraphs': 0,
                'paraphrase_details': [],
                'summary': {}
            }
            
            processed_count = 0
            paraphrased_count = 0
            
            print(f"📄 Processing: {len(doc.paragraphs)} paragraphs")
            
            for para_analysis in combined_analysis['paragraph_analysis']:
                para_index = para_analysis['paragraph_index'] - 1  # 0-based for docx
                
                if para_index >= len(doc.paragraphs):
                    continue
                
                paragraph = doc.paragraphs[para_index]
                para_text = paragraph.text.strip()
                
                # Skip empty or very short paragraphs
                if not para_text or len(para_text.split()) < self.config['min_paragraph_length']:
                    continue
                
                processed_count += 1
                
                # Determine if paraphrasing is needed
                should_paraphrase = False
                reason = ""
                
                combined_score = para_analysis['combined_risk_score']
                recommendation = para_analysis['recommendation']
                
                if recommendation in ['immediate_paraphrase', 'paraphrase_required']:
                    should_paraphrase = True
                    reason = f"High risk (score: {combined_score:.1f})"
                elif recommendation == 'paraphrase_recommended' and combined_score >= self.config['auto_paraphrase_threshold']:
                    should_paraphrase = True
                    reason = f"Above auto-paraphrase threshold ({combined_score:.1f})"
                
                print(f"\n📄 Paragraph {para_analysis['paragraph_index']}: ", end="")
                
                if should_paraphrase:
                    print(f"🤖 Paraphrasing ({reason})")
                    
                    # Choose paraphrasing mode based on risk level
                    if combined_score >= 90:
                        self.paraphraser.switch_mode('aggressive')
                        aggressiveness = 0.9
                    elif combined_score >= 75:
                        self.paraphraser.switch_mode('balanced')
                        aggressiveness = 0.7
                    else:
                        self.paraphraser.switch_mode('smart')
                        aggressiveness = 0.6
                    
                    # Paraphrase the paragraph
                    paraphrase_result = self.paraphraser.process_paragraph_ultimate(
                        para_text,
                        paragraph_index=para_analysis['paragraph_index'],
                        total_paragraphs=len(doc.paragraphs),
                        aggressiveness=aggressiveness
                    )
                    
                    if paraphrase_result and paraphrase_result['paraphrase']:
                        # Replace paragraph text
                        paragraph.text = paraphrase_result['paraphrase']
                        
                        # Highlight the changed paragraph (optional)
                        for run in paragraph.runs:
                            run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                        
                        paraphrased_count += 1
                        
                        # Store paraphrasing details
                        paraphrase_details = {
                            'paragraph_index': para_analysis['paragraph_index'],
                            'original_text': para_text[:100] + '...' if len(para_text) > 100 else para_text,
                            'paraphrased_text': paraphrase_result['paraphrase'][:100] + '...' if len(paraphrase_result['paraphrase']) > 100 else paraphrase_result['paraphrase'],
                            'original_similarity': combined_score,
                            'new_similarity': paraphrase_result['similarity'],
                            'improvement': combined_score - paraphrase_result['similarity'],
                            'method_used': paraphrase_result['method'],
                            'reason': reason
                        }
                        
                        paraphrase_results['paraphrase_details'].append(paraphrase_details)
                        
                        print(f"   ✅ Success: {combined_score:.1f}% → {paraphrase_result['similarity']:.1f}% (-{paraphrase_details['improvement']:.1f}%)")
                        
                    else:
                        print(f"   ❌ Failed to paraphrase")
                else:
                    print(f"✅ Skipped ({combined_score:.1f}% - acceptable)")
            
            # Update results
            paraphrase_results['paragraphs_processed'] = processed_count
            paraphrase_results['paragraphs_paraphrased'] = paraphrased_count
            paraphrase_results['skipped_paragraphs'] = processed_count - paraphrased_count
            
            # Save the modified document
            output_info = self.save_paraphrased_document(doc, file_path)
            if output_info:
                paraphrase_results['output_files'] = output_info
                paraphrase_results['output_file'] = output_info['paraphrased_file']  # Backward compatibility
            
            # Update statistics
            paraphrase_end = time.time()
            self.stats['paraphrasing_time'] = paraphrase_end - paraphrase_start
            self.stats['paraphrased_paragraphs'] = paraphrased_count
            
            # Generate summary
            self.generate_paraphrase_summary(paraphrase_results)
            
            # Save paraphrase report
            self.save_paraphrase_report(paraphrase_results)
            
            return paraphrase_results
            
        except Exception as e:
            print(f"❌ Error during auto-paraphrasing: {e}")
            return None
    
    def save_paraphrased_document(self, doc, original_path):
        """Save the paraphrased document in 'completed' folder"""
        try:
            path_obj = Path(original_path)
            
            # Create 'completed' folder if it doesn't exist
            completed_dir = Path("completed")
            completed_dir.mkdir(exist_ok=True)
            
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            output_filename = f"{path_obj.stem}_paraphrased_{timestamp}{path_obj.suffix}"
            output_path = completed_dir / output_filename
            
            # Save paraphrased document
            doc.save(str(output_path))
            
            # Also copy original document to completed folder for comparison
            original_copy = completed_dir / f"{path_obj.stem}_ORIGINAL_{timestamp}{path_obj.suffix}"
            shutil.copy2(original_path, str(original_copy))
            
            print(f"\n💾 COMPLETED DOCUMENTS:")
            print(f"   📄 Original: {original_copy}")
            print(f"   🤖 Paraphrased: {output_path}")
            print(f"   📁 Location: completed/ folder")
            
            return {
                'paraphrased_file': str(output_path),
                'original_copy': str(original_copy),
                'completed_folder': str(completed_dir)
            }
            
        except Exception as e:
            print(f"❌ Error saving paraphrased document: {e}")
            return None
    
    def generate_paraphrase_summary(self, results):
        """Generate paraphrasing summary"""
        print(f"\n" + "=" * 80)
        print("📊 PARAPHRASING SUMMARY")
        print("=" * 80)
        
        print(f"📄 Document: {results['filename']}")
        print(f"📊 Total paragraphs: {results['total_paragraphs']}")
        print(f"🔍 Processed: {results['paragraphs_processed']}")
        print(f"🤖 Paraphrased: {results['paragraphs_paraphrased']}")
        print(f"⏭️  Skipped: {results['skipped_paragraphs']}")
        
        if results['paragraphs_paraphrased'] > 0:
            print(f"\n🎯 PARAPHRASING RESULTS:")
            
            total_improvement = sum(detail['improvement'] for detail in results['paraphrase_details'] if detail['improvement'] > 0)
            avg_improvement = total_improvement / results['paragraphs_paraphrased']
            
            print(f"   📈 Average improvement: {avg_improvement:.1f}% similarity reduction")
            print(f"   🎉 Total improvement: {total_improvement:.1f}% cumulative reduction")
            
            print(f"\n🔝 TOP IMPROVEMENTS:")
            sorted_details = sorted(results['paraphrase_details'], key=lambda x: x['improvement'], reverse=True)
            for detail in sorted_details[:5]:
                print(f"   Para {detail['paragraph_index']}: {detail['original_similarity']:.1f}% → {detail['new_similarity']:.1f}% (-{detail['improvement']:.1f}%)")
        
        print("=" * 80)
    
    def process_document_complete(self, file_path, auto_paraphrase=True):
        """Complete document processing: Analysis + Paraphrasing + Reporting"""
        print("🚀 STARTING COMPLETE DOCUMENT PROCESSING")
        print("🎯 Workflow: Backup → Analysis → Paraphrasing → Reporting")
        
        # Step 1: Create backup
        backup_path = self.create_backup(file_path)
        
        # Step 2: Comprehensive analysis
        combined_analysis = self.analyze_document_comprehensive(file_path)
        
        if not combined_analysis:
            print("❌ Analysis failed - stopping process")
            return None
        
        # Step 3: Auto-paraphrasing (if enabled)
        paraphrase_results = None
        if auto_paraphrase:
            paraphrase_results = self.auto_paraphrase_document(file_path, combined_analysis)
        
        # Step 4: Final report
        final_report = self.generate_final_report(combined_analysis, paraphrase_results, backup_path)
        
        # Step 5: Update final statistics
        self.stats['processing_end_time'] = datetime.now()
        self.stats['total_processing_time'] = (
            self.stats['processing_end_time'] - self.stats['processing_start_time']
        ).total_seconds()
        
        # Print final summary
        self.print_final_summary(final_report)
        
        return final_report
    
    def generate_final_report(self, analysis_results, paraphrase_results, backup_path):
        """Generate comprehensive final report"""
        report = {
            'document_info': {
                'filename': analysis_results['filename'],
                'processing_timestamp': datetime.now().isoformat(),
                'backup_path': backup_path
            },
            'analysis_results': analysis_results,
            'paraphrase_results': paraphrase_results,
            'processing_statistics': self.stats.copy(),
            'recommendations': analysis_results.get('recommendations', []),
            'next_steps': []
        }
        
        # Generate next steps
        if paraphrase_results:
            if paraphrase_results['paragraphs_paraphrased'] > 0:
                report['next_steps'].append("✅ Document has been automatically paraphrased")
                report['next_steps'].append("🔍 Recommend re-scanning to verify improvement")
                report['next_steps'].append("📝 Review highlighted paragraphs for quality")
            else:
                report['next_steps'].append("ℹ️ No paraphrasing was needed")
        else:
            if analysis_results.get('recommendations'):
                high_risk_recs = [r for r in analysis_results['recommendations'] if r['priority'] in ['critical', 'high']]
                if high_risk_recs:
                    report['next_steps'].append("⚠️ Manual paraphrasing recommended for high-risk content")
                    report['next_steps'].append("🤖 Use ultimate_hybrid_paraphraser.py for specific paragraphs")
        
        # Save comprehensive report
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        report_file = f"integrated_analysis_report_{timestamp}.json"
        
        try:
            with open(report_file, 'w', encoding='utf-8') as f:
                json.dump(report, f, ensure_ascii=False, indent=2, default=str)
            
            report['report_file'] = report_file
            print(f"\n📋 Comprehensive report saved: {report_file}")
            
        except Exception as e:
            print(f"⚠️ Could not save comprehensive report: {e}")
        
        return report
    
    def print_final_summary(self, final_report):
        """Print final processing summary"""
        print("\n" + "=" * 80)
        print("🎉 INTEGRATED PROCESSING COMPLETE")
        print("=" * 80)
        
        # Document info
        print(f"📄 Document: {final_report['document_info']['filename']}")
        if final_report['document_info']['backup_path']:
            print(f"📋 Backup: {os.path.basename(final_report['document_info']['backup_path'])}")
        
        # Processing statistics
        stats = final_report['processing_statistics']
        print(f"\n⏱️ PROCESSING TIME:")
        print(f"   🌐 Online detection: {stats['online_detection_time']:.1f}s")
        print(f"   🔍 Pattern analysis: {stats['pattern_analysis_time']:.1f}s")
        if stats['paraphrasing_time'] > 0:
            print(f"   🤖 Paraphrasing: {stats['paraphrasing_time']:.1f}s")
        print(f"   📊 Total time: {stats['total_processing_time']:.1f}s")
        
        # Analysis results
        analysis = final_report['analysis_results']
        online_summary = analysis['analysis_summary']['online_analysis']
        pattern_summary = analysis['analysis_summary']['pattern_analysis']
        
        print(f"\n📊 ANALYSIS RESULTS:")
        print(f"   🌐 Online detection: {online_summary.get('needs_paraphrasing', 0)} paragraphs need attention")
        print(f"   🔍 Pattern analysis: {pattern_summary.get('needs_paraphrasing', 0)} paragraphs flagged")
        
        # Paraphrasing results
        if final_report['paraphrase_results']:
            paraphrase = final_report['paraphrase_results']
            print(f"   🤖 Auto-paraphrased: {paraphrase['paragraphs_paraphrased']} paragraphs")
            if paraphrase.get('output_files'):
                output_files = paraphrase['output_files']
                print(f"   📁 Results in completed/ folder:")
                print(f"      📄 Original: {os.path.basename(output_files['original_copy'])}")
                print(f"      🤖 Paraphrased: {os.path.basename(output_files['paraphrased_file'])}")
            elif paraphrase.get('output_file'):
                print(f"   💾 Output file: {os.path.basename(paraphrase['output_file'])}")
        
        # Recommendations
        if final_report['recommendations']:
            print(f"\n💡 KEY RECOMMENDATIONS:")
            for rec in final_report['recommendations'][:3]:  # Top 3 recommendations
                print(f"   {rec['description']}")
        
        # Next steps
        if final_report['next_steps']:
            print(f"\n🎯 NEXT STEPS:")
            for step in final_report['next_steps']:
                print(f"   {step}")
        
        print("\n" + "=" * 80)
        print("✨ All processing completed successfully!")
        print("📋 Check the comprehensive report for detailed analysis")
        print("=" * 80)
    
    def update_analysis_statistics(self, combined_results):
        """Update analysis statistics"""
        self.stats['total_paragraphs'] = combined_results['total_paragraphs']
        
        for para in combined_results['paragraph_analysis']:
            risk_level = para['combined_risk_level']
            
            if para['online_analysis']:
                self.stats['online_checked_paragraphs'] += 1
            
            if para['pattern_analysis']:
                self.stats['pattern_checked_paragraphs'] += 1
            
            if risk_level == 'very_high':
                self.stats['high_risk_found'] += 1
            elif risk_level == 'high':
                self.stats['medium_risk_found'] += 1
            elif risk_level == 'medium':
                self.stats['low_risk_found'] += 1
    
    def save_combined_analysis(self, combined_results):
        """Save combined analysis results"""
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        report_file = f"combined_analysis_{timestamp}.json"
        
        try:
            with open(report_file, 'w', encoding='utf-8') as f:
                json.dump(combined_results, f, ensure_ascii=False, indent=2, default=str)
            
            print(f"📋 Combined analysis saved: {report_file}")
            return report_file
            
        except Exception as e:
            print(f"⚠️ Could not save combined analysis: {e}")
            return None
    
    def save_paraphrase_report(self, paraphrase_results):
        """Save paraphrasing report"""
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        report_file = f"paraphrase_report_{timestamp}.json"
        
        try:
            with open(report_file, 'w', encoding='utf-8') as f:
                json.dump(paraphrase_results, f, ensure_ascii=False, indent=2, default=str)
            
            print(f"📋 Paraphrasing report saved: {report_file}")
            return report_file
            
        except Exception as e:
            print(f"⚠️ Could not save paraphrasing report: {e}")
            return None


def find_docx_files():
    """Find all .docx files in the project directory with priority order"""
    docx_files = []
    
    # Priority order: documents/ folder first, then others
    priority_folders = ['./documents', '.']
    other_files = []
    
    # Search in current directory and subdirectories
    for root, dirs, files in os.walk('.'):
        # Skip backup and temporary directories
        if any(skip in root for skip in ['backup', '__pycache__', '.git', 'venv']):
            continue
            
        for file in files:
            if file.endswith('.docx') and not file.startswith('~'):
                full_path = os.path.join(root, file)
                
                # Prioritize files from documents/ folder
                if root == './documents':
                    docx_files.insert(0, full_path)  # Add to beginning
                else:
                    docx_files.append(full_path)  # Add to end
    
    return docx_files

def select_document(auto_select=False):
    """Let user select document to process"""
    print("🔍 Searching for .docx documents...")
    
    docx_files = find_docx_files()
    
    if not docx_files:
        print("❌ No .docx files found in the project directory")
        print("💡 Please add your document to the project folder")
        return None
    
    print(f"📄 Found {len(docx_files)} document(s):")
    print("-" * 50)
    
    for i, file_path in enumerate(docx_files, 1):
        file_size = os.path.getsize(file_path) / 1024  # KB
        mod_time = datetime.fromtimestamp(os.path.getmtime(file_path)).strftime("%Y-%m-%d %H:%M")
        print(f"  {i}. {os.path.basename(file_path)}")
        print(f"     📂 {os.path.dirname(file_path)}")
        print(f"     📏 {file_size:.1f} KB | 🕒 {mod_time}")
        print()
    
    # Auto-select first document if in auto mode or non-interactive
    if auto_select or len(docx_files) == 1:
        selected_file = docx_files[0]
        print(f"✅ Auto-selected: {os.path.basename(selected_file)}")
        return selected_file
    
    # Interactive selection
    while True:
        try:
            choice = input(f"Select document (1-{len(docx_files)}) [default: 1]: ").strip() or '1'
            index = int(choice) - 1
            
            if 0 <= index < len(docx_files):
                selected_file = docx_files[index]
                print(f"✅ Selected: {os.path.basename(selected_file)}")
                return selected_file
            else:
                print(f"❌ Please enter a number between 1 and {len(docx_files)}")
                
        except (ValueError, EOFError):
            # Fallback to auto-select if input fails
            selected_file = docx_files[0]
            print(f"✅ Auto-selected (input unavailable): {os.path.basename(selected_file)}")
            return selected_file

def main():
    """Main function for integrated smart paraphrase system"""
    print("🚀 INTEGRATED SMART PARAPHRASE SYSTEM")
    print("🎯 Complete Solution: Detection → Analysis → Paraphrasing → Reporting")
    print("=" * 80)
    
    # Let user select document to process (with auto-fallback)
    document_path = select_document(auto_select=True)
    
    if not document_path:
        return
    
    # Initialize integrated system
    print("\n🔧 Initializing integrated system...")
    
    # You can change the mode here: 'smart', 'balanced', 'aggressive', 'turnitin_safe'
    system = IntegratedSmartParaphraseSystem(mode='smart')
    
    # Process document completely
    print(f"\n🎯 Processing document: {os.path.basename(document_path)}")
    
    # Auto-processing options (with fallback for non-interactive mode)
    try:
        print(f"\n🤔 Processing options:")
        print(f"   1. Analysis only (no auto-paraphrasing)")
        print(f"   2. Complete processing (analysis + auto-paraphrasing)")
        
        choice = input("Choose option (1/2) [default: 2]: ").strip() or '2'
        auto_paraphrase = choice == '2'
    except EOFError:
        # Auto-fallback for non-interactive environments
        print("🤖 Auto-mode: Running complete processing (analysis + auto-paraphrasing)")
        auto_paraphrase = True
    
    # Start processing
    final_report = system.process_document_complete(document_path, auto_paraphrase=auto_paraphrase)
    
    if final_report:
        print(f"\n🎉 Processing completed successfully!")
        print(f"📋 Check the reports for detailed analysis")
        
        # Show quick stats
        if final_report['paraphrase_results'] and final_report['paraphrase_results']['paragraphs_paraphrased'] > 0:
            print(f"🤖 Auto-paraphrased: {final_report['paraphrase_results']['paragraphs_paraphrased']} paragraphs")
            if 'output_files' in final_report['paraphrase_results']:
                output_files = final_report['paraphrase_results']['output_files']
                print(f"📁 Check 'completed/' folder for:")
                print(f"   📄 Original: {os.path.basename(output_files['original_copy'])}")
                print(f"   🤖 Paraphrased: {os.path.basename(output_files['paraphrased_file'])}")
            elif 'output_file' in final_report['paraphrase_results']:
                print(f"💾 Output file: {os.path.basename(final_report['paraphrase_results']['output_file'])}")
        
    else:
        print("❌ Processing failed. Please check the error messages above.")


if __name__ == "__main__":
    main()