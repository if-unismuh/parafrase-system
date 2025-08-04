# smart_hybrid_paraphraser.py
"""
Smart Hybrid Paraphraser - Local Indonesian NLP + Gemini AI
Hemat biaya tapi tetap akurat dengan AI routing yang cerdas
Author: DevNoLife
Version: 2.0
"""

import os
import json
import re
import time
from datetime import datetime
from pathlib import Path
import docx
from docx.shared import RGBColor
from docx.enum.text import WD_COLOR_INDEX
import shutil
from collections import defaultdict
import requests
import hashlib

# Import sistem parafrase lokal
from main_paraphrase_system import IndonesianParaphraseSystem

class SmartHybridParaphraser:
    def __init__(self, synonym_file='sinonim.json', gemini_api_key=None):
        print("🚀 Initializing Smart Hybrid Paraphraser...")
        
        # Initialize local paraphrase system
        self.local_paraphraser = IndonesianParaphraseSystem(synonym_file)
        
        # Gemini AI configuration
        self.gemini_api_key = "AIzaSyAABOimEqGMOJUnB302ARIB_tzUVxl5JDA" if gemini_api_key else None
        self.gemini_api_url = "https://generativelanguage.googleapis.com/v1beta/models/gemini-2.0-flash:generateContent"
        
        # Cost tracking
        self.cost_tracker = {
            'local_calls': 0,
            'ai_calls': 0,
            'ai_tokens_used': 0,
            'estimated_cost_usd': 0.0,
            'cache_hits': 0
        }
        
        # AI result cache to avoid duplicate calls
        self.ai_cache = {}
        
        # Smart routing thresholds
        self.config = {
            'local_confidence_threshold': 0.25,  # If local achieves >25% reduction, don't use AI
            'complexity_threshold': 0.7,         # High complexity → Use AI
            'plagiarism_risk_threshold': 0.8,    # High risk → Use AI
            'min_paragraph_length': 40,          # Minimum words for AI processing
            'max_batch_size': 5,                 # Max paragraphs per AI batch call
            'ai_retry_attempts': 2               # Retry failed AI calls
        }
        
        # Academic patterns that benefit from AI
        self.ai_priority_patterns = [
            r'menurut\s+\w+\s*\(\d{4}\)',  # Citations
            r'berdasarkan\s+penelitian',    # Research references  
            r'definisi\s+\w+\s+adalah',     # Definitions
            r'konsep\s+\w+\s+merupakan',    # Concepts
            r'teori\s+\w+\s+menyatakan',    # Theories
            r'sistem\s+informasi\s+adalah', # Technical definitions
        ]
        
        print(f"✅ Local system ready with {len(self.local_paraphraser.synonyms)} synonyms")
        print(f"✅ Gemini AI {'configured' if gemini_api_key else 'not configured (will use local only)'}")
        print("✅ Smart Hybrid Paraphraser ready!")
    
    def calculate_paragraph_complexity(self, text):
        """Calculate complexity score for smart routing"""
        words = text.split()
        sentences = text.split('.')
        
        # Length factor
        length_score = min(len(words) / 100, 1.0)  # Normalize to 0-1
        
        # Academic vocabulary density
        academic_words = [
            'penelitian', 'analisis', 'metode', 'sistem', 'implementasi',
            'evaluasi', 'teori', 'konsep', 'pendekatan', 'metodologi'
        ]
        academic_density = sum(1 for word in words if word.lower() in academic_words) / len(words)
        
        # Citation/reference density
        citation_count = len(re.findall(r'\(\d{4}\)|et al|vol\.|no\.', text))
        citation_score = min(citation_count / 3, 1.0)
        
        # Technical pattern density
        pattern_score = sum(1 for pattern in self.ai_priority_patterns if re.search(pattern, text.lower())) / len(self.ai_priority_patterns)
        
        # Average sentence length
        avg_sentence_length = len(words) / len(sentences) if sentences else 0
        sentence_complexity = min(avg_sentence_length / 25, 1.0)
        
        # Weighted complexity score
        complexity = (
            length_score * 0.2 +
            academic_density * 0.3 +
            citation_score * 0.2 +
            pattern_score * 0.2 +
            sentence_complexity * 0.1
        )
        
        return min(complexity, 1.0)
    
    def should_use_ai(self, paragraph_text, local_result):
        """Smart routing decision: Local or AI?"""
        # Always try local first - if it's good enough, don't use AI
        if local_result['plagiarism_reduction'] >= self.config['local_confidence_threshold'] * 100:
            return False, "Local result sufficient"
        
        # Check paragraph length
        word_count = len(paragraph_text.split())
        if word_count < self.config['min_paragraph_length']:
            return False, "Paragraph too short"
        
        # Calculate complexity
        complexity = self.calculate_paragraph_complexity(paragraph_text)
        
        # High complexity → Use AI
        if complexity >= self.config['complexity_threshold']:
            return True, f"High complexity ({complexity:.2f})"
        
        # Check for academic patterns that benefit from AI
        pattern_matches = sum(1 for pattern in self.ai_priority_patterns if re.search(pattern, paragraph_text.lower()))
        if pattern_matches >= 2:
            return True, f"Multiple academic patterns ({pattern_matches})"
        
        # High plagiarism risk based on local analysis
        if local_result['similarity'] >= self.config['plagiarism_risk_threshold'] * 100:
            return True, f"High plagiarism risk ({local_result['similarity']:.1f}%)"
        
        return False, "Local processing adequate"
    
    def create_cache_key(self, text):
        """Create cache key for AI results"""
        return hashlib.md5(text.encode('utf-8')).hexdigest()[:16]
    
    def call_gemini_api(self, paragraphs_batch):
        """Call Gemini API for batch processing"""
        if not self.gemini_api_key:
            return None
        
        try:
            # Create prompt for batch processing
            prompt = self.create_gemini_prompt(paragraphs_batch)
            
            headers = {
                'Content-Type': 'application/json',
            }
            
            data = {
                "contents": [{
                    "parts": [{
                        "text": prompt
                    }]
                }],
                "generationConfig": {
                    "temperature": 0.3,  # Lower temperature for more consistent results
                    "topK": 40,
                    "topP": 0.95,
                    "maxOutputTokens": 4096,
                }
            }
            
            # Make API call with retry logic
            for attempt in range(self.config['ai_retry_attempts']):
                try:
                    response = requests.post(
                        f"{self.gemini_api_url}?key={self.gemini_api_key}",
                        headers=headers,
                        json=data,
                        timeout=30
                    )
                    
                    if response.status_code == 200:
                        result = response.json()
                        if 'candidates' in result and result['candidates']:
                            ai_response = result['candidates'][0]['parts'][0]['text']
                            
                            # Track usage
                            self.cost_tracker['ai_calls'] += 1
                            # Estimate tokens (rough: 1 token ≈ 4 chars)
                            tokens_used = len(prompt + ai_response) // 4
                            self.cost_tracker['ai_tokens_used'] += tokens_used
                            self.cost_tracker['estimated_cost_usd'] += tokens_used * 0.000002  # Rough estimate
                            
                            return self.parse_gemini_response(ai_response, paragraphs_batch)
                    
                    else:
                        print(f"⚠️  Gemini API error {response.status_code}: {response.text}")
                        
                except requests.exceptions.RequestException as e:
                    print(f"⚠️  API call attempt {attempt + 1} failed: {e}")
                    if attempt < self.config['ai_retry_attempts'] - 1:
                        time.sleep(2 ** attempt)  # Exponential backoff
            
            return None
            
        except Exception as e:
            print(f"❌ Gemini API error: {e}")
            return None
    
    def create_gemini_prompt(self, paragraphs_batch):
        """Create optimized prompt for Gemini API"""
        prompt = """Kamu adalah expert paraphrasing untuk teks akademik bahasa Indonesia. 

TUGAS: Parafrase paragraf berikut untuk mengurangi plagiarisme sambil mempertahankan makna dan tone akademik.

ATURAN:
1. Gunakan sinonim yang tepat dan natural
2. Ubah struktur kalimat jika perlu
3. Pertahankan makna asli 100%
4. Pertahankan tone formal akademik
5. Jangan hilangkan informasi penting
6. Hasil harus natural dan mudah dibaca

FORMAT OUTPUT:
Untuk setiap paragraf, berikan:
PARAGRAF_[NOMOR]:
[Hasil parafrase]

INPUT PARAGRAPHS:
"""
        
        for i, paragraph in enumerate(paragraphs_batch, 1):
            prompt += f"\nPARAGRAF_{i}:\n{paragraph['text']}\n"
        
        prompt += "\nMulai parafrase sekarang:"
        
        return prompt
    
    def parse_gemini_response(self, ai_response, original_paragraphs):
        """Parse Gemini response back to structured format"""
        results = []
        
        try:
            # Split response by paragraph markers
            sections = re.split(r'PARAGRAF_(\d+):', ai_response)
            
            for i in range(1, len(sections), 2):
                if i + 1 < len(sections):
                    paragraph_num = int(sections[i]) - 1
                    paraphrased_text = sections[i + 1].strip()
                    
                    if paragraph_num < len(original_paragraphs):
                        original_text = original_paragraphs[paragraph_num]['text']
                        
                        # Calculate similarity
                        similarity = self.local_paraphraser.calculate_similarity(original_text, paraphrased_text)
                        
                        result = {
                            'paraphrase': paraphrased_text,
                            'similarity': round(similarity, 2),
                            'plagiarism_reduction': round(100 - similarity, 2),
                            'changes_made': 1,  # AI made changes
                            'method': 'gemini_ai',
                            'status': self.local_paraphraser._get_plagiarism_status(similarity)
                        }
                        
                        results.append(result)
            
            return results
            
        except Exception as e:
            print(f"⚠️  Error parsing Gemini response: {e}")
            return []
    
    def process_paragraph_smart(self, paragraph_text):
        """Smart processing: Local first, then AI if needed"""
        # Step 1: Try local processing first
        local_result = self.local_paraphraser.generate_paraphrase(paragraph_text, aggressiveness=0.6)
        self.cost_tracker['local_calls'] += 1
        
        # Step 2: Decide if AI is needed
        use_ai, reason = self.should_use_ai(paragraph_text, local_result)
        
        if not use_ai:
            local_result['method'] = 'local_only'
            local_result['routing_reason'] = reason
            return local_result
        
        # Step 3: Check cache first
        cache_key = self.create_cache_key(paragraph_text)
        if cache_key in self.ai_cache:
            self.cost_tracker['cache_hits'] += 1
            cached_result = self.ai_cache[cache_key].copy()
            cached_result['method'] = 'ai_cached'
            cached_result['routing_reason'] = reason
            return cached_result
        
        # Step 4: Use AI
        print(f"    🤖 Using AI: {reason}")
        
        ai_results = self.call_gemini_api([{'text': paragraph_text}])
        
        if ai_results and len(ai_results) > 0:
            ai_result = ai_results[0]
            ai_result['routing_reason'] = reason
            
            # Cache the result
            self.ai_cache[cache_key] = ai_result.copy()
            
            # Choose best result: AI vs Local
            if ai_result['plagiarism_reduction'] > local_result['plagiarism_reduction']:
                return ai_result
            else:
                # If AI doesn't improve much, use local + mark as ai_attempted
                local_result['method'] = 'local_after_ai'
                local_result['routing_reason'] = f"{reason} (AI didn't improve)"
                return local_result
        
        # Step 5: Fallback to local if AI fails
        local_result['method'] = 'local_fallback'
        local_result['routing_reason'] = f"{reason} (AI failed)"
        return local_result
    
    def process_batch_documents(self, input_folder, aggressiveness=0.6, create_backup=True):
        """Process batch with smart hybrid approach"""
        start_time = datetime.now()
        
        print("=" * 80)
        print("🎯 SMART HYBRID PARAPHRASER (Local + Gemini AI)")
        print("=" * 80)
        
        # Setup similar to previous batch processor
        if not os.path.exists(input_folder):
            print(f"❌ Input folder not found: {input_folder}")
            return
        
        docx_files = [f for f in os.listdir(input_folder) if f.endswith('.docx') and not f.startswith('~')]
        
        if not docx_files:
            print(f"❌ No Word documents found in: {input_folder}")
            return
        
        print(f"📁 Input folder: {input_folder}")
        print(f"📄 Found {len(docx_files)} documents")
        print(f"⚙️  Aggressiveness level: {aggressiveness}")
        print(f"🤖 AI Integration: {'Enabled' if self.gemini_api_key else 'Disabled (Local only)'}")
        
        # Create backup
        backup_folder = None
        if create_backup:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            backup_folder = f"{input_folder}_backup_{timestamp}"
            
            try:
                if not os.path.exists(backup_folder):
                    os.makedirs(backup_folder)
                
                for file in docx_files:
                    src = os.path.join(input_folder, file)
                    dst = os.path.join(backup_folder, file)
                    shutil.copy2(src, dst)
                
                print(f"✅ Backup created: {backup_folder}")
            except Exception as e:
                print(f"❌ Error creating backup: {e}")
                return
        
        print(f"\n🚀 Starting smart hybrid processing...")
        
        # Process each document
        total_stats = {
            'documents_processed': 0,
            'paragraphs_processed': 0,
            'local_only': 0,
            'ai_enhanced': 0,
            'total_changes': 0
        }
        
        for i, filename in enumerate(docx_files, 1):
            file_path = os.path.join(input_folder, filename)
            
            print(f"\n{'='*20} Document {i}/{len(docx_files)} {'='*20}")
            
            doc_stats = self.process_single_document(file_path, aggressiveness)
            
            if doc_stats:
                total_stats['documents_processed'] += 1
                total_stats['paragraphs_processed'] += doc_stats['processed_paragraphs']
                total_stats['local_only'] += doc_stats.get('local_only', 0)
                total_stats['ai_enhanced'] += doc_stats.get('ai_enhanced', 0)
                total_stats['total_changes'] += doc_stats['changes_made']
        
        # Calculate processing time
        end_time = datetime.now()
        processing_time = (end_time - start_time).total_seconds()
        
        # Print comprehensive summary
        self.print_hybrid_summary(total_stats, processing_time, backup_folder)
        
        # Generate detailed report
        self.generate_hybrid_report(input_folder, total_stats, processing_time)
    
    def process_single_document(self, file_path, aggressiveness):
        """Process single document with smart routing"""
        try:
            print(f"\n📄 Processing: {os.path.basename(file_path)}")
            
            doc = docx.Document(file_path)
            
            doc_stats = {
                'filename': os.path.basename(file_path),
                'total_paragraphs': 0,
                'processed_paragraphs': 0,
                'changes_made': 0,
                'local_only': 0,
                'ai_enhanced': 0,
                'ai_cached': 0,
                'methods_used': defaultdict(int)
            }
            
            current_section = 'UNKNOWN'
            
            for i, paragraph in enumerate(doc.paragraphs):
                para_text = paragraph.text.strip()
                doc_stats['total_paragraphs'] += 1
                
                if not para_text:
                    continue
                
                # Detect section headers
                if self.is_section_header(para_text):
                    current_section = para_text
                    print(f"  📍 Section: {current_section}")
                    continue
                
                # Check if suitable for processing
                if not self.is_paragraph_suitable_for_paraphrasing(para_text):
                    continue
                
                # Smart hybrid processing
                result = self.process_paragraph_smart(para_text)
                
                # Update document if improvement is significant
                if result['plagiarism_reduction'] > 15:  # Lower threshold than before
                    paragraph.clear()
                    paragraph.add_run(result['paraphrase'])
                    
                    # Color coding based on method
                    if 'ai' in result['method']:
                        # Blue highlight for AI-enhanced
                        for run in paragraph.runs:
                            run.font.highlight_color = WD_COLOR_INDEX.TURQUOISE
                        doc_stats['ai_enhanced'] += 1
                    else:
                        # Yellow highlight for local-only
                        for run in paragraph.runs:
                            run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                        doc_stats['local_only'] += 1
                    
                    doc_stats['processed_paragraphs'] += 1
                    doc_stats['changes_made'] += result.get('changes_made', 1)
                    doc_stats['methods_used'][result['method']] += 1
                    
                    print(f"    ✅ Para {i+1}: {result['plagiarism_reduction']:.1f}% reduction ({result['method']})")
                else:
                    print(f"    ⏭️  Para {i+1}: Skipped ({result['plagiarism_reduction']:.1f}% - {result['method']})")
            
            # Save document
            doc.save(file_path)
            
            print(f"  📊 Document Summary:")
            print(f"     • Total paragraphs: {doc_stats['total_paragraphs']}")
            print(f"     • Processed: {doc_stats['processed_paragraphs']}")
            print(f"     • Local-only: {doc_stats['local_only']}")
            print(f"     • AI-enhanced: {doc_stats['ai_enhanced']}")
            print(f"     • Changes made: {doc_stats['changes_made']}")
            print(f"  ✅ Saved: {file_path}")
            
            return doc_stats
            
        except Exception as e:
            print(f"❌ Error processing {file_path}: {e}")
            return None
    
    def is_section_header(self, text):
        """Check if text is a section header"""
        patterns = [
            r'^BAB\s+[IVX]+',
            r'^\d+\.\d+',
            r'^[A-Z\s]+$',
        ]
        
        for pattern in patterns:
            if re.match(pattern, text.strip()) and len(text.split()) < 10:
                return True
        return False
    
    def is_paragraph_suitable_for_paraphrasing(self, paragraph_text):
        """Check if paragraph should be paraphrased"""
        word_count = len(paragraph_text.split())
        if word_count < 20:  # Lower minimum for hybrid system
            return False
        
        if re.search(r'\d+\.\d+|\btabel\b|\bgambar\b', paragraph_text.lower()):
            return False
        
        if re.search(r'\(\d{4}\).*\(\d{4}\).*\(\d{4}\)', paragraph_text):  # Too many citations
            return False
        
        if paragraph_text.isupper() and word_count < 10:
            return False
        
        return True
    
    def print_hybrid_summary(self, stats, processing_time, backup_folder):
        """Print comprehensive hybrid processing summary"""
        print("\n" + "=" * 80)
        print("📊 SMART HYBRID PROCESSING SUMMARY")
        print("=" * 80)
        
        print(f"📄 Documents processed: {stats['documents_processed']}")
        print(f"📝 Paragraphs processed: {stats['paragraphs_processed']}")
        print(f"🏠 Local-only processing: {stats['local_only']}")
        print(f"🤖 AI-enhanced processing: {stats['ai_enhanced']}")
        print(f"🔄 Total changes made: {stats['total_changes']}")
        print(f"⏱️  Processing time: {processing_time:.1f} seconds")
        
        # Cost summary
        print(f"\n💰 COST SUMMARY:")
        print(f"💻 Local calls: {self.cost_tracker['local_calls']} (FREE)")
        print(f"🤖 AI calls: {self.cost_tracker['ai_calls']}")
        print(f"💾 Cache hits: {self.cost_tracker['cache_hits']} (FREE)")
        print(f"🔢 AI tokens used: {self.cost_tracker['ai_tokens_used']:,}")
        print(f"💵 Estimated cost: ${self.cost_tracker['estimated_cost_usd']:.4f} (~Rp {self.cost_tracker['estimated_cost_usd'] * 15000:.0f})")
        
        # Efficiency metrics
        if stats['paragraphs_processed'] > 0:
            ai_ratio = (stats['ai_enhanced'] / stats['paragraphs_processed']) * 100
            print(f"📈 AI usage efficiency: {ai_ratio:.1f}% of processed paragraphs")
        
        if backup_folder:
            print(f"\n💾 Original files backed up to: {backup_folder}")
        
        print("=" * 80)
    
    def generate_hybrid_report(self, input_folder, stats, processing_time):
        """Generate detailed hybrid processing report"""
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        report_file = os.path.join(input_folder, f"hybrid_paraphrase_report_{timestamp}.json")
        
        try:
            report_data = {
                'timestamp': timestamp,
                'input_folder': input_folder,
                'processing_stats': stats,
                'cost_tracking': self.cost_tracker,
                'configuration': self.config,
                'processing_time_seconds': processing_time,
                'system_info': {
                    'local_synonyms': len(self.local_paraphraser.synonyms),
                    'local_phrases': len(self.local_paraphraser.phrase_replacements),
                    'gemini_enabled': bool(self.gemini_api_key),
                    'version': '2.0'
                }
            }
            
            with open(report_file, 'w', encoding='utf-8') as f:
                json.dump(report_data, f, ensure_ascii=False, indent=2)
            
            print(f"📋 Hybrid processing report saved: {report_file}")
            
        except Exception as e:
            print(f"⚠️  Could not save report: {e}")


def main():
    """Main function for smart hybrid processing"""
    print("🎯 SMART HYBRID PARAPHRASER - LOCAL + GEMINI AI")
    print("Hemat biaya dengan AI routing yang cerdas")
    print("=" * 80)
    
    # Configuration
    SYNONYM_FILE = 'sinonim.json'     # Your large synonym database
    INPUT_FOLDER = 'documents'        # Folder with Word documents
    GEMINI_API_KEY = 'YOUR_API_KEY'   # Your Gemini API key
    AGGRESSIVENESS = 0.6              # Processing aggressiveness
    CREATE_BACKUP = True              # Always backup first
    
    # Initialize hybrid processor
    processor = SmartHybridParaphraser(
        synonym_file=SYNONYM_FILE,
        gemini_api_key=GEMINI_API_KEY
    )
    
    # Process batch with smart hybrid approach
    processor.process_batch_documents(
        input_folder=INPUT_FOLDER,
        aggressiveness=AGGRESSIVENESS,
        create_backup=CREATE_BACKUP
    )
    
    print("\n🎉 Smart hybrid processing completed!")
    print("💡 Check color coding: Yellow = Local only, Blue = AI enhanced")


if __name__ == "__main__":
    main()
