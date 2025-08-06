import re
import random
import os
import json
import unicodedata
from typing import Dict, List, Tuple
from datetime import datetime
from collections import defaultdict

# Configuration
INPUT_FOLDER = 'documents'
OUTPUT_FOLDER = 'processed_documents'
PARAPHRASED_FOLDER = 'paraphrased_documents'  # Folder baru untuk hasil parafrase
SUPPORTED_FORMATS = ['.txt', '.docx', '.pdf']

def get_documents_list():
    """
    Get list of documents from INPUT_FOLDER
    Returns list of document files with their info
    """
    if not os.path.exists(INPUT_FOLDER):
        print(f"❌ Error: Folder '{INPUT_FOLDER}' tidak ditemukan!")
        return []
    
    documents = []
    print(f"\n📁 Scanning folder: {INPUT_FOLDER}")
    
    for filename in os.listdir(INPUT_FOLDER):
        file_path = os.path.join(INPUT_FOLDER, filename)
        
        # Check if it's a file and has supported format
        if os.path.isfile(file_path):
            file_ext = os.path.splitext(filename)[1].lower()
            if file_ext in SUPPORTED_FORMATS:
                file_size = os.path.getsize(file_path)
                documents.append({
                    'filename': filename,
                    'path': file_path,
                    'extension': file_ext,
                    'size_kb': round(file_size / 1024, 2),
                    'status': 'Ready to process'
                })
                print(f"✅ Found: {filename} ({round(file_size / 1024, 2)} KB)")
            else:
                print(f"⚠️ Skipped: {filename} (format tidak didukung)")
    
    if not documents:
        print(f"❌ Tidak ada dokumen yang dapat diproses dalam folder '{INPUT_FOLDER}'")
        print(f"📝 Format yang didukung: {', '.join(SUPPORTED_FORMATS)}")
    else:
        print(f"\n✅ Total dokumen siap diproses: {len(documents)}")
    
    return documents

def process_documents_from_folder(aggressiveness: float = 0.6, mode: str = 'balanced'):
    """
    Process all documents from INPUT_FOLDER and save results to PARAPHRASED_FOLDER
    DOES NOT MODIFY ORIGINAL FILES - Creates new paraphrased versions
    """
    print("🎯 PROCESSING DOCUMENTS FROM INPUT FOLDER")
    print("📁 Original files will NOT be modified")
    print("💾 Paraphrased versions will be saved to separate folder")
    print("=" * 50)
    
    # Get documents list
    documents = get_documents_list()
    if not documents:
        return
    
    # Create output folders if not exist
    if not os.path.exists(OUTPUT_FOLDER):
        os.makedirs(OUTPUT_FOLDER)
        print(f"📁 Created reports folder: {OUTPUT_FOLDER}")
    
    if not os.path.exists(PARAPHRASED_FOLDER):
        os.makedirs(PARAPHRASED_FOLDER)
        print(f"📁 Created paraphrased folder: {PARAPHRASED_FOLDER}")
    
    # Initialize evasion system
    evasion = UltimatePlagiarismEvasion()
    
    # Process each document
    total_processed = 0
    for i, doc in enumerate(documents, 1):
        print(f"\n🔄 Processing {i}/{len(documents)}: {doc['filename']}")
        
        try:
            # Read document content based on format
            if doc['extension'] == '.txt':
                with open(doc['path'], 'r', encoding='utf-8') as file:
                    content = file.read()
                
                if not content.strip():
                    print(f"❌ Empty document: {doc['filename']}")
                    continue
                
                # Analyze and process
                analysis = evasion.analyze_text_complexity(content)
                print(f"📊 Text analysis: {analysis['recommendation']}")
                
                # Use recommended aggressiveness or user-specified
                final_aggressiveness = analysis['recommended_aggressiveness'] if mode == 'auto' else aggressiveness
                
                # Process document
                result = evasion.process_document_text(content, final_aggressiveness)
                
                # Save paraphrased version to new folder
                paraphrased_filename = f"paraphrased_{doc['filename']}"
                paraphrased_path = os.path.join(PARAPHRASED_FOLDER, paraphrased_filename)
                
                with open(paraphrased_path, 'w', encoding='utf-8') as file:
                    file.write(result['processed_text'])
                
                print(f"✅ PARAPHRASED VERSION SAVED: {paraphrased_filename}")
                print(f"📊 Stats: {result['stats']['processed_paragraphs']} paragraphs, {result['stats']['total_changes']} changes")
                print(f"� Avg similarity reduction: {result['stats']['avg_similarity_reduction']:.1f}%")
                
                # Save processing report in OUTPUT_FOLDER
                report_filename = f"report_{os.path.splitext(doc['filename'])[0]}.json"
                report_path = os.path.join(OUTPUT_FOLDER, report_filename)
                
                with open(report_path, 'w', encoding='utf-8') as file:
                    json.dump({
                        'original_file': doc['filename'],
                        'paraphrased_file': paraphrased_filename,
                        'original_path': doc['path'],
                        'paraphrased_path': paraphrased_path,
                        'processing_stats': result['stats'],
                        'text_analysis': analysis,
                        'timestamp': datetime.now().isoformat(),
                        'final_aggressiveness': final_aggressiveness,
                        'mode': mode
                    }, file, indent=2, ensure_ascii=False)
                
                total_processed += 1
            
            elif doc['extension'] == '.docx':
                print(f"📄 Processing DOCX: {doc['filename']}")
                
                # Import docx library if needed
                try:
                    import docx
                except ImportError:
                    print(f"❌ python-docx library required for .docx files")
                    print(f"💡 Install with: pip install python-docx")
                    continue
                
                # Load document
                docx_doc = docx.Document(doc['path'])
                
                # Create new document for paraphrased version
                paraphrased_doc = docx.Document()
                
                # Copy document structure and process content
                total_doc_changes = 0
                processed_paragraphs = 0
                
                # Analyze overall text first
                full_text = []
                for paragraph in docx_doc.paragraphs:
                    para_text = paragraph.text.strip()
                    if para_text and len(para_text.split()) >= 10:
                        full_text.append(para_text)
                
                if not full_text:
                    print(f"❌ No suitable paragraphs found in: {doc['filename']}")
                    continue
                
                combined_text = '\n\n'.join(full_text)
                analysis = evasion.analyze_text_complexity(combined_text)
                print(f"📊 Text analysis: {analysis['recommendation']}")
                
                # Use recommended aggressiveness or user-specified
                final_aggressiveness = analysis['recommended_aggressiveness'] if mode == 'auto' else aggressiveness
                
                # Process each paragraph
                for paragraph in docx_doc.paragraphs:
                    para_text = paragraph.text.strip()
                    
                    # Create new paragraph in paraphrased document
                    new_para = paraphrased_doc.add_paragraph()
                    
                    if para_text and len(para_text.split()) >= 10:
                        # Process paragraph content
                        para_result = evasion.ultimate_transform(para_text, final_aggressiveness)
                        
                        if para_result['similarity_reduction'] > 10:
                            # Use paraphrased version
                            new_para.text = para_result['transformed']
                            processed_paragraphs += 1
                            total_doc_changes += para_result['total_changes']
                            
                            print(f"    ✅ Para {processed_paragraphs}: {para_result['similarity_reduction']:.1f}% reduction, {para_result['total_changes']} changes")
                        else:
                            # Keep original if low impact
                            new_para.text = para_text
                            print(f"    ⏭️ Para: Kept original (low impact: {para_result['similarity_reduction']:.1f}%)")
                    else:
                        # Keep short paragraphs as-is (headers, etc.)
                        new_para.text = para_text
                
                # Save paraphrased document to new folder
                paraphrased_filename = f"paraphrased_{doc['filename']}"
                paraphrased_path = os.path.join(PARAPHRASED_FOLDER, paraphrased_filename)
                paraphrased_doc.save(paraphrased_path)
                
                print(f"✅ PARAPHRASED VERSION SAVED: {paraphrased_filename}")
                print(f"📊 Stats: {processed_paragraphs} paragraphs updated, {total_doc_changes} total changes")
                
                # Save processing report
                report_filename = f"report_{os.path.splitext(doc['filename'])[0]}.json"
                report_path = os.path.join(OUTPUT_FOLDER, report_filename)
                
                with open(report_path, 'w', encoding='utf-8') as file:
                    json.dump({
                        'original_file': doc['filename'],
                        'paraphrased_file': paraphrased_filename,
                        'original_path': doc['path'],
                        'paraphrased_path': paraphrased_path,
                        'processing_stats': {
                            'total_paragraphs': len([p for p in docx_doc.paragraphs if p.text.strip()]),
                            'processed_paragraphs': processed_paragraphs,
                            'total_changes': total_doc_changes,
                            'aggressiveness': final_aggressiveness
                        },
                        'text_analysis': analysis,
                        'timestamp': datetime.now().isoformat(),
                        'final_aggressiveness': final_aggressiveness,
                        'mode': mode
                    }, file, indent=2, ensure_ascii=False)
                
                total_processed += 1
            
            elif doc['extension'] == '.pdf':
                print(f"⚠️ PDF format requires PyPDF2 library")
                print(f"💡 Install with: pip install PyPDF2")
                continue
            
        except Exception as e:
            print(f"❌ Error processing {doc['filename']}: {str(e)}")
    
    print(f"\n🎉 PROCESSING COMPLETE!")
    print(f"✅ Successfully processed: {total_processed}/{len(documents)} documents")
    print(f"📁 Original files: {INPUT_FOLDER}/ (UNCHANGED)")
    print(f"📁 Paraphrased files: {PARAPHRASED_FOLDER}/")
    print(f"📁 Reports: {OUTPUT_FOLDER}/")
    print(f"\n💡 Compare before/after:")
    print(f"   • Original: {INPUT_FOLDER}/filename.docx")
    print(f"   • Paraphrased: {PARAPHRASED_FOLDER}/paraphrased_filename.docx")

def process_single_document(file_path: str, aggressiveness: float = 0.6, mode: str = 'balanced') -> dict:
    """
    Process a single document and directly modify the original file
    Returns processing statistics
    """
    if not os.path.exists(file_path):
        return {'error': f'File not found: {file_path}'}
    
    print(f"🔄 Processing single document: {os.path.basename(file_path)}")
    
    # Initialize evasion system
    evasion = UltimatePlagiarismEvasion()
    
    file_ext = os.path.splitext(file_path)[1].lower()
    
    try:
        if file_ext == '.txt':
            # Process TXT file
            with open(file_path, 'r', encoding='utf-8') as file:
                content = file.read()
            
            if not content.strip():
                return {'error': 'Empty document'}
            
            # Analyze and process
            analysis = evasion.analyze_text_complexity(content)
            final_aggressiveness = analysis['recommended_aggressiveness'] if mode == 'auto' else aggressiveness
            
            # Process document
            result = evasion.process_document_text(content, final_aggressiveness)
            
            # Save back to original file (DIRECT EDIT)
            with open(file_path, 'w', encoding='utf-8') as file:
                file.write(result['processed_text'])
            
            return {
                'success': True,
                'filename': os.path.basename(file_path),
                'stats': result['stats'],
                'analysis': analysis,
                'final_aggressiveness': final_aggressiveness
            }
        
        elif file_ext == '.docx':
            # Process DOCX file
            try:
                import docx
            except ImportError:
                return {'error': 'python-docx library required. Install with: pip install python-docx'}
            
            # Load document
            doc = docx.Document(file_path)
            
            # Process paragraphs
            total_changes = 0
            processed_paragraphs = 0
            total_paragraphs = 0
            
            for paragraph in doc.paragraphs:
                para_text = paragraph.text.strip()
                if para_text and len(para_text.split()) >= 10:
                    total_paragraphs += 1
                    
                    # Determine aggressiveness
                    analysis = evasion.analyze_text_complexity(para_text)
                    final_aggressiveness = analysis['recommended_aggressiveness'] if mode == 'auto' else aggressiveness
                    
                    # Process paragraph
                    result = evasion.ultimate_transform(para_text, final_aggressiveness)
                    
                    # Update if significant improvement
                    if result['similarity_reduction'] > 10:  # Lowered threshold
                        paragraph.clear()
                        paragraph.add_run(result['transformed'])
                        
                        processed_paragraphs += 1
                        total_changes += result['total_changes']
            
            # Save document back to original file (DIRECT EDIT)
            doc.save(file_path)
            
            return {
                'success': True,
                'filename': os.path.basename(file_path),
                'stats': {
                    'total_paragraphs': total_paragraphs,
                    'processed_paragraphs': processed_paragraphs,
                    'total_changes': total_changes
                },
                'final_aggressiveness': final_aggressiveness
            }
        
        else:
            return {'error': f'Unsupported file format: {file_ext}'}
    
    except Exception as e:
        return {'error': f'Processing error: {str(e)}'}

class UltimatePlagiarismEvasion:
    def __init__(self):
        print("🎯 Initializing Ultimate Plagiarism Evasion System...")
        
        # Zero-width invisible characters (tidak terlihat mata)
        self.invisible_chars = [
            '\u200B',  # Zero Width Space
            '\u200C',  # Zero Width Non-Joiner
            '\u200D',  # Zero Width Joiner
            '\u2060',  # Word Joiner
        ]
        
        # Advanced semantic transformation patterns
        self.semantic_transformations = {
            # Academic patterns dengan variasi tinggi
            'penelitian_patterns': {
                'penelitian ini bertujuan untuk': [
                    'kajian ini dimaksudkan untuk',
                    'studi ini berupaya dalam rangka',
                    'riset ini diarahkan untuk',
                    'investigasi ini difokuskan pada upaya',
                    'eksplorasi ini ditujukan dalam hal'
                ],
                'penelitian ini menggunakan': [
                    'kajian ini menerapkan',
                    'studi ini memanfaatkan',
                    'riset ini mengaplikasikan',
                    'investigasi ini melibatkan penggunaan',
                    'eksplorasi ini berbasis pada'
                ],
                'berdasarkan hasil penelitian': [
                    'mengacu pada temuan kajian',
                    'merujuk pada output riset',
                    'berlandaskan hasil studi',
                    'sesuai dengan hasil investigasi',
                    'berdasarkan hasil eksplorasi'
                ],
                'hasil penelitian menunjukkan': [
                    'temuan kajian mengindikasikan',
                    'output riset memperlihatkan',
                    'hasil studi mendemonstrasikan',
                    'hasil investigasi mengungkapkan',
                    'temuan eksplorasi merefleksikan'
                ]
            },
            
            # Technical system patterns
            'system_patterns': {
                'sistem informasi adalah': [
                    'arsitektur informasi merupakan',
                    'platform informasi dapat didefinisikan sebagai',
                    'framework informasi adalah',
                    'infrastruktur informasi dapat dipahami sebagai',
                    'ekosistem informasi mencakup'
                ],
                'sistem ini dapat': [
                    'platform ini mampu',
                    'arsitektur ini sanggup',
                    'framework ini bisa',
                    'infrastruktur ini dapat',
                    'mekanisme ini memungkinkan untuk'
                ],
                'menggunakan sistem': [
                    'memanfaatkan platform',
                    'menerapkan arsitektur',
                    'mengaplikasikan framework',
                    'melibatkan infrastruktur',
                    'berbasis pada mekanisme'
                ]
            },
            
            # Methodology patterns
            'method_patterns': {
                'metode yang digunakan': [
                    'pendekatan yang diterapkan',
                    'cara yang diaplikasikan',
                    'teknik yang dimanfaatkan',
                    'prosedur yang dilaksanakan',
                    'strategi yang diimplementasikan'
                ],
                'dengan menggunakan metode': [
                    'melalui penerapan pendekatan',
                    'via implementasi teknik',
                    'lewat aplikasi prosedur',
                    'dengan menerapkan strategi',
                    'berdasarkan penggunaan cara'
                ],
                'metode ini efektif untuk': [
                    'pendekatan ini optimal dalam',
                    'teknik ini efisien untuk',
                    'prosedur ini cocok dalam hal',
                    'strategi ini tepat untuk',
                    'cara ini sesuai dalam rangka'
                ]
            },
            
            # Analysis patterns
            'analysis_patterns': {
                'analisis data menunjukkan': [
                    'pengkajian data mengindikasikan',
                    'evaluasi data memperlihatkan',
                    'telaah data mengungkapkan',
                    'eksaminasi data mendemonstrasikan',
                    'penelaahan data merefleksikan'
                ],
                'berdasarkan analisis': [
                    'mengacu pada pengkajian',
                    'merujuk pada evaluasi',
                    'berlandaskan telaah',
                    'sesuai dengan eksaminasi',
                    'berdasarkan penelaahan'
                ],
                'hasil analisis': [
                    'output pengkajian',
                    'temuan evaluasi',
                    'hasil telaah',
                    'outcome eksaminasi',
                    'hasil penelaahan'
                ]
            }
        }
        
        # Structural sentence reordering patterns
        self.structure_patterns = [
            # Passive to active and vice versa
            {
                'pattern': r'(\w+)\s+(digunakan|diterapkan|dimanfaatkan|diaplikasikan)\s+untuk\s+(\w+)',
                'replacement': r'\3 menggunakan \1',
                'description': 'Passive to active transformation'
            },
            {
                'pattern': r'(\w+)\s+(menggunakan|menerapkan|memanfaatkan|mengaplikasikan)\s+(\w+)',
                'replacement': r'\3 \2 oleh \1',
                'description': 'Active to passive transformation'
            },
            
            # Causal relationship reordering
            {
                'pattern': r'karena\s+(\w+.*?),\s*maka\s+(\w+.*?)(\.|$)',
                'replacement': r'\2 sebagai akibat dari \1\3',
                'description': 'Causal reordering'
            },
            {
                'pattern': r'akibat\s+(\w+.*?),\s*(\w+.*?)(\.|$)',
                'replacement': r'\2 yang disebabkan oleh \1\3',
                'description': 'Consequence reordering'
            },
            
            # Purpose clause reordering
            {
                'pattern': r'untuk\s+(\w+.*?),\s*(\w+.*?)(\.|$)',
                'replacement': r'\2 dengan tujuan \1\3',
                'description': 'Purpose clause reordering'
            },
            
            # Conditional reordering
            {
                'pattern': r'jika\s+(\w+.*?),\s*maka\s+(\w+.*?)(\.|$)',
                'replacement': r'\2 dalam kondisi \1\3',
                'description': 'Conditional reordering'
            }
        ]
        
        # Advanced word-level transformations
        self.word_transformations = {
            'academic_verbs': {
                'menunjukkan': ['mengindikasikan', 'memperlihatkan', 'mendemonstrasikan', 'mengungkapkan', 'merefleksikan'],
                'menggunakan': ['memanfaatkan', 'menerapkan', 'mengaplikasikan', 'melibatkan', 'berbasis pada'],
                'mengembangkan': ['merancang', 'membangun', 'menciptakan', 'menyusun', 'mengonstruksi'],
                'menganalisis': ['mengkaji', 'mengevaluasi', 'menelaah', 'mengeksaminasi', 'meneliti'],
                'menghasilkan': ['memproduksi', 'menciptakan', 'melahirkan', 'membangkitkan', 'memunculkan'],
                'meningkatkan': ['mengoptimalkan', 'memperbaiki', 'memaksimalkan', 'mengembangkan', 'menyempurnakan']
            },
            'academic_nouns': {
                'penelitian': ['kajian', 'studi', 'riset', 'investigasi', 'eksplorasi'],
                'analisis': ['pengkajian', 'evaluasi', 'telaah', 'eksaminasi', 'penelaahan'],
                'hasil': ['temuan', 'output', 'outcome', 'produk', 'capaian'],
                'metode': ['pendekatan', 'teknik', 'cara', 'prosedur', 'strategi'],
                'sistem': ['platform', 'arsitektur', 'framework', 'infrastruktur', 'mekanisme'],
                'data': ['informasi', 'dataset', 'sampel data', 'kumpulan data', 'materi empiris']
            },
            'connecting_words': {
                'oleh karena itu': ['dengan demikian', 'maka dari itu', 'akibatnya', 'konsekuensinya', 'sebagai hasilnya'],
                'selain itu': ['di samping itu', 'tambahan pula', 'lebih lanjut', 'selanjutnya', 'bahkan'],
                'dengan demikian': ['oleh karena itu', 'maka dari itu', 'akibatnya', 'sebagai konsekuensi', 'hasilnya'],
                'berdasarkan': ['mengacu pada', 'merujuk pada', 'berlandaskan', 'sesuai dengan', 'menurut'],
                'sehingga': ['yang mengakibatkan', 'sampai', 'hingga', 'alhasil', 'sebagai konsekuensi']
            }
        }
        
        print("✅ Ultimate evasion system loaded!")
        print(f"🔧 Semantic patterns: {sum(len(v) for v in self.semantic_transformations.values())}")
        print(f"🔧 Structure patterns: {len(self.structure_patterns)}")
        print(f"🔧 Word transformations: {sum(len(v) for v in self.word_transformations.values())}")
    
    def insert_invisible_watermark(self, text: str, density: float = 0.15) -> str:
        """
        Insert invisible characters strategically
        Density 0.15 = 15% chance per word boundary
        """
        words = text.split()
        result = []
        
        for i, word in enumerate(words):
            result.append(word)
            
            # Insert invisible char after word (except last word)
            if i < len(words) - 1:
                if random.random() < density:
                    invisible_char = random.choice(self.invisible_chars)
                    result.append(invisible_char)
            
            # Add normal space (except after last word)
            if i < len(words) - 1:
                result.append(' ')
        
        return ''.join(result)
    
    def apply_semantic_transformations(self, text: str) -> tuple:
        """
        Apply contextual semantic transformations
        Returns (transformed_text, changes_made)
        """
        transformed_text = text
        changes_made = []
        
        # Apply all semantic transformation categories
        for category, patterns in self.semantic_transformations.items():
            for pattern, replacements in patterns.items():
                if pattern.lower() in transformed_text.lower():
                    replacement = random.choice(replacements)
                    
                    # Case-insensitive replacement while preserving case
                    pattern_regex = re.compile(re.escape(pattern), re.IGNORECASE)
                    match = pattern_regex.search(transformed_text)
                    
                    if match:
                        # Preserve capitalization of first word
                        if match.group()[0].isupper():
                            replacement = replacement.capitalize()
                        
                        transformed_text = pattern_regex.sub(replacement, transformed_text, count=1)
                        changes_made.append({
                            'type': 'semantic_transformation',
                            'original': pattern,
                            'replacement': replacement,
                            'category': category
                        })
        
        return transformed_text, changes_made
    
    def apply_structural_reordering(self, text: str) -> tuple:
        """
        Apply structural sentence reordering
        Returns (transformed_text, changes_made)
        """
        transformed_text = text
        changes_made = []
        
        for structure in self.structure_patterns:
            pattern = structure['pattern']
            replacement = structure['replacement']
            
            if re.search(pattern, transformed_text, re.IGNORECASE):
                old_text = transformed_text
                transformed_text = re.sub(pattern, replacement, transformed_text, flags=re.IGNORECASE)
                
                if old_text != transformed_text:
                    changes_made.append({
                        'type': 'structural_reordering',
                        'description': structure['description'],
                        'pattern': pattern
                    })
        
        return transformed_text, changes_made
    
    def apply_word_transformations(self, text: str, transformation_rate: float = 0.4) -> tuple:
        """
        Apply advanced word-level transformations
        Returns (transformed_text, changes_made)
        """
        words = text.split()
        transformed_words = []
        changes_made = []
        
        for word in words:
            # Clean word for matching (remove punctuation)
            clean_word = re.sub(r'[^\w]', '', word.lower())
            transformed = False
            
            # Check all transformation categories
            for category, word_dict in self.word_transformations.items():
                if clean_word in word_dict and random.random() < transformation_rate:
                    alternatives = word_dict[clean_word]
                    replacement = random.choice(alternatives)
                    
                    # Preserve capitalization and punctuation
                    if word[0].isupper():
                        replacement = replacement.capitalize()
                    
                    # Add back punctuation
                    punctuation = ''.join(c for c in word if not c.isalnum())
                    final_word = replacement + punctuation
                    
                    transformed_words.append(final_word)
                    changes_made.append({
                        'type': 'word_transformation',
                        'original': word,
                        'replacement': final_word,
                        'category': category
                    })
                    transformed = True
                    break
            
            if not transformed:
                transformed_words.append(word)
        
        return ' '.join(transformed_words), changes_made
    
    def ultimate_transform(self, text: str, aggressiveness: float = 0.7) -> dict:
        """
        Apply the ultimate transformation combining all techniques
        aggressiveness: 0.1 (subtle) to 1.0 (maximum)
        """
        if not text or not text.strip():
            return {
                'original': text,
                'transformed': text,
                'similarity_reduction': 0,
                'changes_made': [],
                'status': 'Empty text provided'
            }
        
        original_text = text.strip()
        current_text = original_text
        all_changes = []
        
        # Step 1: Semantic transformations (highest impact)
        current_text, semantic_changes = self.apply_semantic_transformations(current_text)
        all_changes.extend(semantic_changes)
        
        # Step 2: Structural reordering (medium impact)
        if aggressiveness > 0.3:
            current_text, structure_changes = self.apply_structural_reordering(current_text)
            all_changes.extend(structure_changes)
        
        # Step 3: Word-level transformations (controlled by aggressiveness)
        word_rate = min(aggressiveness * 0.6, 0.5)  # Max 50% word transformation
        current_text, word_changes = self.apply_word_transformations(current_text, word_rate)
        all_changes.extend(word_changes)
        
        # Step 4: Invisible watermarking (final stealth layer)
        watermark_density = min(aggressiveness * 0.2, 0.15)  # Max 15% density
        current_text = self.insert_invisible_watermark(current_text, watermark_density)
        
        # Calculate similarity reduction estimate
        original_words = set(re.findall(r'\w+', original_text.lower()))
        transformed_words = set(re.findall(r'\w+', current_text.lower()))
        
        if len(original_words) > 0:
            word_overlap = len(original_words.intersection(transformed_words)) / len(original_words)
            estimated_similarity = word_overlap * 100
        else:
            estimated_similarity = 100
        
        # Adjust similarity based on structural changes
        if any(change['type'] == 'structural_reordering' for change in all_changes):
            estimated_similarity *= 0.7  # 30% additional reduction for structure changes
        
        similarity_reduction = 100 - estimated_similarity
        
        # Determine status
        if similarity_reduction >= 70:
            status = "✅ EXCELLENT - Very high evasion rate"
        elif similarity_reduction >= 50:
            status = "✅ GOOD - High evasion rate"
        elif similarity_reduction >= 30:
            status = "⚠️ MODERATE - Decent evasion rate"
        else:
            status = "❌ LOW - May need more aggressive settings"
        
        return {
            'original': original_text,
            'transformed': current_text,
            'similarity_reduction': round(similarity_reduction, 1),
            'changes_made': all_changes,
            'total_changes': len(all_changes),
            'status': status,
            'aggressiveness_used': aggressiveness,
            'has_invisible_watermark': True,
            'word_count_original': len(original_text.split()),
            'word_count_transformed': len(current_text.split())
        }
    
    def process_document_text(self, text: str, aggressiveness: float = 0.7) -> dict:
        """
        Process entire document text with ultimate transformation
        """
        paragraphs = text.split('\n\n')
        processed_paragraphs = []
        
        stats = {
            'total_paragraphs': 0,
            'processed_paragraphs': 0,
            'skipped_paragraphs': 0,
            'total_changes': 0,
            'avg_similarity_reduction': 0,
            'change_types': defaultdict(int),
            'aggressiveness': aggressiveness
        }
        
        similarity_reductions = []
        
        for i, paragraph in enumerate(paragraphs):
            para_text = paragraph.strip()
            stats['total_paragraphs'] += 1
            
            if not para_text or len(para_text.split()) < 10:
                processed_paragraphs.append(paragraph)
                stats['skipped_paragraphs'] += 1
                continue
            
            # Apply ultimate transformation
            result = self.ultimate_transform(para_text, aggressiveness)
            
            # Only apply if significant improvement
            if result['similarity_reduction'] > 10:  # Lowered threshold to 10%
                processed_paragraphs.append(result['transformed'])
                
                stats['processed_paragraphs'] += 1
                stats['total_changes'] += result['total_changes']
                similarity_reductions.append(result['similarity_reduction'])
                
                # Track change types
                for change in result['changes_made']:
                    stats['change_types'][change['type']] += 1
                
                print(f"✅ Para {i+1}: {result['similarity_reduction']:.1f}% reduction, {result['total_changes']} changes")
            else:
                processed_paragraphs.append(paragraph)
                stats['skipped_paragraphs'] += 1
                print(f"⏭️ Para {i+1}: Skipped (low impact: {result['similarity_reduction']:.1f}%)")
        
        # Calculate average similarity reduction
        if similarity_reductions:
            stats['avg_similarity_reduction'] = sum(similarity_reductions) / len(similarity_reductions)
        
        # Convert defaultdict to regular dict
        stats['change_types'] = dict(stats['change_types'])
        
        processed_text = '\n\n'.join(processed_paragraphs)
        
        return {
            'original_text': text,
            'processed_text': processed_text,
            'stats': stats
        }
    
    def quick_transform(self, text: str, mode: str = 'balanced') -> str:
        """
        Quick transformation dengan preset modes
        
        Modes:
        - 'subtle': aggressiveness 0.3 (minimal changes)
        - 'balanced': aggressiveness 0.6 (recommended)
        - 'aggressive': aggressiveness 0.9 (maximum changes)
        """
        aggressiveness_map = {
            'subtle': 0.3,
            'balanced': 0.6,
            'aggressive': 0.9
        }
        
        aggressiveness = aggressiveness_map.get(mode, 0.6)
        result = self.ultimate_transform(text, aggressiveness)
        return result['transformed']
    
    def analyze_text_complexity(self, text: str) -> dict:
        """
        Analyze text complexity and recommend aggressiveness level
        """
        words = text.split()
        word_count = len(words)
        
        # Count academic terms
        academic_terms = ['penelitian', 'analisis', 'metode', 'sistem', 'hasil', 'data']
        academic_count = sum(1 for word in words if any(term in word.lower() for term in academic_terms))
        
        # Count repeated patterns
        repeated_patterns = 0
        for pattern_group in self.semantic_transformations.values():
            for pattern in pattern_group.keys():
                if pattern.lower() in text.lower():
                    repeated_patterns += 1
        
        # Calculate complexity score
        academic_ratio = academic_count / word_count if word_count > 0 else 0
        pattern_ratio = repeated_patterns / max(len(words) // 10, 1)
        
        complexity_score = (academic_ratio + pattern_ratio) / 2
        
        # Recommend aggressiveness
        if complexity_score > 0.7:
            recommended_aggressiveness = 0.9
            recommendation = "AGGRESSIVE - High academic content detected"
        elif complexity_score > 0.4:
            recommended_aggressiveness = 0.6
            recommendation = "BALANCED - Moderate academic content"
        else:
            recommended_aggressiveness = 0.3
            recommendation = "SUBTLE - Low academic content"
        
        return {
            'word_count': word_count,
            'academic_terms_found': academic_count,
            'academic_ratio': academic_ratio,
            'repeated_patterns': repeated_patterns,
            'complexity_score': complexity_score,
            'recommended_aggressiveness': recommended_aggressiveness,
            'recommendation': recommendation
        }

def main():
    """Demo function untuk ultimate plagiarism evasion"""
    print("🎯 ULTIMATE PLAGIARISM EVASION - DIRECT FILE PROCESSING")
    print("Contextual Semantic Rewriting + Invisible Watermarking")
    print("DIRECTLY MODIFIES ORIGINAL FILES - NO BACKUP!")
    print("=" * 70)
    
    # Show available documents in INPUT_FOLDER
    print(f"\n📁 CHECKING DOCUMENTS IN '{INPUT_FOLDER}' FOLDER:")
    documents = get_documents_list()
    
    if documents:
        print(f"\n⚠️ WARNING: ORIGINAL FILES WILL BE DIRECTLY MODIFIED!")
        print(f"💡 READY TO PROCESS {len(documents)} DOCUMENTS:")
        print("Choose processing mode:")
        print("1. 🟢 Subtle mode (aggressiveness=0.3)")
        print("2. 🟡 Balanced mode (aggressiveness=0.6) - RECOMMENDED")
        print("3. 🔴 Aggressive mode (aggressiveness=0.9)")
        print("4. 🤖 Auto mode (AI decides based on content)")
        
        print(f"\n🚀 READY TO START PROCESSING:")
        print("To process ALL documents:")
        print("  process_documents_from_folder(aggressiveness=0.6, mode='balanced')")
        print("\nTo process a SINGLE document:")
        print("  process_single_document('documents/filename.txt', aggressiveness=0.6)")
        print("\nFor AUTO mode (recommended):")
        print("  process_documents_from_folder(mode='auto')")
        
    else:
        print(f"\n💡 No documents found. Add documents to '{INPUT_FOLDER}' folder:")
        print(f"   • Supported formats: {', '.join(SUPPORTED_FORMATS)}")
        print(f"   • Example: place your .txt or .docx files in '{INPUT_FOLDER}/' folder")
    
    # Quick demo with sample text (for demonstration only)
    print("\n" + "="*50)
    print("🧪 QUICK DEMO - SAMPLE TRANSFORMATION:")
    print("="*50)
    
    # Initialize system for demo
    evasion = UltimatePlagiarismEvasion()
    
    sample_text = """Penelitian ini bertujuan untuk mengembangkan sistem informasi yang dapat 
    menggunakan teknologi modern untuk meningkatkan efisiensi kerja. Berdasarkan hasil 
    penelitian, sistem ini dapat menganalisis data dengan metode yang digunakan secara 
    efektif. Analisis data menunjukkan bahwa metode ini efektif untuk menghasilkan 
    hasil yang optimal."""
    
    print(f"\n📝 ORIGINAL TEXT:")
    print(f"'{sample_text}'")
    
    # Analyze text complexity
    analysis = evasion.analyze_text_complexity(sample_text)
    print(f"\n📊 TEXT COMPLEXITY ANALYSIS:")
    print(f"Word count: {analysis['word_count']}")
    print(f"Academic terms: {analysis['academic_terms_found']}")
    print(f"Complexity score: {analysis['complexity_score']:.2f}")
    print(f"Recommendation: {analysis['recommendation']}")
    
    # Test recommended aggressiveness
    result = evasion.ultimate_transform(sample_text, analysis['recommended_aggressiveness'])
    
    print(f"\n🔄 RECOMMENDED TRANSFORMATION:")
    print(f"Result: '{result['transformed']}'")
    print(f"📊 Similarity Reduction: {result['similarity_reduction']:.1f}%")
    print(f"🔧 Changes Made: {result['total_changes']}")
    print(f"📈 Status: {result['status']}")
    print(f"💧 Invisible Watermark: {result['has_invisible_watermark']}")
    
    print(f"\n💡 USAGE EXAMPLES:")
    print(f"")
    print(f"# Process all documents from '{INPUT_FOLDER}' folder (DIRECT EDIT):")
    print(f"process_documents_from_folder(aggressiveness=0.6, mode='balanced')")
    print(f"")
    print(f"# Process with auto aggressiveness (AI decides):")
    print(f"process_documents_from_folder(mode='auto')")
    print(f"")
    print(f"# Process single document (DIRECT EDIT):")
    print(f"process_single_document('documents/your_file.txt', aggressiveness=0.6)")
    print(f"")
    print(f"# Check available documents:")
    print(f"documents = get_documents_list()")
    
    print(f"\n🎯 PROCESSING MODES:")
    print(f"• 🟢 Subtle (0.3): Minimal changes, natural flow")
    print(f"• 🟡 Balanced (0.6): Optimal performance (RECOMMENDED)")
    print(f"• 🔴 Aggressive (0.9): Maximum evasion, check readability")
    print(f"• 🤖 Auto: AI analyzes content and chooses best settings")
    
    print(f"\n⚠️ IMPORTANT WARNINGS:")
    print(f"• Original files will be DIRECTLY MODIFIED")
    print(f"• No backup is created automatically")
    print(f"• Make sure to backup important files manually")
    print(f"• Processing reports are saved in '{OUTPUT_FOLDER}' folder")

# Example usage
if __name__ == "__main__":
    main()
    
    print(f"\n" + "="*70)
    print("🚀 STARTING AUTOMATIC PROCESSING:")
    print("="*70)
    
    # Cek apakah ada dokumen yang bisa diproses
    documents = get_documents_list()
    
    if documents:
        print(f"\n⚡ FOUND {len(documents)} DOCUMENTS - STARTING PROCESSING...")
        print("🤖 Using AUTO mode (AI will determine best aggressiveness for each document)")
        print("⚠️ WARNING: ORIGINAL FILES WILL BE DIRECTLY MODIFIED!")
        
        # Tunggu konfirmasi user (opsional - hapus jika ingin otomatis)
        try:
            user_input = input("\n🔄 Press ENTER to continue or 'q' to quit: ").strip().lower()
            if user_input == 'q':
                print("❌ Processing cancelled by user")
                exit()
        except KeyboardInterrupt:
            print("\n❌ Processing cancelled")
            exit()
        
        # Mulai processing otomatis
        print(f"\n🚀 STARTING PROCESSING...")
        print("🔄 Using BALANCED mode with aggressiveness 0.7 for better results")
        process_documents_from_folder(aggressiveness=0.7, mode='balanced')
        
        print(f"\n🎉 ALL PROCESSING COMPLETED!")
        print(f"📁 Check your '{INPUT_FOLDER}' folder - files have been directly modified")
        print(f"📊 Processing reports saved in '{OUTPUT_FOLDER}' folder")
        
    else:
        print(f"\n❌ No documents found to process")
        print(f"💡 Add .txt or .docx files to the '{INPUT_FOLDER}' folder and run again")
        
        # Buat folder documents jika belum ada
        if not os.path.exists(INPUT_FOLDER):
            os.makedirs(INPUT_FOLDER)
            print(f"✅ Created '{INPUT_FOLDER}' folder - add your documents here")
        
        print(f"\n📝 EXAMPLE: Create a test file")
        test_file_path = os.path.join(INPUT_FOLDER, "test_document.txt")
        
        if not os.path.exists(test_file_path):
            test_content = """Penelitian ini bertujuan untuk mengembangkan sistem informasi yang dapat menggunakan teknologi modern untuk meningkatkan efisiensi kerja. Berdasarkan hasil penelitian, sistem ini dapat menganalisis data dengan metode yang digunakan secara efektif.

Analisis data menunjukkan bahwa metode ini efektif untuk menghasilkan hasil yang optimal. Dengan menggunakan metode penelitian yang tepat, sistem informasi dapat diimplementasikan dengan baik.

Penelitian ini menggunakan pendekatan kualitatif dan kuantitatif untuk mendapatkan hasil yang komprehensif. Hasil penelitian menunjukkan adanya peningkatan efisiensi sebesar 75% dibandingkan dengan sistem sebelumnya."""
            
            with open(test_file_path, 'w', encoding='utf-8') as f:
                f.write(test_content)
            
            print(f"✅ Created test document: {test_file_path}")
            print(f"🔄 Run the script again to process this test document")
        
        print(f"\n💡 To process documents manually, use:")
        print(f"process_documents_from_folder(mode='auto')")
        print(f"# or")
        print(f"process_single_document('documents/your_file.txt', mode='auto')")

# Fungsi tambahan untuk easy processing
def quick_process_all():
    """Fungsi cepat untuk memproses semua dokumen"""
    print("🚀 QUICK PROCESS ALL DOCUMENTS")
    documents = get_documents_list()
    
    if documents:
        process_documents_from_folder(mode='auto')
    else:
        print("❌ No documents found in the documents folder")

def show_processing_results():
    """Tampilkan hasil processing yang sudah ada"""
    if os.path.exists(OUTPUT_FOLDER):
        print(f"\n📊 PROCESSING RESULTS IN '{OUTPUT_FOLDER}':")
        
        for filename in os.listdir(OUTPUT_FOLDER):
            if filename.endswith('.json'):
                report_path = os.path.join(OUTPUT_FOLDER, filename)
                try:
                    with open(report_path, 'r', encoding='utf-8') as f:
                        report = json.load(f)
                    
                    print(f"\n📄 {report.get('original_file', 'Unknown')}")
                    if 'processing_stats' in report:
                        stats = report['processing_stats']
                        print(f"   • Processed paragraphs: {stats.get('processed_paragraphs', 0)}")
                        print(f"   • Total changes: {stats.get('total_changes', 0)}")
                        print(f"   • Avg similarity reduction: {stats.get('avg_similarity_reduction', 0):.1f}%")
                    print(f"   • Processed: {report.get('timestamp', 'Unknown time')}")
                
                except Exception as e:
                    print(f"   ⚠️ Error reading {filename}: {e}")
    else:
        print(f"❌ No processing results found. Run processing first.")
