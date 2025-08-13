#!/usr/bin/env python3
"""
Demo sederhana untuk memproses core chapters (BAB I - BAB V) saja
"""

import os
import sys
from pathlib import Path

# Add current directory to path
sys.path.append(os.path.dirname(os.path.abspath(__file__)))

def demo_core_chapters():
    """Demo core chapters processing"""
    print("🎯 DEMO: Core Chapters Processing")
    print("📖 Focus: BAB I PENDAHULUAN → BAB V KESIMPULAN")
    print("⏭️  Skipping: Cover, appendices, references")
    print("=" * 60)
    
    # Check available files
    doc_dir = "documents"
    if not os.path.exists(doc_dir):
        print(f"❌ Directory not found: {doc_dir}")
        return
    
    # Find Word documents
    doc_files = []
    for ext in ['.docx', '.doc']:
        doc_files.extend(Path(doc_dir).glob(f"*{ext}"))
    
    if not doc_files:
        print(f"❌ No Word documents found in {doc_dir}")
        return
    
    print(f"📄 Found {len(doc_files)} documents:")
    for i, doc_file in enumerate(doc_files, 1):
        print(f"   {i}. {doc_file.name}")
    
    # Use first file for demo
    test_file = str(doc_files[0])
    print(f"\n🎯 Demo with: {Path(test_file).name}")
    
    # Import after checking files
    try:
        from full_gemini_paraphraser import FullGeminiParaphraser
    except ImportError as e:
        print(f"❌ Import error: {e}")
        return
    
    # Initialize with conservative settings for demo
    print("\n🚀 Initializing paraphraser...")
    paraphraser = FullGeminiParaphraser(
        synonym_file='sinonim.json',
        chunk_size=2000,  # Small chunks for faster demo
        max_retries=2     # Fewer retries for demo
    )
    
    # Test chapter detection first
    print("\n🔍 Detecting chapter boundaries...")
    import docx
    
    try:
        doc = docx.Document(test_file)
        start_index, end_index = paraphraser.detect_chapter_boundaries(doc)
        
        if start_index is not None:
            print(f"✅ Found BAB I at paragraph {start_index}")
            start_text = doc.paragraphs[start_index].text.strip()
            print(f"   📝 Start: {start_text[:60]}...")
        else:
            print("❌ Could not find BAB I PENDAHULUAN")
            return
        
        if end_index is not None:
            print(f"✅ Found end boundary at paragraph {end_index}")
            if end_index < len(doc.paragraphs):
                end_text = doc.paragraphs[end_index].text.strip()
                print(f"   📝 End: {end_text[:60]}...")
        else:
            print("⚠️  No specific end boundary found, will process until end")
            end_index = len(doc.paragraphs)
        
        total_paragraphs = sum(1 for i in range(start_index, end_index) 
                              if i < len(doc.paragraphs) and doc.paragraphs[i].text.strip())
        
        print(f"📊 Core chapters contain {total_paragraphs} paragraphs")
        
    except Exception as e:
        print(f"❌ Error detecting chapters: {e}")
        return
    
    # Ask user if they want to proceed with processing
    if total_paragraphs > 10:
        print(f"\n⚠️  This will process {total_paragraphs} paragraphs.")
        print("💡 For demo purposes, recommend testing with smaller document first.")
        response = input("Continue? (y/n): ").lower().strip()
        if response != 'y':
            print("Demo cancelled by user.")
            return
    
    # Process core chapters
    print(f"\n🔄 Processing core chapters...")
    try:
        result = paraphraser.process_core_chapters_only(test_file)
        
        if result:
            print("\n🎉 DEMO COMPLETED!")
            print(f"✅ Output: {result['output_path']}")
            print(f"📊 Processed: {result['processed_paragraphs']} paragraphs")
            print(f"📍 Range: {result['start_paragraph']} to {result['end_paragraph']}")
            
            # Show brief statistics
            stats = result.get('statistics', {})
            if stats:
                print(f"\n📈 Quick Stats:")
                print(f"   Runtime: {stats.get('runtime', 'N/A')}")
                print(f"   API calls: {stats.get('total_api_calls', 'N/A')}")
                print(f"   Avg quality: {stats.get('average_quality', 'N/A')}")
        else:
            print("❌ Processing failed")
            
    except KeyboardInterrupt:
        print("\n⏸️  Demo interrupted by user")
    except Exception as e:
        print(f"❌ Demo error: {e}")

def show_usage():
    """Show usage examples"""
    print("🎯 Core Chapters Processing - Usage Examples")
    print("=" * 50)
    print()
    print("📖 Process core chapters only (BAB I - BAB V):")
    print("   python full_gemini_paraphraser.py core")
    print()
    print("📄 Process single file core chapters:")
    print('   python full_gemini_paraphraser.py core-single "path/to/file.docx"')
    print()
    print("🧪 Run this demo:")
    print("   python demo_core_chapters.py")
    print()
    print("💡 Benefits:")
    print("   ✅ Faster processing (skip cover pages, appendices)")
    print("   ✅ Focus on main content only")
    print("   ✅ Cleaner output documents")
    print("   ✅ Reduced API usage and costs")

if __name__ == "__main__":
    if len(sys.argv) > 1 and sys.argv[1] == "usage":
        show_usage()
    else:
        demo_core_chapters()