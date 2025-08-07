# smart_plagiarism_checker.py
"""
Smart Plagiarism Checker - Offline Pattern Detection
Deteksi pola plagiarisme tanpa perlu koneksi internet menggunakan analisis pattern dan similarity
Author: DevNoLife
Version: 1.0
"""

import os
import json
import re
import hashlib
from datetime import datetime
from difflib import SequenceMatcher
from collections import defaultdict, Counter
import docx

class SmartPlagiarismChecker:
    def __init__(self):
        print("🔍 Initializing Smart Plagiarism Checker...")
        
        # Common plagiarism patterns in Indonesian academic text
        self.plagiarism_patterns = {
            'copy_paste_indicators': [
                r'menurut\s+\w+\s+\(\d{4}\)\s+menyatakan\s+bahwa',  # Citation patterns
                r'berdasarkan\s+penelitian\s+\w+\s+\(\d{4}\)',
                r'definisi\s+\w+\s+menurut\s+\w+\s+adalah',
                r'teori\s+\w+\s+dikemukakan\s+oleh\s+\w+',
                r'konsep\s+\w+\s+dijelaskan\s+oleh\s+\w+',
            ],
            'academic_cliches': [
                r'penelitian\s+ini\s+bertujuan\s+untuk\s+mengetahui',
                r'berdasarkan\s+latar\s+belakang\s+di\s+atas',
                r'dari\s+uraian\s+di\s+atas\s+dapat\s+disimpulkan',
                r'hasil\s+penelitian\s+menunjukkan\s+bahwa',
                r'sesuai\s+dengan\s+rumusan\s+masalah',
            ],
            'repetitive_structures': [
                r'^(penelitian|kajian|studi)\s+ini\s+',
                r'^(berdasarkan|menurut|sesuai)\s+',
                r'^(dengan\s+demikian|oleh\s+karena\s+itu)',
                r'^(hasil|temuan|output)\s+',
            ]
        }
        
        # Suspicious word sequences that might indicate copying
        self.suspicious_sequences = [
            'menurut para ahli dalam bidang',
            'berdasarkan teori yang dikemukakan oleh',
            'sejalan dengan pendapat yang dikemukakan',
            'sesuai dengan definisi yang diberikan oleh',
            'dalam konteks ini dapat dijelaskan bahwa',
            'dari penjelasan tersebut dapat dipahami',
            'hal ini sejalan dengan teori yang menyatakan',
        ]
        
        # Settings
        self.settings = {
            'min_text_length': 20,          # Minimum words to analyze
            'similarity_threshold': 75,      # High similarity threshold
            'pattern_weight': 30,           # Weight for pattern matching
            'sequence_weight': 25,          # Weight for suspicious sequences
            'structure_weight': 20,         # Weight for repetitive structures
            'redundancy_weight': 25,        # Weight for content redundancy
        }
        
        # Analysis results storage
        self.analysis_cache = {}
        
        print("✅ Smart Plagiarism Checker ready!")
        print(f"✅ Loaded {len(self.plagiarism_patterns)} pattern categories")
        print(f"✅ Monitoring {len(self.suspicious_sequences)} suspicious sequences")
    
    def analyze_text_patterns(self, text):
        """Analyze text for plagiarism patterns"""
        if len(text.split()) < self.settings['min_text_length']:
            return {
                'risk_score': 0,
                'patterns_found': [],
                'analysis': 'Text too short for analysis'
            }
        
        patterns_found = []
        total_score = 0
        
        # Check for copy-paste indicators
        for pattern in self.plagiarism_patterns['copy_paste_indicators']:
            matches = re.findall(pattern, text, re.IGNORECASE)
            if matches:
                score = len(matches) * 15
                total_score += score
                patterns_found.append({
                    'type': 'copy_paste_indicator',
                    'pattern': pattern,
                    'matches': matches,
                    'score': score
                })
        
        # Check for academic clichés
        for pattern in self.plagiarism_patterns['academic_cliches']:
            matches = re.findall(pattern, text, re.IGNORECASE)
            if matches:
                score = len(matches) * 10
                total_score += score
                patterns_found.append({
                    'type': 'academic_cliche',
                    'pattern': pattern,
                    'matches': matches,
                    'score': score
                })
        
        # Check for repetitive structures
        for pattern in self.plagiarism_patterns['repetitive_structures']:
            matches = re.findall(pattern, text, re.IGNORECASE)
            if matches:
                score = len(matches) * 8
                total_score += score
                patterns_found.append({
                    'type': 'repetitive_structure',
                    'pattern': pattern,
                    'matches': matches,
                    'score': score
                })
        
        # Check for suspicious sequences
        for sequence in self.suspicious_sequences:
            if sequence.lower() in text.lower():
                score = 12
                total_score += score
                patterns_found.append({
                    'type': 'suspicious_sequence',
                    'sequence': sequence,
                    'score': score
                })
        
        return {
            'risk_score': min(total_score, 100),
            'patterns_found': patterns_found,
            'analysis': self._get_risk_analysis(total_score)
        }
    
    def analyze_content_redundancy(self, text):
        """Analyze for repetitive content within text"""
        sentences = [s.strip() for s in text.split('.') if s.strip()]
        
        if len(sentences) < 3:
            return {
                'redundancy_score': 0,
                'repeated_phrases': [],
                'analysis': 'Insufficient sentences for redundancy analysis'
            }
        
        # Find repeated phrases
        phrase_counts = Counter()
        repeated_phrases = []
        
        for sentence in sentences:
            words = sentence.split()
            # Check for repeated 3-word phrases
            for i in range(len(words) - 2):
                phrase = ' '.join(words[i:i+3]).lower()
                if len(phrase) > 15:  # Skip very short phrases
                    phrase_counts[phrase] += 1
        
        redundancy_score = 0
        for phrase, count in phrase_counts.items():
            if count > 1:
                score = (count - 1) * 10
                redundancy_score += score
                repeated_phrases.append({
                    'phrase': phrase,
                    'count': count,
                    'score': score
                })
        
        return {
            'redundancy_score': min(redundancy_score, 100),
            'repeated_phrases': repeated_phrases,
            'analysis': f"Found {len(repeated_phrases)} repeated phrases"
        }
    
    def analyze_writing_consistency(self, text):
        """Analyze writing style consistency"""
        sentences = [s.strip() for s in text.split('.') if s.strip()]
        
        if len(sentences) < 5:
            return {
                'consistency_score': 100,
                'issues': [],
                'analysis': 'Text too short for style analysis'
            }
        
        issues = []
        inconsistency_score = 0
        
        # Analyze sentence length variation
        sentence_lengths = [len(s.split()) for s in sentences]
        avg_length = sum(sentence_lengths) / len(sentence_lengths)
        
        # Check for extreme variations
        very_short = [i for i, length in enumerate(sentence_lengths) if length < avg_length * 0.3]
        very_long = [i for i, length in enumerate(sentence_lengths) if length > avg_length * 2.5]
        
        if len(very_short) > len(sentences) * 0.3:
            inconsistency_score += 15
            issues.append({
                'type': 'excessive_short_sentences',
                'count': len(very_short),
                'description': 'Too many very short sentences'
            })
        
        if len(very_long) > len(sentences) * 0.3:
            inconsistency_score += 15
            issues.append({
                'type': 'excessive_long_sentences',
                'count': len(very_long),
                'description': 'Too many very long sentences'
            })
        
        # Check for vocabulary complexity variation
        complex_words = 0
        simple_words = 0
        
        for sentence in sentences:
            words = sentence.split()
            complex_in_sentence = sum(1 for w in words if len(w) > 10)
            simple_in_sentence = sum(1 for w in words if len(w) < 4)
            
            if len(words) > 0:
                complex_ratio = complex_in_sentence / len(words)
                simple_ratio = simple_in_sentence / len(words)
                
                if complex_ratio > 0.3:
                    complex_words += 1
                if simple_ratio > 0.5:
                    simple_words += 1
        
        if complex_words > 0 and simple_words > 0:
            inconsistency_score += 20
            issues.append({
                'type': 'vocabulary_inconsistency',
                'description': 'Mix of very simple and very complex vocabulary'
            })
        
        consistency_score = max(0, 100 - inconsistency_score)
        
        return {
            'consistency_score': consistency_score,
            'issues': issues,
            'analysis': f"Writing consistency: {consistency_score}%"
        }
    
    def calculate_overall_risk(self, pattern_analysis, redundancy_analysis, consistency_analysis):
        """Calculate overall plagiarism risk score"""
        pattern_score = pattern_analysis['risk_score']
        redundancy_score = redundancy_analysis['redundancy_score']
        consistency_penalty = max(0, 100 - consistency_analysis['consistency_score'])
        
        # Weighted calculation
        weights = self.settings
        overall_score = (
            (pattern_score * weights['pattern_weight'] / 100) +
            (redundancy_score * weights['redundancy_weight'] / 100) +
            (consistency_penalty * weights['structure_weight'] / 100)
        )
        
        return min(overall_score, 100)
    
    def _get_risk_analysis(self, score):
        """Get risk analysis based on score"""
        if score >= 70:
            return "🔴 HIGH RISK - Strong indicators of potential plagiarism"
        elif score >= 40:
            return "🟡 MEDIUM RISK - Some suspicious patterns detected"
        elif score >= 20:
            return "🟢 LOW RISK - Minor concerns detected"
        else:
            return "✅ CLEAN - No significant plagiarism indicators"
    
    def _get_overall_risk_status(self, score):
        """Get overall risk status"""
        if score >= 75:
            return "🔴 VERY HIGH RISK - Likely plagiarized content"
        elif score >= 50:
            return "🟡 HIGH RISK - Needs thorough review and paraphrasing"
        elif score >= 25:
            return "🟠 MEDIUM RISK - Some improvements recommended"
        else:
            return "✅ LOW RISK - Content appears original"
    
    def check_paragraph(self, text):
        """Check a single paragraph for plagiarism indicators"""
        # Create cache key
        cache_key = hashlib.md5(text.encode('utf-8')).hexdigest()[:16]
        
        if cache_key in self.analysis_cache:
            return self.analysis_cache[cache_key]
        
        # Perform analyses
        pattern_analysis = self.analyze_text_patterns(text)
        redundancy_analysis = self.analyze_content_redundancy(text)
        consistency_analysis = self.analyze_writing_consistency(text)
        
        # Calculate overall risk
        overall_risk = self.calculate_overall_risk(
            pattern_analysis, redundancy_analysis, consistency_analysis
        )
        
        result = {
            'text_preview': text[:100] + '...' if len(text) > 100 else text,
            'overall_risk_score': round(overall_risk, 1),
            'overall_status': self._get_overall_risk_status(overall_risk),
            'pattern_analysis': pattern_analysis,
            'redundancy_analysis': redundancy_analysis,
            'consistency_analysis': consistency_analysis,
            'recommendation': self._get_recommendation(overall_risk)
        }
        
        # Cache result
        self.analysis_cache[cache_key] = result
        
        return result
    
    def _get_recommendation(self, score):
        """Get recommendation based on risk score"""
        if score >= 75:
            return "🚨 IMMEDIATE ACTION: Completely rewrite this paragraph"
        elif score >= 50:
            return "⚠️ PARAPHRASE REQUIRED: Significant changes needed"
        elif score >= 25:
            return "💡 IMPROVE: Minor paraphrasing recommended"
        else:
            return "✅ GOOD: Content appears original"
    
    def scan_document(self, file_path):
        """Scan entire document for plagiarism indicators"""
        print("=" * 80)
        print("🔍 SMART PLAGIARISM CHECK (OFFLINE)")
        print("=" * 80)
        
        if not os.path.exists(file_path):
            print(f"❌ File not found: {file_path}")
            return None
        
        print(f"📄 Analyzing: {os.path.basename(file_path)}")
        
        try:
            doc = docx.Document(file_path)
            
            scan_results = {
                'filename': os.path.basename(file_path),
                'scan_timestamp': datetime.now().isoformat(),
                'total_paragraphs': len(doc.paragraphs),
                'analyzed_paragraphs': 0,
                'very_high_risk': [],
                'high_risk': [],
                'medium_risk': [],
                'low_risk': [],
                'summary': {}
            }
            
            print(f"📊 Total paragraphs: {len(doc.paragraphs)}")
            print(f"🔍 Starting pattern-based analysis...\n")
            
            for i, paragraph in enumerate(doc.paragraphs):
                para_text = paragraph.text.strip()
                
                if not para_text or len(para_text.split()) < self.settings['min_text_length']:
                    continue
                
                # Skip headers and titles
                if (re.match(r'^(BAB|bab)\s+[IVXLC]+', para_text) or 
                    para_text.isupper() and len(para_text.split()) < 10):
                    continue
                
                print(f"📄 Paragraph {i+1}: ", end="")
                
                # Analyze paragraph
                analysis = self.check_paragraph(para_text)
                
                para_result = {
                    'paragraph_index': i + 1,
                    'text_preview': analysis['text_preview'],
                    'risk_score': analysis['overall_risk_score'],
                    'status': analysis['overall_status'],
                    'recommendation': analysis['recommendation'],
                    'patterns_found': len(analysis['pattern_analysis']['patterns_found']),
                    'detailed_analysis': analysis
                }
                
                # Categorize by risk level
                score = analysis['overall_risk_score']
                if score >= 75:
                    print(f"🔴 VERY HIGH RISK ({score}%) - {analysis['recommendation']}")
                    scan_results['very_high_risk'].append(para_result)
                elif score >= 50:
                    print(f"🟡 HIGH RISK ({score}%) - {analysis['recommendation']}")
                    scan_results['high_risk'].append(para_result)
                elif score >= 25:
                    print(f"🟠 MEDIUM RISK ({score}%) - {analysis['recommendation']}")
                    scan_results['medium_risk'].append(para_result)
                else:
                    print(f"✅ LOW RISK ({score}%) - Content OK")
                    scan_results['low_risk'].append(para_result)
                
                scan_results['analyzed_paragraphs'] += 1
            
            # Generate summary
            self._generate_scan_summary(scan_results)
            
            # Save detailed report
            self._save_analysis_report(scan_results)
            
            return scan_results
            
        except Exception as e:
            print(f"❌ Error scanning document: {e}")
            return None
    
    def _generate_scan_summary(self, results):
        """Generate scan summary"""
        total = results['analyzed_paragraphs']
        very_high = len(results['very_high_risk'])
        high = len(results['high_risk'])
        medium = len(results['medium_risk'])
        low = len(results['low_risk'])
        
        results['summary'] = {
            'total_analyzed': total,
            'very_high_risk_count': very_high,
            'high_risk_count': high,
            'medium_risk_count': medium,
            'low_risk_count': low,
            'needs_immediate_action': very_high,
            'needs_paraphrasing': very_high + high,
            'needs_improvement': very_high + high + medium,
            'risk_percentage': round(((very_high + high) / total) * 100, 1) if total > 0 else 0
        }
        
        print(f"\n" + "=" * 80)
        print("📊 PLAGIARISM ANALYSIS SUMMARY")
        print("=" * 80)
        
        print(f"📄 Document: {results['filename']}")
        print(f"📊 Paragraphs analyzed: {total}")
        print(f"🔴 VERY HIGH RISK (≥75%): {very_high} paragraphs")
        print(f"🟡 HIGH RISK (50-74%): {high} paragraphs") 
        print(f"🟠 MEDIUM RISK (25-49%): {medium} paragraphs")
        print(f"✅ LOW RISK (<25%): {low} paragraphs")
        
        print(f"\n🚨 PRIORITY ACTIONS:")
        print(f"   • Immediate rewrite needed: {very_high} paragraphs")
        print(f"   • Paraphrasing required: {very_high + high} paragraphs")
        print(f"   • Minor improvements: {medium} paragraphs")
        print(f"   • Overall risk level: {results['summary']['risk_percentage']}%")
        
        if very_high > 0:
            print(f"\n🔴 VERY HIGH RISK PARAGRAPHS:")
            for para in results['very_high_risk'][:5]:
                print(f"   Para {para['paragraph_index']}: {para['risk_score']}% - '{para['text_preview']}'")
        
        if high > 0:
            print(f"\n🟡 HIGH RISK PARAGRAPHS:")
            for para in results['high_risk'][:3]:
                print(f"   Para {para['paragraph_index']}: {para['risk_score']}% - '{para['text_preview']}'")
        
        print(f"\n💡 NEXT STEPS:")
        if very_high + high == 0:
            print("   ✅ Document quality is good. Minor improvements may be beneficial.")
        elif results['summary']['risk_percentage'] > 30:
            print("   🚨 High plagiarism risk detected. Extensive paraphrasing required.")
        else:
            print(f"   ⚠️ {very_high + high} paragraphs need attention. Use document_processor.py to paraphrase.")
        
        print("=" * 80)
    
    def _save_analysis_report(self, results):
        """Save detailed analysis report"""
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        report_file = f"plagiarism_analysis_report_{timestamp}.json"
        
        try:
            with open(report_file, 'w', encoding='utf-8') as f:
                json.dump(results, f, ensure_ascii=False, indent=2, default=str)
            
            print(f"\n📋 Detailed analysis report saved: {report_file}")
            
        except Exception as e:
            print(f"⚠️ Could not save report: {e}")


def find_docx_files():
    """Find all .docx files in the project directory with priority for documents/ folder"""
    docx_files = []
    
    # Search in current directory and subdirectories
    for root, dirs, files in os.walk('.'):
        # Skip backup and temporary directories
        if any(skip in root for skip in ['backup', '__pycache__', '.git', 'venv']):
            continue
            
        for file in files:
            if file.endswith('.docx') and not file.startswith('~'):
                full_path = os.path.join(root, file)
                
                # Prioritize files from documents/ folder
                if root == './documents':
                    docx_files.insert(0, full_path)  # Add to beginning
                else:
                    docx_files.append(full_path)  # Add to end
    
    return docx_files

def main():
    """Main function for smart plagiarism checking"""
    print("🔍 SMART PLAGIARISM CHECKER - Pattern-Based Analysis")
    print("🚀 No internet required - Uses advanced pattern detection")
    print("=" * 80)
    
    # Initialize checker
    checker = SmartPlagiarismChecker()
    
    # Find available documents
    print("🔍 Searching for documents...")
    docx_files = find_docx_files()
    
    if not docx_files:
        print("❌ No .docx files found in the project directory")
        print("💡 Please add your document to the project folder")
        return
    
    # Use the first document found
    document_path = docx_files[0]
    
    if len(docx_files) > 1:
        print(f"📄 Found {len(docx_files)} documents. Using: {os.path.basename(document_path)}")
    else:
        print(f"📄 Found document: {os.path.basename(document_path)}")
    
    if os.path.exists(document_path):
        print(f"🎯 Analyzing document for plagiarism patterns...")
        
        results = checker.scan_document(document_path)
        
        if results:
            summary = results['summary']
            
            print(f"\n🎉 Analysis completed!")
            print(f"📈 Risk Assessment: {summary['risk_percentage']}% of content needs attention")
            
            if summary['needs_paraphrasing'] > 0:
                print(f"\n🔧 RECOMMENDED WORKFLOW:")
                print(f"   1. Review {summary['needs_paraphrasing']} high-risk paragraphs")
                print(f"   2. Run: python document_processor.py")
                print(f"   3. Re-analyze with this tool to verify improvement")
                print(f"   4. Focus on aggressive mode for high-risk sections")
        
    else:
        print(f"❌ Document not found: {document_path}")
        print(f"💡 Place your document in the 'documents/' folder")


if __name__ == "__main__":
    main()