# balanced_hybrid_paraphraser_v2.py
"""
Balanced Hybrid Paraphraser v2 - Updated with Google GenAI Package
50% Local + 50% AI menggunakan google-genai package terbaru
Author: DevNoLife
Version: 2.2
"""

import os
import json
import re
import time
import random
from datetime import datetime
from pathlib import Path
import docx
from docx.shared import RGBColor
from docx.enum.text import WD_COLOR_INDEX
import shutil
from collections import defaultdict
import hashlib

# Load environment variables from .env file
try:
    from dotenv import load_dotenv
    load_dotenv()
    print("✅ Environment variables loaded from .env file")
except ImportError:
    print("⚠️  python-dotenv not found. Install with: pip install python-dotenv")
    print("💡 Attempting to load from system environment variables...")

# Updated Gemini AI imports
try:
    import google.generativeai as genai
    GEMINI_AVAILABLE = True
except ImportError:
    print("⚠️  google-generativeai package not found. Install with: pip install google-generativeai")
    GEMINI_AVAILABLE = False

# Import sistem parafrase lokal
from main_paraphrase_system import IndonesianParaphraseSystem

class BalancedHybridParaphraserV2:
    def __init__(self, synonym_file='sinonim.json', gemini_api_key=None):
        print("🚀 Initializing Balanced Hybrid Paraphraser v2 (Google GenAI)...")
        
        # Initialize local paraphrase system
        self.local_paraphraser = IndonesianParaphraseSystem(synonym_file)
        
        # Gemini AI configuration with new API
        # Get API key from parameter, environment, or fallback
        self.gemini_api_key = gemini_api_key or os.environ.get("GEMINI_API_KEY")
        self.gemini_client = None
        
        if self.gemini_api_key and GEMINI_AVAILABLE:
            try:
                # Configure the client properly with explicit API key
                genai.configure(api_key=self.gemini_api_key)
                # Verify the configuration worked
                print(f"🔑 API Key configured: {self.gemini_api_key[:10]}...")
                self.gemini_client = genai.GenerativeModel('gemini-1.5-flash')
                print("✅ Gemini client initialized")
            except Exception as e:
                print(f"⚠️  Error initializing Gemini client: {e}")
                print(f"🔍 API Key length: {len(self.gemini_api_key) if self.gemini_api_key else 'None'}")
                self.gemini_client = None
        elif not self.gemini_api_key:
            print("⚠️  No Gemini API key provided")
        elif not GEMINI_AVAILABLE:
            print("⚠️  Google GenAI package not available")
        
        # Cost tracking
        self.cost_tracker = {
            'local_calls': 0,
            'ai_calls': 0,
            'ai_tokens_used': 0,
            'estimated_cost_usd': 0.0,
            'cache_hits': 0,
            'local_vs_ai_comparison': []
        }
        
        # AI result cache
        self.ai_cache = {}
        
        # Balanced configuration
        self.config = {
            'force_ai_ratio': 0.5,              # Target 50% AI usage
            'local_confidence_threshold': 0.15,  # Lower threshold (15% vs 25%)
            'complexity_threshold': 0.4,         # Lower complexity threshold  
            'min_paragraph_length': 25,          # Lower minimum words
            'ai_retry_attempts': 3,              # More retry attempts
            'random_ai_selection': True,         # Randomly select some for AI
            'comparison_mode': True,             # Compare local vs AI results
            'model_name': 'gemini-1.5-flash'    # Updated model
        }
        
        # Academic patterns - more inclusive
        self.ai_priority_patterns = [
            r'menurut\s+\w+',                   # Any citation
            r'berdasarkan',                     # Based on
            r'definisi\s+\w+',                  # Any definition
            r'konsep\s+\w+',                    # Any concept
            r'teori\s+\w+',                     # Any theory
            r'sistem\s+\w+',                    # Any system
            r'penelitian\s+\w+',                # Any research
            r'analisis\s+\w+',                  # Any analysis
            r'metode\s+\w+',                    # Any method
            r'hasil\s+\w+',                     # Any result
        ]
        
        print(f"✅ Local system ready with {len(self.local_paraphraser.synonyms)} synonyms")
        print(f"✅ Gemini AI {'configured (2.0 Flash)' if self.gemini_client else 'not available'}")
        print(f"🎯 Target AI usage: {self.config['force_ai_ratio']*100}%")
        print("✅ Balanced Hybrid Paraphraser v2 ready!")
    
    def should_use_ai_balanced(self, paragraph_text, local_result, paragraph_index, total_paragraphs):
        """Balanced routing: aim for 50-50 distribution"""
        
        # Force some paragraphs to use AI for comparison
        if self.config['random_ai_selection']:
            # Use paragraph index to create deterministic but distributed AI usage
            ai_probability = 0.5  # 50% base probability
            
            # Increase probability for academic content
            complexity = self.calculate_paragraph_complexity(paragraph_text)
            if complexity > 0.5:
                ai_probability += 0.2
            
            # Pattern matching bonus
            pattern_matches = sum(1 for pattern in self.ai_priority_patterns if re.search(pattern, paragraph_text.lower()))
            if pattern_matches > 0:
                ai_probability += 0.1 * pattern_matches
            
            # Length bonus
            word_count = len(paragraph_text.split())
            if word_count > 50:
                ai_probability += 0.1
            
            # Random selection based on probability
            random_factor = (paragraph_index * 7) % 100 / 100  # Deterministic "random"
            
            if random_factor < ai_probability:
                return True, f"Balanced selection (prob: {ai_probability:.2f})"
        
        # Traditional criteria (lowered thresholds)
        if local_result['plagiarism_reduction'] < self.config['local_confidence_threshold'] * 100:
            return True, f"Low local performance ({local_result['plagiarism_reduction']:.1f}%)"
        
        # Check paragraph length (lowered threshold)
        word_count = len(paragraph_text.split())
        if word_count < self.config['min_paragraph_length']:
            return False, "Paragraph too short"
        
        # Complexity check (lowered threshold)
        complexity = self.calculate_paragraph_complexity(paragraph_text)
        if complexity >= self.config['complexity_threshold']:
            return True, f"Moderate complexity ({complexity:.2f})"
        
        # Academic patterns (more generous)
        pattern_matches = sum(1 for pattern in self.ai_priority_patterns if re.search(pattern, paragraph_text.lower()))
        if pattern_matches >= 1:  # Even 1 pattern is enough
            return True, f"Academic pattern detected ({pattern_matches})"
        
        return False, "Local processing preferred"
    
    def calculate_paragraph_complexity(self, text):
        """Calculate complexity score"""
        words = text.split()
        sentences = text.split('.')
        
        length_score = min(len(words) / 80, 1.0)  # Lower normalization
        
        academic_words = [
            'penelitian', 'analisis', 'metode', 'sistem', 'implementasi',
            'evaluasi', 'teori', 'konsep', 'pendekatan', 'metodologi',
            'definisi', 'klasifikasi', 'kategori', 'karakteristik'
        ]
        academic_density = sum(1 for word in words if word.lower() in academic_words) / len(words)
        
        citation_count = len(re.findall(r'\(\d{4}\)|et al|vol\.|no\.', text))
        citation_score = min(citation_count / 2, 1.0)  # Lower threshold
        
        pattern_score = sum(1 for pattern in self.ai_priority_patterns if re.search(pattern, text.lower())) / len(self.ai_priority_patterns)
        
        avg_sentence_length = len(words) / len(sentences) if sentences else 0
        sentence_complexity = min(avg_sentence_length / 20, 1.0)  # Lower threshold
        
        complexity = (
            length_score * 0.25 +
            academic_density * 0.35 +
            citation_score * 0.15 +
            pattern_score * 0.15 +
            sentence_complexity * 0.1
        )
        
        return min(complexity, 1.0)
    
    def create_cache_key(self, text):
        """Create cache key for AI results"""
        return hashlib.md5(text.encode('utf-8')).hexdigest()[:16]
    
    def call_gemini_ai_v2(self, paragraphs_batch):
        """Enhanced Gemini API call using google-generativeai package"""
        if not self.gemini_client:
            print("❌ Gemini client not initialized")
            return None
        
        if not self.gemini_api_key:
            print("❌ No API key available")
            return None
        
        try:
            prompt = self.create_enhanced_gemini_prompt(paragraphs_batch)
            
            # Configure generation
            generation_config = genai.GenerationConfig(
                temperature=0.4,
                top_k=50,
                top_p=0.9,
                max_output_tokens=6144,
            )
            
            for attempt in range(self.config['ai_retry_attempts']):
                try:
                    # Generate content using standard API
                    response = self.gemini_client.generate_content(
                        prompt,
                        generation_config=generation_config,
                    )
                    
                    if response.text and response.text.strip():
                        # Track usage
                        self.cost_tracker['ai_calls'] += 1
                        tokens_used = len(prompt + response.text) // 4
                        self.cost_tracker['ai_tokens_used'] += tokens_used
                        # Gemini Flash pricing (estimated)
                        self.cost_tracker['estimated_cost_usd'] += tokens_used * 0.000001  # Very rough estimate
                        
                        return self.parse_gemini_response(response.text, paragraphs_batch)
                    
                except Exception as e:
                    error_msg = str(e)
                    print(f"⚠️  Gemini API attempt {attempt + 1} failed: {error_msg}")
                    
                    # Special handling for API key errors
                    if "API_KEY_INVALID" in error_msg or "API Key not found" in error_msg:
                        print("🔍 API Key issue detected. Trying to reconfigure...")
                        try:
                            genai.configure(api_key=self.gemini_api_key)
                            print("✅ API key reconfigured")
                        except Exception as reconfig_error:
                            print(f"❌ Failed to reconfigure API key: {reconfig_error}")
                            return None
                    
                    if attempt < self.config['ai_retry_attempts'] - 1:
                        time.sleep(3 ** attempt)  # Exponential backoff
            
            return None
            
        except Exception as e:
            print(f"❌ Gemini API error: {e}")
            return None
    
    def create_enhanced_gemini_prompt(self, paragraphs_batch):
        """Enhanced prompt for better AI results"""
        prompt = """Kamu adalah ahli paraphrasing untuk teks akademik bahasa Indonesia yang sangat berpengalaman.

MISI: Parafrase paragraf berikut untuk mengurangi plagiarisme secara signifikan sambil mempertahankan makna akademik yang tepat.

STRATEGI PARAPHRASING:
✅ Gunakan sinonim yang tepat dan kontekstual
✅ Ubah struktur kalimat (aktif↔pasif, urutan klausa)
✅ Ganti frasa akademik dengan variasi yang setara
✅ Ubah konstruksi gramatikal tanpa mengubah makna
✅ Pertahankan tone formal dan akademik
✅ PENTING: Hasil harus natural dan mudah dibaca

CONTOH TRANSFORMASI:
❌ Buruk: "Penelitian ini menunjukkan bahwa..." → "Riset ini memperlihatkan jika..."
✅ Baik: "Penelitian ini menunjukkan bahwa..." → "Hasil kajian mengindikasikan bahwa..."

TARGET: Minimal 40-60% pengurangan similarity

FORMAT JAWABAN:
PARAGRAF_[NOMOR]:
[Hasil parafrase yang natural dan akademis]

PARAGRAF YANG PERLU DIPARAFRASE:
"""
        
        for i, paragraph in enumerate(paragraphs_batch, 1):
            prompt += f"\nPARAGRAF_{i}:\n{paragraph['text']}\n"
        
        prompt += "\n🎯 MULAI PARAPHRASING SEKARANG (Target: 40-60% similarity reduction):"
        
        return prompt
    
    def parse_gemini_response(self, ai_response, original_paragraphs):
        """Enhanced response parsing"""
        results = []
        
        try:
            # Split by paragraph markers
            sections = re.split(r'PARAGRAF_(\d+):', ai_response)
            
            for i in range(1, len(sections), 2):
                if i + 1 < len(sections):
                    paragraph_num = int(sections[i]) - 1
                    paraphrased_text = sections[i + 1].strip()
                    
                    # Clean up the response
                    paraphrased_text = re.sub(r'^[-•*]\s*', '', paraphrased_text)  # Remove bullets
                    paraphrased_text = re.sub(r'\n+', ' ', paraphrased_text)        # Single line
                    paraphrased_text = paraphrased_text.strip()
                    
                    if paragraph_num < len(original_paragraphs) and paraphrased_text:
                        original_text = original_paragraphs[paragraph_num]['text']
                        
                        # Calculate similarity
                        similarity = self.local_paraphraser.calculate_similarity(original_text, paraphrased_text)
                        
                        result = {
                            'paraphrase': paraphrased_text,
                            'similarity': round(similarity, 2),
                            'plagiarism_reduction': round(100 - similarity, 2),
                            'changes_made': 1,
                            'method': 'gemini_2.0_flash',
                            'status': self.local_paraphraser._get_plagiarism_status(similarity),
                            'original_length': len(original_text.split()),
                            'paraphrase_length': len(paraphrased_text.split())
                        }
                        
                        results.append(result)
            
            return results
            
        except Exception as e:
            print(f"⚠️  Error parsing Gemini response: {e}")
            return []
    
    def process_paragraph_balanced(self, paragraph_text, paragraph_index, total_paragraphs):
        """Balanced processing with local vs AI comparison"""
        
        # Always get local result first
        local_result = self.local_paraphraser.generate_paraphrase(paragraph_text, aggressiveness=0.6)
        local_result['method'] = 'local'
        self.cost_tracker['local_calls'] += 1
        
        # Decide if we should also try AI
        use_ai, reason = self.should_use_ai_balanced(paragraph_text, local_result, paragraph_index, total_paragraphs)
        
        if not use_ai:
            local_result['routing_reason'] = reason
            return local_result
        
        # Check cache first
        cache_key = self.create_cache_key(paragraph_text)
        if cache_key in self.ai_cache:
            self.cost_tracker['cache_hits'] += 1
            ai_result = self.ai_cache[cache_key].copy()
            ai_result['method'] = 'ai_cached'
        else:
            # Call AI using new API
            print(f"    🤖 Using Gemini 2.0: {reason}")
            ai_results = self.call_gemini_ai_v2([{'text': paragraph_text}])
            
            if ai_results and len(ai_results) > 0:
                ai_result = ai_results[0]
                ai_result['method'] = 'gemini_2.0_flash'
                # Cache the result
                self.ai_cache[cache_key] = ai_result.copy()
            else:
                # AI failed, use local
                local_result['method'] = 'local_fallback'
                local_result['routing_reason'] = f"{reason} (AI failed)"
                return local_result
        
        # Compare local vs AI results
        if self.config['comparison_mode']:
            comparison = {
                'paragraph_index': paragraph_index,
                'local_reduction': local_result['plagiarism_reduction'],
                'ai_reduction': ai_result['plagiarism_reduction'],
                'improvement': ai_result['plagiarism_reduction'] - local_result['plagiarism_reduction'],
                'winner': 'ai' if ai_result['plagiarism_reduction'] > local_result['plagiarism_reduction'] else 'local'
            }
            self.cost_tracker['local_vs_ai_comparison'].append(comparison)
            
            print(f"    📊 Comparison: Local {local_result['plagiarism_reduction']:.1f}% vs AI {ai_result['plagiarism_reduction']:.1f}%")
        
        # Choose the better result
        if ai_result['plagiarism_reduction'] > local_result['plagiarism_reduction']:
            ai_result['routing_reason'] = f"{reason} (AI won: +{ai_result['plagiarism_reduction'] - local_result['plagiarism_reduction']:.1f}%)"
            return ai_result
        else:
            local_result['method'] = 'local_better'
            local_result['routing_reason'] = f"{reason} (Local won: +{local_result['plagiarism_reduction'] - ai_result['plagiarism_reduction']:.1f}%)"
            return local_result
    
    def process_batch_documents(self, input_folder, aggressiveness=0.6, create_backup=True):
        """Process batch with balanced approach using Gemini 2.0"""
        start_time = datetime.now()
        
        print("=" * 80)
        print("🎯 BALANCED HYBRID PARAPHRASER v2 (Gemini 2.0 Flash)")
        print("=" * 80)
        
        if not os.path.exists(input_folder):
            print(f"❌ Input folder not found: {input_folder}")
            return
        
        docx_files = [f for f in os.listdir(input_folder) if f.endswith('.docx') and not f.startswith('~')]
        
        if not docx_files:
            print(f"❌ No Word documents found in: {input_folder}")
            return
        
        print(f"📁 Input folder: {input_folder}")
        print(f"📄 Found {len(docx_files)} documents")
        print(f"⚙️  Aggressiveness level: {aggressiveness}")
        print(f"🤖 AI Integration: {'Enabled (Gemini 2.0 Flash)' if self.gemini_client else 'Disabled'}")
        
        # Create backup
        backup_folder = None
        if create_backup:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            backup_folder = f"{input_folder}_backup_v2_{timestamp}"
            
            try:
                if not os.path.exists(backup_folder):
                    os.makedirs(backup_folder)
                
                for file in docx_files:
                    src = os.path.join(input_folder, file)
                    dst = os.path.join(backup_folder, file)
                    shutil.copy2(src, dst)
                
                print(f"✅ Backup created: {backup_folder}")
            except Exception as e:
                print(f"❌ Error creating backup: {e}")
                return
        
        print(f"\n🚀 Starting balanced processing with Gemini 2.0...")
        
        # Process documents
        total_stats = {
            'documents_processed': 0,
            'paragraphs_processed': 0,
            'local_only': 0,
            'ai_enhanced': 0,
            'ai_cached': 0,
            'local_better': 0,
            'total_changes': 0
        }
        
        for i, filename in enumerate(docx_files, 1):
            file_path = os.path.join(input_folder, filename)
            
            print(f"\n{'='*20} Document {i}/{len(docx_files)} {'='*20}")
            
            doc_stats = self.process_single_document_balanced(file_path, aggressiveness)
            
            if doc_stats:
                total_stats['documents_processed'] += 1
                for key in ['paragraphs_processed', 'local_only', 'ai_enhanced', 'ai_cached', 'local_better', 'total_changes']:
                    total_stats[key] += doc_stats.get(key, 0)
        
        # Calculate processing time
        end_time = datetime.now()
        processing_time = (end_time - start_time).total_seconds()
        
        # Print comprehensive summary
        self.print_balanced_summary(total_stats, processing_time, backup_folder)
        
        # Generate detailed report
        self.generate_balanced_report(input_folder, total_stats, processing_time)
    
    def process_single_document_balanced(self, file_path, aggressiveness):
        """Process single document with balanced approach"""
        try:
            print(f"\n📄 Processing: {os.path.basename(file_path)}")
            
            doc = docx.Document(file_path)
            
            doc_stats = {
                'filename': os.path.basename(file_path),
                'total_paragraphs': 0,
                'paragraphs_processed': 0,
                'local_only': 0,
                'ai_enhanced': 0,
                'ai_cached': 0,
                'local_better': 0,
                'total_changes': 0,
                'methods_used': defaultdict(int)
            }
            
            current_section = 'UNKNOWN'
            processed_paragraphs = []
            
            # First pass: collect suitable paragraphs
            for paragraph in doc.paragraphs:
                para_text = paragraph.text.strip()
                doc_stats['total_paragraphs'] += 1
                
                if not para_text:
                    continue
                
                if self.is_section_header(para_text):
                    current_section = para_text
                    print(f"  📍 Section: {current_section}")
                    continue
                
                if self.is_paragraph_suitable_for_paraphrasing(para_text):
                    processed_paragraphs.append({
                        'paragraph': paragraph,
                        'text': para_text,
                        'index': len(processed_paragraphs)
                    })
            
            total_suitable = len(processed_paragraphs)
            print(f"  📝 Found {total_suitable} suitable paragraphs for processing")
            
            # Second pass: process paragraphs
            for para_data in processed_paragraphs:
                paragraph = para_data['paragraph']
                para_text = para_data['text']
                para_index = para_data['index']
                
                # Balanced processing
                result = self.process_paragraph_balanced(para_text, para_index, total_suitable)
                
                # Apply changes if improvement is significant
                if result['plagiarism_reduction'] > 15:  # Threshold
                    paragraph.clear()
                    paragraph.add_run(result['paraphrase'])
                    
                    # Color coding based on method
                    if 'gemini' in result['method'] or 'ai' in result['method']:
                        # Blue for AI
                        for run in paragraph.runs:
                            run.font.highlight_color = WD_COLOR_INDEX.TURQUOISE
                        if 'cached' in result['method']:
                            doc_stats['ai_cached'] += 1
                        else:
                            doc_stats['ai_enhanced'] += 1
                    elif result['method'] == 'local_better':
                        # Green for local that beat AI
                        for run in paragraph.runs:
                            run.font.highlight_color = WD_COLOR_INDEX.BRIGHT_GREEN
                        doc_stats['local_better'] += 1
                    else:
                        # Yellow for local only
                        for run in paragraph.runs:
                            run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                        doc_stats['local_only'] += 1
                    
                    doc_stats['paragraphs_processed'] += 1
                    doc_stats['total_changes'] += result.get('changes_made', 1)
                    doc_stats['methods_used'][result['method']] += 1
                    
                    print(f"    ✅ Para {para_index+1}: {result['plagiarism_reduction']:.1f}% reduction ({result['method']})")
                else:
                    print(f"    ⏭️  Para {para_index+1}: Skipped ({result['plagiarism_reduction']:.1f}% - {result['method']})")
            
            # Save document
            doc.save(file_path)
            
            print(f"  📊 Document Summary:")
            print(f"     • Total paragraphs: {doc_stats['total_paragraphs']}")
            print(f"     • Processed: {doc_stats['paragraphs_processed']}")
            print(f"     • Local-only: {doc_stats['local_only']}")
            print(f"     • AI-enhanced: {doc_stats['ai_enhanced']}")
            print(f"     • AI-cached: {doc_stats['ai_cached']}")
            print(f"     • Local-better: {doc_stats['local_better']}")
            print(f"  ✅ Saved: {file_path}")
            
            return doc_stats
            
        except Exception as e:
            print(f"❌ Error processing {file_path}: {e}")
            return None
    
    def is_section_header(self, text):
        """Check if text is a section header"""
        patterns = [
            r'^BAB\s+[IVX]+',
            r'^\d+\.\d+',
            r'^[A-Z\s]+$',
        ]
        
        for pattern in patterns:
            if re.match(pattern, text.strip()) and len(text.split()) < 10:
                return True
        return False
    
    def is_paragraph_suitable_for_paraphrasing(self, paragraph_text):
        """More lenient suitability check"""
        word_count = len(paragraph_text.split())
        if word_count < 15:  # Lower minimum
            return False
        
        if re.search(r'\d+\.\d+|\btabel\b|\bgambar\b', paragraph_text.lower()):
            return False
        
        if paragraph_text.isupper() and word_count < 10:
            return False
        
        return True
    
    def print_balanced_summary(self, stats, processing_time, backup_folder):
        """Print balanced processing summary"""
        print("\n" + "=" * 80)
        print("📊 BALANCED HYBRID PROCESSING SUMMARY (Gemini 2.0)")
        print("=" * 80)
        
        print(f"📄 Documents processed: {stats['documents_processed']}")
        print(f"📝 Paragraphs processed: {stats['paragraphs_processed']}")
        print(f"🏠 Local-only: {stats['local_only']}")
        print(f"🤖 AI-enhanced: {stats['ai_enhanced']}")
        print(f"💾 AI-cached: {stats['ai_cached']}")
        print(f"🏆 Local-better: {stats['local_better']}")
        print(f"🔄 Total changes: {stats['total_changes']}")
        print(f"⏱️  Processing time: {processing_time:.1f} seconds")
        
        # Calculate percentages
        if stats['paragraphs_processed'] > 0:
            local_pct = ((stats['local_only'] + stats['local_better']) / stats['paragraphs_processed']) * 100
            ai_pct = ((stats['ai_enhanced'] + stats['ai_cached']) / stats['paragraphs_processed']) * 100
            print(f"📈 Distribution: {local_pct:.1f}% Local, {ai_pct:.1f}% AI")
        
        # Cost summary
        print(f"\n💰 COST SUMMARY (Gemini 2.0 Flash):")
        print(f"💻 Local calls: {self.cost_tracker['local_calls']} (FREE)")
        print(f"🤖 AI calls: {self.cost_tracker['ai_calls']}")
        print(f"💾 Cache hits: {self.cost_tracker['cache_hits']} (FREE)")
        print(f"🔢 Tokens used: {self.cost_tracker['ai_tokens_used']:,}")
        print(f"💵 Estimated cost: ${self.cost_tracker['estimated_cost_usd']:.4f} (~Rp {self.cost_tracker['estimated_cost_usd'] * 15000:.0f})")
        
        # AI vs Local comparison
        if self.cost_tracker['local_vs_ai_comparison']:
            comparisons = self.cost_tracker['local_vs_ai_comparison']
            ai_wins = sum(1 for c in comparisons if c['winner'] == 'ai')
            local_wins = sum(1 for c in comparisons if c['winner'] == 'local')
            print(f"\n🏆 PERFORMANCE COMPARISON:")
            print(f"   • AI wins: {ai_wins}/{len(comparisons)} ({ai_wins/len(comparisons)*100:.1f}%)")
            print(f"   • Local wins: {local_wins}/{len(comparisons)} ({local_wins/len(comparisons)*100:.1f}%)")
            
            avg_improvement = sum(c.get('improvement', 0) for c in comparisons) / len(comparisons)
            print(f"   • Average AI improvement: {avg_improvement:.1f}%")
        
        if backup_folder:
            print(f"\n💾 Backup: {backup_folder}")
        
        print("=" * 80)
    
    def generate_balanced_report(self, input_folder, stats, processing_time):
        """Generate detailed balanced report"""
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        report_file = os.path.join(input_folder, f"balanced_v2_report_{timestamp}.json")
        
        try:
            report_data = {
                'timestamp': timestamp,
                'processing_stats': stats,
                'cost_tracking': self.cost_tracker,
                'configuration': self.config,
                'processing_time_seconds': processing_time,
                'local_vs_ai_comparison': self.cost_tracker['local_vs_ai_comparison'],
                'system_info': {
                    'local_synonyms': len(self.local_paraphraser.synonyms),
                    'local_phrases': len(self.local_paraphraser.phrase_replacements),
                    'gemini_enabled': bool(self.gemini_client),
                    'gemini_model': self.config['model_name'],
                    'google_genai_version': '2.0',
                    'version': '2.2'
                }
            }
            
            with open(report_file, 'w', encoding='utf-8') as f:
                json.dump(report_data, f, ensure_ascii=False, indent=2)
            
            print(f"📋 Balanced v2 report saved: {report_file}")
            
        except Exception as e:
            print(f"⚠️  Could not save report: {e}")


def main():
    """Main function for balanced hybrid processing v2"""
    print("🎯 BALANCED HYBRID PARAPHRASER v2 - GEMINI 2.0 FLASH")
    print("50% Local + 50% AI menggunakan Google GenAI package terbaru")
    print("=" * 80)
    
    # Check dependencies
    if not GEMINI_AVAILABLE:
        print("❌ Required package missing!")
        print("📦 Install with: pip install google-generativeai")
        return
    
    # Configuration
    SYNONYM_FILE = 'sinonim.json'           # Your large synonym database
    INPUT_FOLDER = 'documents'              # Folder with Word documents
    AGGRESSIVENESS = 0.6                    # Processing aggressiveness
    CREATE_BACKUP = True                    # Always backup first
    
    # Get API key from environment or set directly
    GEMINI_API_KEY = os.environ.get("GEMINI_API_KEY")
    
    if not GEMINI_API_KEY:
        print("⚠️  WARNING: GEMINI_API_KEY not found!")
        print("🔑 Set your API key:")
        print("   Method 1: export GEMINI_API_KEY='your-api-key-here'")
        print("   Method 2: Set in environment variables")
        print("   Method 3: Uncomment and set directly in code below")
        print()
        # Uncomment and set your API key here if needed:
        # GEMINI_API_KEY = "your-gemini-api-key-here"
        
        if not GEMINI_API_KEY:
            print("❌ Cannot proceed without API key")
            print("📖 Get your free API key from: https://makersuite.google.com/app/apikey")
            return
    
    print(f"✅ API key configured: {GEMINI_API_KEY[:20]}...")
    
    # Initialize balanced processor v2
    processor = BalancedHybridParaphraserV2(
        synonym_file=SYNONYM_FILE,
        gemini_api_key=GEMINI_API_KEY
    )
    
    # Process batch with balanced approach
    processor.process_batch_documents(
        input_folder=INPUT_FOLDER,
        aggressiveness=AGGRESSIVENESS,
        create_backup=CREATE_BACKUP
    )
    
    print("\n🎉 Balanced hybrid v2 processing completed!")
    print("💡 Color coding:")
    print("   🟡 Yellow = Local only")
    print("   🔵 Blue = AI enhanced (Gemini 2.0)")
    print("   🟢 Green = Local beat AI")
    print("\n📊 Check the detailed report for Local vs AI performance comparison!")
    print("🚀 Powered by Gemini 2.0 Flash - Latest Google AI model!")


if __name__ == "__main__":
    main()
