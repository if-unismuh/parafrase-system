# smart_turnitin_detector.py
"""
Smart Turnitin-Like Detector + Selective Paraphraser
Hanya memarafrase bagian yang benar-benar mirip dengan sumber online/database
Seperti cara kerja Turnitin yang membandingkan dengan database mereka
Author: DevNoLife
Version: 4.0
"""

import os
import json
import re
import time
import random
import hashlib
import requests
from datetime import datetime
from pathlib import Path
import docx
from docx.shared import RGBColor
from docx.enum.text import WD_COLOR_INDEX
import shutil
from collections import defaultdict
from difflib import SequenceMatcher
import threading
from concurrent.futures import ThreadPoolExecutor, as_completed

# Updated Gemini AI imports
try:
    from google import genai
    from google.genai import types
    GEMINI_AVAILABLE = True
except ImportError:
    print("⚠️  google-genai package not found. Install with: pip install google-genai")
    GEMINI_AVAILABLE = False

# Import sistem parafrase lokal
from main_paraphrase_system import IndonesianParaphraseSystem

class SmartTurnitinDetector:
    def __init__(self, synonym_file='sinonim.json', gemini_api_key=None):
        print("🔍 Initializing Smart Turnitin-Like Detector...")
        
        # Initialize local paraphrase system
        self.local_paraphraser = IndonesianParaphraseSystem(synonym_file)
        
        # Gemini AI configuration
        self.gemini_api_key = gemini_api_key
        self.gemini_client = None
        
        if gemini_api_key and GEMINI_AVAILABLE:
            try:
                os.environ["GEMINI_API_KEY"] = gemini_api_key
                self.gemini_client = genai.Client(api_key=gemini_api_key)
                print("✅ Gemini 2.0 Flash configured for selective paraphrasing")
            except Exception as e:
                print(f"⚠️  Error initializing Gemini: {e}")
                self.gemini_client = None
        
        # Similarity detection database (simulasi database Turnitin)
        self.similarity_database = self.load_similarity_database()
        
        # Cost tracking
        self.cost_tracker = {
            'local_calls': 0,
            'ai_calls': 0,
            'similarity_checks': 0,
            'flagged_segments': 0,
            'paraphrased_segments': 0,
            'ai_tokens_used': 0,
            'estimated_cost_usd': 0.0
        }
        
        # Configuration untuk similarity detection
        self.config = {
            'similarity_threshold': 0.75,        # 75% similarity = flagged
            'min_segment_length': 10,            # Minimum 10 kata untuk dicheck
            'max_segment_length': 50,            # Maximum 50 kata per segment
            'chunk_overlap': 5,                  # Overlap antar chunk
            'ngram_size': 5,                     # N-gram untuk matching
            'ai_improvement_threshold': 0.6,     # Gunakan AI jika local < 60% improvement
            'model_name': 'gemini-2.0-flash',
            'ai_retry_attempts': 2
        }
        
        # Pattern library (dari analisis Turnitin hasil Anda)
        self.common_academic_patterns = [
            "penelitian ini bertujuan untuk",
            "berdasarkan hasil penelitian",
            "menurut pendapat para ahli",
            "dalam konteks penelitian ini",
            "hasil analisis menunjukkan bahwa",
            "dapat disimpulkan bahwa",
            "berdasarkan latar belakang",
            "rumusan masalah dalam penelitian",
            "tujuan penelitian ini adalah",
            "manfaat penelitian ini",
            "batasan masalah penelitian",
            "definisi operasional dalam penelitian",
            "kajian teori yang relevan",
            "metodologi penelitian yang digunakan",
            "teknik pengumpulan data",
            "analisis data menggunakan",
            "sistem informasi merupakan",
            "teknologi informasi adalah",
            "implementasi sistem dapat",
            "pengembangan aplikasi berbasis"
        ]
        
        # Cache untuk hasil similarity check
        self.similarity_cache = {}
        
        print(f"✅ Similarity database loaded: {len(self.similarity_database)} patterns")
        print(f"✅ Academic patterns loaded: {len(self.common_academic_patterns)} patterns")
        print("✅ Smart Turnitin Detector ready!")
    
    def load_similarity_database(self):
        """Load atau create similarity database (simulasi database Turnitin)"""
        db_file = "similarity_database.json"
        
        if os.path.exists(db_file):
            try:
                with open(db_file, 'r', encoding='utf-8') as f:
                    return json.load(f)
            except:
                pass
        
        # Create default database berdasarkan pattern umum
        default_db = {
            "academic_boilerplate": [
                "penelitian ini bertujuan untuk mengembangkan sistem yang dapat",
                "berdasarkan latar belakang yang telah diuraikan di atas",
                "rumusan masalah dalam penelitian ini adalah sebagai berikut",
                "tujuan dari penelitian ini adalah untuk mengetahui",
                "manfaat penelitian ini diharapkan dapat memberikan kontribusi",
                "batasan masalah dalam penelitian ini meliputi",
                "definisi operasional yang digunakan dalam penelitian ini",
                "kajian pustaka yang relevan dengan penelitian ini",
                "metodologi penelitian yang digunakan dalam penelitian ini",
                "teknik pengumpulan data dalam penelitian ini menggunakan",
                "analisis data dalam penelitian ini menggunakan metode",
                "hasil penelitian menunjukkan bahwa sistem yang dikembangkan",
                "kesimpulan dari penelitian ini adalah bahwa sistem",
                "saran untuk penelitian selanjutnya adalah agar dapat"
            ],
            "technical_definitions": [
                "sistem informasi adalah kombinasi dari teknologi informasi",
                "basis data merupakan kumpulan data yang tersimpan secara sistematis",
                "aplikasi web adalah aplikasi yang dapat diakses melalui browser",
                "framework adalah kerangka kerja untuk pengembangan aplikasi",
                "algoritma adalah urutan langkah-langkah untuk menyelesaikan masalah",
                "pemrograman berorientasi objek adalah paradigma pemrograman",
                "database management system adalah software yang digunakan",
                "user interface adalah antarmuka yang memungkinkan interaksi",
                "system development life cycle adalah metodologi pengembangan",
                "software engineering adalah disiplin ilmu yang berkaitan"
            ],
            "methodology_patterns": [
                "metode penelitian yang digunakan dalam penelitian ini adalah",
                "populasi dalam penelitian ini adalah seluruh pengguna sistem",
                "sampel penelitian ini diambil dengan teknik purposive sampling",
                "instrumen penelitian yang digunakan berupa kuesioner dan wawancara",
                "teknik analisis data menggunakan analisis deskriptif dan inferensial",
                "validitas instrument diuji menggunakan validitas konstruk",
                "reliabilitas instrument diuji menggunakan cronbach alpha",
                "pengujian hipotesis menggunakan teknik analisis regresi",
                "pengumpulan data dilakukan dengan cara observasi dan dokumentasi",
                "analisis kebutuhan sistem dilakukan melalui wawancara dengan user"
            ]
        }
        
        # Save default database
        try:
            with open(db_file, 'w', encoding='utf-8') as f:
                json.dump(default_db, f, ensure_ascii=False, indent=2)
        except:
            pass
        
        return default_db
    
    def create_text_ngrams(self, text, n=5):
        """Create n-grams from text for similarity matching"""
        words = re.findall(r'\w+', text.lower())
        if len(words) < n:
            return [' '.join(words)]
        
        ngrams = []
        for i in range(len(words) - n + 1):
            ngrams.append(' '.join(words[i:i+n]))
        return ngrams
    
    def check_similarity_with_database(self, text_segment):
        """Check similarity dengan database (seperti Turnitin)"""
        cache_key = hashlib.md5(text_segment.encode()).hexdigest()[:12]
        
        if cache_key in self.similarity_cache:
            return self.similarity_cache[cache_key]
        
        self.cost_tracker['similarity_checks'] += 1
        
        max_similarity = 0.0
        best_match = None
        matched_source = None
        
        # Check against all categories in database
        for category, patterns in self.similarity_database.items():
            for pattern in patterns:
                # Calculate similarity using SequenceMatcher
                similarity = SequenceMatcher(None, text_segment.lower(), pattern.lower()).ratio()
                
                if similarity > max_similarity:
                    max_similarity = similarity
                    best_match = pattern
                    matched_source = category
        
        # Check against common academic patterns
        for pattern in self.common_academic_patterns:
            similarity = SequenceMatcher(None, text_segment.lower(), pattern.lower()).ratio()
            
            if similarity > max_similarity:
                max_similarity = similarity
                best_match = pattern
                matched_source = "academic_patterns"
        
        # Additional check: word-level similarity
        segment_words = set(re.findall(r'\w+', text_segment.lower()))
        for category, patterns in self.similarity_database.items():
            for pattern in patterns:
                pattern_words = set(re.findall(r'\w+', pattern.lower()))
                if segment_words and pattern_words:
                    word_similarity = len(segment_words.intersection(pattern_words)) / len(segment_words.union(pattern_words))
                    if word_similarity > max_similarity:
                        max_similarity = word_similarity
                        best_match = pattern
                        matched_source = f"{category}_words"
        
        result = {
            'similarity': max_similarity,
            'is_flagged': max_similarity >= self.config['similarity_threshold'],
            'best_match': best_match,
            'source': matched_source,
            'flagged': max_similarity >= self.config['similarity_threshold']
        }
        
        # Cache result
        self.similarity_cache[cache_key] = result
        
        return result
    
    def segment_text_for_analysis(self, text):
        """Segment text into chunks for similarity analysis"""
        words = text.split()
        segments = []
        
        # Create overlapping segments
        min_len = self.config['min_segment_length']
        max_len = self.config['max_segment_length']
        overlap = self.config['chunk_overlap']
        
        i = 0
        while i < len(words):
            # Variable segment length
            segment_length = min(max_len, len(words) - i)
            if segment_length < min_len and i > 0:
                break
            
            segment_words = words[i:i + segment_length]
            segment_text = ' '.join(segment_words)
            
            segments.append({
                'text': segment_text,
                'start_word': i,
                'end_word': i + segment_length - 1,
                'word_count': len(segment_words)
            })
            
            # Move forward with overlap
            i += max(1, segment_length - overlap)
        
        return segments
    
    def analyze_paragraph_similarity(self, paragraph_text):
        """Analyze paragraph dan detect bagian yang flagged"""
        segments = self.segment_text_for_analysis(paragraph_text)
        flagged_segments = []
        analysis_results = []
        
        for segment in segments:
            similarity_result = self.check_similarity_with_database(segment['text'])
            
            analysis_results.append({
                'segment': segment,
                'similarity_result': similarity_result
            })
            
            if similarity_result['is_flagged']:
                flagged_segments.append({
                    'text': segment['text'],
                    'similarity': similarity_result['similarity'],
                    'best_match': similarity_result['best_match'],
                    'source': similarity_result['source'],
                    'start_word': segment['start_word'],
                    'end_word': segment['end_word']
                })
        
        # Calculate overall similarity
        if analysis_results:
            avg_similarity = sum(r['similarity_result']['similarity'] for r in analysis_results) / len(analysis_results)
            max_similarity = max(r['similarity_result']['similarity'] for r in analysis_results)
        else:
            avg_similarity = 0.0
            max_similarity = 0.0
        
        return {
            'flagged_segments': flagged_segments,
            'all_segments': analysis_results,
            'avg_similarity': avg_similarity,
            'max_similarity': max_similarity,
            'needs_paraphrasing': len(flagged_segments) > 0,
            'flagged_ratio': len(flagged_segments) / len(segments) if segments else 0
        }
    
    def create_targeted_paraphrase_prompt(self, text, flagged_analysis):
        """Create targeted prompt hanya untuk bagian yang flagged"""
        flagged_segments = flagged_analysis['flagged_segments']
        flagged_texts = [seg['text'] for seg in flagged_segments]
        
        prompt = f"""Kamu adalah expert paraphrasing yang sangat memahami sistem deteksi plagiarisme.

MISI SPESIFIK: Hanya parafrase bagian-bagian teks yang TELAH TERDETEKSI memiliki similarity tinggi dengan sumber lain. Jangan ubah bagian yang tidak bermasalah.

BAGIAN YANG TERDETEKSI SIMILARITY TINGGI:
{chr(10).join(f"- {text} (similarity: {seg['similarity']:.2f})" for seg, text in zip(flagged_segments, flagged_texts))}

STRATEGI TARGETED:
✅ FOKUS HANYA pada bagian yang di-flag di atas
✅ Restructure bagian flagged dengan drastis (ubah struktur, sinonim, urutan)
✅ Pertahankan bagian yang tidak flagged TANPA PERUBAHAN
✅ Pastikan makna keseluruhan tetap sama
✅ Hasil harus natural dan akademik

TARGET: Turunkan similarity bagian flagged dari {flagged_analysis['max_similarity']*100:.0f}% menjadi <25%

TEKS LENGKAP:
{text}

PARAFRASE TARGETED (hanya ubah bagian flagged):"""
        
        return prompt
    
    def process_paragraph_with_detection(self, paragraph_text, paragraph_index):
        """Process paragraph dengan similarity detection dulu"""
        
        # Step 1: Analyze similarity (seperti Turnitin scan)
        similarity_analysis = self.analyze_paragraph_similarity(paragraph_text)
        
        # Step 2: Jika tidak ada yang flagged, return original
        if not similarity_analysis['needs_paraphrasing']:
            return {
                'paraphrase': paragraph_text,
                'similarity': 0.0,
                'plagiarism_reduction': 0.0,
                'changes_made': 0,
                'method': 'no_similarity_detected',
                'status': '✅ Original text - No similarity detected',
                'similarity_analysis': similarity_analysis,
                'flagged_segments': 0
            }
        
        # Step 3: Ada bagian yang flagged - track it
        self.cost_tracker['flagged_segments'] += len(similarity_analysis['flagged_segments'])
        
        print(f"    🚨 Similarity detected: {len(similarity_analysis['flagged_segments'])} flagged segments")
        for i, seg in enumerate(similarity_analysis['flagged_segments'][:3]):  # Show max 3
            print(f"      - Segment {i+1}: {seg['similarity']:.2f} similarity with {seg['source']}")
        
        # Step 4: Try local paraphrasing first
        local_result = self.local_paraphraser.generate_paraphrase(paragraph_text, aggressiveness=0.7)
        self.cost_tracker['local_calls'] += 1
        
        # Step 5: Re-check similarity after local paraphrasing
        local_similarity_check = self.analyze_paragraph_similarity(local_result['paraphrase'])
        local_improvement = similarity_analysis['max_similarity'] - local_similarity_check['max_similarity']
        
        # Step 6: Decide if AI is needed
        if (local_improvement >= self.config['ai_improvement_threshold'] and 
            len(local_similarity_check['flagged_segments']) == 0):
            # Local sudah cukup bagus
            local_result['method'] = 'local_similarity_aware'
            local_result['similarity_analysis'] = local_similarity_check
            local_result['flagged_segments'] = len(local_similarity_check['flagged_segments'])
            return local_result
        
        # Step 7: Use AI untuk targeted paraphrasing
        if self.gemini_client and similarity_analysis['flagged_segments']:
            print(f"    🤖 Using AI for targeted paraphrasing of flagged segments")
            
            ai_result_text = self.call_gemini_for_targeted_paraphrasing(paragraph_text, similarity_analysis)
            
            if ai_result_text:
                # Re-check similarity after AI
                ai_similarity_check = self.analyze_paragraph_similarity(ai_result_text)
                ai_similarity = self.local_paraphraser.calculate_similarity(paragraph_text, ai_result_text)
                
                ai_result = {
                    'paraphrase': ai_result_text,
                    'similarity': round(ai_similarity, 2),
                    'plagiarism_reduction': round(100 - ai_similarity, 2),
                    'changes_made': 1,
                    'method': 'ai_targeted_similarity',
                    'status': self.local_paraphraser._get_plagiarism_status(ai_similarity),
                    'similarity_analysis': ai_similarity_check,
                    'flagged_segments': len(ai_similarity_check['flagged_segments'])
                }
                
                # Choose better result
                if len(ai_similarity_check['flagged_segments']) < len(local_similarity_check['flagged_segments']):
                    self.cost_tracker['paraphrased_segments'] += 1
                    return ai_result
        
        # Step 8: Fallback to local result
        local_result['method'] = 'local_fallback_similarity'
        local_result['similarity_analysis'] = local_similarity_check
        local_result['flagged_segments'] = len(local_similarity_check['flagged_segments'])
        return local_result
    
    def call_gemini_for_targeted_paraphrasing(self, text, similarity_analysis):
        """Call Gemini AI specifically for flagged segments"""
        if not self.gemini_client:
            return None
        
        try:
            prompt = self.create_targeted_paraphrase_prompt(text, similarity_analysis)
            
            contents = [
                types.Content(
                    role="user",
                    parts=[types.Part.from_text(text=prompt)],
                ),
            ]
            
            generate_content_config = types.GenerationConfig(
                temperature=0.4,  # Lower temperature for more focused changes
                top_k=50,
                top_p=0.9,
                max_output_tokens=3072,
                candidate_count=1,
            )
            
            for attempt in range(self.config['ai_retry_attempts']):
                try:
                    response_text = ""
                    for chunk in self.gemini_client.models.generate_content_stream(
                        model=self.config['model_name'],
                        contents=contents,
                        config=generate_content_config,
                    ):
                        if hasattr(chunk, 'text') and chunk.text:
                            response_text += chunk.text
                    
                    if response_text.strip():
                        # Track usage
                        self.cost_tracker['ai_calls'] += 1
                        tokens_used = len(prompt + response_text) // 4
                        self.cost_tracker['ai_tokens_used'] += tokens_used
                        self.cost_tracker['estimated_cost_usd'] += tokens_used * 0.000001
                        
                        return response_text.strip()
                    
                except Exception as e:
                    print(f"⚠️  Targeted AI attempt {attempt + 1} failed: {e}")
                    if attempt < self.config['ai_retry_attempts'] - 1:
                        time.sleep(2 ** attempt)
            
            return None
            
        except Exception as e:
            print(f"❌ Targeted AI error: {e}")
            return None
    
    def process_batch_documents(self, input_folder, create_backup=True):
        """Process batch dengan similarity detection dulu"""
        start_time = datetime.now()
        
        print("=" * 80)
        print("🔍 SMART TURNITIN-LIKE DETECTOR + SELECTIVE PARAPHRASER")
        print("=" * 80)
        
        if not os.path.exists(input_folder):
            print(f"❌ Input folder not found: {input_folder}")
            return
        
        docx_files = [f for f in os.listdir(input_folder) if f.endswith('.docx') and not f.startswith('~')]
        
        if not docx_files:
            print(f"❌ No Word documents found in: {input_folder}")
            return
        
        print(f"📁 Input folder: {input_folder}")
        print(f"📄 Found {len(docx_files)} documents")
        print(f"🔍 Similarity threshold: {self.config['similarity_threshold']*100:.0f}%")
        print(f"🤖 AI Enhancement: {'Enabled for flagged segments only' if self.gemini_client else 'Disabled'}")
        
        # Create backup
        backup_folder = None
        if create_backup:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            backup_folder = f"{input_folder}_backup_smart_{timestamp}"
            
            try:
                if not os.path.exists(backup_folder):
                    os.makedirs(backup_folder)
                
                for file in docx_files:
                    src = os.path.join(input_folder, file)
                    dst = os.path.join(backup_folder, file)
                    shutil.copy2(src, dst)
                
                print(f"✅ Backup created: {backup_folder}")
            except Exception as e:
                print(f"❌ Error creating backup: {e}")
                return
        
        print(f"\n🚀 Starting smart similarity detection...")
        
        # Process documents
        total_stats = {
            'documents_processed': 0,
            'paragraphs_analyzed': 0,
            'paragraphs_flagged': 0,
            'paragraphs_clean': 0,
            'paragraphs_paraphrased': 0,
            'flagged_segments_detected': 0,
            'flagged_segments_fixed': 0,
            'total_changes': 0
        }
        
        for i, filename in enumerate(docx_files, 1):
            file_path = os.path.join(input_folder, filename)
            
            print(f"\n{'='*20} Document {i}/{len(docx_files)} {'='*20}")
            
            doc_stats = self.process_single_document_smart(file_path)
            
            if doc_stats:
                total_stats['documents_processed'] += 1
                for key in doc_stats:
                    if key in total_stats:
                        total_stats[key] += doc_stats[key]
        
        # Calculate processing time
        end_time = datetime.now()
        processing_time = (end_time - start_time).total_seconds()
        
        # Print comprehensive summary
        self.print_smart_detector_summary(total_stats, processing_time, backup_folder)
        
        # Generate detailed report
        self.generate_smart_detector_report(input_folder, total_stats, processing_time)
    
    def process_single_document_smart(self, file_path):
        """Process single document dengan smart similarity detection"""
        try:
            print(f"\n📄 Processing: {os.path.basename(file_path)}")
            
            doc = docx.Document(file_path)
            
            doc_stats = {
                'filename': os.path.basename(file_path),
                'total_paragraphs': 0,
                'paragraphs_analyzed': 0,
                'paragraphs_flagged': 0,
                'paragraphs_clean': 0,
                'paragraphs_paraphrased': 0,
                'flagged_segments_detected': 0,
                'flagged_segments_fixed': 0,
                'total_changes': 0
            }
            
            current_section = 'UNKNOWN'
            processed_paragraphs = []
            
            # Collect suitable paragraphs
            for paragraph in doc.paragraphs:
                para_text = paragraph.text.strip()
                doc_stats['total_paragraphs'] += 1
                
                if not para_text:
                    continue
                
                if self.is_section_header(para_text):
                    current_section = para_text
                    print(f"  📍 Section: {current_section}")
                    continue
                
                if self.is_paragraph_suitable_for_analysis(para_text):
                    processed_paragraphs.append({
                        'paragraph': paragraph,
                        'text': para_text,
                        'index': len(processed_paragraphs)
                    })
            
            total_suitable = len(processed_paragraphs)
            print(f"  📝 Found {total_suitable} suitable paragraphs for similarity analysis")
            
            # Process paragraphs with smart detection
            for para_data in processed_paragraphs:
                paragraph = para_data['paragraph']
                para_text = para_data['text']
                para_index = para_data['index']
                
                # Smart similarity detection + selective paraphrasing
                result = self.process_paragraph_with_detection(para_text, para_index)
                
                doc_stats['paragraphs_analyzed'] += 1
                
                # Track flagged segments
                flagged_count = result.get('flagged_segments', 0)
                doc_stats['flagged_segments_detected'] += flagged_count
                
                if flagged_count > 0:
                    doc_stats['paragraphs_flagged'] += 1
                else:
                    doc_stats['paragraphs_clean'] += 1
                
                # Apply changes only if similarity was detected and fixed
                if result['changes_made'] > 0 and flagged_count > 0:
                    paragraph.clear()
                    paragraph.add_run(result['paraphrase'])
                    
                    # Color coding based on detection results
                    if 'ai_targeted' in result['method']:
                        # Red for AI-targeted similarity fixes
                        for run in paragraph.runs:
                            run.font.highlight_color = WD_COLOR_INDEX.RED
                    elif 'similarity_aware' in result['method']:
                        # Orange for local similarity-aware fixes
                        for run in paragraph.runs:
                            run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                    
                    doc_stats['paragraphs_paraphrased'] += 1
                    doc_stats['total_changes'] += result['changes_made']
                    
                    # Calculate how many flagged segments were fixed
                    original_flagged = result.get('similarity_analysis', {}).get('flagged_segments', [])
                    fixed_segments = max(0, flagged_count - len(original_flagged))
                    doc_stats['flagged_segments_fixed'] += fixed_segments
                    
                    print(f"    ✅ Para {para_index+1}: {result['plagiarism_reduction']:.1f}% reduction, {flagged_count} flagged segments ({result['method']})")
                elif flagged_count == 0:
                    print(f"    ✅ Para {para_index+1}: No similarity detected - kept original")
                else:
                    print(f"    ⏭️  Para {para_index+1}: Flagged but minimal improvement ({result['method']})")
            
            # Save document
            doc.save(file_path)
            
            print(f"  📊 Document Summary:")
            print(f"     • Total paragraphs: {doc_stats['total_paragraphs']}")
            print(f"     • Analyzed: {doc_stats['paragraphs_analyzed']}")
            print(f"     • Clean (no similarity): {doc_stats['paragraphs_clean']}")
            print(f"     • Flagged: {doc_stats['paragraphs_flagged']}")
            print(f"     • Paraphrased: {doc_stats['paragraphs_paraphrased']}")
            print(f"     • Flagged segments detected: {doc_stats['flagged_segments_detected']}")
            print(f"  ✅ Saved: {file_path}")
            
            return doc_stats
            
        except Exception as e:
            print(f"❌ Error processing {file_path}: {e}")
            return None
    
    def is_section_header(self, text):
        """Check if text is a section header"""
        patterns = [
            r'^BAB\s+[IVX]+',
            r'^\d+\.\d+',
            r'^[A-Z\s]+$',
        ]
        
        for pattern in patterns:
            if re.match(pattern, text.strip()) and len(text.split()) < 10:
                return True
        return False
    
    def is_paragraph_suitable_for_analysis(self, paragraph_text):
        """Check if paragraph should be analyzed"""
        word_count = len(paragraph_text.split())
        if word_count < self.config['min_segment_length']:
            return False
        
        if re.search(r'\d+\.\d+|\btabel\b|\bgambar\b', paragraph_text.lower()):
            return False
        
        if paragraph_text.isupper() and word_count < 15:
            return False
        
        return True
    
    def print_smart_detector_summary(self, stats, processing_time, backup_folder):
        """Print comprehensive smart detector summary"""
        print("\n" + "=" * 80)
        print("📊 SMART TURNITIN-LIKE DETECTOR SUMMARY")
        print("=" * 80)
        
        print(f"📄 Documents processed: {stats['documents_processed']}")
        print(f"📝 Paragraphs analyzed: {stats['paragraphs_analyzed']}")
        print(f"✅ Clean paragraphs (no similarity): {stats['paragraphs_clean']}")
        print(f"🚨 Flagged paragraphs: {stats['paragraphs_flagged']}")
        print(f"✏️  Paragraphs paraphrased: {stats['paragraphs_paraphrased']}")
        print(f"🔍 Flagged segments detected: {stats['flagged_segments_detected']}")
        print(f"🔧 Flagged segments fixed: {stats['flagged_segments_fixed']}")
        print(f"🔄 Total changes made: {stats['total_changes']}")
        print(f"⏱️  Processing time: {processing_time:.1f} seconds")
        
        # Efficiency metrics
        if stats['paragraphs_analyzed'] > 0:
            clean_ratio = (stats['paragraphs_clean'] / stats['paragraphs_analyzed']) * 100
            flagged_ratio = (stats['paragraphs_flagged'] / stats['paragraphs_analyzed']) * 100
            print(f"\n📊 SIMILARITY ANALYSIS:")
            print(f"   • Clean content: {clean_ratio:.1f}%")
            print(f"   • Flagged content: {flagged_ratio:.1f}%")
            
            if stats['paragraphs_flagged'] > 0:
                fix_rate = (stats['paragraphs_paraphrased'] / stats['paragraphs_flagged']) * 100
                print(f"   • Fix success rate: {fix_rate:.1f}%")
        
        # Cost analysis
        print(f"\n💰 COST ANALYSIS:")
        print(f"   💻 Local calls: {self.cost_tracker['local_calls']} (FREE)")
        print(f"   🔍 Similarity checks: {self.cost_tracker['similarity_checks']} (FREE)")
        print(f"   🤖 AI calls: {self.cost_tracker['ai_calls']}")
        print(f"   💵 Estimated cost: ${self.cost_tracker['estimated_cost_usd']:.6f}")
        
        # Efficiency comparison
        print(f"\n🎯 EFFICIENCY GAINS:")
        total_paragraphs = stats['paragraphs_analyzed']
        if total_paragraphs > 0:
            efficiency = ((stats['paragraphs_clean']) / total_paragraphs) * 100
            print(f"   • Content preserved unchanged: {efficiency:.1f}%")
            print(f"   • Only flagged content processed: {100-efficiency:.1f}%")
            print(f"   • Selective processing saves: ~{efficiency:.0f}% unnecessary changes")
        
        if backup_folder:
            print(f"\n💾 Original files backed up to: {backup_folder}")
        
        print("=" * 80)
    
    def generate_smart_detector_report(self, input_folder, stats, processing_time):
        """Generate detailed smart detector report"""
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        report_file = os.path.join(input_folder, f"smart_detector_report_{timestamp}.json")
        
        try:
            report_data = {
                'timestamp': timestamp,
                'input_folder': input_folder,
                'processing_time_seconds': processing_time,
                'statistics': stats,
                'cost_tracker': self.cost_tracker,
                'configuration': self.config,
                'similarity_database_size': sum(len(patterns) for patterns in self.similarity_database.values()),
                'system_info': {
                    'local_synonyms': len(self.local_paraphraser.synonyms),
                    'similarity_threshold': self.config['similarity_threshold'],
                    'gemini_enabled': bool(self.gemini_client),
                    'version': '4.0'
                }
            }
            
            with open(report_file, 'w', encoding='utf-8') as f:
                json.dump(report_data, f, ensure_ascii=False, indent=2)
            
            print(f"📋 Smart detector report saved: {report_file}")
            
        except Exception as e:
            print(f"⚠️  Could not save report: {e}")
    
    def add_to_similarity_database(self, text, category="user_added"):
        """Add new pattern to similarity database"""
        if category not in self.similarity_database:
            self.similarity_database[category] = []
        
        self.similarity_database[category].append(text)
        
        # Save updated database
        try:
            with open("similarity_database.json", 'w', encoding='utf-8') as f:
                json.dump(self.similarity_database, f, ensure_ascii=False, indent=2)
            print(f"✅ Added to similarity database: {text[:50]}...")
        except Exception as e:
            print(f"⚠️  Could not save to database: {e}")
    
    def batch_add_patterns_from_file(self, file_path, category="batch_import"):
        """Batch add patterns from text file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                patterns = [line.strip() for line in f if line.strip()]
            
            if category not in self.similarity_database:
                self.similarity_database[category] = []
            
            self.similarity_database[category].extend(patterns)
            
            # Save updated database
            with open("similarity_database.json", 'w', encoding='utf-8') as f:
                json.dump(self.similarity_database, f, ensure_ascii=False, indent=2)
            
            print(f"✅ Added {len(patterns)} patterns from {file_path}")
            
        except Exception as e:
            print(f"❌ Error importing patterns: {e}")
    
    def export_flagged_segments(self, output_file="flagged_segments.txt"):
        """Export all flagged segments for analysis"""
        try:
            flagged_segments = []
            
            # Collect from cache
            for cache_key, result in self.similarity_cache.items():
                if result['is_flagged']:
                    flagged_segments.append({
                        'similarity': result['similarity'],
                        'best_match': result['best_match'],
                        'source': result['source']
                    })
            
            # Write to file
            with open(output_file, 'w', encoding='utf-8') as f:
                f.write("FLAGGED SEGMENTS ANALYSIS\n")
                f.write("=" * 50 + "\n\n")
                
                for i, seg in enumerate(flagged_segments, 1):
                    f.write(f"Segment {i}:\n")
                    f.write(f"Similarity: {seg['similarity']:.3f}\n")
                    f.write(f"Source: {seg['source']}\n")
                    f.write(f"Best Match: {seg['best_match']}\n")
                    f.write("-" * 30 + "\n")
            
            print(f"📋 Exported {len(flagged_segments)} flagged segments to {output_file}")
            
        except Exception as e:
            print(f"⚠️  Could not export flagged segments: {e}")


def main():
    """Main function for smart Turnitin-like detection"""
    print("🔍 SMART TURNITIN-LIKE DETECTOR + SELECTIVE PARAPHRASER")
    print("Hanya memarafrase bagian yang benar-benar terdeteksi similarity tinggi")
    print("=" * 80)
    
    # Check dependencies
    if not GEMINI_AVAILABLE:
        print("⚠️  google-genai package not available, using local processing only")
    
    # Configuration
    SYNONYM_FILE = 'sinonim.json'            # Your synonym database
    INPUT_FOLDER = 'documents'               # Folder with Word documents
    GEMINI_API_KEY = None                    # Set your API key here if available
    CREATE_BACKUP = True                     # Always backup first
    
    # Get API key from environment if available
    GEMINI_API_KEY = os.environ.get("GEMINI_API_KEY", GEMINI_API_KEY)
    
    if not GEMINI_API_KEY:
        print("⚠️  No Gemini API key found - using local processing only")
        print("💡 Set GEMINI_API_KEY environment variable to enable AI enhancement")
    else:
        print(f"✅ API key configured: {GEMINI_API_KEY[:20]}...")
    
    # Initialize smart detector
    detector = SmartTurnitinDetector(
        synonym_file=SYNONYM_FILE,
        gemini_api_key=GEMINI_API_KEY
    )
    
    # Configuration options
    print(f"\n⚙️  CONFIGURATION:")
    print(f"   🎯 Similarity threshold: {detector.config['similarity_threshold']*100:.0f}%")
    print(f"   📏 Min segment length: {detector.config['min_segment_length']} words")
    print(f"   📏 Max segment length: {detector.config['max_segment_length']} words")
    print(f"   🔄 Chunk overlap: {detector.config['chunk_overlap']} words")
    
    # Process batch with smart detection
    detector.process_batch_documents(
        input_folder=INPUT_FOLDER,
        create_backup=CREATE_BACKUP
    )
    
    print("\n🎉 Smart Turnitin-like detection completed!")
    print("💡 Color coding:")
    print("   🟡 Yellow = Local similarity-aware fixes")  
    print("   🔴 Red = AI-targeted similarity fixes")
    print("   ⚪ No highlight = Original text (no similarity detected)")
    
    print("\n📊 Key Benefits:")
    print("   ✅ Only processes content that actually has similarity issues")
    print("   ✅ Preserves original text when no problems detected")
    print("   ✅ Targeted AI usage only for problematic segments")
    print("   ✅ Much more cost-effective than blanket processing")
    
    # Optional: Export analysis for review
    detector.export_flagged_segments("flagged_analysis.txt")
    print(f"\n📋 Flagged segments analysis exported for review")


if __name__ == "__main__":
    main()
