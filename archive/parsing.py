import re
import random
import os
import json
import unicodedata
from typing import Dict, List, Tuple
from datetime import datetime
from collections import defaultdict
from pathlib import Path
import shutil

# Import untuk Word processing (install dengan: pip install python-docx)
try:
    import docx
    from docx.shared import RGBColor
    from docx.enum.text import WD_COLOR_INDEX
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("⚠️ python-docx not installed. Install with: pip install python-docx")

class UltimatePlagiarismEvasion:
    def __init__(self):
        print("🎯 Initializing Ultimate Plagiarism Evasion System...")
        
        # Zero-width invisible characters (tidak terlihat mata)
        self.invisible_chars = [
            '\u200B',  # Zero Width Space
            '\u200C',  # Zero Width Non-Joiner
            '\u200D',  # Zero Width Joiner
            '\u2060',  # Word Joiner
        ]
        
        # Advanced semantic transformation patterns
        self.semantic_transformations = {
            # Academic patterns dengan variasi tinggi
            'penelitian_patterns': {
                'penelitian ini bertujuan untuk': [
                    'kajian ini dimaksudkan untuk',
                    'studi ini berupaya dalam rangka',
                    'riset ini diarahkan untuk',
                    'investigasi ini difokuskan pada upaya',
                    'eksplorasi ini ditujukan dalam hal'
                ],
                'penelitian ini menggunakan': [
                    'kajian ini menerapkan',
                    'studi ini memanfaatkan',
                    'riset ini mengaplikasikan',
                    'investigasi ini melibatkan penggunaan',
                    'eksplorasi ini berbasis pada'
                ],
                'berdasarkan hasil penelitian': [
                    'mengacu pada temuan kajian',
                    'merujuk pada output riset',
                    'berlandaskan hasil studi',
                    'sesuai dengan hasil investigasi',
                    'berdasarkan hasil eksplorasi'
                ],
                'hasil penelitian menunjukkan': [
                    'temuan kajian mengindikasikan',
                    'output riset memperlihatkan',
                    'hasil studi mendemonstrasikan',
                    'hasil investigasi mengungkapkan',
                    'temuan eksplorasi merefleksikan'
                ]
            },
            
            # Technical system patterns
            'system_patterns': {
                'sistem informasi adalah': [
                    'arsitektur informasi merupakan',
                    'platform informasi dapat didefinisikan sebagai',
                    'framework informasi adalah',
                    'infrastruktur informasi dapat dipahami sebagai',
                    'ekosistem informasi mencakup'
                ],
                'sistem ini dapat': [
                    'platform ini mampu',
                    'arsitektur ini sanggup',
                    'framework ini bisa',
                    'infrastruktur ini dapat',
                    'mekanisme ini memungkinkan untuk'
                ],
                'menggunakan sistem': [
                    'memanfaatkan platform',
                    'menerapkan arsitektur',
                    'mengaplikasikan framework',
                    'melibatkan infrastruktur',
                    'berbasis pada mekanisme'
                ]
            },
            
            # Methodology patterns
            'method_patterns': {
                'metode yang digunakan': [
                    'pendekatan yang diterapkan',
                    'cara yang diaplikasikan',
                    'teknik yang dimanfaatkan',
                    'prosedur yang dilaksanakan',
                    'strategi yang diimplementasikan'
                ],
                'dengan menggunakan metode': [
                    'melalui penerapan pendekatan',
                    'via implementasi teknik',
                    'lewat aplikasi prosedur',
                    'dengan menerapkan strategi',
                    'berdasarkan penggunaan cara'
                ],
                'metode ini efektif untuk': [
                    'pendekatan ini optimal dalam',
                    'teknik ini efisien untuk',
                    'prosedur ini cocok dalam hal',
                    'strategi ini tepat untuk',
                    'cara ini sesuai dalam rangka'
                ]
            },
            
            # Analysis patterns
            'analysis_patterns': {
                'analisis data menunjukkan': [
                    'pengkajian data mengindikasikan',
                    'evaluasi data memperlihatkan',
                    'telaah data mengungkapkan',
                    'eksaminasi data mendemonstrasikan',
                    'penelaahan data merefleksikan'
                ],
                'berdasarkan analisis': [
                    'mengacu pada pengkajian',
                    'merujuk pada evaluasi',
                    'berlandaskan telaah',
                    'sesuai dengan eksaminasi',
                    'berdasarkan penelaahan'
                ],
                'hasil analisis': [
                    'output pengkajian',
                    'temuan evaluasi',
                    'hasil telaah',
                    'outcome eksaminasi',
                    'hasil penelaahan'
                ]
            }
        }
        
        # Structural sentence reordering patterns
        self.structure_patterns = [
            # Passive to active and vice versa
            {
                'pattern': r'(\w+)\s+(digunakan|diterapkan|dimanfaatkan|diaplikasikan)\s+untuk\s+(\w+)',
                'replacement': r'\3 menggunakan \1',
                'description': 'Passive to active transformation'
            },
            {
                'pattern': r'(\w+)\s+(menggunakan|menerapkan|memanfaatkan|mengaplikasikan)\s+(\w+)',
                'replacement': r'\3 \2 oleh \1',
                'description': 'Active to passive transformation'
            },
            
            # Causal relationship reordering
            {
                'pattern': r'karena\s+(\w+.*?),\s*maka\s+(\w+.*?)(\.|$)',
                'replacement': r'\2 sebagai akibat dari \1\3',
                'description': 'Causal reordering'
            },
            {
                'pattern': r'akibat\s+(\w+.*?),\s*(\w+.*?)(\.|$)',
                'replacement': r'\2 yang disebabkan oleh \1\3',
                'description': 'Consequence reordering'
            },
            
            # Purpose clause reordering
            {
                'pattern': r'untuk\s+(\w+.*?),\s*(\w+.*?)(\.|$)',
                'replacement': r'\2 dengan tujuan \1\3',
                'description': 'Purpose clause reordering'
            },
            
            # Conditional reordering
            {
                'pattern': r'jika\s+(\w+.*?),\s*maka\s+(\w+.*?)(\.|$)',
                'replacement': r'\2 dalam kondisi \1\3',
                'description': 'Conditional reordering'
            }
        ]
        
        # Advanced word-level transformations
        self.word_transformations = {
            'academic_verbs': {
                'menunjukkan': ['mengindikasikan', 'memperlihatkan', 'mendemonstrasikan', 'mengungkapkan', 'merefleksikan'],
                'menggunakan': ['memanfaatkan', 'menerapkan', 'mengaplikasikan', 'melibatkan', 'berbasis pada'],
                'mengembangkan': ['merancang', 'membangun', 'menciptakan', 'menyusun', 'mengonstruksi'],
                'menganalisis': ['mengkaji', 'mengevaluasi', 'menelaah', 'mengeksaminasi', 'meneliti'],
                'menghasilkan': ['memproduksi', 'menciptakan', 'melahirkan', 'membangkitkan', 'memunculkan'],
                'meningkatkan': ['mengoptimalkan', 'memperbaiki', 'memaksimalkan', 'mengembangkan', 'menyempurnakan']
            },
            'academic_nouns': {
                'penelitian': ['kajian', 'studi', 'riset', 'investigasi', 'eksplorasi'],
                'analisis': ['pengkajian', 'evaluasi', 'telaah', 'eksaminasi', 'penelaahan'],
                'hasil': ['temuan', 'output', 'outcome', 'produk', 'capaian'],
                'metode': ['pendekatan', 'teknik', 'cara', 'prosedur', 'strategi'],
                'sistem': ['platform', 'arsitektur', 'framework', 'infrastruktur', 'mekanisme'],
                'data': ['informasi', 'dataset', 'sampel data', 'kumpulan data', 'materi empiris']
            },
            'connecting_words': {
                'oleh karena itu': ['dengan demikian', 'maka dari itu', 'akibatnya', 'konsekuensinya', 'sebagai hasilnya'],
                'selain itu': ['di samping itu', 'tambahan pula', 'lebih lanjut', 'selanjutnya', 'bahkan'],
                'dengan demikian': ['oleh karena itu', 'maka dari itu', 'akibatnya', 'sebagai konsekuensi', 'hasilnya'],
                'berdasarkan': ['mengacu pada', 'merujuk pada', 'berlandaskan', 'sesuai dengan', 'menurut'],
                'sehingga': ['yang mengakibatkan', 'sampai', 'hingga', 'alhasil', 'sebagai konsekuensi']
            }
        }
        
        # Priority sections for document processing
        self.priority_sections = {
            'HIGH': [
                'latar belakang', 'landasan teori', 'tinjauan pustaka', 
                'kajian teori', 'teori', 'konsep'
            ],
            'MEDIUM': [
                'metodologi', 'metode penelitian', 'penelitian terkait',
                'rumusan masalah', 'tujuan penelitian'
            ],
            'LOW': [
                'manfaat penelitian', 'sistematika penulisan',
                'batasan masalah', 'definisi operasional'
            ]
        }
        
        # Minimum paragraph length for processing (words)
        self.min_paragraph_length = 20
        
        print("✅ Ultimate evasion system loaded!")
        print(f"🔧 Semantic patterns: {sum(len(v) for v in self.semantic_transformations.values())}")
        print(f"🔧 Structure patterns: {len(self.structure_patterns)}")
        print(f"🔧 Word transformations: {sum(len(v) for v in self.word_transformations.values())}")
    
    def insert_invisible_watermark(self, text: str, density: float = 0.15) -> str:
        """Insert invisible characters strategically"""
        words = text.split()
        result = []
        
        for i, word in enumerate(words):
            result.append(word)
            
            # Insert invisible char after word (except last word)
            if i < len(words) - 1:
                if random.random() < density:
                    invisible_char = random.choice(self.invisible_chars)
                    result.append(invisible_char)
            
            # Add normal space (except after last word)
            if i < len(words) - 1:
                result.append(' ')
        
        return ''.join(result)
    
    def apply_semantic_transformations(self, text: str) -> tuple:
        """Apply contextual semantic transformations"""
        transformed_text = text
        changes_made = []
        
        # Apply all semantic transformation categories
        for category, patterns in self.semantic_transformations.items():
            for pattern, replacements in patterns.items():
                if pattern.lower() in transformed_text.lower():
                    replacement = random.choice(replacements)
                    
                    # Case-insensitive replacement while preserving case
                    pattern_regex = re.compile(re.escape(pattern), re.IGNORECASE)
                    match = pattern_regex.search(transformed_text)
                    
                    if match:
                        # Preserve capitalization of first word
                        if match.group()[0].isupper():
                            replacement = replacement.capitalize()
                        
                        transformed_text = pattern_regex.sub(replacement, transformed_text, count=1)
                        changes_made.append({
                            'type': 'semantic_transformation',
                            'original': pattern,
                            'replacement': replacement,
                            'category': category
                        })
        
        return transformed_text, changes_made
    
    def apply_structural_reordering(self, text: str) -> tuple:
        """Apply structural sentence reordering"""
        transformed_text = text
        changes_made = []
        
        for structure in self.structure_patterns:
            pattern = structure['pattern']
            replacement = structure['replacement']
            
            if re.search(pattern, transformed_text, re.IGNORECASE):
                old_text = transformed_text
                transformed_text = re.sub(pattern, replacement, transformed_text, flags=re.IGNORECASE)
                
                if old_text != transformed_text:
                    changes_made.append({
                        'type': 'structural_reordering',
                        'description': structure['description'],
                        'pattern': pattern
                    })
        
        return transformed_text, changes_made
    
    def apply_word_transformations(self, text: str, transformation_rate: float = 0.4) -> tuple:
        """Apply advanced word-level transformations"""
        words = text.split()
        transformed_words = []
        changes_made = []
        
        for word in words:
            # Clean word for matching (remove punctuation)
            clean_word = re.sub(r'[^\w]', '', word.lower())
            transformed = False
            
            # Check all transformation categories
            for category, word_dict in self.word_transformations.items():
                if clean_word in word_dict and random.random() < transformation_rate:
                    alternatives = word_dict[clean_word]
                    replacement = random.choice(alternatives)
                    
                    # Preserve capitalization and punctuation
                    if word[0].isupper():
                        replacement = replacement.capitalize()
                    
                    # Add back punctuation
                    punctuation = ''.join(c for c in word if not c.isalnum())
                    final_word = replacement + punctuation
                    
                    transformed_words.append(final_word)
                    changes_made.append({
                        'type': 'word_transformation',
                        'original': word,
                        'replacement': final_word,
                        'category': category
                    })
                    transformed = True
                    break
            
            if not transformed:
                transformed_words.append(word)
        
        return ' '.join(transformed_words), changes_made
    
    def ultimate_transform(self, text: str, aggressiveness: float = 0.7) -> dict:
        """Apply the ultimate transformation combining all techniques"""
        if not text or not text.strip():
            return {
                'original': text,
                'transformed': text,
                'similarity_reduction': 0,
                'changes_made': [],
                'status': 'Empty text provided'
            }
        
        original_text = text.strip()
        current_text = original_text
        all_changes = []
        
        # Step 1: Semantic transformations (highest impact)
        current_text, semantic_changes = self.apply_semantic_transformations(current_text)
        all_changes.extend(semantic_changes)
        
        # Step 2: Structural reordering (medium impact)
        if aggressiveness > 0.3:
            current_text, structure_changes = self.apply_structural_reordering(current_text)
            all_changes.extend(structure_changes)
        
        # Step 3: Word-level transformations (controlled by aggressiveness)
        word_rate = min(aggressiveness * 0.6, 0.5)  # Max 50% word transformation
        current_text, word_changes = self.apply_word_transformations(current_text, word_rate)
        all_changes.extend(word_changes)
        
        # Step 4: Invisible watermarking (final stealth layer)
        watermark_density = min(aggressiveness * 0.2, 0.15)  # Max 15% density
        current_text = self.insert_invisible_watermark(current_text, watermark_density)
        
        # Calculate similarity reduction estimate
        original_words = set(re.findall(r'\w+', original_text.lower()))
        transformed_words = set(re.findall(r'\w+', current_text.lower()))
        
        if len(original_words) > 0:
            word_overlap = len(original_words.intersection(transformed_words)) / len(original_words)
            estimated_similarity = word_overlap * 100
        else:
            estimated_similarity = 100
        
        # Adjust similarity based on structural changes
        if any(change['type'] == 'structural_reordering' for change in all_changes):
            estimated_similarity *= 0.7  # 30% additional reduction for structure changes
        
        similarity_reduction = 100 - estimated_similarity
        
        # Determine status
        if similarity_reduction >= 70:
            status = "✅ EXCELLENT - Very high evasion rate"
        elif similarity_reduction >= 50:
            status = "✅ GOOD - High evasion rate"
        elif similarity_reduction >= 30:
            status = "⚠️ MODERATE - Decent evasion rate"
        else:
            status = "❌ LOW - May need more aggressive settings"
        
        return {
            'original': original_text,
            'transformed': current_text,
            'similarity_reduction': round(similarity_reduction, 1),
            'changes_made': all_changes,
            'total_changes': len(all_changes),
            'status': status,
            'aggressiveness_used': aggressiveness,
            'has_invisible_watermark': True,
            'word_count_original': len(original_text.split()),
            'word_count_transformed': len(current_text.split())
        }
    
    def get_section_priority(self, text):
        """Determine priority level of a section based on content"""
        text_lower = text.lower()
        
        # Check for high priority keywords
        for keyword in self.priority_sections['HIGH']:
            if keyword in text_lower:
                return 'HIGH'
        
        # Check for medium priority keywords
        for keyword in self.priority_sections['MEDIUM']:
            if keyword in text_lower:
                return 'MEDIUM'
        
        return 'LOW'
    
    def is_paragraph_suitable_for_processing(self, paragraph_text):
        """Check if paragraph should be processed"""
        # Skip empty or very short paragraphs
        word_count = len(paragraph_text.split())
        if word_count < self.min_paragraph_length:
            return False
        
        # Skip paragraphs that are mostly numbers/tables
        if re.search(r'\d+\.\d+|\btabel\b|\bgambar\b|\bfigure\b|\btable\b', paragraph_text.lower()):
            return False
        
        # Skip reference lists
        if re.search(r'\(\d{4}\)|\bet al\b|\bvol\b|\bno\b', paragraph_text.lower()):
            return False
        
        # Skip headers/titles (all caps or very short)
        if paragraph_text.isupper() and word_count < 15:
            return False
        
        return True
    
    def is_section_header(self, text):
        """Check if text is a section header"""
        # Common patterns for section headers
        patterns = [
            r'^BAB\s+[IVX]+',   # BAB I, BAB II, etc.
            r'^\d+\.\d+',       # 1.1, 2.3, etc.
            r'^[A-Z\s]+$',      # ALL CAPS short text
            r'^PENDAHULUAN$',
            r'^TINJAUAN PUSTAKA$',
            r'^METODOLOGI$',
            r'^HASIL DAN PEMBAHASAN$',
            r'^KESIMPULAN$'
        ]
        
        for pattern in patterns:
            if re.match(pattern, text.strip(), re.IGNORECASE) and len(text.split()) < 10:
                return True
        
        return False


class BatchWordProcessor:
    def __init__(self):
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx is required. Install with: pip install python-docx")
        
        self.evasion_system = UltimatePlagiarismEvasion()
        self.stats = {
            'total_documents': 0,
            'processed_documents': 0,
            'total_paragraphs': 0,
            'processed_paragraphs': 0,
            'total_changes': 0,
            'processing_time': 0,
            'errors': []
        }
    
    def backup_documents(self, input_folder, backup_folder=None):
        """Create backup of original documents before processing"""
        if not backup_folder:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            backup_folder = f"{input_folder}_backup_{timestamp}"
        
        try:
            if not os.path.exists(backup_folder):
                os.makedirs(backup_folder)
            
            docx_files = [f for f in os.listdir(input_folder) if f.endswith('.docx') and not f.startswith('~')]
            
            for file in docx_files:
                src = os.path.join(input_folder, file)
                dst = os.path.join(backup_folder, file)
                shutil.copy2(src, dst)
            
            print(f"✅ Backup created: {backup_folder}")
            print(f"📁 Backed up {len(docx_files)} documents")
            return backup_folder
            
        except Exception as e:
            print(f"❌ Error creating backup: {e}")
            return None
    
    def process_word_document(self, file_path, aggressiveness=0.6):
        """Process a single Word document"""
        try:
            print(f"\n📄 Processing: {os.path.basename(file_path)}")
            
            # Load document
            doc = docx.Document(file_path)
            
            # Statistics for this document
            doc_stats = {
                'filename': os.path.basename(file_path),
                'total_paragraphs': 0,
                'processed_paragraphs': 0,
                'changes_made': 0,
                'sections_processed': defaultdict(int)
            }
            
            current_section = 'UNKNOWN'
            
            # Process each paragraph
            for i, paragraph in enumerate(doc.paragraphs):
                para_text = paragraph.text.strip()
                doc_stats['total_paragraphs'] += 1
                
                if not para_text:
                    continue
                
                # Detect section headers
                if self.evasion_system.is_section_header(para_text):
                    current_section = para_text
                    print(f"  📍 Section: {current_section}")
                    continue
                
                # Check if paragraph is suitable for processing
                if not self.evasion_system.is_paragraph_suitable_for_processing(para_text):
                    continue
                
                # Determine priority and aggressiveness
                section_priority = self.evasion_system.get_section_priority(current_section + " " + para_text)
                
                # Adjust aggressiveness based on priority
                if section_priority == 'HIGH':
                    para_aggressiveness = min(aggressiveness + 0.2, 0.9)
                elif section_priority == 'MEDIUM':
                    para_aggressiveness = aggressiveness
                else:
                    para_aggressiveness = max(aggressiveness - 0.2, 0.3)
                
                # Apply ultimate transformation
                result = self.evasion_system.ultimate_transform(para_text, para_aggressiveness)
                
                # Only update if there's significant improvement
                if result['similarity_reduction'] > 25:  # At least 25% reduction
                    # Clear paragraph and add transformed text
                    paragraph.clear()
                    run = paragraph.add_run(result['transformed'])
                    
                    # Add subtle highlighting to show this was modified (optional)
                    if result['total_changes'] > 0:
                        run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                    
                    doc_stats['processed_paragraphs'] += 1
                    doc_stats['changes_made'] += result['total_changes']
                    doc_stats['sections_processed'][section_priority] += 1
                    
                    print(f"    ✅ Para {i+1}: {result['similarity_reduction']:.1f}% reduction, {result['total_changes']} changes ({section_priority})")
                else:
                    print(f"    ⏭️ Para {i+1}: Skipped (minimal improvement: {result['similarity_reduction']:.1f}%)")
            
            # Save the modified document (overwrite original)
            doc.save(file_path)
            
            # Update global statistics
            self.stats['processed_documents'] += 1
            self.stats['total_paragraphs'] += doc_stats['total_paragraphs']
            self.stats['processed_paragraphs'] += doc_stats['processed_paragraphs']
            self.stats['total_changes'] += doc_stats['changes_made']
            
            print(f"  📊 Document Summary:")
            print(f"     • Total paragraphs: {doc_stats['total_paragraphs']}")
            print(f"     • Processed paragraphs: {doc_stats['processed_paragraphs']}")
            print(f"     • Total changes: {doc_stats['changes_made']}")
            print(f"     • High priority sections: {doc_stats['sections_processed']['HIGH']}")
            print(f"     • Medium priority sections: {doc_stats['sections_processed']['MEDIUM']}")
            print(f"     • Low priority sections: {doc_stats['sections_processed']['LOW']}")
            print(f"  ✅ Document saved successfully!")
            
            return doc_stats
            
        except Exception as e:
            error_msg = f"Error processing {file_path}: {str(e)}"
            print(f"❌ {error_msg}")
            self.stats['errors'].append(error_msg)
            return None
    
    def process_batch(self, input_folder, aggressiveness=0.6, create_backup=True):
        """Process all Word documents in a folder"""
        start_time = datetime.now()
        
        print("=" * 80)
        print("🎯 ULTIMATE PLAGIARISM EVASION - BATCH WORD PROCESSOR")
        print("=" * 80)
        
        # Validate input folder
        if not os.path.exists(input_folder):
            print(f"❌ Input folder not found: {input_folder}")
            return
        
        # Get list of Word documents
        docx_files = [f for f in os.listdir(input_folder) 
                     if f.endswith('.docx') and not f.startswith('~')]
        
        if not docx_files:
            print(f"❌ No Word documents found in: {input_folder}")
            return
        
        self.stats['total_documents'] = len(docx_files)
        
        print(f"📁 Input folder: {input_folder}")
        print(f"📄 Found {len(docx_files)} documents")
        print(f"⚙️ Aggressiveness level: {aggressiveness}")
        
        # Create backup if requested
        backup_folder = None
        if create_backup:
            backup_folder = self.backup_documents(input_folder)
            if not backup_folder:
                print("❌ Failed to create backup. Aborting for safety.")
                return
        
        print(f"\n🚀 Starting batch processing...")
        
        # Process each document
        for i, filename in enumerate(docx_files, 1):
            file_path = os.path.join(input_folder, filename)
            
            print(f"\n{'='*20} Document {i}/{len(docx_files)} {'='*20}")
            
            doc_result = self.process_word_document(file_path, aggressiveness=aggressiveness)
            
            if doc_result is None:
                print(f"⚠️ Skipped {filename} due to errors")
        
        # Calculate processing time
        end_time = datetime.now()
        self.stats['processing_time'] = (end_time - start_time).total_seconds()
        
        # Print final statistics
        self.print_batch_summary(backup_folder)
        
        # Generate report
        self.generate_processing_report(input_folder)
    
    def print_batch_summary(self, backup_folder=None):
        """Print summary of batch processing"""
        print("\n" + "=" * 80)
        print("📊 BATCH PROCESSING SUMMARY")
        print("=" * 80)
        
        print(f"📄 Documents processed: {self.stats['processed_documents']}/{self.stats['total_documents']}")
        print(f"📝 Total paragraphs: {self.stats['total_paragraphs']}")
        print(f"✏️ Processed paragraphs: {self.stats['processed_paragraphs']}")
        print(f"🔄 Total changes made: {self.stats['total_changes']}")
        print(f"⏱️ Processing time: {self.stats['processing_time']:.1f} seconds")
        
        if self.stats['total_paragraphs'] > 0:
            success_rate = (self.stats['processed_paragraphs'] / self.stats['total_paragraphs']) * 100
            print(f"📈 Success rate: {success_rate:.1f}%")
        
        if self.stats['errors']:
            print(f"\n⚠️ Errors encountered: {len(self.stats['errors'])}")
            for error in self.stats['errors']:
                print(f"   • {error}")
        
        if backup_folder:
            print(f"\n💾 Original files backed up to: {backup_folder}")
        
        print("=" * 80)
    
    def generate_processing_report(self, input_folder):
        """Generate detailed processing report"""
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        report_file = os.path.join(input_folder, f"evasion_report_{timestamp}.json")
        
        try:
            report_data = {
                'timestamp': timestamp,
                'input_folder': input_folder,
                'statistics': self.stats,
                'system_info': {
                    'semantic_patterns': sum(len(v) for v in self.evasion_system.semantic_transformations.values()),
                    'structure_patterns': len(self.evasion_system.structure_patterns),
                    'word_transformations': sum(len(v) for v in self.evasion_system.word_transformations.values()),
                    'version': '2.0'
                }
            }
            
            with open(report_file, 'w', encoding='utf-8') as f:
                json.dump(report_data, f, ensure_ascii=False, indent=2)
            
            print(f"📋 Processing report saved: {report_file}")
            
        except Exception as e:
            print(f"⚠️ Could not save report: {e}")


def main():
    """Main function for batch Word document processing"""
    print("🎯 ULTIMATE PLAGIARISM EVASION - BATCH WORD PROCESSOR")
    print("Advanced text transformation with document structure preservation")
    print("=" * 80)
    
    # Configuration - SESUAIKAN DENGAN KEBUTUHAN ANDA
    INPUT_FOLDER = 'documents'       # Folder berisi dokumen Word
    AGGRESSIVENESS = 0.6            # 0.3 (subtle) sampai 0.9 (aggressive)
    CREATE_BACKUP = True            # Selalu buat backup sebelum proses
    
    print(f"📁 Target folder: {INPUT_FOLDER}")
    print(f"⚙️ Aggressiveness: {AGGRESSIVENESS}")
    print(f"💾 Create backup: {CREATE_BACKUP}")
    
    # Check if python-docx is available
    if not DOCX_AVAILABLE:
        print("❌ Error: python-docx library is required!")
        print("Install with: pip install python-docx")
        return
    
    # Check if input folder exists
    if not os.path.exists(INPUT_FOLDER):
        print(f"❌ Error: Folder '{INPUT_FOLDER}' tidak ditemukan!")
        print("Buat folder 'documents' dan masukkan file .docx di dalamnya")
        return
    
    try:
        # Initialize the batch processor
        processor = BatchWordProcessor()
        
        # Process the batch
        processor.process_batch(
            input_folder=INPUT_FOLDER,
            aggressiveness=AGGRESSIVENESS,
            create_backup=CREATE_BACKUP
        )
        
        print("\n🎉 Batch processing completed successfully!")
        print("✅ Dokumen telah diproses dan disimpan")
        print("💾 File backup tersedia jika diperlukan")
        
    except Exception as e:
        print(f"❌ Critical error: {e}")
        print("Pastikan semua dependencies terinstall dengan benar")


def demo_single_text():
    """Demo function untuk testing single text"""
    print("\n🧪 DEMO - SINGLE TEXT TRANSFORMATION:")
    
    # Initialize system
    evasion = UltimatePlagiarismEvasion()
    
    # Sample academic text
    sample_text = """Penelitian ini bertujuan untuk mengembangkan sistem informasi yang dapat 
    menggunakan teknologi modern untuk meningkatkan efisiensi kerja. Berdasarkan hasil 
    penelitian, sistem ini dapat menganalisis data dengan metode yang digunakan secara 
    efektif. Analisis data menunjukkan bahwa metode ini efektif untuk menghasilkan 
    hasil yang optimal."""
    
    print(f"\n📝 ORIGINAL TEXT:")
    print(f"'{sample_text}'")
    
    # Test different aggressiveness levels
    for level_name, aggressiveness in [('Subtle', 0.3), ('Balanced', 0.6), ('Aggressive', 0.9)]:
        print(f"\n🔄 {level_name.upper()} TRANSFORMATION (Aggressiveness: {aggressiveness}):")
        
        result = evasion.ultimate_transform(sample_text, aggressiveness)
        
        print(f"Result: '{result['transformed']}'")
        print(f"📊 Similarity Reduction: {result['similarity_reduction']:.1f}%")
        print(f"🔧 Changes Made: {result['total_changes']}")
        print(f"📈 Status: {result['status']}")
        print("-" * 50)


def check_requirements():
    """Check if all requirements are met"""
    print("🔍 Checking requirements...")
    
    requirements_met = True
    
    # Check python-docx
    if not DOCX_AVAILABLE:
        print("❌ python-docx not installed")
        print("   Install with: pip install python-docx")
        requirements_met = False
    else:
        print("✅ python-docx available")
    
    # Check if documents folder exists
    if os.path.exists('documents'):
        docx_files = [f for f in os.listdir('documents') if f.endswith('.docx')]
        print(f"✅ Documents folder found with {len(docx_files)} .docx files")
    else:
        print("⚠️ 'documents' folder not found")
        print("   Create 'documents' folder and put your .docx files there")
    
    return requirements_met


def interactive_mode():
    """Interactive mode untuk konfigurasi"""
    print("\n🎛️ INTERACTIVE CONFIGURATION MODE")
    print("=" * 50)
    
    # Get input folder
    input_folder = input("📁 Input folder path (default: 'documents'): ").strip()
    if not input_folder:
        input_folder = 'documents'
    
    # Get aggressiveness level
    print("\n⚙️ Aggressiveness levels:")
    print("   0.3 = Subtle (minimal changes, natural)")
    print("   0.6 = Balanced (recommended)")
    print("   0.9 = Aggressive (maximum changes)")
    
    aggressiveness_input = input("Aggressiveness level (0.3-0.9, default: 0.6): ").strip()
    try:
        aggressiveness = float(aggressiveness_input) if aggressiveness_input else 0.6
        aggressiveness = max(0.3, min(0.9, aggressiveness))  # Clamp between 0.3-0.9
    except ValueError:
        aggressiveness = 0.6
    
    # Backup option
    backup_input = input("💾 Create backup? (Y/n, default: Y): ").strip().lower()
    create_backup = backup_input not in ['n', 'no']
    
    print(f"\n📋 Configuration:")
    print(f"   📁 Input folder: {input_folder}")
    print(f"   ⚙️ Aggressiveness: {aggressiveness}")
    print(f"   💾 Create backup: {create_backup}")
    
    confirm = input("\n▶️ Start processing? (Y/n): ").strip().lower()
    if confirm in ['n', 'no']:
        print("❌ Processing cancelled")
        return
    
    # Process
    try:
        processor = BatchWordProcessor()
        processor.process_batch(
            input_folder=input_folder,
            aggressiveness=aggressiveness,
            create_backup=create_backup
        )
        print("\n🎉 Processing completed!")
    except Exception as e:
        print(f"❌ Error: {e}")


if __name__ == "__main__":
    print("🎯 ULTIMATE PLAGIARISM EVASION SYSTEM")
    print("Advanced Academic Text Transformation")
    print("=" * 80)
    
    # Check requirements first
    if not check_requirements():
        print("\n❌ Requirements not met. Please install missing dependencies.")
        exit(1)
    
    print("\n🚀 Choose operation mode:")
    print("1. 🤖 Auto Mode (use default settings)")
    print("2. 🎛️ Interactive Mode (configure settings)")
    print("3. 🧪 Demo Mode (test with sample text)")
    print("4. ❌ Exit")
    
    choice = input("\nEnter choice (1-4): ").strip()
    
    if choice == '1':
        main()
    elif choice == '2':
        interactive_mode()
    elif choice == '3':
        demo_single_text()
    elif choice == '4':
        print("👋 Goodbye!")
    else:
        print("❌ Invalid choice. Running auto mode...")
        main()
