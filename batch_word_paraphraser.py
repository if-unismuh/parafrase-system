# batch_word_paraphraser.py
"""
Batch Word Document Paraphraser - Direct Modification
Memproses multiple dokumen Word BAB 1-3 dan langsung memodifikasi file asli
Author: DevNoLife
Version: 1.0
"""

import os
import json
import re
from datetime import datetime
from pathlib import Path
import docx
from docx.shared import RGBColor
from docx.enum.text import WD_COLOR_INDEX
import shutil
from collections import defaultdict

# Import sistem parafrase utama
from main_paraphrase_system import IndonesianParaphraseSystem

class BatchWordParaphraser:
    def __init__(self, synonym_file='sinonim.json'):
        print("🚀 Initializing Batch Word Document Paraphraser...")
        
        # Initialize paraphrase system
        self.paraphraser = IndonesianParaphraseSystem(synonym_file)
        
        # Processing statistics
        self.stats = {
            'total_documents': 0,
            'processed_documents': 0,
            'total_paragraphs': 0,
            'paraphrased_paragraphs': 0,
            'total_changes': 0,
            'processing_time': 0,
            'errors': []
        }
        
        # Priority sections for paraphrasing
        self.priority_sections = {
            'HIGH': [
                'latar belakang', 'landasan teori', 'tinjauan pustaka', 
                'kajian teori', 'teori', 'konsep'
            ],
            'MEDIUM': [
                'metodologi', 'metode penelitian', 'penelitian terkait',
                'rumusan masalah', 'tujuan penelitian'
            ],
            'LOW': [
                'manfaat penelitian', 'sistematika penulisan',
                'batasan masalah', 'definisi operasional'
            ]
        }
        
        # Minimum paragraph length for paraphrasing (words)
        self.min_paragraph_length = 30
        
        print("✅ Batch Word Paraphraser ready!")
    
    def backup_documents(self, input_folder, backup_folder=None):
        """Create backup of original documents before processing"""
        if not backup_folder:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            backup_folder = f"{input_folder}_backup_{timestamp}"
        
        try:
            if not os.path.exists(backup_folder):
                os.makedirs(backup_folder)
            
            docx_files = [f for f in os.listdir(input_folder) if f.endswith('.docx')]
            
            for file in docx_files:
                src = os.path.join(input_folder, file)
                dst = os.path.join(backup_folder, file)
                shutil.copy2(src, dst)
            
            print(f"✅ Backup created: {backup_folder}")
            print(f"📁 Backed up {len(docx_files)} documents")
            return backup_folder
            
        except Exception as e:
            print(f"❌ Error creating backup: {e}")
            return None
    
    def get_section_priority(self, text):
        """Determine priority level of a section based on content"""
        text_lower = text.lower()
        
        # Check for high priority keywords
        for keyword in self.priority_sections['HIGH']:
            if keyword in text_lower:
                return 'HIGH'
        
        # Check for medium priority keywords
        for keyword in self.priority_sections['MEDIUM']:
            if keyword in text_lower:
                return 'MEDIUM'
        
        return 'LOW'
    
    def is_paragraph_suitable_for_paraphrasing(self, paragraph_text):
        """Check if paragraph should be paraphrased"""
        # Skip empty or very short paragraphs
        word_count = len(paragraph_text.split())
        if word_count < self.min_paragraph_length:
            return False
        
        # Skip paragraphs that are mostly numbers/tables
        if re.search(r'\d+\.\d+|\btabel\b|\bgambar\b', paragraph_text.lower()):
            return False
        
        # Skip reference lists
        if re.search(r'\(\d{4}\)|\bet al\b|\bvol\b|\bno\b', paragraph_text.lower()):
            return False
        
        # Skip headers/titles (all caps or very short)
        if paragraph_text.isupper() and word_count < 10:
            return False
        
        return True
    
    def paraphrase_paragraph(self, paragraph_text, aggressiveness=0.5):
        """Paraphrase a single paragraph"""
        try:
            result = self.paraphraser.generate_paraphrase(paragraph_text, aggressiveness)
            return result
        except Exception as e:
            print(f"⚠️  Error paraphrasing paragraph: {e}")
            return {
                'paraphrase': paragraph_text,  # Return original if error
                'similarity': 100.0,
                'plagiarism_reduction': 0.0,
                'changes_made': 0,
                'status': f'Error: {str(e)}'
            }
    
    def process_word_document(self, file_path, output_path=None, aggressiveness=0.5):
        """Process a single Word document"""
        try:
            print(f"\n📄 Processing: {os.path.basename(file_path)}")
            
            # Load document
            doc = docx.Document(file_path)
            
            if not output_path:
                output_path = file_path  # Overwrite original
            
            # Statistics for this document
            doc_stats = {
                'filename': os.path.basename(file_path),
                'total_paragraphs': 0,
                'processed_paragraphs': 0,
                'changes_made': 0,
                'sections_processed': defaultdict(int)
            }
            
            current_section = 'UNKNOWN'
            
            # Process each paragraph
            for i, paragraph in enumerate(doc.paragraphs):
                para_text = paragraph.text.strip()
                doc_stats['total_paragraphs'] += 1
                
                if not para_text:
                    continue
                
                # Detect section headers
                if self.is_section_header(para_text):
                    current_section = para_text
                    print(f"  📍 Section: {current_section}")
                    continue
                
                # Check if paragraph is suitable for paraphrasing
                if not self.is_paragraph_suitable_for_paraphrasing(para_text):
                    continue
                
                # Determine priority and aggressiveness
                section_priority = self.get_section_priority(current_section + " " + para_text)
                
                # Adjust aggressiveness based on priority
                if section_priority == 'HIGH':
                    para_aggressiveness = min(aggressiveness + 0.2, 0.8)
                elif section_priority == 'MEDIUM':
                    para_aggressiveness = aggressiveness
                else:
                    para_aggressiveness = max(aggressiveness - 0.2, 0.3)
                
                # Paraphrase the paragraph
                result = self.paraphrase_paragraph(para_text, para_aggressiveness)
                
                # Only update if there's significant improvement
                if result['plagiarism_reduction'] > 20:  # At least 20% reduction
                    # Update paragraph text
                    paragraph.clear()
                    paragraph.add_run(result['paraphrase'])
                    
                    # Add comment/highlight for tracking (optional)
                    if result['changes_made'] > 0:
                        # Add subtle highlighting to show this was modified
                        for run in paragraph.runs:
                            run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                    
                    doc_stats['processed_paragraphs'] += 1
                    doc_stats['changes_made'] += result['changes_made']
                    doc_stats['sections_processed'][section_priority] += 1
                    
                    print(f"    ✅ Para {i+1}: {result['plagiarism_reduction']:.1f}% reduction, {result['changes_made']} changes")
                else:
                    print(f"    ⏭️  Para {i+1}: Skipped (minimal improvement: {result['plagiarism_reduction']:.1f}%)")
            
            # Save the modified document
            doc.save(output_path)
            
            # Update global statistics
            self.stats['processed_documents'] += 1
            self.stats['total_paragraphs'] += doc_stats['total_paragraphs']
            self.stats['paraphrased_paragraphs'] += doc_stats['processed_paragraphs']
            self.stats['total_changes'] += doc_stats['changes_made']
            
            print(f"  📊 Document Summary:")
            print(f"     • Total paragraphs: {doc_stats['total_paragraphs']}")
            print(f"     • Processed paragraphs: {doc_stats['processed_paragraphs']}")
            print(f"     • Total changes: {doc_stats['changes_made']}")
            print(f"     • High priority sections: {doc_stats['sections_processed']['HIGH']}")
            print(f"     • Medium priority sections: {doc_stats['sections_processed']['MEDIUM']}")
            print(f"  ✅ Saved: {output_path}")
            
            return doc_stats
            
        except Exception as e:
            error_msg = f"Error processing {file_path}: {str(e)}"
            print(f"❌ {error_msg}")
            self.stats['errors'].append(error_msg)
            return None
    
    def is_section_header(self, text):
        """Check if text is a section header"""
        # Common patterns for section headers
        patterns = [
            r'^BAB\s+[IVX]+',  # BAB I, BAB II, etc.
            r'^\d+\.\d+',      # 1.1, 2.3, etc.
            r'^[A-Z\s]+$',     # ALL CAPS short text
        ]
        
        for pattern in patterns:
            if re.match(pattern, text.strip()) and len(text.split()) < 10:
                return True
        
        return False
    
    def process_batch(self, input_folder, aggressiveness=0.5, create_backup=True):
        """Process all Word documents in a folder"""
        start_time = datetime.now()
        
        print("=" * 80)
        print("🎯 BATCH WORD DOCUMENT PARAPHRASER")
        print("=" * 80)
        
        # Validate input folder
        if not os.path.exists(input_folder):
            print(f"❌ Input folder not found: {input_folder}")
            return
        
        # Get list of Word documents
        docx_files = [f for f in os.listdir(input_folder) if f.endswith('.docx') and not f.startswith('~')]
        
        if not docx_files:
            print(f"❌ No Word documents found in: {input_folder}")
            return
        
        self.stats['total_documents'] = len(docx_files)
        
        print(f"📁 Input folder: {input_folder}")
        print(f"📄 Found {len(docx_files)} documents")
        print(f"⚙️  Aggressiveness level: {aggressiveness}")
        
        # Create backup if requested
        backup_folder = None
        if create_backup:
            backup_folder = self.backup_documents(input_folder)
            if not backup_folder:
                print("❌ Failed to create backup. Aborting for safety.")
                return
        
        print(f"\n🚀 Starting batch processing...")
        
        # Process each document
        for i, filename in enumerate(docx_files, 1):
            file_path = os.path.join(input_folder, filename)
            
            print(f"\n{'='*20} Document {i}/{len(docx_files)} {'='*20}")
            
            doc_result = self.process_word_document(file_path, aggressiveness=aggressiveness)
            
            if doc_result is None:
                print(f"⚠️  Skipped {filename} due to errors")
        
        # Calculate processing time
        end_time = datetime.now()
        self.stats['processing_time'] = (end_time - start_time).total_seconds()
        
        # Print final statistics
        self.print_batch_summary(backup_folder)
        
        # Generate report
        self.generate_processing_report(input_folder)
    
    def print_batch_summary(self, backup_folder=None):
        """Print summary of batch processing"""
        print("\n" + "=" * 80)
        print("📊 BATCH PROCESSING SUMMARY")
        print("=" * 80)
        
        print(f"📄 Documents processed: {self.stats['processed_documents']}/{self.stats['total_documents']}")
        print(f"📝 Total paragraphs: {self.stats['total_paragraphs']}")
        print(f"✏️  Paraphrased paragraphs: {self.stats['paraphrased_paragraphs']}")
        print(f"🔄 Total changes made: {self.stats['total_changes']}")
        print(f"⏱️  Processing time: {self.stats['processing_time']:.1f} seconds")
        
        if self.stats['total_paragraphs'] > 0:
            success_rate = (self.stats['paraphrased_paragraphs'] / self.stats['total_paragraphs']) * 100
            print(f"📈 Success rate: {success_rate:.1f}%")
        
        if self.stats['errors']:
            print(f"\n⚠️  Errors encountered: {len(self.stats['errors'])}")
            for error in self.stats['errors']:
                print(f"   • {error}")
        
        if backup_folder:
            print(f"\n💾 Original files backed up to: {backup_folder}")
        
        print("=" * 80)
    
    def generate_processing_report(self, input_folder):
        """Generate detailed processing report"""
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        report_file = os.path.join(input_folder, f"paraphrase_report_{timestamp}.json")
        
        try:
            report_data = {
                'timestamp': timestamp,
                'input_folder': input_folder,
                'statistics': self.stats,
                'system_info': {
                    'synonym_groups': len(self.paraphraser.synonyms),
                    'phrase_patterns': len(self.paraphraser.phrase_replacements),
                    'version': '1.0'
                }
            }
            
            with open(report_file, 'w', encoding='utf-8') as f:
                json.dump(report_data, f, ensure_ascii=False, indent=2)
            
            print(f"📋 Processing report saved: {report_file}")
            
        except Exception as e:
            print(f"⚠️  Could not save report: {e}")


def main():
    """Main function for batch processing"""
    print("🎯 BATCH WORD DOCUMENT PARAPHRASER")
    print("Directly modifies original Word documents with paraphrased content")
    print("=" * 80)
    
    # Configuration
    SYNONYM_FILE = 'sinonim.json'  # Path to your synonym database
    INPUT_FOLDER = 'documents'    # Folder containing Word documents
    AGGRESSIVENESS = 0.6         # 0.3 (conservative) to 0.8 (aggressive)
    CREATE_BACKUP = True         # Always create backup before processing
    
    # Initialize the batch processor
    processor = BatchWordParaphraser(SYNONYM_FILE)
    
    # Process the batch
    processor.process_batch(
        input_folder=INPUT_FOLDER,
        aggressiveness=AGGRESSIVENESS,
        create_backup=CREATE_BACKUP
    )
    
    print("\n🎉 Batch processing completed!")
    print("Check the processed documents and backup folder.")


if __name__ == "__main__":
    main()
