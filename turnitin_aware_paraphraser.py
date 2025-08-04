# turnitin_aware_paraphraser.py
"""
Turnitin-Aware Paraphraser - Targeting Specific Plagiarism Patterns
Based on actual Turnitin detection results to maximize effectiveness
Author: DevNoLife
Version: 3.0
"""

import os
import json
import re
import time
import random
from datetime import datetime
from pathlib import Path
import docx
from docx.shared import RGBColor
from docx.enum.text import WD_COLOR_INDEX
import shutil
from collections import defaultdict
import hashlib

# Updated Gemini AI imports
try:
    from google import genai
    from google.genai import types
    GEMINI_AVAILABLE = True
except ImportError:
    print("⚠️  google-genai package not found. Install with: pip install google-genai")
    GEMINI_AVAILABLE = False

# Import sistem parafrase lokal
from main_paraphrase_system import IndonesianParaphraseSystem

class TurnitinAwareParaphraser:
    def __init__(self, synonym_file='contoh.json', gemini_api_key=None):
        print("🎯 Initializing Turnitin-Aware Paraphraser...")
        
        # Initialize local paraphrase system
        self.local_paraphraser = IndonesianParaphraseSystem(synonym_file)
        
        # Gemini AI configuration
        self.gemini_api_key = gemini_api_key
        self.gemini_client = None
        
        if gemini_api_key and GEMINI_AVAILABLE:
            try:
                os.environ["GEMINI_API_KEY"] = gemini_api_key
                self.gemini_client = genai.Client(api_key=gemini_api_key)
                print("✅ Gemini 2.0 Flash configured for Turnitin targeting")
            except Exception as e:
                print(f"⚠️  Error initializing Gemini: {e}")
                self.gemini_client = None
        
        # Cost tracking
        self.cost_tracker = {
            'local_calls': 0,
            'ai_calls': 0,
            'ai_tokens_used': 0,
            'estimated_cost_usd': 0.0,
            'turnitin_patterns_detected': 0,
            'high_risk_patterns_fixed': 0
        }
        
        # AI result cache
        self.ai_cache = {}
        
        # Turnitin-specific configuration
        self.config = {
            'turnitin_threshold': 0.8,           # High priority for Turnitin patterns
            'academic_template_threshold': 0.9,  # Very high for academic templates
            'technical_definition_threshold': 0.85, # High for technical definitions
            'model_name': 'gemini-2.0-flash',
            'max_batch_size': 2,                 # Smaller batches for quality
            'ai_retry_attempts': 3
        }
        
        # HIGH PRIORITY: Patterns that Turnitin actually detected (from the PDF)
        self.turnitin_detected_patterns = {
            # Academic boilerplate (heavily flagged in Turnitin)
            'academic_templates': [
                r'berdasarkan latar belakang yang telah diuraikan',
                r'maka rumusan masalah dalam penelitian ini adalah',
                r'penelitian ini bertujuan untuk',
                r'menambah wawasan dan pemahaman penulis',
                r'menjadi referensi bagi peneliti',
                r'secara garis besar penulisan.*terbagi menjadi',
                r'bab ini menerangkan secara singkat dan jelas',
                r'pada bab ini membahas tentang teori-teori',
                r'membahas tentang metode penelitian dan alat'
            ],
            
            # Technical definitions (LSB was heavily flagged)
            'technical_definitions': [
                r'least significant bit.*adalah teknik steganografi',
                r'cara kerja metode lsb yaitu mengubah bit',
                r'bit terakhir.*bit yang paling tidak signifikan',
                r'cover image.*yang tidak berpengaruh signifikan',
                r'metode ini banyak digunakan dalam invisible watermarking',
                r'quick response.*qr.*code telah diadopsi',
                r'fleksibilitas dan kapasitas penyimpanan qr code'
            ],
            
            # Research methodology (standard formats)
            'methodology_patterns': [
                r'penelitian ini dibatasi pada',
                r'fokus penelitian adalah pada',
                r'tidak termasuk dalam ruang lingkup penelitian',
                r'penelitian ini tidak mencakup',
                r'karena.*tidak dapat dipertahankan pada'
            ],
            
            # Domain-specific terminology
            'domain_terms': [
                'steganografi', 'watermarking', 'least significant bit', 'lsb',
                'cover image', 'invisible watermarking', 'qr code', 'quick response',
                'imperceptibility', 'ketidakterlihatan', 'citra penampung'
            ]
        }
        
        # ADVANCED: Paraphrase strategies specifically for Turnitin patterns
        self.turnitin_paraphrase_strategies = {
            # Academic template replacements
            'academic_replacements': {
                'berdasarkan latar belakang yang telah diuraikan': [
                    'mengacu pada konteks yang telah dijelaskan sebelumnya',
                    'merujuk pada uraian latar belakang di atas',
                    'berlandaskan pemaparan konteks yang telah dikemukakan'
                ],
                'maka rumusan masalah dalam penelitian ini adalah': [
                    'sehingga permasalahan yang akan dikaji meliputi',
                    'dengan demikian fokus permasalahan penelitian mencakup',
                    'oleh karena itu pertanyaan penelitian yang diajukan yaitu'
                ],
                'penelitian ini bertujuan untuk': [
                    'kajian ini dimaksudkan untuk',
                    'studi ini berupaya untuk',
                    'riset ini dirancang untuk'
                ],
                'menambah wawasan dan pemahaman penulis': [
                    'memperluas pengetahuan dan insight peneliti',
                    'mengembangkan kompetensi dan perspektif penulis',
                    'meningkatkan kapasitas dan pemahaman mendalam penulis'
                ]
            },
            
            # Technical definition restructuring
            'technical_restructuring': {
                'cara kerja metode': 'prinsip operasional teknik',
                'mengubah bit terakhir': 'memodifikasi bit paling tidak signifikan',
                'tidak berpengaruh signifikan': 'berdampak minimal terhadap',
                'banyak digunakan dalam': 'sering diaplikasikan untuk',
                'implementasinya yang relatif mudah': 'penerapannya yang cukup sederhana'
            },
            
            # Domain terminology alternatives
            'domain_alternatives': {
                'steganografi': ['teknik penyembunyian data', 'metode kamuflase informasi'],
                'watermarking': ['penandaan digital', 'pemberian tanda pada konten'],
                'cover image': ['citra penampung', 'gambar kontainer'],
                'invisible watermarking': ['penandaan tak terlihat', 'watermark tersembunyi'],
                'imperceptibility': ['tingkat ketidakterlihatan', 'sifat tidak kasat mata'],
                'least significant bit': ['bit paling tidak signifikan', 'bit dengan bobot terendah']
            }
        }
        
        print(f"✅ Local system ready with {len(self.local_paraphraser.synonyms)} synonyms")
        print(f"🎯 Turnitin patterns loaded: {sum(len(v) for v in self.turnitin_detected_patterns.values())} patterns")
        print(f"🤖 AI Enhancement: {'Enabled' if self.gemini_client else 'Local only'}")
        print("✅ Turnitin-Aware Paraphraser ready!")
    
    def detect_turnitin_risk_level(self, text):
        """Detect how likely text is to be flagged by Turnitin"""
        risk_score = 0.0
        detected_patterns = []
        
        text_lower = text.lower()
        
        # Check for academic templates (highest risk)
        for pattern in self.turnitin_detected_patterns['academic_templates']:
            if re.search(pattern, text_lower):
                risk_score += 0.3
                detected_patterns.append(('academic_template', pattern))
        
        # Check for technical definitions (high risk)
        for pattern in self.turnitin_detected_patterns['technical_definitions']:
            if re.search(pattern, text_lower):
                risk_score += 0.25
                detected_patterns.append(('technical_definition', pattern))
        
        # Check for methodology patterns (medium-high risk)
        for pattern in self.turnitin_detected_patterns['methodology_patterns']:
            if re.search(pattern, text_lower):
                risk_score += 0.2
                detected_patterns.append(('methodology', pattern))
        
        # Check for domain-specific terms (medium risk)
        domain_count = sum(1 for term in self.turnitin_detected_patterns['domain_terms'] 
                          if term in text_lower)
        if domain_count > 0:
            risk_score += min(domain_count * 0.1, 0.4)
            detected_patterns.append(('domain_terms', f"{domain_count} terms"))
        
        # Normalize risk score
        risk_level = min(risk_score, 1.0)
        
        return {
            'risk_level': risk_level,
            'risk_category': self.get_risk_category(risk_level),
            'detected_patterns': detected_patterns,
            'needs_ai': risk_level >= self.config['turnitin_threshold']
        }
    
    def get_risk_category(self, risk_level):
        """Convert risk level to category"""
        if risk_level >= 0.8:
            return "🔴 VERY HIGH - Turnitin will definitely flag this"
        elif risk_level >= 0.6:
            return "🟠 HIGH - Likely to be flagged"
        elif risk_level >= 0.4:
            return "🟡 MEDIUM - May be flagged"
        elif risk_level >= 0.2:
            return "🟢 LOW - Probably safe"
        else:
            return "✅ VERY LOW - Safe"
    
    def apply_turnitin_specific_paraphrasing(self, text):
        """Apply paraphrasing specifically targeting Turnitin patterns"""
        modified_text = text
        changes_made = []
        
        # Strategy 1: Academic template replacement
        for original, alternatives in self.turnitin_paraphrase_strategies['academic_replacements'].items():
            pattern = re.compile(original, re.IGNORECASE)
            if pattern.search(modified_text):
                replacement = random.choice(alternatives)
                modified_text = pattern.sub(replacement, modified_text, count=1)
                changes_made.append({
                    'type': 'academic_template',
                    'original': original,
                    'replacement': replacement,
                    'risk_reduction': 0.3
                })
        
        # Strategy 2: Technical restructuring
        for original, replacement in self.turnitin_paraphrase_strategies['technical_restructuring'].items():
            if original in modified_text.lower():
                pattern = re.compile(re.escape(original), re.IGNORECASE)
                modified_text = pattern.sub(replacement, modified_text)
                changes_made.append({
                    'type': 'technical_restructuring',
                    'original': original,
                    'replacement': replacement,
                    'risk_reduction': 0.2
                })
        
        # Strategy 3: Domain terminology replacement
        for term, alternatives in self.turnitin_paraphrase_strategies['domain_alternatives'].items():
            if term in modified_text.lower():
                replacement = random.choice(alternatives)
                pattern = re.compile(r'\b' + re.escape(term) + r'\b', re.IGNORECASE)
                modified_text = pattern.sub(replacement, modified_text)
                changes_made.append({
                    'type': 'domain_terminology',
                    'original': term,
                    'replacement': replacement,
                    'risk_reduction': 0.15
                })
        
        return modified_text, changes_made
    
    def create_cache_key(self, text):
        """Create cache key for AI results"""
        return hashlib.md5(text.encode('utf-8')).hexdigest()[:16]
    
    def call_gemini_for_turnitin_patterns(self, text, risk_analysis):
        """Call Gemini AI specifically for high-risk Turnitin patterns"""
        if not self.gemini_client:
            return None
        
        try:
            # Create specialized prompt for Turnitin avoidance
            prompt = self.create_turnitin_avoidance_prompt(text, risk_analysis)
            
            contents = [
                types.Content(
                    role="user",
                    parts=[types.Part.from_text(text=prompt)],
                ),
            ]
            
            generate_content_config = types.GenerateContentConfig(
                temperature=0.5,  # Slightly higher for more variation
                top_k=60,
                top_p=0.9,
                max_output_tokens=4096,
                candidate_count=1,
            )
            
            for attempt in range(self.config['ai_retry_attempts']):
                try:
                    response_text = ""
                    for chunk in self.gemini_client.models.generate_content_stream(
                        model=self.config['model_name'],
                        contents=contents,
                        config=generate_content_config,
                    ):
                        if hasattr(chunk, 'text') and chunk.text:
                            response_text += chunk.text
                    
                    if response_text.strip():
                        # Track usage
                        self.cost_tracker['ai_calls'] += 1
                        tokens_used = len(prompt + response_text) // 4
                        self.cost_tracker['ai_tokens_used'] += tokens_used
                        self.cost_tracker['estimated_cost_usd'] += tokens_used * 0.000001
                        
                        return response_text.strip()
                    
                except Exception as e:
                    print(f"⚠️  Turnitin AI attempt {attempt + 1} failed: {e}")
                    if attempt < self.config['ai_retry_attempts'] - 1:
                        time.sleep(2 ** attempt)
            
            return None
            
        except Exception as e:
            print(f"❌ Turnitin AI error: {e}")
            return None
    
    def create_turnitin_avoidance_prompt(self, text, risk_analysis):
        """Create specialized prompt for avoiding Turnitin detection"""
        detected_patterns = [p[1] for p in risk_analysis['detected_patterns']]
        
        prompt = f"""Kamu adalah expert anti-plagiarisme yang sangat memahami cara kerja Turnitin dan sistem deteksi plagiarisme lainnya.

MISI KRITIS: Parafrase teks berikut untuk menghindari deteksi Turnitin. Teks ini memiliki risk level {risk_analysis['risk_level']:.2f} dan mengandung pattern berisiko tinggi.

PATTERN BERISIKO TERDETEKSI:
{chr(10).join(f"- {pattern}" for pattern in detected_patterns[:5])}

STRATEGI ANTI-TURNITIN:
✅ Ubah STRUKTUR KALIMAT secara drastis (aktif↔pasif, urutan klausa)
✅ Ganti academic boilerplate dengan variasi yang sangat berbeda
✅ Restructure definisi teknis dengan approach yang beda
✅ Variasikan terminology domain dengan sinonim yang jarang digunakan
✅ Pecah kalimat panjang atau gabungkan kalimat pendek
✅ Ubah urutan informasi tanpa mengubah makna

PANTANGAN:
❌ Jangan hanya ganti sinonim sederhana
❌ Jangan pertahankan struktur kalimat yang sama
❌ Jangan gunakan template akademik standar

TARGET: Turunkan similarity dari {risk_analysis['risk_level']*100:.0f}% menjadi <20%
FORMAT: Berikan hasil parafrase langsung tanpa penjelasan.

TEKS ASLI:
{text}

PARAFRASE ANTI-TURNITIN:"""
        
        return prompt
    
    def process_paragraph_turnitin_aware(self, paragraph_text, paragraph_index):
        """Process paragraph with Turnitin-specific awareness"""
        
        # Step 1: Analyze Turnitin risk
        risk_analysis = self.detect_turnitin_risk_level(paragraph_text)
        self.cost_tracker['turnitin_patterns_detected'] += len(risk_analysis['detected_patterns'])
        
        # Step 2: Apply local Turnitin-specific paraphrasing first
        local_modified, local_changes = self.apply_turnitin_specific_paraphrasing(paragraph_text)
        
        # Step 3: Get local system result on the modified text
        local_result = self.local_paraphraser.generate_paraphrase(local_modified, aggressiveness=0.7)
        local_result['method'] = 'local_turnitin_aware'
        local_result['risk_analysis'] = risk_analysis
        local_result['turnitin_changes'] = len(local_changes)
        self.cost_tracker['local_calls'] += 1
        
        # Step 4: Decide if AI is needed for high-risk patterns
        if risk_analysis['needs_ai'] and self.gemini_client:
            # Check cache first
            cache_key = self.create_cache_key(paragraph_text)
            if cache_key in self.ai_cache:
                ai_result_text = self.ai_cache[cache_key]
            else:
                print(f"    🎯 High-risk Turnitin pattern detected - Using AI")
                ai_result_text = self.call_gemini_for_turnitin_patterns(paragraph_text, risk_analysis)
                
                if ai_result_text:
                    self.ai_cache[cache_key] = ai_result_text
                else:
                    # AI failed, use local result
                    local_result['method'] = 'local_fallback_turnitin'
                    return local_result
            
            if ai_result_text:
                # Calculate AI result metrics
                ai_similarity = self.local_paraphraser.calculate_similarity(paragraph_text, ai_result_text)
                
                ai_result = {
                    'paraphrase': ai_result_text,
                    'similarity': round(ai_similarity, 2),
                    'plagiarism_reduction': round(100 - ai_similarity, 2),
                    'changes_made': 1,
                    'method': 'turnitin_ai_specialized',
                    'status': self.local_paraphraser._get_plagiarism_status(ai_similarity),
                    'risk_analysis': risk_analysis,
                    'turnitin_changes': 1
                }
                
                # Choose better result
                if ai_result['plagiarism_reduction'] > local_result['plagiarism_reduction']:
                    self.cost_tracker['high_risk_patterns_fixed'] += 1
                    return ai_result
        
        return local_result
    
    def process_batch_documents(self, input_folder, create_backup=True):
        """Process batch with Turnitin-aware approach"""
        start_time = datetime.now()
        
        print("=" * 80)
        print("🎯 TURNITIN-AWARE PARAPHRASER - ANTI-PLAGIARISM TARGETING")
        print("=" * 80)
        
        if not os.path.exists(input_folder):
            print(f"❌ Input folder not found: {input_folder}")
            return
        
        docx_files = [f for f in os.listdir(input_folder) if f.endswith('.docx') and not f.startswith('~')]
        
        if not docx_files:
            print(f"❌ No Word documents found in: {input_folder}")
            return
        
        print(f"📁 Input folder: {input_folder}")
        print(f"📄 Found {len(docx_files)} documents")
        print(f"🎯 Turnitin pattern detection: Enabled")
        print(f"🤖 AI Enhancement: {'Enabled for high-risk patterns' if self.gemini_client else 'Disabled'}")
        
        # Create backup
        backup_folder = None
        if create_backup:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            backup_folder = f"{input_folder}_backup_turnitin_{timestamp}"
            
            try:
                if not os.path.exists(backup_folder):
                    os.makedirs(backup_folder)
                
                for file in docx_files:
                    src = os.path.join(input_folder, file)
                    dst = os.path.join(backup_folder, file)
                    shutil.copy2(src, dst)
                
                print(f"✅ Backup created: {backup_folder}")
            except Exception as e:
                print(f"❌ Error creating backup: {e}")
                return
        
        print(f"\n🚀 Starting Turnitin-aware processing...")
        
        # Process documents
        total_stats = {
            'documents_processed': 0,
            'paragraphs_processed': 0,
            'high_risk_detected': 0,
            'medium_risk_detected': 0,
            'low_risk_detected': 0,
            'ai_enhanced': 0,
            'local_turnitin_aware': 0,
            'total_changes': 0
        }
        
        for i, filename in enumerate(docx_files, 1):
            file_path = os.path.join(input_folder, filename)
            
            print(f"\n{'='*20} Document {i}/{len(docx_files)} {'='*20}")
            
            doc_stats = self.process_single_document_turnitin_aware(file_path)
            
            if doc_stats:
                total_stats['documents_processed'] += 1
                for key in doc_stats:
                    if key in total_stats:
                        total_stats[key] += doc_stats[key]
        
        # Calculate processing time
        end_time = datetime.now()
        processing_time = (end_time - start_time).total_seconds()
        
        # Print comprehensive summary
        self.print_turnitin_aware_summary(total_stats, processing_time, backup_folder)
        
        # Generate detailed report
        self.generate_turnitin_aware_report(input_folder, total_stats, processing_time)
    
    def process_single_document_turnitin_aware(self, file_path):
        """Process single document with Turnitin awareness"""
        try:
            print(f"\n📄 Processing: {os.path.basename(file_path)}")
            
            doc = docx.Document(file_path)
            
            doc_stats = {
                'filename': os.path.basename(file_path),
                'total_paragraphs': 0,
                'paragraphs_processed': 0,
                'high_risk_detected': 0,
                'medium_risk_detected': 0, 
                'low_risk_detected': 0,
                'ai_enhanced': 0,
                'local_turnitin_aware': 0,
                'total_changes': 0
            }
            
            current_section = 'UNKNOWN'
            processed_paragraphs = []
            
            # Collect suitable paragraphs
            for paragraph in doc.paragraphs:
                para_text = paragraph.text.strip()
                doc_stats['total_paragraphs'] += 1
                
                if not para_text:
                    continue
                
                if self.is_section_header(para_text):
                    current_section = para_text
                    print(f"  📍 Section: {current_section}")
                    continue
                
                if self.is_paragraph_suitable_for_paraphrasing(para_text):
                    processed_paragraphs.append({
                        'paragraph': paragraph,
                        'text': para_text,
                        'index': len(processed_paragraphs)
                    })
            
            total_suitable = len(processed_paragraphs)
            print(f"  📝 Found {total_suitable} suitable paragraphs for Turnitin analysis")
            
            # Process paragraphs with Turnitin awareness
            for para_data in processed_paragraphs:
                paragraph = para_data['paragraph']
                para_text = para_data['text']
                para_index = para_data['index']
                
                # Turnitin-aware processing
                result = self.process_paragraph_turnitin_aware(para_text, para_index)
                
                # Track risk levels
                risk_level = result.get('risk_analysis', {}).get('risk_level', 0)
                if risk_level >= 0.6:
                    doc_stats['high_risk_detected'] += 1
                elif risk_level >= 0.3:
                    doc_stats['medium_risk_detected'] += 1
                else:
                    doc_stats['low_risk_detected'] += 1
                
                # Apply changes if improvement is significant
                if result['plagiarism_reduction'] > 12:  # Lower threshold for Turnitin patterns
                    paragraph.clear()
                    paragraph.add_run(result['paraphrase'])
                    
                    # Color coding based on risk and method
                    if 'turnitin_ai' in result['method']:
                        # Red for AI-enhanced high-risk patterns
                        for run in paragraph.runs:
                            run.font.highlight_color = WD_COLOR_INDEX.RED
                        doc_stats['ai_enhanced'] += 1
                    elif 'turnitin_aware' in result['method']:
                        # Orange for local Turnitin-aware
                        for run in paragraph.runs:
                            run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                        doc_stats['local_turnitin_aware'] += 1
                    
                    doc_stats['paragraphs_processed'] += 1
                    doc_stats['total_changes'] += result.get('changes_made', 1)
                    
                    risk_category = result.get('risk_analysis', {}).get('risk_category', 'Unknown')
                    print(f"    ✅ Para {para_index+1}: {result['plagiarism_reduction']:.1f}% reduction ({result['method']}) - {risk_category}")
                else:
                    print(f"    ⏭️  Para {para_index+1}: Skipped ({result['plagiarism_reduction']:.1f}% - {result['method']})")
            
            # Save document
            doc.save(file_path)
            
            print(f"  📊 Document Summary:")
            print(f"     • Total paragraphs: {doc_stats['total_paragraphs']}")
            print(f"     • Processed: {doc_stats['paragraphs_processed']}")
            print(f"     • High-risk detected: {doc_stats['high_risk_detected']}")
            print(f"     • Medium-risk detected: {doc_stats['medium_risk_detected']}")
            print(f"     • AI-enhanced: {doc_stats['ai_enhanced']}")
            print(f"     • Local Turnitin-aware: {doc_stats['local_turnitin_aware']}")
            print(f"  ✅ Saved: {file_path}")
            
            return doc_stats
            
        except Exception as e:
            print(f"❌ Error processing {file_path}: {e}")
            return None
    
    def is_section_header(self, text):
        """Check if text is a section header"""
        patterns = [
            r'^BAB\s+[IVX]+',
            r'^\d+\.\d+',
            r'^[A-Z\s]+$',
        ]
        
        for pattern in patterns:
            if re.match(pattern, text.strip()) and len(text.split()) < 10:
                return True
        return False
    
    def is_paragraph_suitable_for_paraphrasing(self, paragraph_text):
        """Check if paragraph should be processed"""
        word_count = len(paragraph_text.split())
        if word_count < 10:  # Even shorter paragraphs for Turnitin patterns
            return False
        
        if re.search(r'\d+\.\d+|\btabel\b|\bgambar\b', paragraph_text.lower()):
            return False
        
        if paragraph_text.isupper() and word_count < 8:
            return False
        
        return True
    
    def print_turnitin_aware_summary(self, stats, processing_time, backup_folder):
        """Print Turnitin-aware processing summary"""
        print("\n" + "=" * 80)
        print("📊 TURNITIN-AWARE PROCESSING SUMMARY")
        print("=" * 80)
        
        print(f"📄 Documents processed: {stats['documents_processed']}")
        print(f"📝 Paragraphs processed: {stats['paragraphs_processed']}")
        print(f"🔴 High-risk patterns: {stats['high_risk_detected']}")
        print(f"🟠 Medium-risk patterns: {stats['medium_risk_detected']}")
        print(f"🟢 Low-risk patterns: {stats['low_risk_detected']}")
        print(f"🤖 AI-enhanced: {stats['ai_enhanced']}")
        print(f"🎯 Local Turnitin-aware: {stats['local_turnitin_aware']}")
        print(f"🔄 Total changes: {stats['total_changes']}")
        print(f"⏱️  Processing time: {processing_time:.1f} seconds")
        
        # Turnitin-specific metrics
        print(f"\n🎯 TURNITIN ANALYSIS:")
        print(f"🔍 Patterns detected: {self.cost_tracker['turnitin_patterns_detected']}")
        print(f"🛠️  High-risk patterns fixed: {self.cost_tracker['high_risk_patterns_fixed']}")
        
        if stats['paragraphs_processed'] > 0:
            risk_distribution = (
                f"High: {stats['high_risk_detected']}, "
                f"Medium: {stats['medium_risk_detected']}, "
                f"Low: {stats['low_risk_detected']}"
            )
            print(f"📊 Risk distribution: {risk_distribution}")
        else:
            print(f"📊 Risk distribution: No risks detected")
        print(f"💰 Estimated AI cost: ${self.cost_tracker['estimated_cost_usd']:.6f} USD")
        print(f"💬 AI calls made: {self.cost_tracker['ai_calls']}" )

