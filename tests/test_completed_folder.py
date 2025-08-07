# test_completed_folder.py
"""
Test Completed Folder Functionality
Shows how the system saves original and paraphrased documents for comparison
"""

import os
import docx
from datetime import datetime
from pathlib import Path
import shutil

def create_sample_document():
    """Create a small sample document for testing"""
    print("📄 Creating sample document for testing...")
    
    # Create sample document
    doc = docx.Document()
    
    # Add sample academic content
    paragraphs = [
        "Penelitian ini bertujuan untuk menganalisis pengaruh teknologi informasi terhadap efisiensi kerja karyawan.",
        "Menurut Smith (2020), teknologi informasi memiliki peran penting dalam meningkatkan produktivitas organisasi.",
        "Hasil penelitian menunjukkan bahwa implementasi sistem informasi dapat meningkatkan efisiensi hingga 35%.",
        "Berdasarkan data yang diperoleh, terdapat korelasi positif antara penggunaan teknologi dan kinerja karyawan.",
        "Kesimpulan penelitian ini mengindikasikan bahwa investasi teknologi informasi memberikan dampak positif."
    ]
    
    for paragraph_text in paragraphs:
        doc.add_paragraph(paragraph_text)
    
    sample_path = "sample_test_document.docx"
    doc.save(sample_path)
    
    print(f"✅ Sample document created: {sample_path}")
    return sample_path

def demonstrate_completed_folder():
    """Demonstrate the completed folder functionality"""
    print("🚀 TESTING COMPLETED FOLDER FUNCTIONALITY")
    print("=" * 60)
    
    # Create sample document
    sample_path = create_sample_document()
    
    try:
        # Import and initialize our system
        from integrated_smart_paraphrase_system import IntegratedSmartParaphraseSystem
        
        print("\n🔧 Initializing system...")
        system = IntegratedSmartParaphraseSystem(mode='smart')
        
        # Create a paraphrased version manually for demo
        print("\n🤖 Creating paraphrased version...")
        
        # Read original document
        doc = docx.Document(sample_path)
        
        # Create paraphrased version (simple demo)
        paraphrased_doc = docx.Document()
        
        paraphrased_texts = [
            "Studi ini dimaksudkan untuk mengkaji dampak teknologi informasi pada efisiensi kerja pegawai.",
            "Berdasarkan Smith (2020), teknologi informasi memegang peranan krusial dalam meningkatkan produktivitas organisasi.", 
            "Temuan riset memperlihatkan bahwa penerapan sistem informasi mampu meningkatkan efisiensi sampai 35%.",
            "Sesuai data yang didapat, ada hubungan positif antara pemanfaatan teknologi dengan performa karyawan.",
            "Konklusi riset ini menunjukkan bahwa investasi teknologi informasi memberi dampak positif."
        ]
        
        for text in paraphrased_texts:
            para = paraphrased_doc.add_paragraph(text)
            # Highlight paraphrased text
            for run in para.runs:
                run.font.highlight_color = docx.enum.text.WD_COLOR_INDEX.YELLOW
        
        # Use the system's save function to demonstrate completed folder
        print("\n💾 Saving to completed folder...")
        output_info = system.save_paraphrased_document(paraphrased_doc, sample_path)
        
        if output_info:
            print("\n🎉 SUCCESS! Documents saved to completed folder:")
            print(f"   📁 Folder: {output_info['completed_folder']}")
            print(f"   📄 Original: {os.path.basename(output_info['original_copy'])}")
            print(f"   🤖 Paraphrased: {os.path.basename(output_info['paraphrased_file'])}")
            
            # Show completed folder contents
            print(f"\n📋 COMPLETED FOLDER CONTENTS:")
            completed_files = os.listdir("completed")
            for file in sorted(completed_files):
                if file.endswith('.docx'):
                    file_path = os.path.join("completed", file)
                    size = os.path.getsize(file_path) / 1024
                    mod_time = datetime.fromtimestamp(os.path.getmtime(file_path)).strftime("%H:%M:%S")
                    
                    if "ORIGINAL" in file:
                        print(f"   📄 {file}")
                        print(f"      📏 {size:.1f} KB | 🕒 {mod_time} | 🔖 Original Document")
                    else:
                        print(f"   🤖 {file}")
                        print(f"      📏 {size:.1f} KB | 🕒 {mod_time} | 🎯 Paraphrased Document")
                    print()
            
            print("🔍 COMPARISON GUIDE:")
            print("   1. Open both files in Word/LibreOffice")
            print("   2. Compare side-by-side to see differences")  
            print("   3. Paraphrased version has yellow highlights")
            print("   4. Original version shows source content")
        
        # Cleanup
        if os.path.exists(sample_path):
            os.remove(sample_path)
            
    except Exception as e:
        print(f"❌ Error: {e}")
        # Cleanup on error
        if os.path.exists(sample_path):
            os.remove(sample_path)
    
    print(f"\n" + "=" * 60)
    print("✅ COMPLETED FOLDER TEST FINISHED!")
    print("💡 Your main documents will be saved the same way")
    print("=" * 60)

def show_folder_structure():
    """Show the folder structure for clarity"""
    print("\n📁 FOLDER STRUCTURE AFTER PROCESSING:")
    print("-" * 50)
    print("📂 parafrase-system/")
    print("├── 📄 SKRIPSI  FAHRISAL FADLI-2.docx (original)")
    print("├── 📂 documents/ (input documents)")
    print("├── 📂 backups/ (automatic backups)")
    print("├── 📂 completed/ (comparison ready)")
    print("│   ├── 📄 SKRIPSI_ORIGINAL_20250807_123456.docx")  
    print("│   └── 🤖 SKRIPSI_paraphrased_20250807_123456.docx")
    print("├── 📂 reports/ (JSON analysis files)")
    print("└── 🤖 integrated_smart_paraphrase_system.py")
    print()
    print("🎯 WORKFLOW:")
    print("1. 🔍 System scans original document")
    print("2. 🌐 Checks plagiarism online")
    print("3. 🤖 Auto-paraphrases risky content")
    print("4. 💾 Saves BOTH original + paraphrased in completed/")
    print("5. 📊 Generates comparison reports")

if __name__ == "__main__":
    demonstrate_completed_folder()
    show_folder_structure()