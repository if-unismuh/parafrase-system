# quick_demo_system.py
"""
Quick Demo for Integrated System with Completed Folder
Demonstrates the complete workflow with smaller documents for testing
"""

import os
import json
import shutil
from datetime import datetime
from pathlib import Path

# Import our integrated system
from integrated_smart_paraphrase_system import IntegratedSmartParaphraseSystem

def demo_with_smaller_document():
    """Demo with one of the smaller paraphrased documents"""
    print("🚀 QUICK DEMO - INTEGRATED SMART PARAPHRASE SYSTEM")
    print("🎯 Testing with smaller document for demonstration")
    print("=" * 80)
    
    # Find smaller documents for demo
    demo_files = [
        "./paraphrased_documents/paraphrased_BAB I - 2025-08-05T085341.331.docx",
        "./paraphrased_documents/paraphrased_BAB II - 2025-08-05T085342.216.docx",
        "./paraphrased_documents/paraphrased_BAB III - 2025-08-05T085343.369.docx"
    ]
    
    # Find available demo file
    demo_file = None
    for file_path in demo_files:
        if os.path.exists(file_path):
            demo_file = file_path
            break
    
    if not demo_file:
        print("❌ No demo files found. Using main document...")
        demo_file = "SKRIPSI  FAHRISAL FADLI-2.docx"
    
    print(f"📄 Demo file: {os.path.basename(demo_file)}")
    print(f"📏 File size: {os.path.getsize(demo_file) / 1024:.1f} KB")
    
    # Initialize system
    print("\n🔧 Initializing system...")
    system = IntegratedSmartParaphraseSystem(mode='smart')
    
    # Demo: Analysis only first
    print(f"\n🔍 DEMO 1: ANALYSIS ONLY")
    print("-" * 40)
    
    try:
        # Just analysis
        analysis_results = system.analyze_document_comprehensive(demo_file)
        
        if analysis_results:
            print(f"✅ Analysis completed!")
            print(f"📊 Total paragraphs analyzed: {analysis_results['total_paragraphs']}")
            
            # Show sample recommendations
            if analysis_results.get('recommendations'):
                print(f"\n💡 Sample recommendations:")
                for rec in analysis_results['recommendations'][:2]:
                    print(f"   • {rec['description']}")
        
        # Demo: Complete processing
        print(f"\n🤖 DEMO 2: COMPLETE PROCESSING")
        print("-" * 40)
        
        # Complete processing on first few paragraphs only (for demo)
        print("📝 Note: Processing first 20 paragraphs for demo purposes...")
        
        # Create a smaller version for demo
        import docx
        doc = docx.Document(demo_file)
        
        # Take only first 20 paragraphs for quick demo
        demo_paragraphs = []
        for i, para in enumerate(doc.paragraphs[:20]):
            if para.text.strip() and len(para.text.split()) >= 10:
                demo_paragraphs.append(para)
        
        if demo_paragraphs:
            # Create temporary demo document
            demo_doc = docx.Document()
            for para in demo_paragraphs:
                demo_doc.add_paragraph(para.text)
            
            demo_temp_path = "temp_demo_document.docx"
            demo_doc.save(demo_temp_path)
            
            # Process the demo document
            final_report = system.process_document_complete(demo_temp_path, auto_paraphrase=True)
            
            if final_report:
                print(f"\n🎉 DEMO COMPLETED!")
                
                # Show results
                if final_report.get('paraphrase_results'):
                    paraphrase = final_report['paraphrase_results']
                    if paraphrase.get('output_files'):
                        output_files = paraphrase['output_files']
                        print(f"\n📁 CHECK 'completed/' FOLDER:")
                        print(f"   📄 Original: {os.path.basename(output_files['original_copy'])}")
                        print(f"   🤖 Paraphrased: {os.path.basename(output_files['paraphrased_file'])}")
                        print(f"   🔍 Compare the two files to see the differences!")
                        
                        # Show folder contents
                        print(f"\n📋 COMPLETED FOLDER CONTENTS:")
                        completed_files = os.listdir("completed")
                        for file in sorted(completed_files):
                            if file.endswith('.docx'):
                                file_path = os.path.join("completed", file)
                                size = os.path.getsize(file_path) / 1024
                                mod_time = datetime.fromtimestamp(os.path.getmtime(file_path)).strftime("%H:%M")
                                
                                if "ORIGINAL" in file:
                                    print(f"   📄 {file} ({size:.1f} KB) - {mod_time}")
                                else:
                                    print(f"   🤖 {file} ({size:.1f} KB) - {mod_time}")
            
            # Cleanup temp file
            if os.path.exists(demo_temp_path):
                os.remove(demo_temp_path)
        
    except Exception as e:
        print(f"❌ Demo error: {e}")
    
    print(f"\n" + "=" * 80)
    print("🎉 DEMO COMPLETED!")
    print("💡 Key features demonstrated:")
    print("   ✅ Auto-document detection")
    print("   ✅ Online plagiarism checking") 
    print("   ✅ Auto-paraphrasing")
    print("   ✅ Completed folder with original + paraphrased files")
    print("   ✅ Side-by-side comparison ready")
    print("=" * 80)


def show_system_structure():
    """Show the complete system file structure"""
    print("\n📁 SYSTEM FILE STRUCTURE:")
    print("-" * 40)
    
    structure = {
        "Input Documents": [
            "SKRIPSI  FAHRISAL FADLI-2.docx (main document)",
            "documents/ folder (additional docs)",
            "paraphrased_documents/ (existing processed)"
        ],
        "Core System Files": [
            "integrated_smart_paraphrase_system.py (main system)",
            "plagiarism_detector.py (online detection)",
            "smart_plagiarism_checker.py (pattern analysis)", 
            "ultimate_hybrid_paraphraser.py (AI paraphrasing)"
        ],
        "Output Folders": [
            "completed/ (original + paraphrased comparison)",
            "backups/ (automatic document backups)",
            "reports/ (JSON analysis reports)"
        ]
    }
    
    for category, files in structure.items():
        print(f"\n📂 {category}:")
        for file in files:
            print(f"   • {file}")


if __name__ == "__main__":
    demo_with_smaller_document()
    show_system_structure()