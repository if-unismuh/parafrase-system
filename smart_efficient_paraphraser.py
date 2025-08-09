# smart_efficient_paraphraser.py
"""
Smart Efficient Paraphraser - Cost-Effective System
Strategy: Local paraphrase first + Gemini AI quality checker for poor results
Features:
- Local synonym-based paraphrasing (primary method)
- AI quality assessment to detect ngawur/kacau sentences
- Gemini AI refinement only for problematic results
- Cost tracking and optimization
- Content protection (titles, headings, etc.)

Author: DevNoLife
Version: 1.0 - Efficient Edition
"""

import os
import json
import re
import time
from datetime import datetime
from pathlib import Path
import docx
from collections import defaultdict
import requests
from difflib import SequenceMatcher
import glob
import random

# DDGS imports
try:
    from ddgs import DDGS
    DDGS_AVAILABLE = True
    print("✅ DuckDuckGo Search (DDGS) imported successfully")
except ImportError:
    try:
        from duckduckgo_search import DDGS
        DDGS_AVAILABLE = True
        print("✅ DuckDuckGo Search (legacy) imported successfully")
    except ImportError:
        DDGS_AVAILABLE = False
        print("⚠️  DuckDuckGo Search not available. Install with: pip install ddgs")

# Load environment variables
try:
    from dotenv import load_dotenv
    load_dotenv()
    print("✅ Environment variables loaded from .env file")
except ImportError:
    print("⚠️  python-dotenv not found. Install with: pip install python-dotenv")
    print("💡 Using system environment variables...")

# Gemini AI imports
try:
    import google.generativeai as genai
    GEMINI_AVAILABLE = True
except ImportError:
    try:
        # Fallback to requests-based API
        GEMINI_AVAILABLE = "requests"
    except:
        GEMINI_AVAILABLE = False

class SmartEfficientParaphraser:
    def __init__(self, synonym_file='sinonim.json', gemini_api_key=None):
        print("🚀 Initializing Smart Efficient Paraphraser v1.0...")
        print("💡 Strategy: Local First + AI Quality Check")
        
        # API Configuration (only for quality checking)
        self.gemini_api_key = gemini_api_key or os.getenv('GEMINI_API_KEY')
        if not self.gemini_api_key:
            print("⚠️  GEMINI_API_KEY not found - AI quality check disabled")
            self.gemini_available = False
        else:
            self._initialize_gemini_client()
            self.gemini_available = True
        
        # Load enhanced synonym database
        self.synonyms, self.academic_words, self.pos_tags = self.load_synonyms(synonym_file)
        print(f"📚 Total synonym groups: {len(self.synonyms):,}")
        print(f"🎯 Academic word groups: {len(self.academic_words):,}")
        
        # Initialize Indonesian paraphrase patterns
        self._initialize_paraphrase_patterns()
        
        # Initialize DDGS search capability
        self.ddgs_available = DDGS_AVAILABLE
        if self.ddgs_available:
            self.ddgs = DDGS()
            print("🔍 DuckDuckGo Search initialized for content discovery")
        
        # Search configuration
        self.search_config = {
            'max_results': 5,
            'region': 'id-id',  # Indonesia region
            'language': 'id',   # Indonesian language
            'safesearch': 'moderate',
            'min_content_length': 100,  # Minimum content length to consider
            'max_content_length': 2000  # Maximum content length to use
        }
        
        # Content protection patterns
        self.protection_patterns = {
            'titles': [
                r'^BAB\s+[IVX]+.*',  # Chapter titles
                r'^CHAPTER\s+\d+.*',
                r'^Bab\s+\d+.*',
                r'^\d+\.\s*[A-Z][^.]*$',  # Numbered titles without ending punctuation
                r'^[A-Z][A-Z\s]{5,}$',  # ALL CAPS titles
            ],
            'author_names': [
                r'Author:\s*(.+)',
                r'Penulis:\s*(.+)', 
                r'Oleh:\s*(.+)',
                r'^([A-Z][a-z]+\s+[A-Z][a-z]+(\s+[A-Z][a-z]+)*)$',  # Proper names
            ],
            'headings': [
                r'^\d+\.\d+.*',  # 1.1 Heading format
                r'^\d+\.\d+\.\d+.*',  # 1.1.1 Subheading format
                r'^[A-Z]\.\s+[A-Z].*',  # A. Heading format
                r'^ABSTRAK$|^ABSTRACT$|^PENDAHULUAN$|^KESIMPULAN$',
                r'^DAFTAR\s+(ISI|PUSTAKA|GAMBAR|TABEL)',
            ],
            'citations': [
                r'\(.+?\d{4}.+?\)',  # (Author, 2024)
                r'\[.+?\]',  # [1], [Author, 2024]
                r'(Menurut|According to).+?\(\d{4}\)',
            ]
        }
        
        # Quality assessment criteria
        self.quality_criteria = {
            'min_readability_score': 60,  # Minimum readability
            'max_repetition_ratio': 0.3,  # Max word repetition
            'min_word_variety': 0.4,      # Minimum word diversity
            'max_weird_combinations': 2   # Max strange word combinations
        }
        
        # Cost tracking
        self.cost_tracker = {
            'local_paraphrases': 0,
            'ai_quality_checks': 0,
            'ai_refinements': 0,
            'total_api_calls': 0,
            'protected_content': 0,
            'search_queries': 0,
            'context_enhanced_paraphrases': 0,
            'cost_savings_percent': 0
        }
        
        # Processing statistics
        self.stats = {
            'start_time': datetime.now(),
            'quality_improvements': 0,
            'perfect_local_results': 0,
            'ai_refinement_needed': 0
        }
        
        print("✅ Smart Efficient Paraphraser initialized!")

    def _initialize_gemini_client(self):
        """Initialize Gemini AI client for quality checking only"""
        if GEMINI_AVAILABLE == True:
            try:
                genai.configure(api_key=self.gemini_api_key)
                self.gemini_client = genai.GenerativeModel('gemini-1.5-flash')  # Use lighter model
                self.api_method = 'genai'
                print(f"✅ Gemini AI quality checker initialized")
            except Exception as e:
                print(f"⚠️  Gemini initialization failed: {e}")
                self.gemini_available = False
        else:
            self.gemini_available = False

    def load_synonyms(self, filename):
        """Load enhanced synonym database from JSON file"""
        if not filename or not os.path.exists(filename):
            # Try components directory
            components_path = os.path.join('components', filename)
            if os.path.exists(components_path):
                filename = components_path
            else:
                print("⚠️  No synonym file found")
                return {}, {}, set()
        
        try:
            print(f"📁 Loading comprehensive synonym database: {filename}")
            with open(filename, 'r', encoding='utf-8') as f:
                data = json.load(f)
            
            synonyms = defaultdict(list)
            academic_words = defaultdict(list)
            pos_tags = set()
            
            # Academic word categories for better filtering
            academic_categories = {
                'research': ['penelitian', 'analisis', 'studi', 'riset', 'investigasi', 'eksplorasi', 'kajian', 'telaah'],
                'method': ['menggunakan', 'menerapkan', 'memanfaatkan', 'mengaplikasikan', 'mengimplementasikan'],
                'result': ['menunjukkan', 'mengindikasikan', 'membuktikan', 'menghasilkan', 'menciptakan'],
                'academic_verbs': ['bertujuan', 'bermaksud', 'dimaksudkan', 'ditujukan', 'berupaya'],
                'academic_nouns': ['pendekatan', 'metode', 'teknik', 'prosedur', 'sistem', 'struktur']
            }
            
            processed_count = 0
            academic_count = 0
            
            for word, info in data.items():
                if isinstance(info, dict) and 'sinonim' in info:
                    synonym_list = info['sinonim']
                    pos_tag = info.get('tag', 'unknown')
                    pos_tags.add(pos_tag)
                    
                    if isinstance(synonym_list, list) and len(synonym_list) > 0:
                        # Clean synonyms and remove unwanted entries
                        clean_synonyms = []
                        for synonym in synonym_list:
                            clean_syn = synonym.lower().strip()
                            # Filter out very short words, numbers, and weird entries
                            if (len(clean_syn) >= 3 and 
                                not clean_syn.isdigit() and 
                                not any(char in clean_syn for char in ['(', ')', '[', ']', 'cak', 'v ', 'n ', 'a '])):
                                clean_synonyms.append(clean_syn)
                        
                        if clean_synonyms:
                            clean_word = word.lower().strip()
                            all_words = [clean_word] + clean_synonyms
                            
                            # Build synonym mapping
                            for w in all_words:
                                synonyms[w] = [s for s in all_words if s != w]
                            
                            # Check if it's an academic word
                            is_academic = False
                            for category, words in academic_categories.items():
                                if clean_word in words or any(aw in clean_word for aw in words):
                                    academic_words[clean_word] = clean_synonyms[:5]  # Limit to best 5 synonyms
                                    is_academic = True
                                    academic_count += 1
                                    break
                            
                            processed_count += 1
            
            print(f"✅ Loaded {processed_count:,} words with synonyms")
            print(f"🎯 Academic words identified: {academic_count:,}")
            print(f"🏷️ POS tags found: {', '.join(sorted(pos_tags))}")
            
            return dict(synonyms), dict(academic_words), pos_tags
            
        except Exception as e:
            print(f"❌ Error loading synonyms: {e}")
            return {}, {}, set()

    def _initialize_paraphrase_patterns(self):
        """Initialize Indonesian paraphrase patterns"""
        # Academic phrase replacements
        self.phrase_replacements = {
            'penelitian ini': ['studi ini', 'kajian ini', 'riset ini', 'investigasi ini'],
            'hasil penelitian': ['temuan riset', 'output kajian', 'hasil studi', 'temuan penelitian'],
            'bertujuan untuk': ['dimaksudkan untuk', 'ditujukan untuk', 'bermaksud untuk', 'berupaya untuk'],
            'dapat': ['mampu', 'bisa', 'sanggup'],
            'menggunakan': ['memakai', 'memanfaatkan', 'mempergunakan', 'mengaplikasikan'],
            'menunjukkan': ['memperlihatkan', 'mengindikasikan', 'mendemonstrasikan', 'membuktikan'],
            'mengurangi': ['menurunkan', 'meminimalisir', 'meminimalisasi', 'memperkecil'],
            'meningkatkan': ['menaikkan', 'memperbesar', 'mengoptimalkan', 'memaksimalkan'],
            'tingkat': ['level', 'taraf', 'derajat', 'kadar'],
            'sistem': ['metode', 'cara', 'teknik', 'prosedur', 'mekanisme'],
            'pendekatan': ['metode', 'cara', 'teknik', 'strategi'],
            'menghasilkan': ['menciptakan', 'membuat', 'memproduksi', 'melahirkan'],
            'berbeda': ['tidak sama', 'berlainan', 'bervariasi', 'beragam'],
            'struktur': ['susunan', 'bentuk', 'format', 'konstruksi'],
            'mempertahankan': ['menjaga', 'memelihara', 'melestarikan']
        }
        
        # Sentence structure patterns
        self.structure_patterns = [
            {'pattern': r'^(.+?) adalah (.+?)$', 'replacement': r'\2 merupakan \1'},
            {'pattern': r'^Dengan (.+?), (.+?)$', 'replacement': r'Melalui \1, \2'},
            {'pattern': r'^Hal ini (.+?) karena (.+?)$', 'replacement': r'Kondisi tersebut \1 disebabkan \2'},
            {'pattern': r'^Oleh karena itu, (.+?)$', 'replacement': r'Dengan demikian, \1'},
            {'pattern': r'^Selain itu, (.+?)$', 'replacement': r'Di samping itu, \1'}
        ]

    def is_protected_content(self, text):
        """Check if content should be protected from paraphrasing"""
        text = text.strip()
        
        for category, patterns in self.protection_patterns.items():
            for pattern in patterns:
                if re.match(pattern, text, re.IGNORECASE):
                    return True, category
        
        return False, None

    def search_related_content(self, query, search_type='text'):
        """Search for related content using DuckDuckGo"""
        if not self.ddgs_available:
            print("⚠️  DuckDuckGo Search not available")
            return []
        
        try:
            print(f"🔍 Searching for: {query[:50]}...")
            
            search_results = []
            
            if search_type == 'text':
                # Search for text content
                results = self.ddgs.text(
                    query,
                    region=self.search_config['region'],
                    safesearch=self.search_config['safesearch'],
                    max_results=self.search_config['max_results']
                )
                
                for result in results:
                    if result and 'body' in result and result['body']:
                        content_length = len(result['body'])
                        if (self.search_config['min_content_length'] <= content_length <= 
                            self.search_config['max_content_length']):
                            
                            search_results.append({
                                'title': result.get('title', 'No title'),
                                'content': result['body'],
                                'url': result.get('href', ''),
                                'content_length': content_length,
                                'relevance_score': self._calculate_relevance(query, result['body'])
                            })
            
            elif search_type == 'news':
                # Search for news content
                results = self.ddgs.news(
                    query,
                    region=self.search_config['region'],
                    safesearch=self.search_config['safesearch'],
                    max_results=self.search_config['max_results']
                )
                
                for result in results:
                    if result and 'body' in result and result['body']:
                        content_length = len(result['body'])
                        if (self.search_config['min_content_length'] <= content_length <= 
                            self.search_config['max_content_length']):
                            
                            search_results.append({
                                'title': result.get('title', 'No title'),
                                'content': result['body'],
                                'url': result.get('url', ''),
                                'date': result.get('date', ''),
                                'content_length': content_length,
                                'relevance_score': self._calculate_relevance(query, result['body'])
                            })
            
            # Sort by relevance score
            search_results.sort(key=lambda x: x['relevance_score'], reverse=True)
            
            print(f"✅ Found {len(search_results)} relevant results")
            return search_results
            
        except Exception as e:
            print(f"❌ Search error: {e}")
            return []

    def _calculate_relevance(self, query, content):
        """Calculate relevance score between query and content"""
        if not query or not content:
            return 0
        
        query_words = set(re.findall(r'\w+', query.lower()))
        content_words = set(re.findall(r'\w+', content.lower()))
        
        if not query_words or not content_words:
            return 0
        
        # Calculate word overlap
        overlap = len(query_words.intersection(content_words))
        relevance = (overlap / len(query_words)) * 100
        
        return round(relevance, 2)

    def extract_keywords_from_text(self, text):
        """Extract key terms from text for searching"""
        if not text or not text.strip():
            return []
        
        # Remove common stop words in Indonesian
        stop_words = {
            'dan', 'atau', 'yang', 'ini', 'itu', 'pada', 'untuk', 'dengan', 'dalam', 'dari', 
            'ke', 'di', 'adalah', 'akan', 'dapat', 'telah', 'sudah', 'belum', 'masih',
            'hanya', 'juga', 'tidak', 'bukan', 'ada', 'menjadi', 'membuat', 'seperti'
        }
        
        # Extract words longer than 3 characters
        words = re.findall(r'\w+', text.lower())
        keywords = [word for word in words if len(word) > 3 and word not in stop_words]
        
        # Count frequency and get most common
        word_freq = defaultdict(int)
        for word in keywords:
            word_freq[word] += 1
        
        # Sort by frequency and take top keywords
        sorted_words = sorted(word_freq.items(), key=lambda x: x[1], reverse=True)
        top_keywords = [word for word, freq in sorted_words[:5]]  # Top 5 keywords
        
        return top_keywords

    def search_and_analyze_content(self, original_text):
        """Search for related content and analyze for paraphrasing potential"""
        if not self.ddgs_available:
            return None
        
        # Extract keywords from original text
        keywords = self.extract_keywords_from_text(original_text)
        if not keywords:
            print("⚠️  No keywords found in text")
            return None
        
        search_query = ' '.join(keywords[:3])  # Use top 3 keywords
        print(f"🎯 Search query: {search_query}")
        
        # Search for content
        search_results = self.search_related_content(search_query, 'text')
        self.cost_tracker['search_queries'] += 1
        
        if not search_results:
            print("❌ No search results found")
            return None
        
        # Analyze and select best content for paraphrasing context
        best_content = self._select_best_content_for_paraphrasing(
            original_text, 
            search_results
        )
        
        return best_content

    def _select_best_content_for_paraphrasing(self, original_text, search_results):
        """Select the best search result for paraphrasing context"""
        if not search_results:
            return None
        
        scored_results = []
        
        for result in search_results:
            content = result['content']
            
            # Calculate various scores
            relevance = result['relevance_score']
            length_score = min(100, (len(content) / 500) * 100)  # Prefer moderate length
            similarity = self.calculate_similarity(original_text, content)
            
            # We want high relevance, good length, but not too similar (to avoid plagiarism)
            final_score = (relevance * 0.5) + (length_score * 0.2) + ((100 - similarity) * 0.3)
            
            scored_results.append({
                'result': result,
                'score': final_score,
                'similarity_to_original': similarity
            })
        
        # Sort by score and select best
        scored_results.sort(key=lambda x: x['score'], reverse=True)
        best_result = scored_results[0]
        
        print(f"📊 Best content selected:")
        print(f"   Relevance: {best_result['result']['relevance_score']}%")
        print(f"   Similarity to original: {best_result['similarity_to_original']}%")
        print(f"   Final score: {round(best_result['score'], 1)}")
        
        return best_result['result']

    def local_paraphrase(self, text):
        """Enhanced local paraphrasing using comprehensive synonym database"""
        if not text.strip():
            return text, 0
        
        paraphrased = text
        changes_count = 0
        
        # Step 1: Replace academic phrases (higher priority)
        for phrase, replacements in self.phrase_replacements.items():
            if phrase in paraphrased.lower():
                replacement = random.choice(replacements)
                paraphrased = re.sub(
                    re.escape(phrase), 
                    replacement, 
                    paraphrased, 
                    flags=re.IGNORECASE
                )
                changes_count += 1
        
        # Step 2: Enhanced synonym replacement with academic word priority
        words = paraphrased.split()
        new_words = []
        
        for word in words:
            clean_word = re.sub(r'[^\w]', '', word.lower())
            word_replaced = False
            
            # Priority 1: Academic words (higher replacement chance)
            if clean_word in self.academic_words:
                academic_synonyms = self.academic_words[clean_word]
                if academic_synonyms and random.random() < 0.6:  # 60% chance for academic words
                    new_word = random.choice(academic_synonyms)
                    # Preserve original case
                    if word.isupper():
                        new_word = new_word.upper()
                    elif word.istitle():
                        new_word = new_word.capitalize()
                    
                    # Replace while preserving punctuation
                    word_with_punct = re.sub(re.escape(clean_word), new_word, word, flags=re.IGNORECASE)
                    new_words.append(word_with_punct)
                    changes_count += 1
                    word_replaced = True
            
            # Priority 2: Regular synonyms (if not academic word)
            if not word_replaced and clean_word in self.synonyms:
                synonyms = self.synonyms[clean_word]
                # Filter for quality synonyms (length >= 3, not too obscure)
                quality_synonyms = [s for s in synonyms if len(s) >= 3 and len(s) <= 15]
                
                if quality_synonyms and random.random() < 0.35:  # 35% chance for regular words
                    new_word = random.choice(quality_synonyms[:10])  # Use top 10 synonyms only
                    # Preserve original case
                    if word.isupper():
                        new_word = new_word.upper()
                    elif word.istitle():
                        new_word = new_word.capitalize()
                    
                    # Replace while preserving punctuation
                    word_with_punct = re.sub(re.escape(clean_word), new_word, word, flags=re.IGNORECASE)
                    new_words.append(word_with_punct)
                    changes_count += 1
                    word_replaced = True
            
            # Keep original word if not replaced
            if not word_replaced:
                new_words.append(word)
        
        paraphrased = ' '.join(new_words)
        
        # Step 3: Apply sentence structure patterns
        for pattern_info in self.structure_patterns:
            if re.match(pattern_info['pattern'], paraphrased):
                new_structure = re.sub(
                    pattern_info['pattern'], 
                    pattern_info['replacement'], 
                    paraphrased
                )
                if new_structure != paraphrased:
                    paraphrased = new_structure
                    changes_count += 1
                break
        
        return paraphrased, changes_count

    def enhanced_local_paraphrase(self, text, search_context=None):
        """Enhanced local paraphrasing that can use search context for better synonyms"""
        if not text.strip():
            return text, 0
        
        # Start with standard local paraphrasing
        paraphrased, changes_count = self.local_paraphrase(text)
        
        # If we have search context, try to improve with context-aware synonyms
        if search_context and search_context.get('content'):
            context_synonyms = self._extract_context_synonyms(search_context['content'])
            if context_synonyms:
                enhanced_paraphrased, additional_changes = self._apply_context_synonyms(
                    paraphrased, context_synonyms
                )
                paraphrased = enhanced_paraphrased
                changes_count += additional_changes
                self.cost_tracker['context_enhanced_paraphrases'] += 1
                print(f"📈 Applied {additional_changes} context-enhanced changes")
        
        return paraphrased, changes_count

    def _extract_context_synonyms(self, context_content):
        """Extract useful synonyms from search context"""
        if not context_content:
            return {}
        
        # Extract words from context that might serve as alternative expressions
        context_words = re.findall(r'\w+', context_content.lower())
        
        # Filter for meaningful words (length > 3, not common stop words)
        meaningful_words = [
            word for word in context_words 
            if len(word) > 3 and word not in {
                'dengan', 'untuk', 'dalam', 'adalah', 'yang', 'akan', 'dapat', 'telah',
                'sudah', 'hanya', 'juga', 'tidak', 'bukan', 'menjadi', 'membuat'
            }
        ]
        
        # Count frequency and get common context words
        word_freq = defaultdict(int)
        for word in meaningful_words:
            word_freq[word] += 1
        
        # Get words that appear frequently (potential good alternatives)
        context_synonyms = {}
        for word, freq in word_freq.items():
            if freq >= 2 and word in self.synonyms:  # Appears multiple times and has synonyms
                context_synonyms[word] = self.synonyms[word][:3]  # Top 3 synonyms
        
        return context_synonyms

    def _apply_context_synonyms(self, text, context_synonyms):
        """Apply context-aware synonyms to improve paraphrasing"""
        if not context_synonyms:
            return text, 0
        
        enhanced_text = text
        changes_count = 0
        
        words = enhanced_text.split()
        new_words = []
        
        for word in words:
            clean_word = re.sub(r'[^\w]', '', word.lower())
            
            # Check if word can be replaced with context-aware synonym
            if clean_word in context_synonyms:
                context_syns = context_synonyms[clean_word]
                if context_syns and random.random() < 0.4:  # 40% chance for context synonyms
                    new_word = random.choice(context_syns)
                    # Preserve original case
                    if word.isupper():
                        new_word = new_word.upper()
                    elif word.istitle():
                        new_word = new_word.capitalize()
                    
                    # Replace while preserving punctuation
                    word_with_punct = re.sub(re.escape(clean_word), new_word, word, flags=re.IGNORECASE)
                    new_words.append(word_with_punct)
                    changes_count += 1
                else:
                    new_words.append(word)
            else:
                new_words.append(word)
        
        enhanced_text = ' '.join(new_words)
        return enhanced_text, changes_count

    def assess_quality(self, original, paraphrased):
        """Assess quality of paraphrased text to detect if AI refinement needed"""
        issues = []
        score = 100
        
        if not paraphrased or not paraphrased.strip():
            return 0, ["Empty result"]
        
        # Check 1: Readability (avoid overly complex/weird sentences)
        words = paraphrased.split()
        if not words:
            return 0, ["No words found"]
        
        avg_word_length = sum(len(word) for word in words) / len(words)
        if avg_word_length > 8:  # Too many long words
            score -= 20
            issues.append("Overly complex words")
        
        # Check 2: Word repetition
        word_freq = defaultdict(int)
        for word in words:
            clean_word = re.sub(r'[^\w]', '', word.lower())
            if len(clean_word) > 3:  # Only count significant words
                word_freq[clean_word] += 1
        
        if word_freq:
            max_repetition = max(word_freq.values())
            repetition_ratio = max_repetition / len(words)
            if repetition_ratio > self.quality_criteria['max_repetition_ratio']:
                score -= 25
                issues.append("Too much word repetition")
        
        # Check 3: Word variety
        unique_words = len(set(word.lower() for word in words if len(word) > 3))
        total_words = len([word for word in words if len(word) > 3])
        if total_words > 0:
            variety_ratio = unique_words / total_words
            if variety_ratio < self.quality_criteria['min_word_variety']:
                score -= 20
                issues.append("Low word variety")
        
        # Check 4: Detect weird combinations (simple heuristic)
        weird_patterns = [
            r'\b(\w+)\s+\1\b',  # Same word repeated
            r'\bmenjadi adalah\b',  # Weird combinations
            r'\badalah merupakan\b',
            r'\bsangat sekali\b',
            r'\bmembuat menciptakan\b'
        ]
        
        weird_count = 0
        for pattern in weird_patterns:
            if re.search(pattern, paraphrased.lower()):
                weird_count += 1
        
        if weird_count > self.quality_criteria['max_weird_combinations']:
            score -= 30
            issues.append(f"Weird word combinations detected ({weird_count})")
        
        # Check 5: Grammar basic checks
        if re.search(r'\b(di|ke|dari)\s+(di|ke|dari)\b', paraphrased):
            score -= 15
            issues.append("Grammar issues detected")
        
        # Check 6: Sentence completeness
        if not paraphrased.strip().endswith(('.', '!', '?')):
            if original.strip().endswith(('.', '!', '?')):
                score -= 10
                issues.append("Incomplete sentence")
        
        return max(0, score), issues

    def refine_with_ai(self, original_text, poor_paraphrase, quality_issues):
        """Use AI to refine poor quality paraphrased text"""
        if not self.gemini_available:
            print("⚠️  AI refinement not available")
            return poor_paraphrase
        
        prompt = f"""Perbaiki hasil parafrase bahasa Indonesia berikut yang memiliki masalah kualitas.

TEKS ASLI:
"{original_text}"

HASIL PARAFRASE BERMASALAH:
"{poor_paraphrase}"

MASALAH YANG TERDETEKSI:
{', '.join(quality_issues)}

TUGAS:
Perbaiki hasil parafrase agar:
✅ Natural dan mudah dibaca
✅ Tata bahasa Indonesia yang benar
✅ Tidak ada pengulangan kata berlebihan
✅ Kombinasi kata yang wajar
✅ Mempertahankan makna asli
✅ Gaya bahasa akademis

HANYA berikan hasil perbaikan tanpa penjelasan tambahan:"""

        try:
            if self.api_method == 'genai':
                response = self.gemini_client.generate_content(
                    prompt,
                    generation_config=genai.types.GenerationConfig(
                        temperature=0.3,  # Lower temperature for more focused refinement
                        max_output_tokens=1000,
                        top_p=0.9
                    )
                )
                
                if response and response.text:
                    self.cost_tracker['total_api_calls'] += 1
                    self.cost_tracker['ai_refinements'] += 1
                    return response.text.strip()
        
        except Exception as e:
            print(f"⚠️  AI refinement failed: {e}")
        
        return poor_paraphrase

    def calculate_similarity(self, text1, text2):
        """Calculate similarity percentage between two texts"""
        if not text1 or not text2:
            return 100.0 if not text1 and not text2 else 0.0
        
        # Word-based Jaccard similarity
        words1 = set(re.findall(r'\w+', text1.lower()))
        words2 = set(re.findall(r'\w+', text2.lower()))
        
        if not words1 and not words2:
            return 100.0
        if not words1 or not words2:
            return 0.0
        
        jaccard_sim = len(words1.intersection(words2)) / len(words1.union(words2)) * 100
        
        # SequenceMatcher for structural similarity
        sequence_sim = SequenceMatcher(None, text1.lower(), text2.lower()).ratio() * 100
        
        # Weighted combination
        final_similarity = (jaccard_sim * 0.6) + (sequence_sim * 0.4)
        
        return round(final_similarity, 2)

    def paraphrase_text(self, text, use_search=True):
        """Main paraphrasing method with smart AI usage and optional search integration"""
        
        # Check if content should be protected
        is_protected, category = self.is_protected_content(text)
        if is_protected:
            print(f"🛡️  Protected content: {text[:50]}...")
            self.cost_tracker['protected_content'] += 1
            return {
                'paraphrase': text,
                'original': text,
                'similarity': 100.0,
                'method': 'protected_content',
                'category': category,
                'status': 'PROTECTED'
            }
        
        # Optional: Search for related content first for context
        search_context = None
        if use_search and self.ddgs_available:
            print(f"🔍 Searching for context: {text[:50]}...")
            search_context = self.search_and_analyze_content(text)
            if search_context:
                print(f"✅ Found contextual information from: {search_context['title'][:30]}...")
            else:
                print("ℹ️  No relevant context found, proceeding with standard paraphrasing")
        
        # Step 1: Enhanced local paraphrasing (with optional context)
        local_result, changes_count = self.enhanced_local_paraphrase(text, search_context)
        self.cost_tracker['local_paraphrases'] += 1
        
        # Step 2: Quality assessment
        quality_score, quality_issues = self.assess_quality(text, local_result)
        
        # Step 3: Decide if AI refinement is needed
        if quality_score >= self.quality_criteria['min_readability_score']:
            # Good enough, no AI needed
            similarity = self.calculate_similarity(text, local_result)
            self.stats['perfect_local_results'] += 1
            
            method_used = 'local_with_search_context' if search_context else 'local_only'
            return {
                'paraphrase': local_result,
                'original': text,
                'similarity': similarity,
                'plagiarism_reduction': 100 - similarity,
                'method': method_used,
                'quality_score': quality_score,
                'changes_made': changes_count,
                'search_context': search_context['title'] if search_context else None,
                'status': 'HIGH_QUALITY_LOCAL'
            }
        
        else:
            # Quality issues detected, use AI refinement
            print(f"⚠️  Quality issues detected (score: {quality_score}): {quality_issues}")
            refined_result = self.refine_with_ai(text, local_result, quality_issues)
            
            # Re-assess refined quality
            final_quality_score, _ = self.assess_quality(text, refined_result)
            similarity = self.calculate_similarity(text, refined_result)
            
            self.stats['ai_refinement_needed'] += 1
            self.stats['quality_improvements'] += (final_quality_score - quality_score)
            
            method_used = 'local_plus_ai_refinement_with_search' if search_context else 'local_plus_ai_refinement'
            return {
                'paraphrase': refined_result,
                'original': text,
                'similarity': similarity,
                'plagiarism_reduction': 100 - similarity,
                'method': method_used,
                'quality_score': final_quality_score,
                'original_quality_score': quality_score,
                'quality_issues': quality_issues,
                'changes_made': changes_count,
                'search_context': search_context['title'] if search_context else None,
                'status': 'AI_REFINED'
            }

    def process_document(self, input_path, output_path=None, use_search=True):
        """Process entire document with smart efficient paraphrasing"""
        print(f"🎯 Processing document: {input_path}")
        search_status = "enabled" if use_search and self.ddgs_available else "disabled"
        print(f"💡 Strategy: Local paraphrase + AI quality check + Search context ({search_status})")
        
        if not os.path.exists(input_path):
            print(f"❌ File not found: {input_path}")
            return None
        
        # Generate output path if not provided
        if not output_path:
            base_name = Path(input_path).stem
            ext = Path(input_path).suffix
            completed_dir = "completed"
            os.makedirs(completed_dir, exist_ok=True)
            output_path = os.path.join(completed_dir, f"{base_name}_efficient_paraphrased{ext}")
        
        try:
            doc = docx.Document(input_path)
            results = []
            total_paragraphs = len([p for p in doc.paragraphs if p.text.strip()])
            
            print(f"📄 Found {total_paragraphs} paragraphs to process")
            
            ai_used_count = 0
            local_only_count = 0
            
            for i, paragraph in enumerate(doc.paragraphs):
                if not paragraph.text.strip():
                    continue
                
                print(f"🔄 Processing paragraph {i+1}/{total_paragraphs}")
                
                result = self.paraphrase_text(paragraph.text.strip(), use_search=use_search)
                results.append(result)
                
                # Count method usage
                if 'ai' in result['method']:
                    ai_used_count += 1
                    print(f"🤖 AI refinement used")
                else:
                    local_only_count += 1
                    print(f"✅ Local result sufficient")
                
                # Update paragraph
                if result['method'] != 'protected_content':
                    paragraph.text = result['paraphrase']
                
                time.sleep(0.1)  # Small delay
            
            # Calculate cost savings
            total_processed = local_only_count + ai_used_count
            if total_processed > 0:
                savings_percent = (local_only_count / total_processed) * 100
                self.cost_tracker['cost_savings_percent'] = round(savings_percent, 2)
            else:
                self.cost_tracker['cost_savings_percent'] = 0
            
            # Save processed document
            doc.save(output_path)
            print(f"✅ Document saved: {output_path}")
            
            # Report efficiency
            print(f"\n📊 Efficiency Report:")
            if total_processed > 0:
                print(f"   Local only: {local_only_count}/{total_processed} ({round((local_only_count/total_processed)*100, 1)}%)")
                print(f"   AI refinement: {ai_used_count}/{total_processed} ({round((ai_used_count/total_processed)*100, 1)}%)")
                print(f"   🔍 Search queries: {self.cost_tracker['search_queries']}")
                print(f"   📈 Context-enhanced: {self.cost_tracker['context_enhanced_paraphrases']}")
                print(f"   💰 Cost savings: {self.cost_tracker['cost_savings_percent']}%")
            else:
                print(f"   No paragraphs processed in this document")
            
            # Generate report
            report = self.generate_report(results, input_path, output_path)
            
            return {
                'output_path': output_path,
                'results': results,
                'report': report,
                'efficiency_stats': {
                    'local_only': local_only_count,
                    'ai_refined': ai_used_count,
                    'cost_savings_percent': self.cost_tracker['cost_savings_percent']
                }
            }
            
        except Exception as e:
            print(f"❌ Error processing document: {e}")
            return None

    def process_all_documents(self, input_dir="documents", output_dir="completed", use_search=True):
        """Process all documents with efficient strategy"""
        print(f"🚀 Smart Efficient Processing from: {input_dir}")
        search_status = "enabled" if use_search and self.ddgs_available else "disabled"
        print(f"💡 Strategy: Local first + AI quality check + Search context ({search_status})")
        
        os.makedirs(output_dir, exist_ok=True)
        
        # Find documents
        doc_patterns = ["*.docx", "*.doc"]
        doc_files = []
        
        for pattern in doc_patterns:
            doc_files.extend(glob.glob(os.path.join(input_dir, pattern)))
        
        if not doc_files:
            print(f"❌ No documents found in {input_dir}")
            return []
        
        print(f"📄 Found {len(doc_files)} documents to process")
        
        results = []
        total_local_only = 0
        total_ai_refined = 0
        
        for i, doc_path in enumerate(doc_files, 1):
            print(f"\n🔄 Processing document {i}/{len(doc_files)}: {Path(doc_path).name}")
            
            base_name = Path(doc_path).stem
            ext = Path(doc_path).suffix
            output_path = os.path.join(output_dir, f"{base_name}_efficient{ext}")
            
            result = self.process_document(doc_path, output_path, use_search=use_search)
            
            if result:
                results.append(result)
                total_local_only += result['efficiency_stats']['local_only']
                total_ai_refined += result['efficiency_stats']['ai_refined']
                print(f"✅ Completed: {Path(output_path).name}")
            
            if i < len(doc_files):
                print("⏳ Brief pause...")
                time.sleep(2)
        
        # Overall efficiency report
        total_processed = total_local_only + total_ai_refined
        overall_savings = (total_local_only / total_processed * 100) if total_processed > 0 else 0
        
        print(f"\n🎉 Batch Processing Complete!")
        print(f"📊 Overall Efficiency:")
        print(f"   Total paragraphs: {total_processed}")
        
        if total_processed > 0:
            print(f"   Local only: {total_local_only} ({round((total_local_only/total_processed)*100, 1)}%)")
            print(f"   AI refined: {total_ai_refined} ({round((total_ai_refined/total_processed)*100, 1)}%)")
            print(f"   🔍 Total search queries: {self.cost_tracker['search_queries']}")
            print(f"   📈 Total context-enhanced: {self.cost_tracker['context_enhanced_paraphrases']}")
            print(f"   💰 Overall cost savings: {round(overall_savings, 1)}%")
        else:
            print("   No paragraphs processed - check document content")
        
        return results

    def generate_report(self, results, input_path, output_path):
        """Generate processing report"""
        report = {
            'processing_info': {
                'input_file': input_path,
                'output_file': output_path,
                'timestamp': datetime.now().isoformat(),
                'processing_time': str(datetime.now() - self.stats['start_time'])
            },
            'efficiency_summary': {
                'total_paragraphs': len(results),
                'local_only_success': sum(1 for r in results if 'local' in r.get('method', '')),
                'ai_refinement_used': sum(1 for r in results if 'ai' in r.get('method', '')),
                'protected_content': sum(1 for r in results if r.get('method') == 'protected_content'),
                'cost_savings_percent': self.cost_tracker['cost_savings_percent']
            },
            'quality_metrics': {
                'average_similarity_reduction': round(
                    sum(r.get('plagiarism_reduction', 0) for r in results) / len(results), 2
                ) if results else 0,
                'average_quality_score': round(
                    sum(r.get('quality_score', 0) for r in results if 'quality_score' in r) /
                    len([r for r in results if 'quality_score' in r]), 2
                ) if any('quality_score' in r for r in results) else 0
            },
            'cost_tracking': self.cost_tracker,
            'processing_stats': {
                'start_time': self.stats['start_time'].isoformat(),
                'quality_improvements': self.stats['quality_improvements'],
                'perfect_local_results': self.stats['perfect_local_results'],
                'ai_refinement_needed': self.stats['ai_refinement_needed']
            }
        }
        
        # Save report
        report_path = f"report_efficient_{Path(output_path).stem}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json"
        with open(report_path, 'w', encoding='utf-8') as f:
            json.dump(report, f, indent=2, ensure_ascii=False)
        
        return report

def main():
    """Main processing function"""
    try:
        print("🚀 Smart Efficient Paraphraser System with DDGS")
        print("💡 Strategy: Search Context + Local First + AI Quality Check")
        print("=" * 60)
        
        paraphraser = SmartEfficientParaphraser(synonym_file='sinonim.json')
        
        # Process all documents with search enabled by default
        results = paraphraser.process_all_documents(
            input_dir="documents",
            output_dir="completed",
            use_search=True  # Enable search by default
        )
        
        if results:
            print("\n🎉 All documents processed with smart efficiency!")
        else:
            print("❌ No documents were processed successfully")
        
    except Exception as e:
        print(f"❌ Processing error: {e}")
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    main()