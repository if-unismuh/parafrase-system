#!/usr/bin/env python3
"""
Test script untuk memproses core chapters (BAB I - BAB V) saja
"""

from full_gemini_paraphraser import FullGeminiParaphraser
import os

def test_core_chapters():
    """Test processing core chapters only"""
    print("🧪 Testing Core Chapters Processing")
    print("=" * 50)
    
    # Initialize paraphraser
    paraphraser = FullGeminiParaphraser(
        synonym_file='sinonim.json',
        chunk_size=3000,  # Smaller chunks for testing
        max_retries=3
    )
    
    # Check if test file exists
    test_file = "documents/SKRIPSI  FAHRISAL FADLI-2.docx"
    
    if not os.path.exists(test_file):
        print(f"❌ Test file not found: {test_file}")
        print("📁 Available files in documents/:")
        if os.path.exists("documents"):
            for file in os.listdir("documents"):
                if file.endswith(('.docx', '.doc')):
                    print(f"   📄 {file}")
        return
    
    print(f"🎯 Testing with: {test_file}")
    
    # Process core chapters only
    result = paraphraser.process_core_chapters_only(test_file)
    
    if result:
        print("\n✅ Core Chapters Processing Completed!")
        print(f"📄 Output file: {result['output_path']}")
        print(f"📊 Paragraphs processed: {result['processed_paragraphs']}")
        print(f"📍 Range: {result['start_paragraph']} to {result['end_paragraph']}")
        
        # Show statistics
        stats = result['statistics']
        print("\n📈 Statistics:")
        for key, value in stats.items():
            print(f"   {key}: {value}")
    else:
        print("❌ Core chapters processing failed")

def test_chapter_detection():
    """Test chapter boundary detection"""
    print("\n🔍 Testing Chapter Detection")
    print("=" * 30)
    
    import docx
    
    test_file = "documents/SKRIPSI  FAHRISAL FADLI-2.docx"
    
    if not os.path.exists(test_file):
        print("❌ Test file not found")
        return
    
    paraphraser = FullGeminiParaphraser(synonym_file='sinonim.json')
    doc = docx.Document(test_file)
    
    start_index, end_index = paraphraser.detect_chapter_boundaries(doc)
    
    print(f"📍 Start index: {start_index}")
    print(f"📍 End index: {end_index}")
    
    if start_index is not None:
        print(f"🔸 Start text: {doc.paragraphs[start_index].text[:100]}...")
    
    if end_index is not None and end_index < len(doc.paragraphs):
        print(f"🔸 End text: {doc.paragraphs[end_index].text[:100]}...")
    
    # Show some context around boundaries
    if start_index is not None:
        print("\n📖 Context around start:")
        for i in range(max(0, start_index-2), min(len(doc.paragraphs), start_index+3)):
            marker = "👉" if i == start_index else "  "
            text = doc.paragraphs[i].text.strip()[:80]
            if text:
                print(f"{marker} [{i}] {text}...")

if __name__ == "__main__":
    # Test chapter detection first
    test_chapter_detection()
    
    # Then test core chapters processing
    test_core_chapters()