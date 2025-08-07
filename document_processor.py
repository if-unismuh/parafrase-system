# document_processor.py
"""
Document Processor untuk Ultimate Hybrid Paraphraser
Khusus untuk memproses dokumen skripsi dengan perlindungan judul dan bagian penting
Author: DevNoLife
Version: 1.0
"""

import os
import json
import re
from datetime import datetime
import docx
from docx.shared import RGBColor
from docx.enum.text import WD_COLOR_INDEX
import shutil
from ultimate_hybrid_paraphraser import UltimateHybridParaphraser

class DocumentProcessor:
    def __init__(self, synonym_file='sinonim.json', gemini_api_key=None):
        print("🚀 Initializing Document Processor for Skripsi...")
        
        # Initialize paraphraser
        self.paraphraser = UltimateHybridParaphraser(
            synonym_file=synonym_file,
            gemini_api_key=gemini_api_key,
            mode='smart'
        )
        
        # Patterns for protected content (DO NOT PARAPHRASE)
        self.protected_patterns = [
            r'^(BAB\s+[IVXLC]+)', # Chapter headers
            r'^(JUDUL|TITLE)', # Titles
            r'^(\d+\.\d+)', # Section numbers  
            r'^(DAFTAR\s+ISI|DAFTAR\s+PUSTAKA|BIBLIOGRAPHY)', # Table of contents, bibliography
            r'^(ABSTRAK|ABSTRACT)', # Abstract headers
            r'^(KATA\s+PENGANTAR|FOREWORD)', # Preface
            r'^(HALAMAN\s+PENGESAHAN)', # Approval page
            r'^\s*$', # Empty lines
            r'^(Gambar\s+\d+)', # Figure captions
            r'^(Tabel\s+\d+)', # Table captions
            r'^\s*\d+\s*$', # Page numbers
            r'https?://', # URLs
            r'www\.', # Web addresses
        ]
        
        # Academic section headers to protect
        self.academic_headers = [
            'PENDAHULUAN', 'TINJAUAN PUSTAKA', 'METODOLOGI', 'PEMBAHASAN',
            'KESIMPULAN', 'SARAN', 'HASIL', 'DISKUSI', 'LANDASAN TEORI'
        ]
        
        print(f"✅ Document Processor ready!")
        print(f"✅ Protected patterns: {len(self.protected_patterns)}")
        print(f"✅ Academic headers: {len(self.academic_headers)}")
    
    def is_protected_content(self, text):
        """Check if content should be protected from paraphrasing"""
        if not text or not text.strip():
            return True
            
        text_clean = text.strip()
        
        # Check protected patterns
        for pattern in self.protected_patterns:
            if re.match(pattern, text_clean, re.IGNORECASE):
                return True
        
        # Check academic headers
        for header in self.academic_headers:
            if header.lower() in text_clean.lower() and len(text_clean.split()) <= 5:
                return True
        
        # Protect very short text (likely headers/titles)
        if len(text_clean.split()) <= 3:
            return True
            
        # Protect text that's all uppercase (likely headers)
        if text_clean.isupper() and len(text_clean.split()) <= 8:
            return True
            
        return False
    
    def is_suitable_for_paraphrasing(self, text):
        """Enhanced suitability check for academic documents"""
        if self.is_protected_content(text):
            return False
        
        # Minimum length check
        word_count = len(text.strip().split())
        if word_count < 15:
            return False
        
        # Skip references and citations heavy paragraphs
        citation_ratio = len(re.findall(r'\(\d{4}\)', text)) / word_count if word_count > 0 else 0
        if citation_ratio > 0.3:  # More than 30% citations
            return False
        
        # Skip figure/table references
        if re.search(r'\b(gambar|tabel|figure|table)\s+\d+', text.lower()):
            return False
        
        return True
    
    def process_document(self, file_path, mode='smart', aggressiveness=0.6, create_backup=True):
        """Process single document with protection for important content"""
        try:
            print(f"\n📄 Processing: {os.path.basename(file_path)}")
            
            if not os.path.exists(file_path):
                print(f"❌ File not found: {file_path}")
                return None
            
            # Create backup if requested
            if create_backup:
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                backup_name = f"{os.path.splitext(file_path)[0]}_backup_{timestamp}.docx"
                shutil.copy2(file_path, backup_name)
                print(f"✅ Backup created: {backup_name}")
            
            # Load document
            doc = docx.Document(file_path)
            
            # Switch to requested mode
            if mode != self.paraphraser.mode:
                self.paraphraser.switch_mode(mode)
            
            stats = {
                'filename': os.path.basename(file_path),
                'total_paragraphs': len(doc.paragraphs),
                'processed_paragraphs': 0,
                'protected_paragraphs': 0,
                'skipped_paragraphs': 0,
                'changes_made': 0,
                'processing_mode': mode
            }
            
            print(f"📊 Document has {stats['total_paragraphs']} paragraphs")
            print(f"🔒 Using protection mode for academic content")
            
            # Process each paragraph
            for i, paragraph in enumerate(doc.paragraphs):
                para_text = paragraph.text.strip()
                
                if not para_text:
                    continue
                
                # Check if content is protected
                if self.is_protected_content(para_text):
                    stats['protected_paragraphs'] += 1
                    print(f"  🔒 Para {i+1}: PROTECTED - '{para_text[:50]}...'")
                    continue
                
                # Check if suitable for paraphrasing
                if not self.is_suitable_for_paraphrasing(para_text):
                    stats['skipped_paragraphs'] += 1
                    print(f"  ⏭️  Para {i+1}: SKIPPED - Not suitable")
                    continue
                
                # Process with paraphraser
                result = self.paraphraser.process_paragraph_ultimate(
                    para_text, 
                    paragraph_index=i, 
                    total_paragraphs=len(doc.paragraphs),
                    aggressiveness=aggressiveness
                )
                
                # Apply changes if improvement is significant
                improvement_threshold = 20  # 20% minimum improvement
                if result['plagiarism_reduction'] >= improvement_threshold:
                    # Replace paragraph content
                    paragraph.clear()
                    paragraph.add_run(result['paraphrase'])
                    
                    # Color coding based on method
                    if 'ai' in result['method'].lower() or 'gemini' in result['method'].lower():
                        # Blue for AI processing
                        for run in paragraph.runs:
                            run.font.highlight_color = WD_COLOR_INDEX.TURQUOISE
                    else:
                        # Yellow for local processing
                        for run in paragraph.runs:
                            run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                    
                    stats['processed_paragraphs'] += 1
                    stats['changes_made'] += result.get('changes_made', 1)
                    
                    print(f"  ✅ Para {i+1}: {result['plagiarism_reduction']:.1f}% reduction ({result['method']})")
                else:
                    stats['skipped_paragraphs'] += 1
                    print(f"  ⏭️  Para {i+1}: Insufficient improvement ({result['plagiarism_reduction']:.1f}%)")
            
            # Save processed document
            doc.save(file_path)
            
            # Print summary
            print(f"\n📊 PROCESSING SUMMARY:")
            print(f"   • Total paragraphs: {stats['total_paragraphs']}")
            print(f"   • Protected (not changed): {stats['protected_paragraphs']}")
            print(f"   • Processed successfully: {stats['processed_paragraphs']}")
            print(f"   • Skipped: {stats['skipped_paragraphs']}")
            print(f"   • Changes made: {stats['changes_made']}")
            print(f"   • Mode used: {stats['processing_mode']}")
            print(f"  ✅ Saved: {file_path}")
            
            return stats
            
        except Exception as e:
            print(f"❌ Error processing document: {e}")
            return None
    
    def process_documents_folder(self, documents_folder, mode='smart', aggressiveness=0.6, create_backup=True):
        """Process all documents in a folder"""
        print("=" * 80)
        print("📁 DOCUMENT FOLDER PROCESSOR")  
        print("🔒 Academic Document Protection Enabled")
        print("=" * 80)
        
        if not os.path.exists(documents_folder):
            print(f"❌ Documents folder not found: {documents_folder}")
            return
        
        # Find all Word documents
        docx_files = [f for f in os.listdir(documents_folder) 
                     if f.endswith('.docx') and not f.startswith('~')]
        
        if not docx_files:
            print(f"❌ No Word documents found in: {documents_folder}")
            return
        
        print(f"📄 Found {len(docx_files)} documents:")
        for doc in docx_files:
            print(f"   • {doc}")
        
        print(f"\n⚙️  Processing mode: {mode.upper()}")
        print(f"⚙️  Aggressiveness: {aggressiveness}")
        print(f"💾 Create backup: {create_backup}")
        
        # Process each document
        total_stats = {
            'documents_processed': 0,
            'total_paragraphs': 0,
            'processed_paragraphs': 0,
            'protected_paragraphs': 0,
            'skipped_paragraphs': 0,
            'total_changes': 0
        }
        
        for i, filename in enumerate(docx_files, 1):
            file_path = os.path.join(documents_folder, filename)
            
            print(f"\n{'='*50}")
            print(f"📄 Document {i}/{len(docx_files)}: {filename}")
            print(f"{'='*50}")
            
            doc_stats = self.process_document(file_path, mode, aggressiveness, create_backup)
            
            if doc_stats:
                total_stats['documents_processed'] += 1
                total_stats['total_paragraphs'] += doc_stats['total_paragraphs']
                total_stats['processed_paragraphs'] += doc_stats['processed_paragraphs']
                total_stats['protected_paragraphs'] += doc_stats['protected_paragraphs']
                total_stats['skipped_paragraphs'] += doc_stats['skipped_paragraphs']
                total_stats['total_changes'] += doc_stats['changes_made']
        
        # Generate final report
        self.generate_processing_report(documents_folder, total_stats, mode)
        
        print(f"\n" + "="*80)
        print("🎉 DOCUMENT PROCESSING COMPLETED!")
        print("="*80)
        print(f"📄 Documents processed: {total_stats['documents_processed']}")
        print(f"📝 Total paragraphs: {total_stats['total_paragraphs']}")
        print(f"✅ Processed: {total_stats['processed_paragraphs']}")
        print(f"🔒 Protected: {total_stats['protected_paragraphs']} (judul, header, dll)")
        print(f"⏭️  Skipped: {total_stats['skipped_paragraphs']}")
        print(f"🔄 Total changes: {total_stats['total_changes']}")
        
        # Show paraphraser statistics
        paraphraser_stats = self.paraphraser.get_processing_statistics()
        print(f"\n💰 COST SUMMARY:")
        print(f"💻 Local calls: {paraphraser_stats['cost_tracker']['local_calls']}")
        print(f"🤖 AI calls: {paraphraser_stats['cost_tracker']['ai_calls']}")
        print(f"💾 Cache hits: {paraphraser_stats['cost_tracker']['cache_hits']}")
        print(f"💵 Estimated cost: ${paraphraser_stats['cost_tracker']['estimated_cost_usd']:.4f}")
        
        print("="*80)
    
    def generate_processing_report(self, documents_folder, stats, mode):
        """Generate detailed processing report"""
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        report_file = os.path.join(documents_folder, f"document_processing_report_{timestamp}.json")
        
        try:
            paraphraser_stats = self.paraphraser.get_processing_statistics()
            
            report_data = {
                'timestamp': timestamp,
                'documents_folder': documents_folder,
                'processing_mode': mode,
                'document_stats': stats,
                'paraphraser_stats': paraphraser_stats,
                'protection_enabled': True,
                'protected_patterns_count': len(self.protected_patterns),
                'academic_headers_count': len(self.academic_headers)
            }
            
            with open(report_file, 'w', encoding='utf-8') as f:
                json.dump(report_data, f, ensure_ascii=False, indent=2)
            
            print(f"📋 Processing report saved: {report_file}")
            
        except Exception as e:
            print(f"⚠️  Could not save report: {e}")


def main():
    """Main function untuk memproses dokumen skripsi"""
    print("🎓 DOCUMENT PROCESSOR FOR SKRIPSI")
    print("🔒 Protected Academic Content Processing")
    print("=" * 80)
    
    # Configuration
    DOCUMENTS_FOLDER = 'documents'  # Folder dengan dokumen skripsi
    PROCESSING_MODE = 'smart'       # smart, balanced, aggressive, turnitin_safe
    AGGRESSIVENESS = 0.6           # Level aggressiveness (0.1-0.9)
    CREATE_BACKUP = True           # Selalu backup original
    
    # Cek Gemini API Key (opsional)
    GEMINI_API_KEY = os.environ.get("GEMINI_API_KEY")
    if not GEMINI_API_KEY:
        print("⚠️  GEMINI_API_KEY not found in environment variables")
        print("💡 Running in LOCAL-ONLY mode (still very effective!)")
        GEMINI_API_KEY = None
    
    # Initialize processor
    processor = DocumentProcessor(
        synonym_file='sinonim.json',
        gemini_api_key=GEMINI_API_KEY
    )
    
    # Process documents folder
    processor.process_documents_folder(
        documents_folder=DOCUMENTS_FOLDER,
        mode=PROCESSING_MODE,
        aggressiveness=AGGRESSIVENESS,
        create_backup=CREATE_BACKUP
    )
    
    print("\n🎉 Document processing completed!")
    print("💡 Color coding in documents:")
    print("   🟡 Yellow highlight = Local paraphrasing")
    print("   🔵 Blue highlight = AI paraphrasing") 
    print("   No highlight = Protected content (unchanged)")


if __name__ == "__main__":
    main()