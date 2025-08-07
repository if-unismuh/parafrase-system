# plagiarism_detector.py
"""
Plagiarism Detector - Online Source Checking System
Deteksi plagiarisme dengan pengecekan ke sumber online dan database publik
Author: DevNoLife
Version: 1.0
"""

import os
import json
import re
import time
import requests
from datetime import datetime
import hashlib
from urllib.parse import quote_plus
from bs4 import BeautifulSoup
import difflib
from collections import defaultdict

class PlagiarismDetector:
    def __init__(self):
        print("🔍 Initializing Plagiarism Detector...")
        
        # Search engines and APIs
        self.search_engines = {
            'google': {
                'url': 'https://www.googleapis.com/customsearch/v1',
                'enabled': False,  # Requires API key
                'free_searches': 0
            },
            'bing': {
                'url': 'https://api.bing.microsoft.com/v7.0/search',
                'enabled': False,  # Requires API key
                'free_searches': 0
            },
            'duckduckgo': {
                'url': 'https://duckduckgo.com/html/',
                'enabled': True,  # Free to use
                'free_searches': 100
            }
        }
        
        # Academic databases to check
        self.academic_sources = [
            'scholar.google.com',
            'researchgate.net',
            'academia.edu',
            'arxiv.org',
            'pubmed.ncbi.nlm.nih.gov',
            'jstor.org',
            'doaj.org'
        ]
        
        # Common Indonesian academic repositories
        self.indonesian_repos = [
            'repository.usu.ac.id',
            'digilib.unila.ac.id', 
            'repository.uin-suka.ac.id',
            'repository.unair.ac.id',
            'repository.ui.ac.id',
            'repository.ipb.ac.id',
            'repository.ugm.ac.id',
            'eprints.undip.ac.id'
        ]
        
        # Detection settings
        self.settings = {
            'min_text_length': 15,      # Minimum words to check
            'similarity_threshold': 70,  # 70% similarity = potential plagiarism
            'chunk_size': 50,           # Words per search chunk
            'max_chunks_per_paragraph': 3,  # Limit API calls
            'search_delay': 2,          # Delay between searches (seconds)
            'max_results_per_search': 5, # Limit results to process
            'timeout': 10               # Request timeout
        }
        
        # Results storage
        self.detection_results = []
        self.search_cache = {}
        
        print("✅ Plagiarism Detector ready!")
        print(f"✅ Search engines: {len(self.search_engines)} configured")
        print(f"✅ Academic sources: {len(self.academic_sources)} databases")
        print(f"✅ Indonesian repos: {len(self.indonesian_repos)} repositories")
    
    def create_search_query(self, text):
        """Create optimized search query from text"""
        # Clean text
        text = re.sub(r'[^\w\s]', ' ', text)
        words = text.split()
        
        # Remove common Indonesian stopwords
        stopwords = {
            'yang', 'dan', 'di', 'ke', 'dari', 'dalam', 'untuk', 'pada',
            'dengan', 'adalah', 'akan', 'atau', 'juga', 'telah', 'dapat',
            'tidak', 'ada', 'ini', 'itu', 'saya', 'kami', 'kita', 'mereka'
        }
        
        # Keep important words
        important_words = [w for w in words if w.lower() not in stopwords and len(w) > 3]
        
        # Create query with most distinctive words
        if len(important_words) >= 6:
            # Use first and last few words for better matching
            query_words = important_words[:3] + important_words[-3:]
        else:
            query_words = important_words
        
        return ' '.join(query_words[:8])  # Max 8 words for effective search
    
    def search_duckduckgo(self, query):
        """Search using DuckDuckGo (free)"""
        try:
            # DuckDuckGo search with academic site restriction
            params = {
                'q': f'"{query}" site:(.edu OR .ac.id OR scholar.google OR researchgate OR academia)',
                'kl': 'wt-wt'  # No region restriction
            }
            
            headers = {
                'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'
            }
            
            response = requests.get(
                'https://duckduckgo.com/html/',
                params=params,
                headers=headers,
                timeout=self.settings['timeout']
            )
            
            if response.status_code != 200:
                return []
            
            soup = BeautifulSoup(response.text, 'html.parser')
            results = []
            
            # Parse DuckDuckGo results
            for result in soup.find_all('a', class_='result__a')[:self.settings['max_results_per_search']]:
                title = result.get_text(strip=True)
                url = result.get('href', '')
                
                if url and title:
                    results.append({
                        'title': title,
                        'url': url,
                        'source': 'duckduckgo'
                    })
            
            return results
            
        except Exception as e:
            print(f"⚠️  DuckDuckGo search error: {e}")
            return []
    
    def search_google_custom(self, query, api_key, search_engine_id):
        """Search using Google Custom Search API (requires API key)"""
        try:
            params = {
                'key': api_key,
                'cx': search_engine_id,
                'q': f'"{query}"',
                'num': self.settings['max_results_per_search']
            }
            
            response = requests.get(
                self.search_engines['google']['url'],
                params=params,
                timeout=self.settings['timeout']
            )
            
            if response.status_code != 200:
                return []
            
            data = response.json()
            results = []
            
            for item in data.get('items', []):
                results.append({
                    'title': item.get('title', ''),
                    'url': item.get('link', ''),
                    'snippet': item.get('snippet', ''),
                    'source': 'google_custom'
                })
            
            return results
            
        except Exception as e:
            print(f"⚠️  Google Custom Search error: {e}")
            return []
    
    def fetch_page_content(self, url):
        """Fetch and extract text content from URL"""
        try:
            headers = {
                'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'
            }
            
            response = requests.get(url, headers=headers, timeout=self.settings['timeout'])
            
            if response.status_code != 200:
                return ""
            
            soup = BeautifulSoup(response.text, 'html.parser')
            
            # Remove script and style elements
            for script in soup(["script", "style"]):
                script.decompose()
            
            # Extract text content
            text = soup.get_text()
            
            # Clean up whitespace
            lines = (line.strip() for line in text.splitlines())
            chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
            text = ' '.join(chunk for chunk in chunks if chunk)
            
            return text
            
        except Exception as e:
            print(f"⚠️  Error fetching {url}: {e}")
            return ""
    
    def calculate_text_similarity(self, text1, text2):
        """Calculate similarity between two texts"""
        # Use SequenceMatcher for accurate similarity
        similarity = difflib.SequenceMatcher(None, text1.lower(), text2.lower()).ratio()
        return similarity * 100
    
    def check_text_against_sources(self, text):
        """Check text against online sources"""
        if len(text.split()) < self.settings['min_text_length']:
            return {
                'similarity': 0,
                'sources': [],
                'status': 'too_short'
            }
        
        # Create cache key
        cache_key = hashlib.md5(text.encode('utf-8')).hexdigest()[:16]
        
        if cache_key in self.search_cache:
            return self.search_cache[cache_key]
        
        # Create search query
        query = self.create_search_query(text)
        if not query:
            return {
                'similarity': 0,
                'sources': [],
                'status': 'no_query'
            }
        
        print(f"  🔍 Searching: '{query[:50]}...'")
        
        # Search for similar content
        search_results = self.search_duckduckgo(query)
        
        if not search_results:
            result = {
                'similarity': 0,
                'sources': [],
                'status': 'no_results'
            }
            self.search_cache[cache_key] = result
            return result
        
        # Check similarity with found sources
        max_similarity = 0
        similar_sources = []
        
        for result in search_results[:3]:  # Check top 3 results
            url = result['url']
            title = result['title']
            
            print(f"    📄 Checking: {title[:50]}...")
            
            # Fetch page content
            content = self.fetch_page_content(url)
            
            if content:
                # Find best matching passage
                content_chunks = [content[i:i+len(text)*2] for i in range(0, len(content), len(text))]
                
                for chunk in content_chunks[:5]:  # Check first 5 chunks
                    similarity = self.calculate_text_similarity(text, chunk)
                    
                    if similarity > max_similarity:
                        max_similarity = similarity
                    
                    if similarity >= self.settings['similarity_threshold']:
                        similar_sources.append({
                            'url': url,
                            'title': title,
                            'similarity': round(similarity, 2),
                            'matched_text': chunk[:200] + '...' if len(chunk) > 200 else chunk
                        })
            
            # Add delay to be respectful
            time.sleep(self.settings['search_delay'])
        
        result = {
            'similarity': round(max_similarity, 2),
            'sources': similar_sources,
            'status': 'checked',
            'query_used': query
        }
        
        # Cache result
        self.search_cache[cache_key] = result
        
        return result
    
    def detect_plagiarism_in_text(self, text):
        """Detect plagiarism in a single text passage"""
        print(f"\n🔍 PLAGIARISM DETECTION")
        print(f"📝 Text: '{text[:100]}...'")
        
        # Split text into checkable chunks if too long
        words = text.split()
        
        if len(words) <= self.settings['chunk_size']:
            # Check entire text
            return self.check_text_against_sources(text)
        else:
            # Check in chunks
            chunk_size = self.settings['chunk_size']
            chunks = [' '.join(words[i:i+chunk_size]) for i in range(0, len(words), chunk_size//2)]
            
            max_similarity = 0
            all_sources = []
            
            for i, chunk in enumerate(chunks[:self.settings['max_chunks_per_paragraph']]):
                print(f"  📄 Checking chunk {i+1}/{min(len(chunks), self.settings['max_chunks_per_paragraph'])}")
                
                chunk_result = self.check_text_against_sources(chunk)
                
                if chunk_result['similarity'] > max_similarity:
                    max_similarity = chunk_result['similarity']
                
                all_sources.extend(chunk_result['sources'])
            
            return {
                'similarity': max_similarity,
                'sources': all_sources,
                'status': 'checked_chunks',
                'chunks_checked': min(len(chunks), self.settings['max_chunks_per_paragraph'])
            }
    
    def scan_document_for_plagiarism(self, file_path):
        """Scan entire document for plagiarism"""
        print("=" * 80)
        print("🔍 PLAGIARISM DETECTION SCAN")
        print("=" * 80)
        
        if not os.path.exists(file_path):
            print(f"❌ File not found: {file_path}")
            return None
        
        print(f"📄 Scanning: {os.path.basename(file_path)}")
        
        try:
            import docx
            doc = docx.Document(file_path)
            
            scan_results = {
                'filename': os.path.basename(file_path),
                'scan_timestamp': datetime.now().isoformat(),
                'total_paragraphs': len(doc.paragraphs),
                'paragraphs_scanned': 0,
                'high_risk_paragraphs': [],
                'medium_risk_paragraphs': [],
                'low_risk_paragraphs': [],
                'summary': {}
            }
            
            print(f"📊 Total paragraphs: {len(doc.paragraphs)}")
            print(f"🔍 Starting plagiarism scan...")
            
            for i, paragraph in enumerate(doc.paragraphs):
                para_text = paragraph.text.strip()
                
                if not para_text or len(para_text.split()) < self.settings['min_text_length']:
                    continue
                
                print(f"\n📄 Paragraph {i+1}/{len(doc.paragraphs)}")
                
                # Detect plagiarism
                detection = self.detect_plagiarism_in_text(para_text)
                
                similarity = detection['similarity']
                sources = detection['sources']
                
                para_result = {
                    'paragraph_index': i + 1,
                    'text_preview': para_text[:100] + '...' if len(para_text) > 100 else para_text,
                    'similarity': similarity,
                    'sources_found': len(sources),
                    'sources': sources,
                    'detection_status': detection['status']
                }
                
                # Categorize risk level
                if similarity >= 80:
                    print(f"  🔴 HIGH RISK: {similarity}% similarity - {len(sources)} sources found")
                    scan_results['high_risk_paragraphs'].append(para_result)
                elif similarity >= 50:
                    print(f"  🟡 MEDIUM RISK: {similarity}% similarity - {len(sources)} sources found")
                    scan_results['medium_risk_paragraphs'].append(para_result)
                elif similarity >= 20:
                    print(f"  🟢 LOW RISK: {similarity}% similarity - {len(sources)} sources found")
                    scan_results['low_risk_paragraphs'].append(para_result)
                else:
                    print(f"  ✅ CLEAN: {similarity}% similarity - No significant matches")
                
                scan_results['paragraphs_scanned'] += 1
            
            # Generate summary
            total_scanned = scan_results['paragraphs_scanned']
            high_risk = len(scan_results['high_risk_paragraphs'])
            medium_risk = len(scan_results['medium_risk_paragraphs'])
            low_risk = len(scan_results['low_risk_paragraphs'])
            clean = total_scanned - high_risk - medium_risk - low_risk
            
            scan_results['summary'] = {
                'total_scanned': total_scanned,
                'high_risk_count': high_risk,
                'medium_risk_count': medium_risk,
                'low_risk_count': low_risk,
                'clean_count': clean,
                'high_risk_percentage': round((high_risk / total_scanned) * 100, 1) if total_scanned > 0 else 0,
                'needs_paraphrasing': high_risk + medium_risk
            }
            
            # Save detailed report
            self.save_plagiarism_report(scan_results)
            
            # Print summary
            self.print_scan_summary(scan_results)
            
            return scan_results
            
        except Exception as e:
            print(f"❌ Error scanning document: {e}")
            return None
    
    def save_plagiarism_report(self, scan_results):
        """Save detailed plagiarism report"""
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        report_file = f"plagiarism_report_{timestamp}.json"
        
        try:
            with open(report_file, 'w', encoding='utf-8') as f:
                json.dump(scan_results, f, ensure_ascii=False, indent=2)
            
            print(f"\n📋 Detailed report saved: {report_file}")
            
        except Exception as e:
            print(f"⚠️  Could not save report: {e}")
    
    def print_scan_summary(self, results):
        """Print plagiarism scan summary"""
        summary = results['summary']
        
        print("\n" + "=" * 80)
        print("📊 PLAGIARISM SCAN SUMMARY")
        print("=" * 80)
        
        print(f"📄 Document: {results['filename']}")
        print(f"📊 Paragraphs scanned: {summary['total_scanned']}")
        print(f"🔴 HIGH RISK (≥80%): {summary['high_risk_count']} paragraphs")
        print(f"🟡 MEDIUM RISK (50-79%): {summary['medium_risk_count']} paragraphs")
        print(f"🟢 LOW RISK (20-49%): {summary['low_risk_count']} paragraphs")
        print(f"✅ CLEAN (<20%): {summary['clean_count']} paragraphs")
        
        print(f"\n🚨 NEEDS PARAPHRASING: {summary['needs_paraphrasing']} paragraphs")
        print(f"📈 High risk percentage: {summary['high_risk_percentage']}%")
        
        if summary['high_risk_count'] > 0:
            print(f"\n🔴 HIGH RISK PARAGRAPHS:")
            for para in results['high_risk_paragraphs'][:5]:  # Show first 5
                print(f"   Para {para['paragraph_index']}: {para['similarity']}% - '{para['text_preview']}'")
        
        if summary['medium_risk_count'] > 0:
            print(f"\n🟡 MEDIUM RISK PARAGRAPHS:")
            for para in results['medium_risk_paragraphs'][:3]:  # Show first 3
                print(f"   Para {para['paragraph_index']}: {para['similarity']}% - '{para['text_preview']}'")
        
        print("\n💡 RECOMMENDATION:")
        if summary['needs_paraphrasing'] == 0:
            print("   ✅ Document appears to be original. No paraphrasing needed.")
        elif summary['high_risk_percentage'] > 30:
            print("   🚨 High plagiarism risk detected. Immediate paraphrasing required.")
        else:
            print(f"   ⚠️  {summary['needs_paraphrasing']} paragraphs need paraphrasing.")
        
        print("=" * 80)


def find_docx_files():
    """Find all .docx files in the project directory with priority for documents/ folder"""
    docx_files = []
    
    # Search in current directory and subdirectories
    for root, dirs, files in os.walk('.'):
        # Skip backup and temporary directories
        if any(skip in root for skip in ['backup', '__pycache__', '.git', 'venv']):
            continue
            
        for file in files:
            if file.endswith('.docx') and not file.startswith('~'):
                full_path = os.path.join(root, file)
                
                # Prioritize files from documents/ folder
                if root == './documents':
                    docx_files.insert(0, full_path)  # Add to beginning
                else:
                    docx_files.append(full_path)  # Add to end
    
    return docx_files

def main():
    """Main function untuk plagiarism detection"""
    print("🔍 PLAGIARISM DETECTOR - Online Source Checker")
    print("=" * 80)
    
    # Initialize detector
    detector = PlagiarismDetector()
    
    # Find available documents
    print("🔍 Searching for documents...")
    docx_files = find_docx_files()
    
    if not docx_files:
        print("❌ No .docx files found in the project directory")
        print("💡 Please add your document to the project folder")
        return
    
    # Use the first document found, or let user choose
    document_path = docx_files[0]
    
    if len(docx_files) > 1:
        print(f"📄 Found {len(docx_files)} documents. Using: {os.path.basename(document_path)}")
    else:
        print(f"📄 Found document: {os.path.basename(document_path)}")
    
    if os.path.exists(document_path):
        scan_results = detector.scan_document_for_plagiarism(document_path)
        
        if scan_results:
            print(f"\n🎉 Plagiarism scan completed!")
            print(f"📋 Check the detailed report for complete analysis")
            
            # Recommendation for paraphrasing
            needs_paraphrasing = scan_results['summary']['needs_paraphrasing']
            if needs_paraphrasing > 0:
                print(f"\n💡 NEXT STEPS:")
                print(f"   1. Review {needs_paraphrasing} paragraphs marked as HIGH/MEDIUM risk")
                print(f"   2. Use document_processor.py to paraphrase these sections")
                print(f"   3. Re-scan after paraphrasing to verify improvement")
        else:
            print("❌ Scan failed. Please check the document and try again.")
    else:
        print(f"❌ Document not found: {document_path}")
        print(f"💡 Please place your document in the 'documents/' folder")


if __name__ == "__main__":
    main()